#!/usr/bin/env python3
"""
Test script to verify that DOCX vector search is now working.
"""

import requests
import tempfile
import os
from pathlib import Path


def create_test_docx():
    """Create a test DOCX file for testing."""
    try:
        from docx import Document

        # Create a test .docx file
        doc = Document()
        doc.add_heading("Leadership Skills in AI", 0)
        doc.add_paragraph(
            "Effective leadership skills are crucial in AI development and implementation."
        )
        doc.add_paragraph("Key leadership qualities include:")

        # Add a list
        doc.add_paragraph("• Strategic thinking and vision", style="List Bullet")
        doc.add_paragraph(
            "• Technical expertise combined with business acumen", style="List Bullet"
        )
        doc.add_paragraph("• Ethical decision-making capabilities", style="List Bullet")
        doc.add_paragraph("• Strong communication skills", style="List Bullet")

        doc.add_heading("Risk Management in AI Leadership", level=1)
        doc.add_paragraph(
            "AI leaders must understand and mitigate various risks including:"
        )
        doc.add_paragraph("• Model bias and fairness issues")
        doc.add_paragraph("• Data privacy and security concerns")
        doc.add_paragraph("• Regulatory compliance requirements")

        # Save to a temporary file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            doc.save(temp_file.name)
            return temp_file.name

    except ImportError:
        print("python-docx not available, skipping DOCX creation test")
        return None


def test_docx_vector_search():
    """Test DOCX vector search through the API."""

    # Create test DOCX file
    docx_path = create_test_docx()
    if not docx_path:
        print("❌ Could not create test DOCX file")
        return False

    try:
        print("📄 Created test DOCX file")

        # Test the vector search endpoint
        api_url = "http://localhost:8000/api/v1/chat/document"

        # Read the DOCX file
        with open(docx_path, "rb") as f:
            files = {
                "files": (
                    "test_leadership.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }

            params = {
                "question": "What does this document say about leadership skills?",
                "use_default_models": "true",
                "session_id": "test_session",
                "is_follow_up": "false",
                "search_mode": "vector",  # This should now work with DOCX!
            }

            print("🔍 Testing vector search with DOCX file...")
            response = requests.post(api_url, files=files, params=params)

            print(f"Response Status: {response.status_code}")

            if response.status_code == 200:
                result = response.json()
                print("✅ Vector search with DOCX SUCCESSFUL!")
                print(f"Answer: {result.get('answer', 'No answer')[:200]}...")
                print(f"Sources found: {len(result.get('sources', []))}")
                return True
            else:
                print(f"❌ Vector search FAILED with status {response.status_code}")
                print(f"Error: {response.text}")
                return False

    except Exception as e:
        print(f"❌ Error during test: {e}")
        return False
    finally:
        # Clean up
        if docx_path and os.path.exists(docx_path):
            os.unlink(docx_path)


if __name__ == "__main__":
    print("Testing DOCX Vector Search Fix...")
    success = test_docx_vector_search()

    if success:
        print("\n🎉 DOCX Vector Search test PASSED!")
        print("The fix is working correctly.")
    else:
        print("\n💥 DOCX Vector Search test FAILED!")
        print("The issue may still need to be addressed.")
