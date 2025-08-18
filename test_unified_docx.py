#!/usr/bin/env python3
"""
Test script to verify that the unified document processing fixes DOCX extraction.
"""

import sys
import os
from pathlib import Path

# Add the backend app to the Python path
backend_path = Path(__file__).parent / "backend"
sys.path.insert(0, str(backend_path))

try:
    from app.services.document_utils import (
        extract_text_from_file_unified,
        extract_documents_from_file_unified,
    )
    from docx import Document
    import tempfile

    def test_unified_docx_processing():
        """Test unified DOCX processing functions."""
        print("Testing unified DOCX processing functions...")

        # Create a test .docx file
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test paragraph in a Word document.")
        doc.add_paragraph("This is another paragraph with some text.")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1,1"
        table.cell(0, 1).text = "Cell 1,2"
        table.cell(1, 0).text = "Cell 2,1"
        table.cell(1, 1).text = "Cell 2,2"

        # Save to a temporary file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            doc.save(temp_file.name)
            temp_file_path = temp_file.name

        try:
            # Read the file content
            with open(temp_file_path, "rb") as f:
                file_content = f.read()

            print("\n=== Testing extract_text_from_file_unified ===")

            # Test unified text extraction
            extracted_text = extract_text_from_file_unified(
                file_content, "test_document.docx"
            )

            print("Extracted text using unified function:")
            print("-" * 50)
            print(extracted_text)
            print("-" * 50)

            # Check if the extraction was successful
            expected_texts = [
                "Test Document",
                "This is a test paragraph",
                "This is another paragraph",
                "Cell 1,1",
                "Cell 2,2",
            ]

            success = True
            for expected in expected_texts:
                if expected not in extracted_text:
                    print(
                        f"ERROR: Expected text '{expected}' not found in extracted text!"
                    )
                    success = False
                else:
                    print(f"✓ Found expected text: '{expected}'")

            if success:
                print("\n✅ extract_text_from_file_unified test PASSED!")
            else:
                print("\n❌ extract_text_from_file_unified test FAILED!")

            print("\n=== Testing extract_documents_from_file_unified ===")

            # Test unified document extraction (for LangChain compatibility)
            documents = extract_documents_from_file_unified(
                file_content, "test_document.docx"
            )

            print(f"Number of documents returned: {len(documents)}")

            if documents:
                doc_text = documents[0].page_content
                print("Document content:")
                print("-" * 50)
                print(doc_text)
                print("-" * 50)

                # Check if the extraction was successful
                doc_success = True
                for expected in expected_texts:
                    if expected not in doc_text:
                        print(
                            f"ERROR: Expected text '{expected}' not found in document content!"
                        )
                        doc_success = False
                    else:
                        print(f"✓ Found expected text: '{expected}'")

                if doc_success:
                    print("\n✅ extract_documents_from_file_unified test PASSED!")
                else:
                    print("\n❌ extract_documents_from_file_unified test FAILED!")
            else:
                print(
                    "\n❌ extract_documents_from_file_unified test FAILED - No documents returned!"
                )

            return success and (doc_success if documents else False)

        finally:
            # Clean up the temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)

    if __name__ == "__main__":
        result = test_unified_docx_processing()
        if result:
            print("\n🎉 All unified DOCX processing tests PASSED!")
            print(
                "The unified document processing should now correctly handle DOCX files."
            )
        else:
            print("\n💥 Some tests FAILED!")
            print("Please check the implementation.")

except ImportError as e:
    print(f"Import error: {e}")
    print(
        "Make sure you're running this from the correct directory and all dependencies are installed."
    )
except Exception as e:
    print(f"Error: {e}")
    import traceback

    traceback.print_exc()
