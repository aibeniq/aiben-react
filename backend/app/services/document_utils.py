"""
Unified document processing utilities.
Centralizes all document text extraction logic for consistent handling across the application.
"""

import tempfile
import os
from pathlib import Path
from typing import List, Union, Any, Tuple
from langchain_core.documents import Document
from langchain_community.document_loaders import TextLoader


def extract_text_from_docx_bytes(file_content: bytes, filename: str) -> str:
    """
    Extract text from DOCX file content (bytes).

    Args:
        file_content: Raw bytes of the DOCX file
        filename: Name of the file (for metadata/error reporting)

    Returns:
        Extracted text content as string
    """
    try:
        from docx import Document as DocxDocument

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        # File is now closed and ready to be read by docx

        try:
            doc = DocxDocument(temp_file_path)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n\n".join(text_parts)

        finally:
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)

    except Exception as e:
        print(f"Error extracting text from DOCX {filename}: {e}")
        return f"Failed to extract text from {filename}: {str(e)}"


def extract_text_from_docx_path(file_path: str, filename: str) -> str:
    """
    Extract text from DOCX file path.

    Args:
        file_path: Path to the DOCX file
        filename: Name of the file (for metadata/error reporting)

    Returns:
        Extracted text content as string
    """
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        text_parts = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n\n".join(text_parts)

    except Exception as e:
        print(f"Error extracting text from DOCX {filename}: {e}")
        return f"Failed to extract text from {filename}: {str(e)}"


def extract_text_from_docx_langchain(file_path: str, filename: str) -> List[Document]:
    """
    Extract text from DOCX file and return as LangChain Document objects.
    Compatible with existing knowledge base processing.

    Args:
        file_path: Path to the DOCX file
        filename: Name of the file (for metadata)

    Returns:
        List of LangChain Document objects
    """
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        text_parts = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        combined_text = "\n\n".join(text_parts)

        # Create metadata
        metadata = {
            "source": filename,
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }

        # Try to get document properties
        try:
            core_properties = doc.core_properties
            if core_properties.title:
                metadata["title"] = core_properties.title
            if core_properties.author:
                metadata["author"] = core_properties.author
            if core_properties.created:
                metadata["created"] = str(core_properties.created)
            if core_properties.modified:
                metadata["modified"] = str(core_properties.modified)
        except Exception as e:
            print(f"Could not extract document properties: {str(e)}")

        return [Document(page_content=combined_text, metadata=metadata)]

    except Exception as e:
        print(f"Error extracting text from DOCX {filename}: {e}")
        error_doc = Document(
            page_content=f"Failed to extract text from {filename}: {str(e)}",
            metadata={"source": filename, "error": True},
        )
        return [error_doc]


def extract_text_from_csv_bytes(file_content: bytes, filename: str) -> str:
    """
    Extract text from CSV file content (bytes).

    Args:
        file_content: Raw bytes of the CSV file
        filename: Name of the file (for metadata/error reporting)

    Returns:
        Extracted text content as string
    """
    try:
        import pandas as pd
        from io import BytesIO

        # Try to read the CSV with different encodings
        try:
            # Try UTF-8 first
            csv_string = file_content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                # Fall back to latin-1
                csv_string = file_content.decode("latin-1")
            except UnicodeDecodeError:
                # Fall back to cp1252 (Windows encoding)
                csv_string = file_content.decode("cp1252")

        # Read CSV into DataFrame
        csv_io = BytesIO(csv_string.encode("utf-8"))
        df = pd.read_csv(csv_io)

        # Convert DataFrame to readable text format
        text_parts = []

        # Add column headers
        text_parts.append("Column Headers:")
        text_parts.append(" | ".join(df.columns.astype(str)))
        text_parts.append("")  # Empty line for separation

        # Add data rows (limit to reasonable number for text extraction)
        max_rows = min(1000, len(df))  # Limit to 1000 rows for performance
        text_parts.append(f"Data ({max_rows} of {len(df)} rows):")

        for index, row in df.head(max_rows).iterrows():
            row_text = " | ".join(row.astype(str).fillna(""))
            text_parts.append(row_text)

        # Add summary information
        text_parts.append("")
        text_parts.append(f"CSV Summary: {len(df)} rows, {len(df.columns)} columns")

        return "\n".join(text_parts)

    except Exception as e:
        print(f"Error extracting text from CSV {filename}: {e}")
        return f"Failed to extract text from CSV {filename}: {str(e)}"


def extract_text_from_xlsx_bytes(file_content: bytes, filename: str) -> str:
    """
    Extract text from XLSX file content (bytes).

    Args:
        file_content: Raw bytes of the XLSX file
        filename: Name of the file (for metadata/error reporting)

    Returns:
        Extracted text content as string
    """
    try:
        import pandas as pd
        from io import BytesIO

        # Read XLSX into DataFrame
        xlsx_io = BytesIO(file_content)

        # Read all sheets
        excel_file = pd.ExcelFile(xlsx_io)
        text_parts = []

        text_parts.append(f"Excel file with {len(excel_file.sheet_names)} sheet(s)")
        text_parts.append("")

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(xlsx_io, sheet_name=sheet_name)

            text_parts.append(f"=== Sheet: {sheet_name} ===")

            # Add column headers
            text_parts.append("Column Headers:")
            text_parts.append(" | ".join(df.columns.astype(str)))
            text_parts.append("")

            # Add data rows (limit to reasonable number for text extraction)
            max_rows = min(500, len(df))  # Limit to 500 rows per sheet for performance
            text_parts.append(f"Data ({max_rows} of {len(df)} rows):")

            for index, row in df.head(max_rows).iterrows():
                row_text = " | ".join(row.astype(str).fillna(""))
                text_parts.append(row_text)

            # Add summary for this sheet
            text_parts.append("")
            text_parts.append(
                f"Sheet Summary: {len(df)} rows, {len(df.columns)} columns"
            )
            text_parts.append("")

        return "\n".join(text_parts)

    except Exception as e:
        print(f"Error extracting text from XLSX {filename}: {e}")
        return f"Failed to extract text from XLSX {filename}: {str(e)}"


def extract_text_from_file_unified(
    file_content: bytes, filename: str, pdf_parsing_mode: str = None, current_user=None
) -> str:
    """
    Unified file text extraction function that handles multiple file types.
    This is the main entry point for document text extraction.

    Args:
        file_content: Raw bytes of the file
        filename: Name of the file
        pdf_parsing_mode: PDF parsing mode ('auto', 'enhanced', 'basic').
                         Priority order:
                         1. Explicit pdf_parsing_mode parameter
                         2. User's preference (current_user.pdf_parsing_preference)
                         3. Global setting (settings.PDF_PARSING_MODE)
        current_user: Current user object (optional)

    Returns:
        Extracted text content as string
    """
    try:
        # Determine file type from extension
        file_ext = Path(filename).suffix.lower()

        if file_ext == ".pdf":
            # Handle PDF files
            from app.services.pdf_utils import extract_text_from_pdf_bytes
            from app.core.config import settings

            # Priority order for mode selection
            if pdf_parsing_mode is not None:
                # Explicit parameter takes highest priority
                mode = pdf_parsing_mode
                print(
                    f"[DOCUMENT_UTILS] Using explicit pdf_parsing_mode parameter: {mode}"
                )
            elif current_user is not None:
                # User preference takes second priority
                mode = getattr(
                    current_user, "pdf_parsing_preference", settings.PDF_PARSING_MODE
                )
                print(
                    f"[DOCUMENT_UTILS] Using user {current_user.id} PDF parsing preference: {mode}"
                )
            else:
                # Global setting is fallback
                mode = settings.PDF_PARSING_MODE
                print(f"[DOCUMENT_UTILS] Using global PDF_PARSING_MODE setting: {mode}")

            return extract_text_from_pdf_bytes(
                file_content, filename, parsing_mode=mode
            )

        elif file_ext in [".docx", ".doc"]:
            # Handle Word documents using our unified DOCX function
            if file_ext == ".docx":
                return extract_text_from_docx_bytes(file_content, filename)
            else:
                # For .doc files, fall back to textloader approach
                with tempfile.NamedTemporaryFile(
                    delete=False, suffix=file_ext
                ) as temp_file:
                    temp_file.write(file_content)
                    temp_file_path = temp_file.name

                try:
                    loader = TextLoader(temp_file_path, encoding="utf-8")
                    documents = loader.load()
                    return "\n\n".join([doc.page_content for doc in documents])
                except UnicodeDecodeError:
                    try:
                        loader = TextLoader(temp_file_path, encoding="latin-1")
                        documents = loader.load()
                        return "\n\n".join([doc.page_content for doc in documents])
                    except Exception as e:
                        return f"Failed to extract text from {filename}: {str(e)}"
                finally:
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)

        elif file_ext in [".txt", ".md"]:
            # Handle text files
            try:
                return file_content.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    return file_content.decode("latin-1")
                except UnicodeDecodeError:
                    return f"Unable to extract text from {filename} - encoding issue"

        elif file_ext == ".csv":
            # Handle CSV files
            return extract_text_from_csv_bytes(file_content, filename)

        elif file_ext in [".xlsx", ".xls"]:
            # Handle Excel files
            return extract_text_from_xlsx_bytes(file_content, filename)

        else:
            # Try to decode as text for unknown file types
            try:
                return file_content.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    return file_content.decode("latin-1")
                except UnicodeDecodeError:
                    return f"Unable to extract text from {filename} - unsupported file format"

    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
        return f"Failed to extract text from {filename}: {str(e)}"


async def extract_documents_from_file_unified(
    file_content: bytes, filename: str, pdf_parsing_mode: str = None, current_user=None
) -> List[Document]:
    """
    Unified file extraction function that returns LangChain Document objects.
    Used by vector search and knowledge base processing.

    This function is async to support non-blocking PDF processing with PyMuPDF4LLM.

    Args:
        file_content: Raw bytes of the file
        filename: Name of the file
        pdf_parsing_mode: PDF parsing mode ('auto', 'enhanced', 'basic').
                         Priority order:
                         1. Explicit pdf_parsing_mode parameter
                         2. User's preference (current_user.pdf_parsing_preference)
                         3. Global setting (settings.PDF_PARSING_MODE)
        current_user: Current user object (optional)

    Returns:
        List of LangChain Document objects
    """
    try:
        # Determine file type from extension
        file_ext = Path(filename).suffix.lower()

        if file_ext == ".pdf":
            # Handle PDF files
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            try:
                from app.services.pdf_utils import load_pdf_with_pypdf
                from app.core.config import settings

                # Priority order for mode selection
                if pdf_parsing_mode is not None:
                    # Explicit parameter takes highest priority
                    mode = pdf_parsing_mode
                    print(
                        f"[DOCUMENT_UTILS] Using explicit pdf_parsing_mode parameter: {mode}"
                    )
                elif current_user is not None:
                    # User preference takes second priority
                    mode = getattr(
                        current_user,
                        "pdf_parsing_preference",
                        settings.PDF_PARSING_MODE,
                    )
                    print(
                        f"[DOCUMENT_UTILS] Using user {current_user.id} PDF parsing preference: {mode}"
                    )
                else:
                    # Global setting is fallback
                    mode = settings.PDF_PARSING_MODE
                    print(
                        f"[DOCUMENT_UTILS] Using global PDF_PARSING_MODE setting: {mode}"
                    )

                return await load_pdf_with_pypdf(
                    temp_file_path,
                    filename,
                    parsing_mode=mode,
                )
            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

        elif file_ext in [".docx", ".doc"]:
            # Handle Word documents
            if file_ext == ".docx":
                with tempfile.NamedTemporaryFile(
                    delete=False, suffix=".docx"
                ) as temp_file:
                    temp_file.write(file_content)
                    temp_file_path = temp_file.name

                try:
                    return extract_text_from_docx_langchain(temp_file_path, filename)
                finally:
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)
            else:
                # For .doc files, fall back to textloader
                with tempfile.NamedTemporaryFile(
                    delete=False, suffix=file_ext
                ) as temp_file:
                    temp_file.write(file_content)
                    temp_file_path = temp_file.name

                try:
                    loader = TextLoader(temp_file_path, encoding="utf-8")
                    documents = loader.load()
                    return documents
                except UnicodeDecodeError:
                    try:
                        loader = TextLoader(temp_file_path, encoding="latin-1")
                        documents = loader.load()
                        return documents
                    except Exception as e:
                        error_doc = Document(
                            page_content=f"Failed to extract text from {filename}: {str(e)}",
                            metadata={"source": filename, "error": True},
                        )
                        return [error_doc]
                finally:
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)

        elif file_ext in [".txt", ".md"]:
            # Handle text files
            try:
                text_content = file_content.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    text_content = file_content.decode("latin-1")
                except UnicodeDecodeError:
                    text_content = (
                        f"Unable to extract text from {filename} - encoding issue"
                    )

            return [
                Document(
                    page_content=text_content,
                    metadata={"source": filename, "content_type": "text/plain"},
                )
            ]

        elif file_ext == ".csv":
            # Handle CSV files
            text_content = extract_text_from_csv_bytes(file_content, filename)
            return [
                Document(
                    page_content=text_content,
                    metadata={"source": filename, "content_type": "text/csv"},
                )
            ]

        elif file_ext in [".xlsx", ".xls"]:
            # Handle Excel files
            text_content = extract_text_from_xlsx_bytes(file_content, filename)
            return [
                Document(
                    page_content=text_content,
                    metadata={
                        "source": filename,
                        "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    },
                )
            ]

        else:
            # Try to decode as text for unknown file types
            try:
                text_content = file_content.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    text_content = file_content.decode("latin-1")
                except UnicodeDecodeError:
                    text_content = f"Unable to extract text from {filename} - unsupported file format"

            return [
                Document(
                    page_content=text_content,
                    metadata={"source": filename, "content_type": "unknown"},
                )
            ]

    except Exception as e:
        print(f"Error extracting documents from {filename}: {e}")
        error_doc = Document(
            page_content=f"Failed to extract text from {filename}: {str(e)}",
            metadata={"source": filename, "error": True},
        )
        return [error_doc]


def extract_documents_and_images_from_file_unified(
    file_content: bytes, filename: str, current_user=None
) -> Tuple[List[Document], List[str]]:
    """
    Enhanced unified extraction that returns both text documents and images.

    Args:
        file_content: Raw bytes of the file
        filename: Name of the file
        current_user: Current user object (optional) for PDF parsing preference

    Returns:
        Tuple of (text_documents, image_list_base64)
    """
    import base64
    import logging

    logger = logging.getLogger(__name__)

    # Get existing text extraction
    documents = extract_documents_from_file_unified(
        file_content, filename, current_user=current_user
    )

    # Extract images based on file type
    images = []
    file_ext = Path(filename).suffix.lower() if "." in filename else ""

    try:
        if file_ext == ".pdf":
            images = extract_images_from_pdf_bytes(file_content)
        elif file_ext in [".docx", ".doc"]:
            images = extract_images_from_docx_bytes(file_content)
        elif file_ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"]:
            # Single image file
            img_base64 = base64.b64encode(file_content).decode()
            images = [img_base64]
    except Exception as e:
        logger.error(f"Error extracting images from {filename}: {e}")

    return documents, images


def extract_images_from_pdf_bytes(file_content: bytes) -> List[str]:
    """Extract images from PDF bytes using PyMuPDF or similar library."""
    import base64
    import logging

    logger = logging.getLogger(__name__)
    images = []

    try:
        # Try to import fitz (PyMuPDF)
        import fitz

        doc = fitz.open("pdf", file_content)

        for page_num in range(doc.page_count):  # Limit pages
            page = doc[page_num]

            # REMOVED FOR REDUNDANCY
            # Convert page to image for comprehensive analysis
            # pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
            # img_data = pix.tobytes("png")
            # img_base64 = base64.b64encode(img_data).decode()
            # images.append(img_base64)

            # Extract embedded images
            image_list = page.get_images()
            # Extract all embedded images (remove arbitrary per-page cap)
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    img_base64 = base64.b64encode(image_bytes).decode()
                    images.append(img_base64)
                except:
                    continue

        doc.close()

    except ImportError:
        logger.warning("PyMuPDF not available for PDF image extraction")
        # Fallback: Convert entire PDF to page images using pdf2image
        try:
            from pdf2image import convert_from_bytes

            pages = convert_from_bytes(file_content, dpi=150, fmt="PNG")
            for page in pages:
                import io

                img_buffer = io.BytesIO()
                page.save(img_buffer, format="PNG")
                img_data = img_buffer.getvalue()
                img_base64 = base64.b64encode(img_data).decode()
                images.append(img_base64)

        except ImportError:
            logger.warning("pdf2image not available for PDF image extraction")
    except Exception as e:
        logger.error(f"PDF image extraction error: {e}")

    return images


def extract_images_from_docx_bytes(file_content: bytes) -> List[str]:
    """Extract images from DOCX bytes."""
    import base64
    import tempfile
    import logging

    logger = logging.getLogger(__name__)
    images = []

    try:
        from docx import Document as DocxDocument
        import io

        # Create document from bytes
        doc = DocxDocument(io.BytesIO(file_content))

        # Extract embedded images from document relationships
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    img_data = rel.target_part.blob
                    img_base64 = base64.b64encode(img_data).decode()
                    images.append(img_base64)
                except:
                    continue

    except Exception as e:
        logger.error(f"DOCX image extraction error: {e}")

    return images


def create_fallback_document_for_vision(
    images: List[str], source_filename: str
) -> Document:
    """
    Create a fallback document for vision processing when no text content is available.
    This allows documents with only images to still be processed by vector search.

    Args:
        images: List of base64 encoded images
        source_filename: Name of the source file

    Returns:
        Document object with metadata indicating it's for vision processing
    """
    image_count = len(images) if images else 0

    fallback_content = f"This document '{source_filename}' contains {image_count} image(s) with visual content that can be analyzed using vision-enabled models."

    return Document(
        page_content=fallback_content,
        metadata={
            "source_filename": source_filename,
            "is_vision_fallback": True,
            "image_count": image_count,
            "content_type": "images_only",
        },
    )


def ensure_documents_for_vector_search(
    documents: List[Document], source_filename: str, images: List[str] = None
) -> List[Document]:
    """
    Ensure there are documents available for vector search, creating fallback documents if needed.
    This prevents vector search failures when documents contain only images or no extractable content.

    Args:
        documents: List of extracted documents (may be empty)
        source_filename: Name of the source file
        images: Optional list of base64 encoded images

    Returns:
        List of documents guaranteed to have at least one document for vector search
    """
    # Filter out empty documents
    valid_documents = [doc for doc in documents if doc.page_content.strip()]

    # If we have valid text documents, return them
    if valid_documents:
        return valid_documents

    # If no text documents but we have images, create fallback for vision processing
    if images and len(images) > 0:
        fallback_doc = create_fallback_document_for_vision(images, source_filename)
        return [fallback_doc]

    # Last resort: create a minimal document to prevent vector search failure
    fallback_content = f"This document '{source_filename}' appears to be empty or contains no extractable content."
    fallback_doc = Document(
        page_content=fallback_content,
        metadata={
            "source_filename": source_filename,
            "is_empty_fallback": True,
            "content_type": "empty",
        },
    )

    return [fallback_doc]


async def extract_text_with_vision_enhancement(
    file_content: bytes,
    filename: str,
    llm,
    purpose: str = "analysis",
    current_user=None,
) -> str:
    """
    Enhanced document processing that combines text extraction with visual processing.
    Used for suggestion endpoints (generate questions, outline, form fields, etc.)

    Args:
        file_content: Raw bytes of the document
        filename: Original filename
        llm: Language model instance
        purpose: Purpose of extraction (e.g., "checklist generation", "outline creation")
        current_user: Current user object for vision analysis permission check

    Returns:
        Combined text content from both text extraction and vision analysis
    """
    from app.services.vision_service import VisionService

    # Check file extension
    file_ext = Path(filename).suffix.lower() if filename else ""

    # Always try text extraction first
    print(f"📄 Extracting text from {filename} for {purpose}")
    text_content = extract_text_from_file_unified(file_content, filename)

    # Check if vision processing should be attempted
    vision_enabled = VisionService.is_vision_enabled(llm, current_user)

    if not vision_enabled:
        print(f"ℹ️ Vision not enabled, using text-only extraction for {filename}")
        return text_content

    # Try vision enhancement for PDFs and DOCX files (both can contain images)
    if file_ext not in [".pdf", ".docx"]:
        print(
            f"ℹ️ File {filename} format doesn't support embedded images, skipping vision enhancement"
        )
        return text_content

    # Try to extract images from PDF or DOCX
    try:
        print(f"🖼️ Attempting to extract images from {file_ext.upper()}: {filename}")
        _, document_images = extract_documents_and_images_from_file_unified(
            file_content, filename
        )

        if not document_images:
            print(f"ℹ️ No images found in {filename}, using text-only content")
            return text_content

        print(
            f"✅ Found {len(document_images)} images in {filename}, performing vision analysis"
        )

        # Convert images for vision analysis
        vision_images = []
        for i, img_b64 in enumerate(document_images):
            vision_images.append(
                {
                    "image_data": img_b64,
                    "metadata": {"source": filename, "page": i + 1},
                }
            )

        # Create purpose-specific vision prompt
        vision_prompt = f"""Analyze the visual content in these document images to extract information relevant for {purpose}.

Document: {filename}
Images: {len(document_images)} pages

Please extract any visual information that would be helpful for {purpose}, including:
1. Charts, graphs, diagrams, and visual data
2. Form fields, tables, and structured layouts  
3. Images, photos, and visual elements
4. Text that may not have been captured in standard text extraction
5. Visual indicators, symbols, and formatting

Provide a detailed description of the visual content that would be useful for {purpose}:"""

        # Perform vision analysis
        vision_result = VisionService.safe_vision_analysis(
            llm=llm,
            prompt_template=vision_prompt,
            variables={},
            images=vision_images,
        )

        if (
            vision_result
            and isinstance(vision_result, str)
            and len(vision_result.strip()) > 50
        ):
            # Combine text and vision results
            combined_content = text_content
            if combined_content.strip():
                combined_content += (
                    f"\n\n--- VISUAL ANALYSIS OF {filename.upper()} ---\n"
                )
            else:
                combined_content = f"--- VISUAL ANALYSIS OF {filename.upper()} ---\n"
            combined_content += vision_result.strip()

            print(
                f"✅ Successfully enhanced {filename} with vision analysis (+{len(vision_result)} chars)"
            )
            return combined_content
        else:
            print(
                f"⚠️ Vision analysis of {filename} produced minimal content, using text only"
            )
            return text_content

    except Exception as e:
        print(f"⚠️ Vision enhancement failed for {filename}: {str(e)}")
        return text_content
