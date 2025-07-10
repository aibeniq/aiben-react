import logging
import mimetypes
import os
import re
import tempfile
import time
from typing import Any

import docx
from fastapi import HTTPException, UploadFile
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pymilvus import (
    AnnSearchRequest,
    CollectionSchema,
    DataType,
    FieldSchema,
    MilvusClient,
    WeightedRanker,
)

from app.core.config import settings
from app.services.embeddings.main import EmbeddingService
from app.services.vectordb.config import (
    BM25_FUNCTION,
    MILVUS_SCHEMA_TEMPLATE,
    MILVUS_URL,
)
from app.services.vectordb.types import (
    ChunkData,
    EmbeddedChunkData,
    SearchEntity,
    SearchHit,
    SearchResults,
    VectorDBError,
    EmbeddingModelError,
)

logger = logging.getLogger(__name__)


def _extract_text_from_docx(file_path: str, filename: str) -> list[Any]:
    doc = docx.Document(file_path)

    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():  # Skip empty paragraphs
            full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                full_text.append(" | ".join(row_text))

    combined_text = "\n\n".join(full_text)

    metadata = {
        "source": filename,
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    # Try to get document properties
    try:
        core_properties = doc.core_properties
        if core_properties.title:
            metadata["title"] = core_properties.title
        if core_properties.author:
            metadata["author"] = core_properties.author
        if core_properties.created:
            metadata["created"] = str(core_properties.created)
        if core_properties.modified:
            metadata["modified"] = str(core_properties.modified)
    except Exception as e:
        print(f"Could not extract document properties: {str(e)}")

    # Create a Document object compatible with langchain
    return [Document(page_content=combined_text, metadata=metadata)]


def _load_uploaded_file(file: UploadFile) -> list[Any]:
    """
    Load an uploaded file based on its type (e.g., PDF, text file).

    Args:
        file (UploadFile): The uploaded file to process.

    Returns:
        List[Any]: A list of loaded documents from the file.
    """
    print(f"Processing file: {file.filename}")
    if not file.filename:
        raise ValueError("File has no filename")
    content_type = file.content_type or mimetypes.guess_type(file.filename)[0]
    if not content_type:
        raise ValueError("File has no content type")
    print(f"Detected content type: {content_type}")

    with tempfile.NamedTemporaryFile(
        delete=False, suffix=f"_{file.filename}"
    ) as temp_file:
        temp_file.write(
            file.file.read()
        )  # Write the file content to the temporary file
        temp_file_path = temp_file.name

    try:
        if content_type == "application/pdf" or file.filename.lower().endswith(".pdf"):
            print("Loading PDF with PyPDFLoader...")
            pypdf_loader = PyPDFLoader(temp_file_path)
            loaded_documents = pypdf_loader.load()
        elif (
            content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or file.filename.lower().endswith(".docx")
        ):
            print("Loading DOCX with python-docx library...")
            loaded_documents = _extract_text_from_docx(temp_file_path, file.filename)
        else:
            print("Loading text with TextLoader...")
            # Try with different encodings if utf-8 fails
            try:
                text_loader = TextLoader(temp_file_path, encoding="utf-8")
                loaded_documents = text_loader.load()
            except UnicodeDecodeError:
                print("UTF-8 decoding failed. Retrying with Latin-1 encoding...")
                text_loader = TextLoader(temp_file_path, encoding="latin-1")
                loaded_documents = text_loader.load()

        return loaded_documents
    except Exception as e:
        print(f"Error processing file {file.filename}: {str(e)}")
        raise HTTPException(
            status_code=400, detail=f"Error processing file {file.filename}: {str(e)}"
        )
    finally:
        # Clean up the temporary file
        os.unlink(temp_file_path)


class VectorDBService:
    """Service for interacting with the vector database."""

    def __init__(self) -> None:
        """Initialize the vector database service."""
        try:
            # initialize Milvus client
            self.client = MilvusClient(MILVUS_URL)
            logger.info(f"Connected to Milvus at {MILVUS_URL}")

            # cache for embedding models and initialized collections
            self._embedding_models: dict[str, Embeddings] = {}
            self._initialized_collections: set[str] = set()

            self.default_output_fields = [
                "content",
                "tags",
                "title",
                "url",
                "knowledge_base_id",
                "user_id",
                "source_id",
            ]

        except Exception as e:
            logger.error(f"Failed to initialize VectorDB service: {e}")
            raise

    def _get_collection_name(self, embedding_model_id: str) -> str:
        """Generate collection name from embedding model ID.

        Returns:
            str: The collection name

        Raises:
            EmbeddingModelError: If embedding model is invalid
        """
        try:
            spec = EmbeddingService.get_model_spec(embedding_model_id)
        except ValueError as e:
            raise EmbeddingModelError(
                f"Failed to get collection name for {embedding_model_id}: {str(e)}"
            ) from e

        # create safe collection name (only alphanumerics and underscores): provider_model_dimensions
        safe_name = re.sub(
            r"[^a-zA-Z0-9_]", "", f"{spec.provider}_{spec.model_name}_{spec.dimensions}"
        )
        return safe_name

    def _get_embedding_model(self, embedding_model_id: str) -> Embeddings:
        """
        Get or load embedding model by ID.

        Args:
            embedding_model_id: The ID of the embedding model to get

        Returns:
            Embeddings: The embedding model

        Raises:
            EmbeddingModelError: If embedding model is invalid
        """
        if embedding_model_id not in self._embedding_models:
            try:
                self._embedding_models[embedding_model_id] = EmbeddingService.get_model(
                    embedding_model_id
                )
            except Exception as e:
                raise EmbeddingModelError(
                    f"Failed to load embedding model {embedding_model_id}: {str(e)}"
                ) from e

        return self._embedding_models[embedding_model_id]

    def _create_schema(self, embedding_model_id: str) -> CollectionSchema:
        """
        Create Milvus schema for the specified embedding model.

        Args:
            embedding_model_id: The ID of the embedding model to create a schema for

        Returns:
            CollectionSchema: The schema for the embedding model

        Raises:
            EmbeddingModelError: If embedding model is invalid
            VectorDBError: If schema creation fails
        """

        try:
            spec = EmbeddingService.get_model_spec(embedding_model_id)
        except ValueError as e:
            raise EmbeddingModelError(
                f"Unknown embedding model: {embedding_model_id}"
            ) from e

        try:
            # create schema fields with appropriate embedding dimension
            fields = MILVUS_SCHEMA_TEMPLATE + [
                FieldSchema(
                    name="dense",
                    dtype=DataType.FLOAT_VECTOR,
                    description="dense embedding vector",
                    dim=spec.dimensions,
                    nullable=False,
                )
            ]

            # create schema
            schema = CollectionSchema(
                fields=fields,
                description=f"schema for {embedding_model_id}",
            )

            # add BM25 function for keyword search
            schema.add_function(BM25_FUNCTION)

            return schema
        except Exception as e:
            raise VectorDBError(
                f"Failed to create schema for {embedding_model_id}: {str(e)}"
            ) from e

    def _init_collection(self, embedding_model_id: str) -> str:
        """
        Initialize a collection for the specified embedding model if it doesn't exist.

        Args:
            embedding_model_id: The ID of the embedding model to initialize a collection for

        Returns:
            str: The name of the collection

        Raises:
            EmbeddingModelError: If embedding model is invalid
            VectorDBError: If collection initialization fails
        """

        try:
            collection_name = self._get_collection_name(embedding_model_id)

            # check if already initialized
            if collection_name in self._initialized_collections:
                return collection_name

            # check if collection exists
            if self.client.has_collection(collection_name=collection_name):
                logger.info(f"Collection '{collection_name}' already exists.")
                self._initialized_collections.add(collection_name)
                return collection_name

            logger.info(f"Creating collection {collection_name}...")

            # create schema
            schema = self._create_schema(embedding_model_id)

            # create collection
            self.client.create_collection(
                collection_name=collection_name,
                schema=schema,
            )

            # create index parameters
            index_params = self.client.prepare_index_params()

            # add indexes
            index_params.add_index(field_name="id", index_type="STL_SORT")

            index_params.add_index(
                field_name="dense",
                index_type="HNSW",
                metric_type="COSINE",
                params={"M": 16, "efConstruction": 500},
            )

            index_params.add_index(
                field_name="sparse",
                index_type="SPARSE_INVERTED_INDEX",
                metric_type="BM25",
                params={"inverted_index_algo": "DAAT_MAXSCORE"},
            )

            index_params.add_index(field_name="knowledge_base_id")
            index_params.add_index(field_name="user_id")
            index_params.add_index(field_name="source_id")

            # create indexes
            self.client.create_index(
                collection_name=collection_name,
                index_params=index_params,
                sync=True,
            )

            logger.info(f"Collection '{collection_name}' created successfully.")
            self._initialized_collections.add(collection_name)

            return collection_name

        except EmbeddingModelError:
            raise
        except Exception as e:
            logger.error(f"Error initializing collection '{collection_name}': {e}")
            raise VectorDBError(
                f"Failed to initialize collection '{collection_name}': {str(e)}"
            ) from e

    def add_source(
        self,
        file: UploadFile,
        knowledge_base_id: str,
        user_id: str,
        source_id: str,
        embedding_model_id: str,
    ) -> None:
        """Add files to the collection.

        Raises:
            EmbeddingModelError: If embedding model is invalid
            VectorDBError: If file processing or insertion fails
        """

        try:
            # validate embedding model
            EmbeddingService.get_model_spec(embedding_model_id)

            documents = []

            # Load documents from file
            loaded_documents = _load_uploaded_file(file)  # TODO: switch to docling
            documents.extend(loaded_documents)

            # Split documents into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=settings.DOCUMENT_CHUNK_SIZE,
                chunk_overlap=settings.DOCUMENT_CHUNK_OVERLAP,
            )
            splits = text_splitter.split_documents(documents)

            chunks_to_add: list[ChunkData] = []
            for split in splits:
                chunk_data = ChunkData(
                    knowledge_base_id=str(knowledge_base_id),
                    source_id=str(source_id),
                    user_id=str(user_id),
                    content=split.page_content,
                    tags=[],
                    title=split.metadata.get("title", ""),
                    summary="",
                    author=split.metadata.get("author", ""),
                    url=split.metadata.get("source", ""),
                    created_at=int(time.time()),
                    updated_at=int(time.time()),
                )
                chunks_to_add.append(chunk_data)

            self._add_chunks(
                chunks_to_add,
                embedding_model_id=embedding_model_id,
            )
        except EmbeddingModelError:
            raise
        except Exception as e:
            raise VectorDBError(f"Failed to add source {source_id}: {str(e)}") from e

    def delete_source(self, source_id: str, embedding_model_id: str) -> None:
        """Delete a file from the collection.

        Raises:
            EmbeddingModelError: If embedding model is invalid
            VectorDBError: If deletion fails
        """
        try:
            # validate embedding model
            EmbeddingService.get_model_spec(embedding_model_id)

        except Exception as e:
            raise EmbeddingModelError(
                f"Failed to delete source {source_id}: {str(e)}"
            ) from e

        try:
            self._delete_chunks(
                source_id=source_id,
                embedding_model_id=embedding_model_id,
            )
        except Exception as e:
            raise VectorDBError(f"Failed to delete source {source_id}: {str(e)}") from e

    def delete_knowledge_base(
        self, knowledge_base_id: str, embedding_model_id: str
    ) -> None:
        """Delete a knowledge base from the collection.

        Raises:
            EmbeddingModelError: If embedding model is invalid
            VectorDBError: If deletion fails
        """
        try:
            # validate embedding model
            EmbeddingService.get_model_spec(embedding_model_id)
        except Exception as e:
            raise EmbeddingModelError(
                f"Failed to delete knowledge base {knowledge_base_id}: {str(e)}"
            ) from e

        try:
            self._delete_chunks(
                knowledge_base_id=knowledge_base_id,
                embedding_model_id=embedding_model_id,
            )
        except Exception as e:
            raise VectorDBError(
                f"Failed to delete knowledge base {knowledge_base_id}: {str(e)}"
            ) from e

    def _add_chunks(
        self,
        chunks: list[ChunkData],
        embedding_model_id: str,
    ) -> None:
        """
        Add chunks to the collection.

        Args:
            chunks: List of ChunkData objects containing chunk data
            embedding_model_id: ID of the embedding model to use

        Raises:
            EmbeddingModelError: If embedding model is invalid
            VectorDBError: If insertion fails
        """

        try:

            # ensure collection exists and is loaded
            collection_name = self._init_collection(embedding_model_id)

            # load collection if not already loaded
            self.client.load_collection(collection_name=collection_name)

            # get embedding model and generate embeddings
            embedding_model = self._get_embedding_model(embedding_model_id)
            embeddings = embedding_model.embed_documents(
                [chunk.content for chunk in chunks]
            )

            # prepare data for insertion
            data_to_insert: list[dict[str, Any]] = []
            for chunk, embedding in zip(chunks, embeddings, strict=False):
                chunk_data = EmbeddedChunkData(**chunk.model_dump(), dense=embedding)
                # convert pydantic model to dictionary for milvus insertion
                data_to_insert.append(chunk_data.model_dump())

            # insert data
            self.client.insert(
                collection_name=collection_name,
                data=data_to_insert,
            )

            logger.info(
                f"Successfully inserted {len(data_to_insert)} chunks into {collection_name}"
            )

        except EmbeddingModelError:
            raise
        except Exception as e:
            raise VectorDBError(
                f"Failed to add chunks to {collection_name}: {str(e)}"
            ) from e

    def _delete_chunks(
        self,
        embedding_model_id: str,
        knowledge_base_id: str | None = None,
        user_id: str | None = None,
        source_id: str | None = None,
    ) -> None:
        """Delete chunks from the collection.

        Raises:
            VectorDBError: If deletion fails
        """

        try:
            if not (knowledge_base_id or user_id or source_id):
                raise VectorDBError(
                    "At least one of knowledge_base_id, user_id, or source_id must be provided"
                )

            # ensure collection exists and is loaded
            collection_name = self._init_collection(embedding_model_id)

            # prepare filter
            filter_expr: list[str] = []
            filter_params: dict[str, str] = {}
            if knowledge_base_id:
                filter_expr.append("knowledge_base_id == {kb_id}")
                filter_params["kb_id"] = knowledge_base_id
            if user_id:
                filter_expr.append("user_id == {user_id}")
                filter_params["user_id"] = user_id
            if source_id:
                filter_expr.append("source_id == {source_id}")
                filter_params["source_id"] = source_id
            filter = " AND ".join(filter_expr) if filter_expr else None

            # delete chunks
            self.client.delete(
                collection_name=collection_name,
                filter=filter,
                filter_params=filter_params if filter_params else None,
            )

            logger.info(f"Successfully deleted chunks from {collection_name}")

        except EmbeddingModelError:
            raise
        except Exception as e:
            raise VectorDBError(
                f"Failed to delete chunks from {collection_name}: {str(e)}"
            ) from e

    def search_semantic(
        self,
        query: str,
        embedding_model_id: str,
        knowledge_base_id: str | None = None,
        user_id: str | None = None,
        source_id: str | None = None,
        limit: int = 10,
        output_fields: list[str] | None = None,
    ) -> SearchResults:
        """
        Semantic similarity search.

        Args:
            query: Search query text
            embedding_model_id: ID of the embedding model to use
            knowledge_base_id: Optional filter by knowledge base ID
            user_id: Optional filter by user ID
            source_id: Optional filter by source ID
            limit: Maximum number of results to return
            output_fields: List of fields to return in results

        Returns:
            SearchResults

        Raises:
            VectorDBError: If search operation fails
            EmbeddingModelError: If embedding model is invalid
        """

        try:
            if not (knowledge_base_id or user_id or source_id):
                raise VectorDBError(
                    "At least one of knowledge_base_id, user_id, or source_id must be provided"
                )

            # ensure collection exists and is loaded
            collection_name = self._init_collection(embedding_model_id)

            # get embedding model and generate query embedding
            embedding_model = self._get_embedding_model(embedding_model_id)
            query_embedding = embedding_model.embed_query(query)

            # prepare filter
            filter_expr = []
            filter_params = {}
            if knowledge_base_id:
                filter_expr.append("knowledge_base_id == {kb_id}")
                filter_params["kb_id"] = knowledge_base_id
            if user_id:
                filter_expr.append("user_id == {user_id}")
                filter_params["user_id"] = user_id
            if source_id:
                filter_expr.append("source_id == {source_id}")
                filter_params["source_id"] = source_id
            filter = " AND ".join(filter_expr) if filter_expr else None

            # perform search
            results = self.client.search(
                collection_name=collection_name,
                data=[query_embedding],
                anns_field="dense",
                filter=filter,
                filter_params=filter_params if filter_params else None,
                limit=limit,
                output_fields=(
                    output_fields
                    if output_fields is not None
                    else self.default_output_fields
                ),
            )

            # convert raw results to typed hits
            hits = []
            if results and results[0]:
                for hit in results[0]:
                    # convert entity dict to SearchEntity
                    entity_data = hit.get("entity", {})
                    search_entity = SearchEntity(
                        content=entity_data.get("content", ""),
                        knowledge_base_id=entity_data.get("knowledge_base_id", ""),
                        source_id=entity_data.get("source_id", ""),
                        user_id=entity_data.get("user_id", ""),
                        title=entity_data.get("title", ""),
                        author=entity_data.get("author", ""),
                        url=entity_data.get("url", ""),
                        tags=entity_data.get("tags", []),
                        summary=entity_data.get("summary", ""),
                        created_at=entity_data.get("created_at", 0),
                        updated_at=entity_data.get("updated_at", 0),
                    )
                    search_hit = SearchHit(
                        id=str(hit.get("id", "")),
                        distance=float(hit.get("distance", 0.0)),
                        entity=search_entity,
                    )
                    hits.append(search_hit)

            search_results = SearchResults(hits=hits)

            logger.info(f"Semantic search completed with {len(hits)} results")
            return search_results

        except EmbeddingModelError:
            raise
        except Exception as e:
            logger.error(f"Error searching chunks in {collection_name}: {e}")
            raise VectorDBError(f"Semantic search failed: {str(e)}") from e

    def search_keyword(
        self,
        query: str,
        embedding_model_id: str,
        knowledge_base_id: str | None = None,
        user_id: str | None = None,
        source_id: str | None = None,
        limit: int = 10,
        output_fields: list[str] | None = None,
    ) -> SearchResults:
        """
        Keyword search.

        Args:
            query: Search query text
            embedding_model_id: ID of the embedding model to use
            knowledge_base_id: Optional filter by knowledge base ID
            user_id: Optional filter by user ID
            source_id: Optional filter by source ID
            limit: Maximum number of results to return
            output_fields: List of fields to return in results

        Returns:
            SearchResults

        Raises:
            VectorDBError: If search operation fails
            EmbeddingModelError: If embedding model is invalid
        """

        try:

            if not (knowledge_base_id or user_id or source_id):
                raise VectorDBError(
                    "At least one of knowledge_base_id, user_id, or source_id must be provided"
                )

            # ensure collection exists and is loaded
            collection_name = self._init_collection(embedding_model_id)

            # prepare filter
            filter_expr: list[str] = []
            filter_params: dict[str, str] = {}
            if knowledge_base_id:
                filter_expr.append("knowledge_base_id == {kb_id}")
                filter_params["kb_id"] = knowledge_base_id
            if user_id:
                filter_expr.append("user_id == {user_id}")
                filter_params["user_id"] = user_id
            if source_id:
                filter_expr.append("source_id == {source_id}")
                filter_params["source_id"] = source_id
            filter = " AND ".join(filter_expr) if filter_expr else None

            # perform search
            results = self.client.search(
                collection_name=collection_name,
                data=[query],
                anns_field="sparse",
                filter=filter,
                filter_params=filter_params if filter_params else None,
                limit=limit,
                output_fields=(
                    output_fields
                    if output_fields is not None
                    else self.default_output_fields
                ),
            )

            # convert raw results to typed hits
            hits = []
            if results and results[0]:
                for hit in results[0]:
                    # convert entity dict to SearchEntity
                    entity_data = hit.get("entity", {})
                    search_entity = SearchEntity(
                        content=entity_data.get("content", ""),
                        knowledge_base_id=entity_data.get("knowledge_base_id", ""),
                        source_id=entity_data.get("source_id", ""),
                        user_id=entity_data.get("user_id", ""),
                        title=entity_data.get("title", ""),
                        author=entity_data.get("author", ""),
                        url=entity_data.get("url", ""),
                        tags=entity_data.get("tags", []),
                        summary=entity_data.get("summary", ""),
                        created_at=entity_data.get("created_at", 0),
                        updated_at=entity_data.get("updated_at", 0),
                    )
                    search_hit = SearchHit(
                        id=str(hit.get("id", "")),
                        distance=float(hit.get("distance", 0.0)),
                        entity=search_entity,
                    )
                    hits.append(search_hit)

            search_results = SearchResults(hits=hits)

            logger.info(f"Keyword search completed with {len(hits)}")
            return search_results

        except EmbeddingModelError:
            raise
        except Exception as e:
            logger.error(f"Error searching chunks in {collection_name}: {e}")
            raise VectorDBError(f"Keyword search failed: {str(e)}") from e

    def search_hybrid(
        self,
        query: str,
        embedding_model_id: str,
        knowledge_base_id: str | None = None,
        user_id: str | None = None,
        source_id: str | None = None,
        limit: int = 10,
        output_fields: list[str] | None = None,
        alpha: float = 0.5,
        rerank_k: int = 20,
    ) -> SearchResults:
        """
        Hybrid search combining semantic and keyword search.

        Args:
            query: Search query text
            embedding_model_id: ID of the embedding model to use
            knowledge_base_id: Optional filter by knowledge base ID
            user_id: Optional filter by user ID
            source_id: Optional filter by source ID
            limit: Maximum number of results to return
            output_fields: List of fields to return in results
            alpha: Weight for semantic search results (0.0-1.0)
            rerank_k: Number of candidates to retrieve before reranking

        Returns:
            SearchResults

        Raises:
            VectorDBError: If search operation fails
            EmbeddingModelError: If embedding model is invalid
        """

        try:
            if not (knowledge_base_id or user_id or source_id):
                raise VectorDBError(
                    "At least one of knowledge_base_id, user_id, or source_id must be provided"
                )

            # ensure collection exists and is loaded
            collection_name = self._init_collection(embedding_model_id)

            if not 0.0 <= alpha <= 1.0:
                raise VectorDBError("alpha must be between 0.0 and 1.0")

            # get embedding model and generate query embedding
            embedding_model = self._get_embedding_model(embedding_model_id)
            embedded_query = embedding_model.embed_query(query)

            # prepare filter
            filter_expr: list[str] = []
            filter_params: dict[str, str] = {}
            if knowledge_base_id:
                filter_expr.append("knowledge_base_id == {kb_id}")
                filter_params["kb_id"] = knowledge_base_id
            if user_id:
                filter_expr.append("user_id == {user_id}")
                filter_params["user_id"] = user_id
            if source_id:
                filter_expr.append("source_id == {source_id}")
                filter_params["source_id"] = source_id
            filter = " AND ".join(filter_expr) if filter_expr else None

            # perform hybrid search using Milvus native hybrid search
            dense_search_params = {
                "data": [embedded_query],
                "anns_field": "dense",
                "param": {"efSearch": 500},
                "limit": rerank_k,
                "expr": filter,
                "expr_params": filter_params if filter_params else None,
            }
            sparse_search_params = {
                "data": [query],
                "anns_field": "sparse",
                "param": {"drop_ratio_search": 0.2},
                "limit": rerank_k,
                "expr": filter,
                "expr_params": filter_params if filter_params else None,
            }

            dense_request = AnnSearchRequest(**dense_search_params)
            sparse_request = AnnSearchRequest(**sparse_search_params)

            ranker = WeightedRanker(alpha, 1 - alpha)

            # execute hybrid search
            results = self.client.hybrid_search(
                collection_name=collection_name,
                reqs=[dense_request, sparse_request],
                ranker=ranker,
                limit=limit,
                output_fields=(
                    output_fields
                    if output_fields is not None
                    else self.default_output_fields
                ),
            )

            # convert raw results to typed hits
            hits = []
            if results and results[0]:
                for hit in results[0]:
                    # convert entity dict to SearchEntity
                    entity_data = hit.get("entity", {})
                    search_entity = SearchEntity(
                        content=entity_data.get("content", ""),
                        knowledge_base_id=entity_data.get("knowledge_base_id", ""),
                        source_id=entity_data.get("source_id", ""),
                        user_id=entity_data.get("user_id", ""),
                        title=entity_data.get("title", ""),
                        author=entity_data.get("author", ""),
                        url=entity_data.get("url", ""),
                        tags=entity_data.get("tags", []),
                        summary=entity_data.get("summary", ""),
                        created_at=entity_data.get("created_at", 0),
                        updated_at=entity_data.get("updated_at", 0),
                    )
                    search_hit = SearchHit(
                        id=str(hit.get("id", "")),
                        distance=float(hit.get("distance", 0.0)),
                        entity=search_entity,
                    )
                    hits.append(search_hit)

            search_results = SearchResults(hits=hits)

            logger.info(f"Hybrid search completed with {len(hits)} results")
            return search_results

        except EmbeddingModelError:
            raise
        except Exception as e:
            logger.error(f"Error in hybrid search for {collection_name}: {e}")
            raise VectorDBError(f"Hybrid search failed: {str(e)}") from e
