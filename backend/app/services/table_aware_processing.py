"""
Table-aware document processing utilities.
Provides enhanced table extraction and structured formatting for better vector search.
"""

import json
import re
import tempfile
import os
from typing import List, Dict, Any, Optional, Tuple, Union
from pathlib import Path
import pandas as pd
from io import BytesIO, StringIO
import logging

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


class TableAwareProcessor:
    """
    Enhanced document processor that preserves table structure and relationships.
    Outputs structured formats (JSON) rather than flattened text for better embeddings.
    """

    def __init__(
        self,
        preserve_headers: bool = True,
        max_table_rows: int = 1000,
        max_row_documents: int = 50,
        enable_json_format: bool = True,
        enable_structured_format: bool = True,
        enable_row_documents: bool = True,
    ):
        """
        Initialize the table-aware processor.

        Args:
            preserve_headers: Whether to preserve column headers in each row
            max_table_rows: Maximum number of rows to process per table (performance)
            max_row_documents: Maximum individual row documents to create
            enable_json_format: Whether to create JSON representation of tables
            enable_structured_format: Whether to create structured text format
            enable_row_documents: Whether to create individual row documents
        """
        self.preserve_headers = preserve_headers
        self.max_table_rows = max_table_rows
        self.max_row_documents = max_row_documents
        self.enable_json_format = enable_json_format
        self.enable_structured_format = enable_structured_format
        self.enable_row_documents = enable_row_documents

    def extract_tables_from_pdf_bytes(
        self, file_content: bytes, filename: str
    ) -> List[Dict[str, Any]]:
        """
        Extract tables from PDF using pdfplumber for better table detection.

        Args:
            file_content: PDF bytes
            filename: Name of the file

        Returns:
            List of table dictionaries with structure preserved
        """
        tables = []

        try:
            # Use pdfplumber for better table extraction
            import pdfplumber

            with pdfplumber.open(BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_tables = page.extract_tables()

                        for table_idx, table in enumerate(page_tables):
                            if not table or len(table) == 0:
                                continue

                            processed_table = self._process_table_structure(
                                table,
                                f"{filename}_page_{page_num + 1}_table_{table_idx + 1}",
                            )

                            if processed_table:
                                processed_table["metadata"] = {
                                    "source": filename,
                                    "page": page_num + 1,
                                    "table_index": table_idx + 1,
                                    "type": "table",
                                }
                                tables.append(processed_table)

                    except Exception as e:
                        logger.warning(
                            f"Error extracting tables from page {page_num + 1}: {e}"
                        )
                        continue

        except ImportError:
            logger.warning(
                "pdfplumber not available, falling back to basic PDF processing"
            )
            # Fallback to basic pypdf extraction (no table structure)
            return self._extract_tables_fallback_pdf(file_content, filename)
        except Exception as e:
            logger.error(f"Error in PDF table extraction for {filename}: {e}")
            return []

        return tables

    def extract_tables_from_docx_bytes(
        self, file_content: bytes, filename: str
    ) -> List[Dict[str, Any]]:
        """
        Extract tables from DOCX with structure preserved.

        Args:
            file_content: DOCX bytes
            filename: Name of the file

        Returns:
            List of table dictionaries with structure preserved
        """
        tables = []

        try:
            from docx import Document as DocxDocument

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            try:
                doc = DocxDocument(temp_file_path)

                for table_idx, table in enumerate(doc.tables):
                    try:
                        # Extract table data as 2D array
                        table_data = []
                        for row in table.rows:
                            row_data = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                row_data.append(cell_text)
                            table_data.append(row_data)

                        if len(table_data) > 0:
                            processed_table = self._process_table_structure(
                                table_data, f"{filename}_table_{table_idx + 1}"
                            )

                            if processed_table:
                                processed_table["metadata"] = {
                                    "source": filename,
                                    "table_index": table_idx + 1,
                                    "type": "table",
                                }
                                tables.append(processed_table)

                    except Exception as e:
                        logger.warning(
                            f"Error extracting table {table_idx + 1} from DOCX: {e}"
                        )
                        continue

            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

        except Exception as e:
            logger.error(f"Error in DOCX table extraction for {filename}: {e}")
            return []

        return tables

    def extract_tables_from_csv_bytes(
        self, file_content: bytes, filename: str
    ) -> List[Dict[str, Any]]:
        """
        Extract table structure from CSV file.

        Args:
            file_content: CSV bytes
            filename: Name of the file

        Returns:
            List containing single table dictionary
        """
        try:
            # Try different encodings
            encodings = ["utf-8", "latin-1", "cp1252"]
            df = None

            for encoding in encodings:
                try:
                    csv_string = file_content.decode(encoding)
                    df = pd.read_csv(StringIO(csv_string))
                    break
                except (UnicodeDecodeError, pd.errors.EmptyDataError):
                    continue

            if df is None or df.empty:
                return []

            # Convert DataFrame to table structure
            table_data = [df.columns.tolist()]  # Headers
            table_data.extend(df.values.tolist())  # Data rows

            processed_table = self._process_table_structure(table_data, filename)

            if processed_table:
                processed_table["metadata"] = {
                    "source": filename,
                    "table_index": 1,
                    "type": "table",
                }
                return [processed_table]

        except Exception as e:
            logger.error(f"Error in CSV table extraction for {filename}: {e}")

        return []

    def extract_tables_from_xlsx_bytes(
        self, file_content: bytes, filename: str
    ) -> List[Dict[str, Any]]:
        """
        Extract tables from XLSX with structure preserved for each sheet.

        Args:
            file_content: XLSX bytes
            filename: Name of the file

        Returns:
            List of table dictionaries (one per sheet)
        """
        tables = []

        try:
            xlsx_io = BytesIO(file_content)
            excel_file = pd.ExcelFile(xlsx_io)

            for sheet_idx, sheet_name in enumerate(excel_file.sheet_names):
                try:
                    df = pd.read_excel(xlsx_io, sheet_name=sheet_name)

                    if df.empty:
                        continue

                    # Convert DataFrame to table structure
                    table_data = [df.columns.tolist()]  # Headers
                    table_data.extend(df.values.tolist())  # Data rows

                    processed_table = self._process_table_structure(
                        table_data, f"{filename}_sheet_{sheet_name}"
                    )

                    if processed_table:
                        processed_table["metadata"] = {
                            "source": filename,
                            "sheet_name": sheet_name,
                            "table_index": sheet_idx + 1,
                            "type": "table",
                        }
                        tables.append(processed_table)

                except Exception as e:
                    logger.warning(
                        f"Error extracting sheet '{sheet_name}' from XLSX: {e}"
                    )
                    continue

        except Exception as e:
            logger.error(f"Error in XLSX table extraction for {filename}: {e}")
            return []

        return tables

    def _process_table_structure(
        self, table_data: List[List[str]], table_id: str
    ) -> Optional[Dict[str, Any]]:
        """
        Process raw table data into structured format for better embeddings.

        Args:
            table_data: 2D array of table cells
            table_id: Identifier for the table

        Returns:
            Structured table dictionary or None if invalid
        """
        if not table_data or len(table_data) == 0:
            return None

        # Remove empty rows and clean data
        clean_data = []
        for row in table_data:
            clean_row = [str(cell).strip() if cell is not None else "" for cell in row]
            # Skip completely empty rows
            if any(cell for cell in clean_row):
                clean_data.append(clean_row)

        if len(clean_data) < 2:  # Need at least header + 1 data row
            return None

        # Assume first row is headers
        headers = clean_data[0]
        data_rows = clean_data[1:]

        # Limit rows for performance
        if len(data_rows) > self.max_table_rows:
            data_rows = data_rows[: self.max_table_rows]

        # Create structured representation
        structured_table = {
            "table_id": table_id,
            "headers": headers,
            "rows": [],
            "structured_text": "",
            "json_representation": "",
        }

        # Process each row
        for row_idx, row in enumerate(data_rows):
            # Pad row if shorter than headers
            while len(row) < len(headers):
                row.append("")

            # Truncate row if longer than headers
            row = row[: len(headers)]

            # Create row dictionary
            row_dict = {}
            for col_idx, header in enumerate(headers):
                if col_idx < len(row):
                    row_dict[header] = row[col_idx]
                else:
                    row_dict[header] = ""

            structured_table["rows"].append(row_dict)

        # Create multiple text representations for embedding diversity

        # 1. Structured text with clear column associations
        text_parts = []
        text_parts.append(f"Table: {table_id}")
        text_parts.append(f"Headers: {' | '.join(headers)}")
        text_parts.append("")

        for row_idx, row_dict in enumerate(structured_table["rows"]):
            row_text = []
            for header in headers:
                value = row_dict.get(header, "")
                if value:
                    row_text.append(f"{header}: {value}")

            if row_text:
                text_parts.append(f"Row {row_idx + 1}: {' | '.join(row_text)}")

        structured_table["structured_text"] = "\n".join(text_parts)

        # 2. JSON representation for precise structure preservation
        json_repr = {
            "table_id": table_id,
            "headers": headers,
            "data": structured_table["rows"],
        }
        structured_table["json_representation"] = json.dumps(json_repr, indent=2)

        return structured_table

    def _extract_tables_fallback_pdf(
        self, file_content: bytes, filename: str
    ) -> List[Dict[str, Any]]:
        """
        Fallback PDF table extraction when pdfplumber is not available.
        Uses basic text parsing to identify potential tables.
        """
        tables = []

        try:
            from app.services.pdf_utils import extract_text_from_pdf_bytes

            text = extract_text_from_pdf_bytes(file_content, filename)

            # Simple heuristic: look for lines with multiple delimiters
            potential_tables = self._extract_tables_from_text(text, filename)
            tables.extend(potential_tables)

        except Exception as e:
            logger.error(f"Error in fallback PDF table extraction for {filename}: {e}")

        return tables

    def _extract_tables_from_text(
        self, text: str, filename: str
    ) -> List[Dict[str, Any]]:
        """
        Extract potential tables from plain text using heuristics.
        """
        tables = []

        lines = text.split("\n")
        potential_table_lines = []

        # Look for lines that might be table rows (multiple separators)
        separators = ["|", "\t", "  "]

        for line_idx, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            # Count separators
            separator_counts = [line.count(sep) for sep in separators]
            max_separators = max(separator_counts)

            if max_separators >= 2:  # Potential table row
                potential_table_lines.append((line_idx, line, max_separators))

        if len(potential_table_lines) < 2:
            return tables

        # Group consecutive potential table lines
        table_groups = []
        current_group = []

        for i, (line_idx, line, sep_count) in enumerate(potential_table_lines):
            if (
                not current_group or line_idx - current_group[-1][0] <= 2
            ):  # Allow 1-2 line gaps
                current_group.append((line_idx, line, sep_count))
            else:
                if len(current_group) >= 2:
                    table_groups.append(current_group)
                current_group = [(line_idx, line, sep_count)]

        if len(current_group) >= 2:
            table_groups.append(current_group)

        # Process each table group
        for table_idx, group in enumerate(table_groups):
            try:
                table_lines = [line for _, line, _ in group]

                # Determine the best separator
                best_sep = self._find_best_separator(table_lines)

                # Parse table
                table_data = []
                for line in table_lines:
                    row = [cell.strip() for cell in line.split(best_sep)]
                    row = [cell for cell in row if cell]  # Remove empty cells
                    if row:
                        table_data.append(row)

                if len(table_data) >= 2:
                    processed_table = self._process_table_structure(
                        table_data, f"{filename}_text_table_{table_idx + 1}"
                    )

                    if processed_table:
                        processed_table["metadata"] = {
                            "source": filename,
                            "table_index": table_idx + 1,
                            "type": "text_table",
                            "extraction_method": "heuristic",
                        }
                        tables.append(processed_table)

            except Exception as e:
                logger.warning(f"Error processing text table {table_idx + 1}: {e}")
                continue

        return tables

    def _find_best_separator(self, lines: List[str]) -> str:
        """Find the most consistent separator in table lines."""
        separators = ["|", "\t", "  ", " "]
        separator_scores = {}

        for sep in separators:
            scores = []
            for line in lines:
                count = line.count(sep)
                scores.append(count)

            # Consistency score (lower variance is better)
            if len(scores) > 1:
                mean_score = sum(scores) / len(scores)
                variance = sum((x - mean_score) ** 2 for x in scores) / len(scores)
                separator_scores[sep] = (mean_score, variance)
            else:
                separator_scores[sep] = (scores[0] if scores else 0, 0)

        # Choose separator with highest mean count and lowest variance
        best_sep = "|"
        best_score = (-1, float("inf"))

        for sep, (mean, variance) in separator_scores.items():
            if mean > best_score[0] or (
                mean == best_score[0] and variance < best_score[1]
            ):
                best_score = (mean, variance)
                best_sep = sep

        return best_sep

    def create_table_aware_documents(
        self, file_content: bytes, filename: str
    ) -> List[Document]:
        """
        Create LangChain documents with table-aware processing.

        Args:
            file_content: Raw file bytes
            filename: Name of the file

        Returns:
            List of Document objects with enhanced table representation
        """
        documents = []

        # Determine file type
        file_ext = Path(filename).suffix.lower()

        # Extract tables based on file type
        tables = []

        if file_ext == ".pdf":
            tables = self.extract_tables_from_pdf_bytes(file_content, filename)
        elif file_ext in [".docx", ".doc"]:
            tables = self.extract_tables_from_docx_bytes(file_content, filename)
        elif file_ext == ".csv":
            tables = self.extract_tables_from_csv_bytes(file_content, filename)
        elif file_ext in [".xlsx", ".xls"]:
            tables = self.extract_tables_from_xlsx_bytes(file_content, filename)
        else:
            # For other formats, try basic text extraction
            try:
                from app.services.document_utils import extract_text_from_file_unified

                text = extract_text_from_file_unified(file_content, filename)
                tables = self._extract_tables_from_text(text, filename)
            except Exception as e:
                logger.warning(f"Could not extract tables from {filename}: {e}")

        # Create documents for each table
        for table in tables:
            # Create multiple document representations for better search coverage

            # 1. Structured text representation (if enabled)
            if self.enable_structured_format:
                doc_structured = Document(
                    page_content=table["structured_text"],
                    metadata={
                        **table["metadata"],
                        "content_type": "table_structured",
                        "table_id": table["table_id"],
                        "headers": (
                            " | ".join(table["headers"]) if table["headers"] else ""
                        ),
                    },
                )
                documents.append(doc_structured)

            # 2. JSON representation for exact structure matching (if enabled)
            if self.enable_json_format:
                doc_json = Document(
                    page_content=table["json_representation"],
                    metadata={
                        **table["metadata"],
                        "content_type": "table_json",
                        "table_id": table["table_id"],
                        "headers": (
                            " | ".join(table["headers"]) if table["headers"] else ""
                        ),
                    },
                )
                documents.append(doc_json)

            # 3. Individual row documents for granular matching (if enabled)
            if self.enable_row_documents:
                max_rows = min(self.max_row_documents, len(table["rows"]))
                for row_idx, row_dict in enumerate(table["rows"][:max_rows]):
                    row_text_parts = []

                    # Include headers in each row for context (if preserve_headers is enabled)
                    if self.preserve_headers:
                        row_text_parts.append(
                            f"Table Headers: {' | '.join(table['headers'])}"
                        )

                    # Add row data with headers
                    row_data = []
                    for header in table["headers"]:
                        value = row_dict.get(header, "")
                        if value:
                            if self.preserve_headers:
                                row_data.append(f"{header}: {value}")
                            else:
                                row_data.append(value)

                    if row_data:
                        if self.preserve_headers:
                            row_text_parts.append(f"Row Data: {' | '.join(row_data)}")
                        else:
                            row_text_parts.append(" | ".join(row_data))

                        row_doc = Document(
                            page_content="\n".join(row_text_parts),
                            metadata={
                                **table["metadata"],
                                "content_type": "table_row",
                                "table_id": table["table_id"],
                                "row_index": row_idx,
                                "headers": (
                                    " | ".join(table["headers"])
                                    if table["headers"]
                                    else ""
                                ),
                            },
                        )
                        documents.append(row_doc)

        return documents


# Global instance - will be initialized with settings when first used
table_processor = None


def get_table_processor():
    """Get configured table processor instance."""
    global table_processor
    if table_processor is None:
        try:
            from app.core.config import settings

            table_processor = TableAwareProcessor(
                preserve_headers=settings.TABLE_PRESERVE_HEADERS,
                max_table_rows=settings.TABLE_MAX_ROWS_PER_TABLE,
                max_row_documents=settings.TABLE_MAX_ROW_DOCUMENTS,
                enable_json_format=settings.TABLE_ENABLE_JSON_FORMAT,
                enable_structured_format=settings.TABLE_ENABLE_STRUCTURED_FORMAT,
                enable_row_documents=settings.TABLE_ENABLE_ROW_DOCUMENTS,
            )
        except ImportError:
            # Fallback to defaults if settings not available
            table_processor = TableAwareProcessor()
    return table_processor


def extract_tables_from_file(
    file_content: bytes, filename: str
) -> List[Dict[str, Any]]:
    """
    Convenience function to extract tables from any supported file type.

    Args:
        file_content: Raw file bytes
        filename: Name of the file

    Returns:
        List of structured table dictionaries
    """
    processor = get_table_processor()
    return processor.create_table_aware_documents(file_content, filename)


def enhance_document_with_tables(file_content: bytes, filename: str) -> List[Document]:
    """
    Enhanced document processing that includes both regular text and table-aware content.

    Args:
        file_content: Raw file bytes
        filename: Name of the file

    Returns:
        List of Document objects including table-enhanced content
    """
    documents = []

    try:
        # Get regular documents
        from app.services.document_utils import extract_documents_from_file_unified

        regular_docs = extract_documents_from_file_unified(file_content, filename)

        # Get table-aware documents
        processor = get_table_processor()
        table_docs = processor.create_table_aware_documents(file_content, filename)

        # Combine both types
        documents.extend(regular_docs)
        documents.extend(table_docs)

    except Exception as e:
        logger.error(f"Error in enhanced document processing for {filename}: {e}")
        # Fallback to regular processing
        try:
            from app.services.document_utils import extract_documents_from_file_unified

            documents = extract_documents_from_file_unified(file_content, filename)
        except Exception as e2:
            logger.error(f"Fallback processing also failed for {filename}: {e2}")

    return documents
