"""
Page counting utilities for different document types.
"""
import tempfile
import os
from pathlib import Path
from typing import Tuple
import pypdf
from docx import Document as DocxDocument
from io import BytesIO


class PageCounter:
    @staticmethod
    def count_pages_from_bytes(file_content: bytes, filename: str) -> int:
        """
        Count pages in a document from its byte content.
        
        Args:
            file_content: Raw bytes of the document
            filename: Original filename to determine document type
            
        Returns:
            Number of pages in the document
        """
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == ".pdf":
            return PageCounter._count_pdf_pages(file_content, filename)
        elif file_ext in [".docx"]:
            return PageCounter._count_docx_pages(file_content, filename)
        elif file_ext in [".txt", ".md"]:
            return PageCounter._count_text_pages(file_content, filename)
        else:
            # For unknown file types, default to 1 page
            return 1
    
    @staticmethod
    def _count_pdf_pages(file_content: bytes, filename: str) -> int:
        """Count pages in a PDF document."""
        try:
            pdf_reader = pypdf.PdfReader(BytesIO(file_content))
            return len(pdf_reader.pages)
        except Exception as e:
            print(f"Error counting PDF pages for {filename}: {e}")
            return 1  # Default fallback
    
    @staticmethod
    def _count_docx_pages(file_content: bytes, filename: str) -> int:
        """
        Count pages in a DOCX document.
        Note: DOCX doesn't have explicit page breaks, so we estimate based on content.
        """
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                doc = DocxDocument(temp_file_path)
                
                # Method 1: Use page break elements (most accurate)
                page_breaks = 0
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                            page_breaks += 1
                
                # If explicit page breaks found, use them (+1 for first page)
                if page_breaks > 0:
                    return page_breaks + 1
                
                # Method 2: Estimate based on content length
                total_chars = sum(len(p.text) for p in doc.paragraphs)
                
                # Rough estimation: ~2000 characters per page (adjustable)
                estimated_pages = max(1, (total_chars + 1999) // 2000)
                
                return estimated_pages
                
            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            print(f"Error counting DOCX pages for {filename}: {e}")
            return 1  # Default fallback
    
    @staticmethod
    def _count_text_pages(file_content: bytes, filename: str) -> int:
        """Count pages in a text document based on content length."""
        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                text = file_content.decode('latin-1')
            
            # Count explicit page breaks (form feed characters)
            explicit_breaks = text.count('\f')
            if explicit_breaks > 0:
                return explicit_breaks + 1
            
            # Estimate based on line count (rough: ~50 lines per page)
            lines = text.count('\n') + 1
            estimated_pages = max(1, (lines + 49) // 50)
            
            return estimated_pages
            
        except Exception as e:
            print(f"Error counting text pages for {filename}: {e}")
            return 1  # Default fallback
