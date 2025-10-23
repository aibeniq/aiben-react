"""
Test fixtures for document processing and file handling.
Provides sample data for testing document extraction functions.
"""

import pytest
import io
from pathlib import Path


@pytest.fixture
def sample_docx_bytes():
    """Return bytes representing a simple DOCX file for testing."""
    try:
        from docx import Document
        import io

        # Create a simple DOCX document
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test paragraph with some content.")
        doc.add_paragraph("This is another paragraph.")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"

        # Save to bytes
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        return docx_buffer.getvalue()
    except ImportError:
        # Fallback if python-docx is not available
        return b""


@pytest.fixture
def sample_pdf_bytes():
    """Return bytes representing a simple PDF file for testing."""
    # Create a minimal PDF structure
    # For now, return empty bytes - in real implementation, create actual PDF
    return b"%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n2 0 obj\n<<\n/Type /Pages\n/Kids [3 0 R]\n/Count 1\n>>\nendobj\n3 0 obj\n<<\n/Type /Page\n/Parent 2 0 R\n/MediaBox [0 0 612 792]\n/Contents 4 0 R\n>>\nendobj\n4 0 obj\n<<\n/Length 44\n>>\nstream\nBT\n/F1 12 Tf\n100 700 Td\n(Hello World) Tj\nET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f\n0000000009 00000 n\n0000000058 00000 n\n0000000115 00000 n\n0000000200 00000 n\ntrailer\n<<\n/Size 5\n/Root 1 0 R\n>>\nstartxref\n284\n%%EOF"


@pytest.fixture
def sample_csv_bytes():
    """Return bytes representing a CSV file for testing."""
    return b"name,age,city\nJohn,25,New York\nJane,30,London\n"


@pytest.fixture
def sample_xlsx_bytes():
    """Return bytes representing an XLSX file for testing."""
    try:
        import pandas as pd
        import io

        # Create a simple DataFrame
        data = {
            "Name": ["Alice", "Bob", "Charlie"],
            "Age": [25, 30, 35],
            "City": ["New York", "London", "Paris"],
        }
        df = pd.DataFrame(data)

        # Save to Excel bytes
        excel_buffer = io.BytesIO()
        df.to_excel(excel_buffer, index=False, sheet_name="Sheet1")
        return excel_buffer.getvalue()
    except ImportError:
        # Fallback if pandas/openpyxl is not available
        return b""


@pytest.fixture
def sample_text_bytes():
    """Return bytes representing a plain text file."""
    return b"This is a sample text file.\nIt contains multiple lines.\nFor testing purposes."


@pytest.fixture
def corrupted_file_bytes():
    """Return bytes representing a corrupted file."""
    return b"This is not a valid file format for any document type."


@pytest.fixture
def empty_file_bytes():
    """Return empty bytes representing an empty file."""
    return b""


@pytest.fixture
def large_text_bytes():
    """Return bytes representing a large text file for performance testing."""
    # Create a large text file (1MB)
    content = "This is a test line that will be repeated.\n" * 20000
    return content.encode("utf-8")


@pytest.fixture
def mock_llm_response():
    """Mock response from LLM API for testing."""
    return {
        "choices": [
            {
                "message": {
                    "content": "This is a mock LLM response for testing purposes."
                }
            }
        ]
    }


@pytest.fixture
def temp_file_path(tmp_path):
    """Create a temporary file path for testing."""
    return tmp_path / "test_file.txt"


@pytest.fixture
def temp_directory(tmp_path):
    """Create a temporary directory for testing."""
    test_dir = tmp_path / "test_dir"
    test_dir.mkdir()
    return test_dir
