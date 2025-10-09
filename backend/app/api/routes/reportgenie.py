import re
import json
import csv
import uuid
import zipfile
import tempfile
import traceback
import asyncio
import os
import shutil
import markdown
import uuid
from pathlib import Path
from io import BytesIO, StringIO
from datetime import datetime
from fastapi.responses import StreamingResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup

from app.models import (
    ReportGenieRequest,
    ReportGenieResponse,
    ReportGenieOutline,
    ReportGenieDetailResponse,
    Source,
    SourceData,
    KnowledgeBase,
    EmbeddingModel,
    DocxRequest,
    LlmInteraction,
    Message,
    GenerateOutlineRequest,
    GenerateOutlineResponse,
    OptimizeOutlineRequest,
    OutlineSuggestion,
    OptimizedOutlineResponse,
)
from app.api.deps import CurrentUser, SessionDep
from app.core.config import settings
from app.services.knowledgebases import get_embedding_model
from app.services.embeddings import load_embeddings_model
from app.services.llms import get_default_llm, invoke_llm, record_llm_interaction
from app.services.translation import translate_text_if_needed
from app.services.retrievers import (
    create_ensemble_retriever,
)  # Import the ensemble retriever
from app.services.enhanced_retrieval import SmartRetrieverFactory
from app.services.text_processing import chunk_text, estimate_tokens
from app.services.progress_tracker import progress_tracker

from sqlmodel import select
from fastapi import (
    APIRouter,
    Depends,
    HTTPException,
    UploadFile,
    File,
    Form,
    Request as FastAPIRequest,
)
from typing import List, Dict, Any, Optional

from langchain_community.vectorstores import Chroma

# from langchain_community.document_loaders import PyPDFLoader, TextLoader  # Removed - using pypdf instead
from langchain_community.document_loaders import TextLoader
from app.services.pdf_utils import load_pdf_with_pypdf

router = APIRouter(prefix="/reportgenie", tags=["reportgenie"])


@router.post("/generate/task")
async def create_generate_task():
    """
    Create a progress tracking task for report generation and return task_id immediately.
    This allows frontend to start progress polling before form submission.
    """
    task_id = progress_tracker.create_task(
        "Generating report",
        {"setup": 0.1, "generating": 0.8, "finalizing": 0.1}
    )
    progress_tracker.update_stage_progress(
        task_id, "setup", 0, 1, "Waiting to start report generation..."
    )
    return {"task_id": task_id}


@router.post("/generate-outline/task")
async def create_generate_outline_task():
    """
    Create a progress tracking task for outline generation and return task_id immediately.
    This allows frontend to start progress polling before form submission.
    """
    task_id = progress_tracker.create_task(
        "Generating outline",
        {"processing_files": 0.2, "generating": 0.7, "finalizing": 0.1}
    )
    progress_tracker.update_stage_progress(
        task_id, "processing_files", 0, 1, "Waiting to start outline generation..."
    )
    return {"task_id": task_id}


@router.post("/optimize-outline/task")
async def create_optimize_outline_task():
    """
    Create a progress tracking task for outline optimization and return task_id immediately.
    This allows frontend to start progress polling before form submission.
    """
    task_id = progress_tracker.create_task(
        "Optimizing outline",
        {
            "setup": 0.1,
            "processing_document": 0.1,
            "generating": 0.4,
            "matching": 0.2,
            "comparing": 0.15,
            "finalizing": 0.05
        }
    )
    progress_tracker.update_stage_progress(
        task_id, "setup", 0, 1, "Waiting to start outline optimization..."
    )
    return {"task_id": task_id}


@router.get("/progress/{task_id}")
async def get_reportgenie_progress(
    task_id: str,
    current_user: CurrentUser,
) -> Any:
    """
    Get progress information for a reportgenie task (generate, generate-outline, or optimize-outline).
    """
    # Make this async to prevent blocking during intensive operations
    progress_data = progress_tracker.get_progress(task_id)
    if not progress_data:
        raise HTTPException(status_code=404, detail="Task not found")

    # Debug logging to see what's actually being returned
    print(f"🔍 REPORTGENIE API RETURNING PROGRESS: task_id={task_id}")
    print(f"🔍 PROGRESS DATA: status={progress_data.get('status')}, percentage={progress_data.get('percentage')}, current_stage={progress_data.get('current_stage')}")
    print(f"🔍 PROGRESS MESSAGE: {progress_data.get('message')}")
    print(f"🔍 PROGRESS STAGES: {list(progress_data.get('stages', {}).keys())}")
    
    # Check each stage completion status
    stages = progress_data.get('stages', {})
    for stage_name, stage_data in stages.items():
        completed = stage_data.get('completed', False) if isinstance(stage_data, dict) else False
        print(f"🔍 STAGE {stage_name}: completed={completed}")

    # Yield control to allow other async operations (like this API call) to run
    await asyncio.sleep(0)

    return progress_data



class KnowledgeBaseCache:
    """Cache for knowledge base retrievers to avoid reloading large databases multiple times."""

    def __init__(self):
        self.cached_retrievers = {}
        self.temp_dirs = {}
        self.cached_chroma_dbs = {}

    def get_retriever(self, kb_id: str, kb: "KnowledgeBase", session, current_user):
        """Get or create a cached retriever for the knowledge base."""
        cache_key = f"{kb_id}_{kb.embedding_model_id}"

        if cache_key not in self.cached_retrievers:
            # Create temporary directory for this knowledge base
            temp_dir = tempfile.mkdtemp()
            self.temp_dirs[cache_key] = temp_dir

            print(f"Loading knowledge base {kb_id} into cache (first time)")

            # Extract ChromaDB to temp directory
            if kb.storage_type == "file" and kb.file_path:
                if os.path.exists(kb.file_path):
                    with zipfile.ZipFile(kb.file_path, "r") as zip_ref:
                        zip_ref.extractall(temp_dir)
                else:
                    raise HTTPException(
                        status_code=400, detail="Knowledge base file not found on disk"
                    )
            elif kb.data:
                with zipfile.ZipFile(BytesIO(kb.data), "r") as zip_ref:
                    zip_ref.extractall(temp_dir)
            else:
                raise HTTPException(
                    status_code=400,
                    detail="Knowledge base has no vector database data",
                )

            # Load embeddings model
            if kb.embedding_model_id:
                embedding_model = session.get(EmbeddingModel, kb.embedding_model_id)
                if embedding_model:
                    model_id = embedding_model.model_id
                    provider = embedding_model.provider
                else:
                    embedding_info = get_embedding_model(session, current_user)
                    model_id = embedding_info["model_id"]
                    provider = embedding_info["provider"]
            else:
                embedding_info = get_embedding_model(session, current_user)
                model_id = embedding_info["model_id"]
                provider = embedding_info["provider"]

            embeddings = load_embeddings_model(provider=provider, model_id=model_id)
            chroma_db = Chroma(
                persist_directory=temp_dir, embedding_function=embeddings
            )

            # Create enhanced retriever with content filtering
            retriever = SmartRetrieverFactory.create_general_document_retriever(
                chroma_db=chroma_db,
                search_kwargs={"k": settings.RAG_NUM_CHUNKS},
            )

            # Cache both retriever and chroma_db
            self.cached_retrievers[cache_key] = retriever
            self.cached_chroma_dbs[cache_key] = chroma_db

            print(f"✅ Knowledge base {kb_id} loaded and cached successfully")
        else:
            print(f"♻️  Using cached knowledge base {kb_id}")

        return self.cached_retrievers[cache_key]

    def cleanup(self):
        """Clean up all temporary directories and cached resources."""
        print(f"🧹 Cleaning up knowledge base cache ({len(self.temp_dirs)} temp dirs)")
        for temp_dir in self.temp_dirs.values():
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception as e:
                print(f"Warning: Failed to cleanup temp dir {temp_dir}: {e}")
        self.cached_retrievers.clear()
        self.cached_chroma_dbs.clear()
        self.temp_dirs.clear()


def sanitize_text_for_json(text: str) -> str:
    """Sanitize text to prevent JSON parsing issues with control characters."""
    # Replace smart quotes and apostrophes with regular ones
    text = text.replace(""", "'").replace(""", "'")
    text = text.replace('"', '"').replace('"', '"')
    text = text.replace("–", "-").replace("—", "-")
    text = text.replace("ʼ", "'")  # This specific character from the logs

    # Remove control characters (characters 0-31 except tab, newline, carriage return)
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)

    # Replace any remaining problematic Unicode characters
    text = text.encode("ascii", errors="ignore").decode("ascii")

    return text


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text from various file formats using unified document processing."""
    from app.services.document_utils import extract_text_from_file_unified

    try:
        return extract_text_from_file_unified(file_content, filename)
    except Exception as e:
        raise HTTPException(
            status_code=400, detail=f"Error extracting text from {filename}: {str(e)}"
        )


@router.post("/generate", response_model=ReportGenieResponse)
async def generate_report(
    session: SessionDep,
    current_user: CurrentUser,
    knowledge_base_id: str = Form(...),
    sections: str = Form(...),
    outline_id: str = Form(...),
    search_mode: str = Form("vector"),  # Default to vector search
    custom_instructions: Optional[str] = Form(None),
    request: FastAPIRequest = None,
    task_id: Optional[str] = Form(None),
):
    """
    Generate a report based on sections outline and knowledge base search results.
    Includes real-time progress tracking.
    """
    try:
        # Debug: Log custom instructions if provided
        if custom_instructions:
            print(f"Custom instructions received for generate: {custom_instructions}")

        # 1. Retrieve knowledge base from database
        kb = session.get(KnowledgeBase, knowledge_base_id)
        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # Initialize the LLM
        llm = get_default_llm(session, current_user)

        # Parse sections
        try:
            sections_data = json.loads(sections)
            if isinstance(sections_data, list) and all(
                isinstance(item, dict) and "text" in item and "consultDocuments" in item
                for item in sections_data
            ):
                section_items = sections_data
            else:
                raise ValueError("Invalid sections format")
        except (json.JSONDecodeError, TypeError, ValueError):
            section_list = sections.strip().split("\n")
            section_items = [
                {"text": section.strip(), "consultDocuments": True}
                for section in section_list
                if section.strip()
            ]

        # Create progress tracking task
        if not task_id:
            task_id = progress_tracker.create_task(
                f"Generating report",
                {"setup": 0.1, "generating": 0.8, "finalizing": 0.1}
            )
        
        progress_tracker.update_stage_progress(
            task_id, "setup", 0, 1, "Initializing report generation..."
        )

        # Process each section
        sections = []
        draft_report = ""

        # Initialize knowledge base cache for this report generation
        kb_cache = KnowledgeBaseCache()
        
        progress_tracker.complete_stage(task_id, "setup", "Setup complete")
        progress_tracker.update_stage_progress(
            task_id, "generating", 0, len(section_items), "Starting section generation..."
        )

        try:
            for idx, section_item in enumerate(section_items):
                # Update progress for each section
                section_preview = section_item["text"][:50] + "..." if len(section_item["text"]) > 50 else section_item["text"]
                progress_tracker.update_stage_progress(
                    task_id, "generating", idx, len(section_items),
                    f"Processing section {idx + 1}/{len(section_items)}: {section_preview}"
                )
                
                await asyncio.sleep(0.01)  # Allow progress API to respond
                # CRITICAL: Check if client has disconnected before processing each section
                try:
                    if request and await request.is_disconnected():
                        print(f"❌ CLIENT DISCONNECTED - Stopping report generation")
                        return ReportGenieResponse(
                            results={
                                "status": "cancelled",
                                "message": "Request cancelled - client disconnected"
                            }
                        )
                except Exception as e:
                    print(f"Warning: Could not check disconnect status: {e}")
                
                section_description = section_item["text"]
                consult_documents = section_item.get("consultDocuments", True)
                # Use search_mode from form parameter, not from section_item (which doesn't have searchType)
                # search_type = section_item.get("searchType", "vector")  # ❌ WRONG - this field doesn't exist

                if not section_description:
                    continue

                # Initialize variables for this section
                section_content = ""
                source_citations = []

                if consult_documents:
                    if search_mode == "full_text":
                        # Full Text Scan Logic
                        print(f"Performing Full Text Scan for: {section_description}")
                        all_source_text = ""
                        sources = session.exec(
                            select(Source).where(Source.knowledge_base_id == kb.id)
                        ).all()
                        for source in sources:
                            # Get source data
                            source_data = session.get(SourceData, source.source_data_id)
                            if not source_data:
                                print(f"No source data found for source {source.name}")
                                continue

                            try:
                                # Extract text from the source data
                                if not source_data.data.startswith(b"PK"):
                                    # Direct file extraction
                                    file_content = extract_text_from_file(
                                        source_data.data, source.name
                                    )
                                else:
                                    # Extract from ZIP file
                                    zip_data = BytesIO(source_data.data)
                                    with zipfile.ZipFile(zip_data, "r") as zip_file:
                                        file_info = zip_file.infolist()[0]
                                        raw_file_content = zip_file.read(
                                            file_info.filename
                                        )
                                        file_content = extract_text_from_file(
                                            raw_file_content, source.name
                                        )

                                all_source_text += f"\n\n--- Source: {source.name} ---\n\n{file_content}"
                            except Exception as e:
                                print(
                                    f"Error extracting content from {source.name}: {e}"
                                )
                                # Continue with other sources instead of failing completely
                                continue

                        text_chunks = chunk_text(
                            all_source_text,
                            max_tokens=settings.FULL_SCAN_DOCUMENT_CHUNK_SIZE,
                        )
                        chunk_analyses = []
                        relevant_chunk_indices = []
                        
                        for i, chunk in enumerate(text_chunks):
                            # Add delay between chunk processing to prevent rate limit exhaustion
                            if i > 0 and settings.REPORTGENIE_ENABLE_PROCESSING_DELAYS:
                                await asyncio.sleep(settings.PROCESSING_DELAY_BETWEEN_CHUNKS)
                                
                            # CRITICAL: Check if client has disconnected before processing each chunk
                            try:
                                if request and await request.is_disconnected():
                                    print(f"❌ CLIENT DISCONNECTED - Stopping at chunk {i + 1}")
                                    return ReportGenieResponse(
                                        results={
                                            "status": "cancelled",
                                            "message": "Request cancelled - client disconnected during chunk processing"
                                        }
                                    )
                            except Exception as e:
                                print(f"Warning: Could not check disconnect status: {e}")
                            
                            # Use relevance filter to check if chunk is relevant
                            analysis = invoke_llm(
                                llm,
                                settings.VERADOC_RELEVANCE_FILTER_PROMPT_TEMPLATE,
                                {"chunk": chunk, "question": section_description},
                            )
                            
                            # CRITICAL: Check if client disconnected after LLM call
                            try:
                                if request and await request.is_disconnected():
                                    print(f"❌ CLIENT DISCONNECTED - After LLM call for chunk {i + 1}")
                                    return ReportGenieResponse(
                                        results={
                                            "status": "cancelled",
                                            "message": "Request cancelled - client disconnected after LLM call"
                                        }
                                    )
                            except Exception as e:
                                print(f"Warning: Could not check disconnect status: {e}")
                            
                            # Only include relevant chunks
                            if "No relevant information found" not in analysis:
                                chunk_analyses.append(analysis)
                                relevant_chunk_indices.append(i)

                        # Synthesize the chunk analyses
                        print(
                            f"📊 Relevance filtering: {len(chunk_analyses)} relevant chunks from {len(text_chunks)} total chunks"
                        )

                        if not chunk_analyses:
                            print("No chunk analyses found - using fallback message")
                            section_content = "No relevant information found in the knowledge base to answer this question."
                            source_citations = []
                        else:
                            chunk_analyses_text = "\n\n".join(chunk_analyses)
                            print(
                                f"Template variables: chunk_analyses={len(chunk_analyses_text)} chars, question={len(section_description)} chars"
                            )

                            try:
                                synthesized_answer = invoke_llm(
                                    llm,
                                    settings.CHATBOT_FULL_TEXT_SYNTHESIS_PROMPT_TEMPLATE,
                                    {
                                        "chunk_analyses": "\n\n".join(chunk_analyses),
                                        "question": section_description,
                                    },
                                )
                                
                                # CRITICAL: Check if client disconnected after synthesis LLM call
                                try:
                                    if request and await request.is_disconnected():
                                        print(f"❌ CLIENT DISCONNECTED - After synthesis LLM call")
                                        return ReportGenieResponse(
                                            results={
                                                "status": "cancelled",
                                                "message": "Request cancelled - client disconnected after synthesis"
                                            }
                                        )
                                except Exception as e:
                                    print(f"Warning: Could not check disconnect status: {e}")
                                
                                # Translate the synthesized answer if needed
                                section_content = await translate_text_if_needed(
                                    synthesized_answer, session, current_user, llm
                                )

                                # Create source citations from relevant chunks only
                                source_citations = []
                                for idx in relevant_chunk_indices:
                                    chunk_content = text_chunks[idx]
                                    # Truncate to 500 chars for display
                                    display_content = chunk_content[:500] + ("..." if len(chunk_content) > 500 else "")
                                    
                                    source_citations.append(
                                        {
                                            "content": display_content,
                                            "metadata": {
                                                "source": "Full Document Scan",
                                                "chunk_index": idx,
                                                "scan_type": "full_text",
                                            },
                                        }
                                    )
                            except Exception as e:
                                print(f"Error in synthesis: {e}")
                                print(
                                    f"Template: {settings.CHATBOT_FULL_TEXT_SYNTHESIS_PROMPT_TEMPLATE}"
                                )
                                raise
                    else:
                        # Vector Search Logic
                        print(f"Performing Vector Search for: {section_description}")

                        # Use cached retriever instead of creating new temp directory each time
                        retriever = kb_cache.get_retriever(
                            knowledge_base_id, kb, session, current_user
                        )

                        # Use the cached retriever's get_relevant_documents method
                        search_results = retriever.get_relevant_documents(
                            section_description
                        )

                        # Format search results for the synthesis prompt
                        context = "\n\n".join(
                            [doc.page_content for doc in search_results]
                        )

                        synthesized_answer = invoke_llm(
                            llm,
                            settings.REPORT_GENIE_PROMPT_TEMPLATE,
                            {
                                "report_draft": draft_report,
                                "context": context,
                                "question": section_description,
                                "custom_instructions": (
                                    f"\nADDITIONAL CUSTOM INSTRUCTIONS:\n{custom_instructions}\n"
                                    if custom_instructions
                                    else ""
                                ),
                            },
                        )

                        # CRITICAL: Check if client disconnected after vector search LLM call
                        try:
                            if request and await request.is_disconnected():
                                print(f"❌ CLIENT DISCONNECTED - After vector search LLM call")
                                return ReportGenieResponse(
                                    results={
                                        "status": "cancelled",
                                        "message": "Request cancelled - client disconnected after LLM call"
                                    }
                                )
                        except Exception as e:
                            print(f"Warning: Could not check disconnect status: {e}")

                        # Translate the synthesized answer if needed
                        section_content = await translate_text_if_needed(
                            synthesized_answer, session, current_user, llm
                        )

                        # Extract source citations from search results
                        source_citations = []
                        for doc in search_results:
                            # Extract metadata from the document
                            metadata = doc.metadata if hasattr(doc, "metadata") else {}

                            # Create citation entry
                            citation = {
                                "content": (
                                    doc.page_content[:500] + "..."
                                    if len(doc.page_content) > 500
                                    else doc.page_content
                                ),
                                "metadata": {
                                    "source": metadata.get("source", "Unknown"),
                                    "source_data_id": metadata.get(
                                        "source_data_id", ""
                                    ),
                                    "page": metadata.get("page", ""),
                                    "chunk_index": metadata.get("chunk_index", ""),
                                    "scan_type": "vector_search",
                                },
                            }
                            source_citations.append(citation)
                else:
                    # Use raw text directly without consulting knowledge base
                    section_content = section_description
                    source_citations = []

                section_title = section_description

                # Store the section with its content and sources
                sections.append(
                    {
                        "title": section_title,
                        "content": section_content,
                        "source_citations": source_citations,
                        "consult_documents": consult_documents,
                    }
                )
                draft_report += f"\n\n## {section_title}\n\n{section_content}"

            # 7. Compile the final report
            progress_tracker.complete_stage(task_id, "generating", "All sections generated successfully")
            progress_tracker.update_stage_progress(
                task_id, "finalizing", 0, 1, "Compiling final report..."
            )
            
            full_report = "\n\n\n\n".join(
                [section["content"].strip() for section in sections]
            )

        finally:
            # Always cleanup the knowledge base cache
            kb_cache.cleanup()

        result = {"full_report": full_report, "sections": sections, "task_id": task_id}

        # Debug logging to verify citations are being saved
        print(f"🔍 REPORTGENIE SAVE DEBUG: Saving {len(sections)} sections")
        for i, section in enumerate(sections):
            citations_count = len(section.get("source_citations", []))
            print(
                f"🔍 Section {i+1}: '{section.get('title', 'No title')}' has {citations_count} citations"
            )
            if citations_count > 0:
                print(f"🔍 Sample citation: {section['source_citations'][0]}")

        # Get outline name if outline_id is provided
        outline_name = None
        if outline_id:
            try:
                outline = session.get(ReportGenieOutline, outline_id)
                if outline:
                    outline_name = outline.name
            except Exception as e:
                print(f"Warning: Error retrieving outline name: {e}")

        # If no outline_id but we have sections, use first line or "Custom Outline"
        if not outline_name and sections:
            first_line = sections.strip().split("\n")[0]
            if first_line:
                # Extract a name from the first section (limited to 30 chars)
                outline_name = first_line[:30] + ("..." if len(first_line) > 30 else "")
            else:
                outline_name = "Custom Outline"

        # Store the full report and sections data in extra_data for retrieval later
        detailed_extra_data = {
            "kb_name": kb.title,
            "full_report": full_report,
            "sections": sections,  # This includes section content and sources
            "outline_name": outline_name,  # Add the outline name here
        }

        interaction_id = record_llm_interaction(
            session=session,
            user_id=current_user.id,
            functionality="reportgenie",
            input_data={
                "knowledge_base_id": knowledge_base_id,
                "sections": sections_data,
                "outline_id": outline_id,
                "search_mode": search_mode,
                "kb_name": kb.title,
            },
            output_data=result,  # ← Use result directly, not nested under "results"
            metadata={
                "kb_name": kb.title,
                "kb_id": knowledge_base_id,
                "sections": sections_data,
                "search_mode": search_mode,
                "outline_id": outline_id,
                "full_report": full_report,
                "section_count": len(
                    result.get("sections", [])
                ),  # Move metrics to metadata
                "total_length": len(full_report),
            },
        )

        print(f"[DEBUG] ReportGenie interaction_id returned: {interaction_id}")
        # Add interaction_id to the result
        result["interaction_id"] = str(interaction_id) if interaction_id else None
        print(
            f"[DEBUG] ReportGenie result with interaction_id: {result.get('interaction_id')}"
        )
        
        # Complete the progress tracking
        progress_tracker.complete_stage(task_id, "finalizing", "Report generation complete!")

        return ReportGenieResponse(results=result)

    except Exception as e:
        import traceback

        traceback.print_exc()
        
        # Mark progress as failed if task_id exists
        if 'task_id' in locals() and task_id:
            progress_tracker.fail_task(task_id, f"Report generation failed: {str(e)}")
        
        raise HTTPException(
            status_code=500, detail=f"Error generating report: {str(e)}"
        )


@router.delete("/reports/{report_id}", response_model=Message)
def delete_report(report_id: uuid.UUID, session: SessionDep, current_user: CurrentUser):
    """
    Delete a report by ID.
    """
    report = session.get(LlmInteraction, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="Report not found.")

    # Ensure the current user is the owner of the report
    if report.user_id != current_user.id:
        raise HTTPException(
            status_code=403, detail="Not authorized to delete this report."
        )

    # Only allow deletion of reportgenie reports
    if report.functionality != "reportgenie":
        raise HTTPException(status_code=400, detail="Invalid report type.")

    session.delete(report)
    session.commit()
    return Message(message="Report deleted successfully.")


@router.get("/history", response_model=List[Dict[str, Any]])
async def get_report_history(
    session: SessionDep,
    current_user: CurrentUser,
    skip: int = 0,
    limit: int = 20,
    show_all: bool = False,
):
    """Retrieve past ReportGenie generation history for the current user or all users."""
    print("Retrieving ReportGenie history. Show all:", show_all)

    try:
        # Start with base query
        query = select(LlmInteraction).where(
            LlmInteraction.functionality == "reportgenie"
        )

        # Only filter by user if not showing all users
        if not show_all:
            query = query.where(LlmInteraction.user_id == current_user.id)

        # Add ordering and pagination
        interactions = session.exec(
            query.order_by(LlmInteraction.date_created.desc()).offset(skip).limit(limit)
        ).all()

        result = []
        for interaction in interactions:
            try:
                # Safe parsing of JSON fields with better error handling
                input_data = {}
                output_data = {}
                metadata = {}

                # Parse input_data safely
                if interaction.input_data:
                    if isinstance(interaction.input_data, str):
                        input_data = json.loads(interaction.input_data)
                    elif isinstance(interaction.input_data, dict):
                        input_data = interaction.input_data
                    else:
                        print(
                            f"Unexpected input_data type: {type(interaction.input_data)}"
                        )
                        input_data = {}

                # Parse output_data safely
                if interaction.output_data:
                    if isinstance(interaction.output_data, str):
                        output_data = json.loads(interaction.output_data)
                    elif isinstance(interaction.output_data, dict):
                        output_data = interaction.output_data
                    else:
                        print(
                            f"Unexpected output_data type: {type(interaction.output_data)}"
                        )
                        output_data = {}

                # Parse metadata safely with enhanced error handling
                if interaction.metadata:
                    if isinstance(interaction.metadata, str):
                        metadata = json.loads(interaction.metadata)
                    elif isinstance(interaction.metadata, dict):
                        metadata = interaction.metadata
                    else:
                        # Handle the MetaData object case
                        print(f"Metadata is of type: {type(interaction.metadata)}")
                        if hasattr(interaction.metadata, "__dict__"):
                            # Convert object to dict if possible
                            metadata = interaction.metadata.__dict__
                        else:
                            # Try to serialize it as string and then parse
                            try:
                                metadata_str = str(interaction.metadata)
                                if metadata_str.startswith(
                                    "{"
                                ) and metadata_str.endswith("}"):
                                    metadata = json.loads(metadata_str)
                                else:
                                    metadata = {}
                            except:
                                print(
                                    f"Failed to parse metadata: {interaction.metadata}"
                                )
                                metadata = {}

                # Handle sections data - ensure it's always a string for API response
                sections_data = input_data.get("sections", "")
                if not isinstance(sections_data, str):
                    sections_data = json.dumps(sections_data) if sections_data else ""

                # Create result item with proper field names for archive display
                result_item = {
                    "id": str(interaction.id),
                    "date_created": interaction.date_created,
                    "knowledge_base_id": input_data.get("knowledge_base_id", ""),
                    "sections": sections_data,
                    "kb_name": metadata.get(
                        "kb_name", input_data.get("kb_name", "Unknown")
                    ),
                    "outline_id": input_data.get("outline_id", ""),
                    "search_mode": input_data.get("search_mode", "vector"),
                    "full_report": metadata.get(
                        "full_report", output_data.get("full_report", "")
                    ),
                    "section_count": output_data.get("section_count", 0),
                    "total_length": output_data.get("total_length", 0),
                    "has_feedback": interaction.feedback is not None,
                }

                # Add feedback information if exists
                if interaction.feedback:
                    result_item["feedback"] = {
                        "feedback": interaction.feedback,
                        "feedbackText": interaction.feedback_text,
                        "feedbackDate": interaction.feedback_date,
                    }

                # Add user info for all-users view
                if show_all:
                    from app.models import User  # Import here to avoid circular imports

                    user = session.get(User, interaction.user_id)
                    user_name = (
                        f"{user.full_name or 'User'} ({user.email})"
                        if user
                        else "Unknown User"
                    )
                    result_item["user_name"] = user_name

                result.append(result_item)

            except json.JSONDecodeError:
                # Handle legacy entries or malformed JSON
                result_item = {
                    "id": str(interaction.id),
                    "date_created": interaction.date_created,
                    "knowledge_base_id": "",
                    "sections": "",
                    "kb_name": "Unknown",
                    "outline_id": "",
                    "search_mode": "vector",
                    "full_report": "",
                    "section_count": 0,
                    "total_length": 0,
                    "has_feedback": interaction.feedback is not None,
                }

                # Add user info for all-users view
                if show_all:
                    from app.models import User

                    user = session.get(User, interaction.user_id)
                    user_name = (
                        f"{user.full_name or 'User'} ({user.email})"
                        if user
                        else "Unknown User"
                    )
                    result_item["user_name"] = user_name

                result.append(result_item)

        return result

    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Error retrieving ReportGenie history: {str(e)}",
        )


@router.get("/detail/{report_id}", response_model=Dict[str, Any])
async def get_report_detail(
    report_id: str,
    session: SessionDep,
    # current_user: CurrentUser,
):
    """Retrieve a specific ReportGenie report's full content by ID."""
    try:
        # Convert string ID to UUID for database query
        interaction_id = uuid.UUID(report_id)

        # Get the specific interaction
        interaction = session.exec(
            select(LlmInteraction)
            .where(LlmInteraction.id == interaction_id)
            .where(LlmInteraction.functionality == "reportgenie")
            # .where(LlmInteraction.user_id == current_user.id)
        ).first()

        if not interaction:
            raise HTTPException(status_code=404, detail="Report not found")

        # Safe parsing of JSON fields with better error handling
        input_data = {}
        output_data = {}
        metadata = {}

        # Parse input_data safely
        if interaction.input_data:
            if isinstance(interaction.input_data, str):
                input_data = json.loads(interaction.input_data)
            elif isinstance(interaction.input_data, dict):
                input_data = interaction.input_data
            else:
                print(f"Unexpected input_data type: {type(interaction.input_data)}")
                input_data = {}

        # Parse output_data safely
        if interaction.output_data:
            if isinstance(interaction.output_data, str):
                output_data = json.loads(interaction.output_data)
            elif isinstance(interaction.output_data, dict):
                output_data = interaction.output_data
            else:
                print(f"Unexpected output_data type: {type(interaction.output_data)}")
                output_data = {}

        # Parse metadata safely with enhanced error handling
        if interaction.metadata:
            if isinstance(interaction.metadata, str):
                metadata = json.loads(interaction.metadata)
            elif isinstance(interaction.metadata, dict):
                metadata = interaction.metadata
            else:
                # Handle the MetaData object case
                print(f"Metadata is of type: {type(interaction.metadata)}")
                if hasattr(interaction.metadata, "__dict__"):
                    # Convert object to dict if possible
                    metadata = interaction.metadata.__dict__
                else:
                    # Try to serialize it as string and then parse
                    try:
                        metadata_str = str(interaction.metadata)
                        if metadata_str.startswith("{") and metadata_str.endswith("}"):
                            metadata = json.loads(metadata_str)
                        else:
                            metadata = {}
                    except:
                        print(f"Failed to parse metadata: {interaction.metadata}")
                        metadata = {}

        # Handle sections data - ensure it's always a string for API response
        sections_data = input_data.get("sections", "")
        if not isinstance(sections_data, str):
            sections_data = json.dumps(sections_data) if sections_data else ""

        result_data = {
            "id": str(interaction.id),
            "date_created": interaction.date_created,
            "knowledge_base_id": input_data.get("knowledge_base_id", ""),
            "sections": sections_data,
            "kb_name": metadata.get("kb_name", input_data.get("kb_name", "Unknown")),
            "outline_id": input_data.get("outline_id", ""),
            "search_mode": input_data.get("search_mode", "vector"),
            "full_report": metadata.get(
                "full_report", output_data.get("full_report", "")
            ),
            "results": output_data,
            "section_count": output_data.get("section_count", 0),
            "total_length": output_data.get("total_length", 0),
            "has_feedback": interaction.feedback is not None,
        }

        # Debug logging for citation retrieval
        sections_with_citations = output_data.get("sections", [])
        if sections_with_citations:
            print(
                f"🔍 REPORTGENIE RETRIEVE DEBUG: Retrieved {len(sections_with_citations)} sections"
            )
            for i, section in enumerate(sections_with_citations):
                if isinstance(section, dict):
                    citations_count = len(section.get("source_citations", []))
                    print(
                        f"🔍 Retrieved Section {i+1}: '{section.get('title', 'No title')}' has {citations_count} citations"
                    )
        else:
            print("🔍 REPORTGENIE RETRIEVE DEBUG: No sections found in output_data")

        return result_data

    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid report ID format")
    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Error retrieving report detail: {str(e)}",
        )


# Functions related to Outlines
@router.post("/outlines", response_model=ReportGenieOutline)
def create_outline(
    outline: ReportGenieOutline, session: SessionDep, current_user: CurrentUser
):
    """
    Save a new outline to the database.
    """
    existing_outline = session.exec(
        select(ReportGenieOutline).where(ReportGenieOutline.name == outline.name)
    ).first()
    if existing_outline:
        raise HTTPException(
            status_code=400, detail="An outline with this name already exists."
        )

    outline.owner_id = current_user.id
    session.add(outline)
    session.commit()
    session.refresh(outline)
    return outline


@router.get("/outlines", response_model=List[ReportGenieOutline])
def get_outlines(session: SessionDep, current_user: CurrentUser):
    """
    Retrieve all outlines from the database for this user.
    """
    print(f"Retrieving outlines for user {current_user.id}")
    if not current_user:
        raise HTTPException(status_code=401, detail="Not authenticated.")

    try:
        outlines = session.exec(
            select(ReportGenieOutline).where(
                ReportGenieOutline.owner_id == current_user.id
            )
        ).all()

        # Print the retrieved outlines for debugging
        print(f"Found {len(outlines)} outlines for user {current_user.id}:")
        for i, outline in enumerate(outlines):
            try:
                section_count = (
                    len(outline.sections.split("\n")) if outline.sections else 0
                )
                print(
                    f"  {i+1}. ID: {outline.id}, Name: {outline.name}, Sections: {section_count} sections"
                )
            except Exception as e:
                print(f"  {i+1}. ID: {outline.id}, Error processing outline: {str(e)}")

        return outlines
    except Exception as e:
        print(f"Error retrieving outlines: {str(e)}")
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error retrieving outlines: {str(e)}"
        )


@router.get("/outlines/{outline_id}", response_model=ReportGenieOutline)
def get_outline(outline_id: uuid.UUID, session: SessionDep):
    """
    Retrieve a specific outline by ID.
    """
    outline = session.get(ReportGenieOutline, outline_id)
    if not outline:
        raise HTTPException(status_code=404, detail="Outline not found.")
    return outline


@router.put("/outlines/{outline_id}", response_model=ReportGenieOutline)
def update_outline(
    outline_id: uuid.UUID,
    updated_outline: ReportGenieOutline,
    session: SessionDep,
    current_user: CurrentUser,
):
    """
    Update an existing outline.
    """
    outline = session.get(ReportGenieOutline, outline_id)
    if not outline:
        raise HTTPException(status_code=404, detail="Outline not found.")

    # Ensure the current user is the owner of the outline
    if outline.owner_id != current_user.id:
        raise HTTPException(
            status_code=403, detail="Not authorized to update this outline."
        )

    outline.name = updated_outline.name
    outline.description = updated_outline.description
    outline.sections = updated_outline.sections
    outline.date_modified = datetime.utcnow()

    session.add(outline)
    session.commit()
    session.refresh(outline)
    return outline


@router.delete("/outlines/{outline_id}", response_model=Message)
def delete_outline(
    outline_id: uuid.UUID, session: SessionDep, current_user: CurrentUser
):
    """
    Delete an outline by ID.
    """
    outline = session.get(ReportGenieOutline, outline_id)
    if not outline:
        raise HTTPException(status_code=404, detail="Outline not found.")

    # Ensure the current user is the owner of the outline
    if outline.owner_id != current_user.id:
        raise HTTPException(
            status_code=403, detail="Not authorized to delete this outline."
        )

    session.delete(outline)
    session.commit()
    return Message(message="Outline deleted successfully.")


@router.post("/generate-outline-json", response_model=GenerateOutlineResponse)
async def generate_outline_json(
    session: SessionDep, current_user: CurrentUser, request: GenerateOutlineRequest
):
    """
    Generate outline sections based on a description using LLM, with optional knowledge base reference (JSON version).
    """
    try:
        # Get the default LLM
        llm = get_default_llm(session, current_user)

        # Handle optional description
        description = request.description or ""

        # Prepare variables for the prompt
        prompt_variables = {
            "description": description,
            "report_type": request.report_type or "general",
            "example_document": "",
            "example_instruction": "",
            "example_analysis_instruction": "",
            "knowledge_base_instruction": "",
            "knowledge_base_content": "",
        }

        # If knowledge base is specified, retrieve content using selected search mode
        if request.knowledge_base_id:
            try:
                from app.services.content_retrieval import (
                    retrieve_knowledge_base_content,
                )

                content, instruction = await retrieve_knowledge_base_content(
                    session=session,
                    current_user=current_user,
                    knowledge_base_id=request.knowledge_base_id,
                    search_mode=request.search_mode,
                    query=description,
                )

                if content:
                    prompt_variables["knowledge_base_content"] = (
                        f"REFERENCE DOCUMENTS FROM KNOWLEDGE BASE:\n{content}"
                    )
                    prompt_variables["knowledge_base_instruction"] = (
                        f"\n12. {instruction} Use them as examples for the type of structure, "
                        f"sections, and content organization that would be appropriate. "
                        f"Search mode used: {request.search_mode}"
                    )
                    prompt_variables["example_analysis_instruction"] = (
                        f". Briefly mention how the knowledge base content (using {request.search_mode}) influenced the outline structure"
                    )

            except Exception as e:
                print(f"Error retrieving from knowledge base: {str(e)}")
                # Continue without knowledge base content rather than failing
                pass

        # Generate outline using the LLM
        outline_response = invoke_llm(
            llm,
            settings.REPORTGENIE_GENERATE_OUTLINE_PROMPT_TEMPLATE,
            prompt_variables,
        )

        # Parse the response to extract sections and analysis
        sections = []
        analysis = ""

        lines = outline_response.strip().split("\n")
        in_sections_section = False
        in_analysis_section = False

        for line in lines:
            line = line.strip()
            if line.startswith("SECTIONS:"):
                in_sections_section = True
                in_analysis_section = False
                continue
            elif line.startswith("ANALYSIS:"):
                in_sections_section = False
                in_analysis_section = True
                continue

            if in_sections_section:
                # Extract sections (numbered list)
                if re.match(r"^\d+\.\s+", line):
                    section = re.sub(r"^\d+\.\s+", "", line)
                    if section.strip():
                        sections.append(section.strip())
            elif in_analysis_section:
                if line:
                    if analysis:
                        analysis += " " + line
                    else:
                        analysis = line

        # If parsing failed, try simpler approach
        if not sections:
            # Split by lines and look for numbered items
            for line in lines:
                line = line.strip()
                if re.match(r"^\d+\.\s+", line):
                    section = re.sub(r"^\d+\.\s+", "", line)
                    if section.strip():
                        sections.append(section.strip())

        # Ensure we have some sections
        if not sections:
            raise HTTPException(
                status_code=500,
                detail="Failed to generate outline from the description. Please try with a more detailed description.",
            )

        # Apply user-specified limit if provided, otherwise use all generated sections
        if request.num_sections:
            sections = sections[: request.num_sections]

        if not analysis:
            search_method = (
                "vector search"
                if request.search_mode == "vector"
                else "full document scan"
            )
            analysis = f"Generated {len(sections)} outline sections based on the provided description using {search_method}"
            if request.knowledge_base_id:
                analysis += " with knowledge base reference."

        # Record the interaction
        record_llm_interaction(
            session=session,
            user_id=current_user.id,
            functionality="generate_outline",
            input_data={
                "description": request.description,
                "report_type": request.report_type,
                "requested_sections": request.num_sections,
                "knowledge_base_id": request.knowledge_base_id,
                "search_mode": request.search_mode,
            },
            output_data={
                "sections_count": len(sections),
                "analysis": analysis,
            },
            metadata={},
        )

        return GenerateOutlineResponse(sections=sections, description_analysis=analysis)

    except Exception as e:
        print(f"Error generating outline: {e}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error generating outline: {str(e)}"
        )


@router.post("/generate-outline", response_model=GenerateOutlineResponse)
async def generate_outline(
    session: SessionDep,
    current_user: CurrentUser,
    description: str = Form(...),
    report_type: str = Form(default="general"),
    num_sections: Optional[int] = Form(default=None),
    files: List[UploadFile] = File(default=[]),
    task_id: Optional[str] = Form(None),
):
    """
    Generate outline sections based on a description using LLM, with optional example document.
    Includes real-time progress tracking.
    """
    try:
        # Create progress tracking task
        if not task_id:
            task_id = progress_tracker.create_task(
                f"Generating outline",
                {"processing_files": 0.2, "generating": 0.7, "finalizing": 0.1}
            )
        
        progress_tracker.update_stage_progress(
            task_id, "processing_files", 0, 1, "Initializing outline generation..."
        )
        
        # Get the default LLM
        llm = get_default_llm(session, current_user)

        # Process uploaded files to extract example document content
        example_document_content = ""
        if files:
            print(f"Processing {len(files)} uploaded files for outline generation")
            progress_tracker.update_stage_progress(
                task_id, "processing_files", 0, len(files), f"Processing {len(files)} uploaded files..."
            )

            for idx, file in enumerate(files):
                progress_tracker.update_stage_progress(
                    task_id, "processing_files", idx, len(files), f"Processing file {idx+1}/{len(files)}: {file.filename}..."
                )
                if file.size > 0:
                    try:
                        # Read and process the file content with visual enhancement
                        file_content = await file.read()

                        # Use enhanced processing with vision capabilities
                        from app.services.document_utils import (
                            extract_text_with_vision_enhancement,
                        )

                        file_text = await extract_text_with_vision_enhancement(
                            file_content,
                            file.filename or "unknown",
                            llm,
                            purpose="outline generation",
                        )

                        if file_text.strip():
                            example_document_content += f"\n\n--- Content from {file.filename} ---\n{file_text.strip()}"
                            print(
                                f"Extracted {len(file_text)} characters from {file.filename} (with vision enhancement)"
                            )

                    except Exception as e:
                        print(f"Error processing file {file.filename}: {str(e)}")
                        # Add the error info to the document content so user knows what happened
                        example_document_content += f"\n\n--- Error processing {file.filename} ---\nError: {str(e)}\n"
                        continue
            
            progress_tracker.complete_stage(task_id, "processing_files", f"Processed {len(files)} files")
        else:
            progress_tracker.complete_stage(task_id, "processing_files", "No files to process")

        # Check if content exceeds token limits and chunk if necessary
        if example_document_content:
            print(
                f"Total example document content: {len(example_document_content)} characters"
            )

            # Using conservative chunking similar to TWINCHECK settings
            max_chunk_size = 80000  # Conservative chunk size for 128K context limit

            if len(example_document_content) > max_chunk_size:
                print(
                    f"Example document too large ({len(example_document_content)} chars), chunking for processing"
                )

                from app.services.text_processing import chunk_text

                # Chunk the document content
                chunks = chunk_text(example_document_content, max_tokens=max_chunk_size)
                
                progress_tracker.update_stage_progress(
                    task_id, "generating", 0, len(chunks), f"Processing {len(chunks)} document chunks..."
                )

                # Process each chunk to generate sections
                all_chunk_sections = []

                for i, chunk in enumerate(chunks):
                    progress_tracker.update_stage_progress(
                        task_id, "generating", i, len(chunks), f"Analyzing chunk {i+1}/{len(chunks)}..."
                    )
                    # Add delay between chunk processing to prevent rate limit exhaustion
                    if i > 0 and settings.REPORTGENIE_ENABLE_PROCESSING_DELAYS:
                        await asyncio.sleep(settings.PROCESSING_DELAY_BETWEEN_CHUNKS)
                        
                    print(f"Processing chunk {i+1}/{len(chunks)}")

                    # Generate sections for this chunk
                    chunk_prompt_variables = {
                        "description": description,
                        "report_type": report_type,
                        "example_document": f"EXAMPLE DOCUMENT FOR REFERENCE:\n{chunk}",
                        "example_instruction": "\n12. Use the example document provided above as inspiration for the type of content organization and structure, but adapt the sections to match the specific requirements in the outline description",
                        "example_analysis_instruction": ". Briefly mention how the example document influenced the structure",
                        "knowledge_base_content": "",
                        "knowledge_base_instruction": "",
                    }

                    try:
                        chunk_response = invoke_llm(
                            llm,
                            settings.REPORTGENIE_GENERATE_OUTLINE_PROMPT_TEMPLATE,
                            chunk_prompt_variables,
                        )

                        # Parse sections from chunk response
                        chunk_sections = []
                        lines = chunk_response.strip().split("\n")
                        in_sections_section = False

                        for line in lines:
                            line = line.strip()
                            if line.startswith("SECTIONS:"):
                                in_sections_section = True
                                continue
                            elif line.startswith("ANALYSIS:"):
                                in_sections_section = False
                                continue

                            if in_sections_section:
                                if re.match(r"^\d+\.\s+", line):
                                    section = re.sub(r"^\d+\.\s+", "", line)
                                    if section.strip():
                                        chunk_sections.append(section.strip())

                        # If parsing failed, try simpler approach
                        if not chunk_sections:
                            for line in lines:
                                line = line.strip()
                                if re.match(r"^\d+\.\s+", line):
                                    section = re.sub(r"^\d+\.\s+", "", line)
                                    if section.strip():
                                        chunk_sections.append(section.strip())

                        all_chunk_sections.extend(chunk_sections)

                    except Exception as e:
                        print(f"Error processing chunk {i+1}: {e}")
                        continue

                # Deduplicate and refine sections across all chunks
                if all_chunk_sections:
                    # Remove duplicates while preserving order
                    seen = set()
                    unique_sections = []
                    for s in all_chunk_sections:
                        if s.lower() not in seen:
                            seen.add(s.lower())
                            unique_sections.append(s)

                    # If we have too many sections, synthesize and prioritize
                    if len(unique_sections) > (num_sections or 15):
                        synthesis_prompt = f"""From the following list of outline sections, select and refine the {num_sections or 8} most important and relevant sections for a {report_type} report based on: {description}

Sections to review:
{chr(10).join([f"{i+1}. {s}" for i, s in enumerate(unique_sections)])}

Requirements:
1. Select the most critical and comprehensive sections
2. Ensure logical flow and organization
3. Maintain clarity and specificity
4. Focus on sections most relevant to the description

Return only the final selected sections, one per line, numbered."""

                        try:
                            refined_response = invoke_llm(llm, synthesis_prompt, {})
                            sections = []
                            for line in refined_response.strip().split("\n"):
                                line = line.strip()
                                if line and (
                                    line[0].isdigit()
                                    or line.startswith("-")
                                    or line.startswith("*")
                                ):
                                    section = re.sub(r"^\d+\.\s+", "", line)
                                    section = re.sub(r"^[-*]\s+", "", section)
                                    if section.strip():
                                        sections.append(section.strip())
                        except Exception as e:
                            print(f"Error in section synthesis: {e}")
                            sections = unique_sections[: num_sections or 8]
                    else:
                        sections = unique_sections[: num_sections or 15]

                    # For analysis, show chunked processing was used
                    analysis = f"Generated {len(sections)} sections from chunked example document analysis ({len(chunks)} chunks processed) based on the provided description to ensure comprehensive report structure coverage."

                    # Record the interaction
                    record_llm_interaction(
                        session=session,
                        user_id=current_user.id,
                        functionality="generate_outline_sections",
                        input_data={
                            "description": description,
                            "requested_sections": num_sections,
                            "report_type": report_type,
                            "has_example_document": True,
                            "chunked_processing": True,
                            "chunk_count": len(chunks),
                        },
                        output_data={
                            "sections_count": len(sections),
                            "analysis": analysis,
                        },
                        metadata={},
                    )

                    return GenerateOutlineResponse(
                        sections=sections, description_analysis=analysis
                    )
                else:
                    # Fallback to description-only generation if chunk processing failed
                    example_document_content = ""

        # Continue with existing logic for small documents or when no files provided
        if example_document_content:
            example_document_section = (
                f"EXAMPLE DOCUMENT FOR REFERENCE:\n{example_document_content}"
            )
            example_instruction = "\n12. Use the example document provided above as inspiration for the type of content organization and structure, but adapt the sections to match the specific requirements in the outline description"
            example_analysis_instruction = (
                ". Briefly mention how the example document influenced the structure"
            )
        else:
            example_document_section = ""
            example_instruction = ""
            example_analysis_instruction = ""

        prompt_variables = {
            "description": description,
            "report_type": report_type,
            "example_document": example_document_section,
            "example_instruction": example_instruction,
            "example_analysis_instruction": example_analysis_instruction,
            "knowledge_base_content": "",
            "knowledge_base_instruction": "",
        }

        # Format the prompt with variables to show what exactly is being sent to the LLM
        formatted_prompt = settings.REPORTGENIE_GENERATE_OUTLINE_PROMPT_TEMPLATE.format(
            **prompt_variables
        )
        print("=== OUTLINE GENERATION PROMPT SENT TO LLM ===")
        print(formatted_prompt)
        print("=== END OF PROMPT ===")
        
        progress_tracker.update_stage_progress(
            task_id, "generating", 0, 1, "Generating outline sections with LLM..."
        )

        # Generate outline sections using the LLM
        outline_response = invoke_llm(
            llm,
            settings.REPORTGENIE_GENERATE_OUTLINE_PROMPT_TEMPLATE,
            prompt_variables,
        )
        
        progress_tracker.complete_stage(task_id, "generating", "Outline generated successfully")
        progress_tracker.update_stage_progress(
            task_id, "finalizing", 0, 1, "Parsing and finalizing sections..."
        )

        # Parse the response to extract sections and analysis
        sections = []
        analysis = ""

        lines = outline_response.strip().split("\n")
        in_sections_section = False
        in_analysis_section = False

        for line in lines:
            line = line.strip()
            if line.startswith("SECTIONS:"):
                in_sections_section = True
                in_analysis_section = False
                continue
            elif line.startswith("ANALYSIS:"):
                in_sections_section = False
                in_analysis_section = True
                continue

            if in_sections_section:
                # Extract sections (numbered list)
                if re.match(r"^\d+\.\s+", line):
                    section = re.sub(r"^\d+\.\s+", "", line)
                    if section.strip():
                        sections.append(section.strip())
            elif in_analysis_section:
                if line:
                    if analysis:
                        analysis += " " + line
                    else:
                        analysis = line

        # If parsing failed, try simpler approach
        if not sections:
            # Split by lines and look for numbered items
            for line in lines:
                line = line.strip()
                if re.match(r"^\d+\.\s+", line):
                    section = re.sub(r"^\d+\.\s+", "", line)
                    if section.strip():
                        sections.append(section.strip())

        # Ensure we have some sections
        if not sections:
            raise HTTPException(
                status_code=500,
                detail="Failed to generate sections from the description. Please try with a more detailed description.",
            )

        # Apply user-specified limit if provided, otherwise use all generated sections
        if num_sections:
            sections = sections[:num_sections]

        if not analysis:
            analysis = f"Generated {len(sections)} sections based on the provided description to ensure comprehensive report structure coverage."

        # Record the interaction
        record_llm_interaction(
            session=session,
            user_id=current_user.id,
            functionality="generate_outline_sections",
            input_data={
                "description": description,
                "requested_sections": num_sections,
                "report_type": report_type,
                "has_example_document": bool(example_document_content),
            },
            output_data={
                "sections_count": len(sections),
                "analysis": analysis,
            },
            metadata={},
        )
        
        progress_tracker.complete_stage(task_id, "finalizing", "Outline generation complete!")

        return GenerateOutlineResponse(sections=sections, description_analysis=analysis, task_id=task_id)

    except Exception as e:
        print(f"Error generating outline: {e}")
        traceback.print_exc()
        
        # Mark progress as failed if task_id exists
        if 'task_id' in locals() and task_id:
            progress_tracker.fail_task(task_id, f"Outline generation failed: {str(e)}")
        
        raise HTTPException(
            status_code=500, detail=f"Error generating outline: {str(e)}"
        )


@router.post("/optimize-outline", response_model=OptimizedOutlineResponse)
async def optimize_outline(
    session: SessionDep,
    current_user: CurrentUser,
    knowledge_base_id: str = Form(...),
    outline_id: str = Form(...),
    sections: str = Form(...),
    custom_instructions: Optional[str] = Form(None),
    search_mode: str = Form("vector"),  # Add search_mode as Form parameter
    files: List[UploadFile] = File(...),
    request: FastAPIRequest = None,
    task_id: Optional[str] = Form(None),
):
    """
    Optimize outline sections by testing them against a ground-truth document.
    Generates a report with current outline and compares it to the ground-truth to suggest improvements.
    Includes real-time progress tracking.
    """
    print("optimize_outline function invoked!")
    cancellation_requested = False

    try:
        # Create progress tracking task
        if not task_id:
            task_id = progress_tracker.create_task(
                f"Optimizing outline",
                {
                    "setup": 0.1,
                    "processing_document": 0.1,
                    "generating": 0.4,
                    "matching": 0.2,
                    "comparing": 0.15,
                    "finalizing": 0.05
                }
            )
        
        progress_tracker.update_stage_progress(
            task_id, "setup", 0, 1, "Initializing outline optimization..."
        )
        
        print("Setting up disconnect monitor for outline optimization...")
        disconnect_monitor = None
        if request:

            async def monitor_client_disconnect():
                nonlocal cancellation_requested
                try:
                    await request.is_disconnected()
                    print("Client disconnected, canceling optimization...")
                    cancellation_requested = True
                except asyncio.CancelledError:
                    print("Disconnect monitor cancelled because main task completed")
                except Exception as e:
                    print(f"Error in disconnect monitoring: {str(e)}")

            disconnect_monitor = asyncio.create_task(monitor_client_disconnect())

        print("Starting outline optimization...")

        # Debug: Log custom instructions
        if custom_instructions:
            print(f"Custom instructions received: {custom_instructions}")
            print(
                "✓ Custom instructions will be applied to content generation and optimization analysis"
            )
        else:
            print("No custom instructions provided - using default prompts")

        # 1. Retrieve knowledge base
        kb = session.get(KnowledgeBase, knowledge_base_id)
        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")
        
        progress_tracker.complete_stage(task_id, "setup", "Setup complete")
        progress_tracker.update_stage_progress(
            task_id, "processing_document", 0, 1, "Processing ground-truth document..."
        )

        # 2. Set up the same infrastructure as generate_report
        with tempfile.TemporaryDirectory() as temp_dir:
            # Extract ChromaDB
            if kb.storage_type == "file" and kb.file_path:
                # File-based storage: extract from file path
                if os.path.exists(kb.file_path):
                    with zipfile.ZipFile(kb.file_path, "r") as zip_ref:
                        zip_ref.extractall(temp_dir)
                else:
                    raise HTTPException(
                        status_code=400, detail="Knowledge base file not found on disk"
                    )
            elif kb.data:
                # Database storage: extract from data field
                with zipfile.ZipFile(BytesIO(kb.data), "r") as zip_ref:
                    zip_ref.extractall(temp_dir)
            else:
                raise HTTPException(
                    status_code=400, detail="Knowledge base has no vector database data"
                )

            # Load embeddings and vector database
            if kb.embedding_model_id:
                embedding_model = session.get(EmbeddingModel, kb.embedding_model_id)
                if embedding_model:
                    model_id = embedding_model.model_id
                    provider = embedding_model.provider
                else:
                    embedding_info = get_embedding_model(session, current_user)
                    model_id = embedding_info["model_id"]
                    provider = embedding_info["provider"]
            else:
                embedding_info = get_embedding_model(session, current_user)
                model_id = embedding_info["model_id"]
                provider = embedding_info["provider"]

            embeddings = load_embeddings_model(provider=provider, model_id=model_id)
            chroma_db = Chroma(
                persist_directory=temp_dir, embedding_function=embeddings
            )

            # Create retriever
            retriever = create_ensemble_retriever(
                chroma_db=chroma_db,
                vector_weight=0.7,
                keyword_weight=0.3,
                search_kwargs={"k": settings.RAG_NUM_CHUNKS},
            )

            # Initialize LLM
            llm = get_default_llm(session, current_user)

            # 3. Process the ground-truth document
            file = files[0]
            content = await file.read()
            ground_truth_text = extract_text_from_file(content, file.filename)

            # NEW: Sanitize the ground-truth text to prevent JSON parsing issues
            ground_truth_text = sanitize_text_for_json(ground_truth_text)

            print(f"DIAGNOSTIC INFO:")
            print(
                f"- Number of sections: {len(current_sections) if 'current_sections' in locals() else 'Not yet parsed'}"
            )
            print(f"- Ground truth document size: {len(ground_truth_text)} characters")
            print(f"- Sanitization applied to prevent JSON parsing issues")

            print(
                f"Processing ground-truth document: {file.filename} ({len(ground_truth_text)} characters)"
            )
            
            progress_tracker.complete_stage(task_id, "processing_document", "Ground-truth document processed")

            # 4. Parse current outline sections
            current_sections = json.loads(sections)
            print(f"Optimizing {len(current_sections)} outline sections...")
            
            progress_tracker.update_stage_progress(
                task_id, "generating", 0, len(current_sections), "Starting section generation..."
            )

            # Add limits to prevent oversized processing
            # MAX_SECTIONS = 50  # Adjust based on testing
            MAX_DOCUMENT_SIZE = 500000  # 500KB limit

            # if len(current_sections) > MAX_SECTIONS:
            #    raise HTTPException(
            #        status_code=413,
            #        detail=f"Too many sections. Maximum {MAX_SECTIONS} allowed.",
            #    )

            if len(ground_truth_text) > MAX_DOCUMENT_SIZE:
                raise HTTPException(
                    status_code=413,
                    detail=f"Document too large. Maximum {MAX_DOCUMENT_SIZE} characters allowed.",
                )

            print(
                f"✓ Size validation passed: {len(current_sections)} sections, {len(ground_truth_text)} chars"
            )

            # 5. Generate report content for each section using current outline
            generated_sections = {}
            section_consult_settings = {}  # Track which sections consult documents

            for i, section in enumerate(current_sections):
                # Update progress for section generation
                section_preview = section["text"][:40] + "..." if len(section["text"]) > 40 else section["text"]
                progress_tracker.update_stage_progress(
                    task_id, "generating", i, len(current_sections),
                    f"Generating section {i + 1}/{len(current_sections)}: {section_preview}"
                )
                
                # Add delay between section processing to prevent rate limit exhaustion
                if i > 0 and settings.REPORTGENIE_ENABLE_PROCESSING_DELAYS:
                    await asyncio.sleep(settings.PROCESSING_DELAY_BETWEEN_REQUESTS)
                    
                if cancellation_requested:
                    print("Operation cancelled by client disconnect")
                    raise HTTPException(
                        status_code=408, detail="Operation cancelled by user"
                    )

                section_description = section["text"].strip()
                consult_documents = section.get("consultDocuments", True)

                # Store the consult documents setting for this section
                section_consult_settings[section_description] = consult_documents

                if not section_description:
                    continue

                print(f"Generating content for section: {section_description[:50]}...")

                if consult_documents:
                    # Check search mode to determine which method to use
                    # Use the search_mode parameter directly

                    if search_mode == "full_text":
                        # Full Text Scan Logic (similar to main generate functionality)
                        print(
                            f"Performing Full Text Scan for section: {section_description[:50]}..."
                        )

                        # Get all source documents
                        all_source_text = ""
                        sources = session.exec(
                            select(Source).where(Source.knowledge_base_id == kb.id)
                        ).all()
                        for source in sources:
                            # Get source data
                            source_data = session.get(SourceData, source.source_data_id)
                            if not source_data:
                                print(f"No source data found for source {source.name}")
                                continue

                            try:
                                # Extract text from the source data
                                if not source_data.data.startswith(b"PK"):
                                    # Direct file extraction
                                    file_content = extract_text_from_file(
                                        source_data.data, source.name
                                    )
                                else:
                                    # Extract from ZIP file
                                    zip_data = BytesIO(source_data.data)
                                    with zipfile.ZipFile(zip_data, "r") as zip_file:
                                        file_info = zip_file.infolist()[0]
                                        raw_file_content = zip_file.read(
                                            file_info.filename
                                        )
                                        file_content = extract_text_from_file(
                                            raw_file_content, source.name
                                        )

                                all_source_text += f"\n\n--- Source: {source.name} ---\n\n{file_content}"
                            except Exception as e:
                                print(
                                    f"Error extracting content from {source.name}: {e}"
                                )
                                # Continue with other sources instead of failing completely
                                continue

                        # Split into chunks and analyze each chunk
                        text_chunks = chunk_text(
                            all_source_text,
                            max_tokens=settings.FULL_SCAN_DOCUMENT_CHUNK_SIZE,
                        )
                        
                        # LLM-based relevance filtering for ReportGenie full text scan
                        # Filter chunks before analysis to avoid processing irrelevant content
                        if text_chunks:
                            print(f"🔍 Full Text Scan: Filtering {len(text_chunks)} chunks for relevance to section: {section_description[:50]}...")
                            
                            import asyncio
                            
                            # Batch / concurrency settings with sensible defaults
                            BATCH_SIZE = getattr(settings, "VERADOC_FULL_SCAN_FILTER_BATCH_SIZE", 10)
                            REQUEST_DELAY = getattr(settings, "PROCESSING_DELAY_BETWEEN_REQUESTS", 0.02)

                            loop = asyncio.get_running_loop()
                            filtered_chunks = []

                            # Process chunks in batches to reduce total runtime while still being rate-limit friendly
                            for start in range(0, len(text_chunks), BATCH_SIZE):
                                batch = text_chunks[start : start + BATCH_SIZE]
                                tasks = []

                                # Create async tasks that run the blocking invoke_llm in executor
                                for j, chunk in enumerate(batch):
                                    chunk_idx = start + j

                                    async def _check(chunk=chunk, chunk_idx=chunk_idx):
                                        # run invoke_llm (sync) in a thread pool to allow concurrency
                                        try:
                                            relevance = await loop.run_in_executor(
                                                None,
                                                lambda: invoke_llm(
                                                    llm,
                                                    settings.VERADOC_RELEVANCE_FILTER_PROMPT_TEMPLATE,
                                                    {"chunk": chunk, "question": section_description},
                                                ),
                                            )
                                        except Exception as e:
                                            # bubble up error to be handled by caller
                                            raise e

                                        return chunk_idx, chunk, relevance

                                    tasks.append(asyncio.create_task(_check()))

                                # Await this batch and handle results
                                results = await asyncio.gather(*tasks, return_exceptions=True)

                                for res in results:
                                    if isinstance(res, Exception):
                                        # On error, include the chunk to be safe (preserve previous behavior)
                                        print(f"Warning: error during batch relevance check: {res}")
                                        continue

                                    chunk_idx, chunk, relevance_check = res
                                    # Filter based on LLM response (same logic as before)
                                    if "No relevant information found" not in (relevance_check or ""):
                                        print(f"✅ Chunk {chunk_idx + 1} is relevant")
                                        filtered_chunks.append(chunk)
                                    else:
                                        print(f"❌ Chunk {chunk_idx + 1} is not relevant - excluding from analysis")

                                # Small sleep between batches to help with rate limits and to yield loop
                                if REQUEST_DELAY:
                                    await asyncio.sleep(REQUEST_DELAY)
                            
                            # Use filtered_chunks for analysis
                            text_chunks = filtered_chunks
                            print(f"📊 Relevance filtering: {len(filtered_chunks)}/{len(text_chunks)} chunks are relevant for analysis")
                        
                        chunk_analyses = []
                        for chunk in text_chunks:
                            analysis = invoke_llm(
                                llm,
                                settings.CHATBOT_FULL_TEXT_CHUNK_PROMPT_TEMPLATE,
                                {"chunk": chunk, "question": section_description},
                            )
                            chunk_analyses.append(analysis)

                        # Synthesize the chunk analyses
                        if not chunk_analyses:
                            generated_content = "No relevant information found in the knowledge base to answer this question."
                        else:
                            chunk_analyses_text = "\n\n".join(chunk_analyses)
                            synthesized_answer = invoke_llm(
                                llm,
                                settings.CHATBOT_FULL_TEXT_SYNTHESIS_PROMPT_TEMPLATE,
                                {
                                    "chunk_analyses": chunk_analyses_text,
                                    "question": section_description,
                                },
                            )
                            # Translate the synthesized answer if needed
                            generated_content = await translate_text_if_needed(
                                synthesized_answer, session, current_user, llm
                            )

                        print(
                            f"Generated {len(generated_content)} characters for section using full text scan"
                        )
                    else:
                        # Vector Search Logic (existing code)
                        print(
                            f"Performing Vector Search for section: {section_description[:50]}..."
                        )

                        # Get relevant context for this section from knowledge base
                        docs = retriever.get_relevant_documents(section_description)
                        
                        # LLM-based relevance filtering for ReportGenie section generation (similar to VeraDoc full scan)
                        # This prevents irrelevant chunks from being included in report sections
                        if docs:
                            print(f"🔍 Filtering {len(docs)} retrieved chunks for relevance to section: {section_description[:50]}...")
                            
                            import asyncio
                            
                            # Batch / concurrency settings with sensible defaults
                            BATCH_SIZE = getattr(settings, "VERADOC_FULL_SCAN_FILTER_BATCH_SIZE", 10)
                            REQUEST_DELAY = getattr(settings, "PROCESSING_DELAY_BETWEEN_REQUESTS", 0.02)

                            loop = asyncio.get_running_loop()
                            filtered_docs = []

                            # Process docs in batches to reduce total runtime while still being rate-limit friendly
                            for start in range(0, len(docs), BATCH_SIZE):
                                batch = docs[start : start + BATCH_SIZE]
                                tasks = []

                                # Create async tasks that run the blocking invoke_llm in executor
                                for j, doc in enumerate(batch):
                                    doc_idx = start + j

                                    async def _check(doc=doc, doc_idx=doc_idx):
                                        # run invoke_llm (sync) in a thread pool to allow concurrency
                                        try:
                                            relevance = await loop.run_in_executor(
                                                None,
                                                lambda: invoke_llm(
                                                    llm,
                                                    settings.VERADOC_RELEVANCE_FILTER_PROMPT_TEMPLATE,
                                                    {"chunk": doc.page_content or "", "question": section_description},
                                                ),
                                            )
                                        except Exception as e:
                                            # bubble up error to be handled by caller
                                            raise e

                                        return doc_idx, doc, relevance

                                    tasks.append(asyncio.create_task(_check()))

                                # Await this batch and handle results
                                results = await asyncio.gather(*tasks, return_exceptions=True)

                                for res in results:
                                    if isinstance(res, Exception):
                                        # On error, include the chunk to be safe (preserve previous behavior)
                                        print(f"Warning: error during batch relevance check: {res}")
                                        continue

                                    doc_idx, doc, relevance_check = res
                                    # Filter based on LLM response (same logic as before)
                                    if "No relevant information found" not in (relevance_check or ""):
                                        print(f"✅ Chunk {doc_idx + 1} is relevant")
                                        filtered_docs.append(doc)
                                    else:
                                        print(f"❌ Chunk {doc_idx + 1} is not relevant - excluding from section")

                                # Small sleep between batches to help with rate limits and to yield loop
                                if REQUEST_DELAY:
                                    await asyncio.sleep(REQUEST_DELAY)
                            
                            # Use filtered_docs for the rest of the flow
                            docs = filtered_docs
                            print(f"📊 Relevance filtering: {len(filtered_docs)}/{len(docs)} chunks are relevant for section")
                        
                        context = "\n\n".join([doc.page_content for doc in docs])

                        # Build the report draft so far (all previous sections)
                        report_draft = ""
                        for prev_section, prev_content in generated_sections.items():
                            if prev_content:
                                report_draft += (
                                    f"\n\n## {prev_section}\n\n{prev_content}"
                                )

                        # Generate content for this section using LLM
                        template_vars = {
                            "report_draft": report_draft,
                            "context": context,
                            "question": section_description,
                        }

                        # Add custom instructions if provided
                        if custom_instructions:
                            template_vars["custom_instructions"] = (
                                f"\nADDITIONAL CUSTOM INSTRUCTIONS:\n{custom_instructions}\n"
                            )
                            print(
                                f"✓ Applying custom instructions to content generation for section: {section_description[:30]}..."
                            )
                        else:
                            template_vars["custom_instructions"] = ""

                        generated_content = invoke_llm(
                            llm,
                            settings.REPORT_GENIE_PROMPT_TEMPLATE,
                            template_vars,
                        )
                        # Translate the generated content if needed
                        generated_content = await translate_text_if_needed(
                            generated_content, session, current_user, llm
                        )
                        print(
                            f"Generated {len(generated_content)} characters for section using vector search"
                        )
                else:
                    # Use the section description directly as content (no document consultation)
                    generated_content = section_description
                    print(
                        f"Using section description directly as content (no document consultation)"
                    )

                generated_sections[section_description] = generated_content

            # 6. Map ground-truth document to outline sections sequentially with large chunks
            print(
                "Mapping ground-truth document to outline sections using large-chunk strategy..."
            )

            # Get list of section descriptions in order with their consult_documents flag
            # Create two separate lists to handle the numbering mismatch issue:
            # 1. all_section_descriptions: For LLM mapping (includes ALL sections)
            # 2. consulting_section_descriptions: For optimization results (only consultDocuments: true)

            all_section_descriptions = []  # For LLM mapping (includes ALL sections)
            all_section_consult_flags = []
            consulting_section_descriptions = (
                []
            )  # For optimization results (only consultDocuments: true)
            section_index_mapping = (
                {}
            )  # Maps section_desc to original index for debugging

            for i, section in enumerate(current_sections):
                section_desc = section["text"].strip()
                consult_docs = section.get("consultDocuments", True)

                # Add to complete list for LLM mapping
                all_section_descriptions.append(section_desc)
                all_section_consult_flags.append(consult_docs)
                section_index_mapping[section_desc] = i + 1

                # Add to consulting list for optimization results
                if consult_docs:
                    consulting_section_descriptions.append(section_desc)

            # Use all sections for LLM mapping to ensure proper boundary awareness
            section_descriptions = all_section_descriptions
            section_consult_flags = all_section_consult_flags

            print(f"Section mapping setup:")
            print(f"- Total sections for LLM mapping: {len(all_section_descriptions)}")
            print(
                f"- Sections for optimization results: {len(consulting_section_descriptions)}"
            )

            # Log section types for debugging
            topic_sections = sum(1 for flag in section_consult_flags if flag)
            literal_sections = len(section_consult_flags) - topic_sections
            print(
                f"Section types: {topic_sections} TOPIC/CONCEPT sections, {literal_sections} LITERAL TEXT sections"
            )
            for i, (section, consult_docs) in enumerate(
                zip(section_descriptions, section_consult_flags)
            ):
                section_type = "TOPIC" if consult_docs else "LITERAL"
                print(
                    f"  {i+1}. [{section_type}] {section[:50]}{'...' if len(section) > 50 else ''}"
                )

            # Split ground-truth into large chunks (up to 100,000 characters)
            max_chunk_size = 50000  # 100,000 characters per chunk
            ground_truth_chunks = []

            # Split by paragraphs first to avoid breaking sentences/paragraphs
            paragraphs = ground_truth_text.split("\n\n")
            current_chunk = ""

            for paragraph in paragraphs:
                # If adding this paragraph would exceed the limit, start a new chunk
                if (
                    current_chunk
                    and len(current_chunk) + len(paragraph) + 2 > max_chunk_size
                ):
                    if current_chunk.strip():
                        ground_truth_chunks.append(current_chunk.strip())
                    current_chunk = paragraph
                else:
                    if current_chunk:
                        current_chunk += "\n\n" + paragraph
                    else:
                        current_chunk = paragraph

            # Add the last chunk if it has content
            if current_chunk.strip():
                ground_truth_chunks.append(current_chunk.strip())

            print(
                f"Split ground-truth into {len(ground_truth_chunks)} large chunks (avg {len(ground_truth_text)//len(ground_truth_chunks) if ground_truth_chunks else 0} chars per chunk)"
            )
            
            # Complete generating stage and start matching stage
            progress_tracker.complete_stage(task_id, "generating", "All sections generated")
            progress_tracker.update_stage_progress(
                task_id, "matching", 0, len(ground_truth_chunks), "Starting document matching..."
            )

            # Map each chunk to the most appropriate section
            section_to_chunks = {section: [] for section in section_descriptions}
            section_to_content = {
                section: [] for section in section_descriptions
            }  # NEW: For actual section content

            # Track previous mapping decisions for context
            previous_mappings = []
            chunk_mapping_stats = {
                "total": 0,
            }

            for chunk_idx, chunk in enumerate(ground_truth_chunks):
                # Update matching progress
                progress_tracker.update_stage_progress(
                    task_id, "matching", chunk_idx, len(ground_truth_chunks),
                    f"Matching chunk {chunk_idx + 1}/{len(ground_truth_chunks)}..."
                )
                
                if cancellation_requested:
                    print("Operation cancelled by client disconnect")
                    raise HTTPException(
                        status_code=408, detail="Operation cancelled by user"
                    )

                print(
                    f"Mapping large chunk {chunk_idx + 1}/{len(ground_truth_chunks)} ({len(chunk)} chars)..."
                )

                # Build enhanced context from previous mappings
                context_info = ""
                if previous_mappings:
                    # Show the last 3 mappings and their outcomes
                    recent_mappings = previous_mappings[-3:]
                    context_summary = []

                    for pm in recent_mappings:
                        section_names_with_types = []
                        for i in pm["assigned_sections"]:
                            if 1 <= i <= len(section_descriptions):
                                section_name = section_descriptions[i - 1]
                                section_type = (
                                    "[TOPIC]"
                                    if section_consult_flags[i - 1]
                                    else "[LITERAL]"
                                )
                                section_names_with_types.append(
                                    f"{section_type} {section_name}"
                                )

                        context_summary.append(
                            f"Chunk {pm['chunk_idx']}: Mapped to sections {pm['assigned_sections']} ({', '.join(section_names_with_types[:2])}{'...' if len(section_names_with_types) > 2 else ''})"
                        )

                    context_info = f"""
PREVIOUS MAPPING DECISIONS (Last {len(recent_mappings)} chunks):
{chr(10).join(context_summary)}

MAPPING PROGRESS: {len(previous_mappings)}/{len(ground_truth_chunks)} chunks processed

This context shows how the document has been mapped so far. Consider:
- Sequential flow: Later chunks typically map to later sections
- Content continuity: Related content often spans adjacent chunks
- Section types: [TOPIC] sections need conceptual matches, [LITERAL] sections need exact text matches
- Previous patterns: Maintain consistency with established mapping logic
"""

                # Pre-calculate values for the prompt to avoid f-string issues
                chunk_size = len(chunk)
                # Use the full chunk content instead of truncating
                chunk_content = sanitize_text_for_json(chunk)

                # No truncation - we send the full chunk to the LLM
                truncation_note = ""
                document_position_percent = (chunk_idx + 1) / len(ground_truth_chunks)
                mapping_context_note = (
                    "- Mapping context: Previous chunks have established patterns - maintain logical consistency"
                    if previous_mappings
                    else "- First chunk: Set the mapping foundation for subsequent chunks"
                )

                # Build outline sections list with boundary awareness (next section context)
                outline_sections_with_boundaries = []
                for i, (section, consult_docs) in enumerate(
                    zip(section_descriptions, section_consult_flags)
                ):
                    if consult_docs:
                        section_type = "[TOPIC/CONCEPT]"
                        instruction = (
                            "Look for content that relates to this topic/concept"
                        )
                    else:
                        section_type = "[LITERAL TEXT]"
                        instruction = "Look for content that matches this text exactly or nearly exactly"

                    # Add next section context for boundary awareness
                    next_section_info = ""
                    if i < len(section_descriptions) - 1:
                        next_section = section_descriptions[i + 1]
                        next_consult_docs = section_consult_flags[i + 1]
                        next_type = (
                            "[TOPIC/CONCEPT]" if next_consult_docs else "[LITERAL TEXT]"
                        )
                        next_section_info = (
                            f" | NEXT SECTION: {i+2}. {next_type} {next_section}"
                        )
                    else:
                        next_section_info = " | NEXT SECTION: [END OF OUTLINE]"

                    outline_sections_with_boundaries.append(
                        f"{i+1}. {section_type} {section} - {instruction}{next_section_info}"
                    )

                # Enhanced mapping prompt with boundary awareness
                mapping_prompt = f"""
Analyze this document chunk and identify the individual sections within it. Extract the actual text content for each section and map it to exactly one outline section.

IMPORTANT BOUNDARY RULES:
- When mapping content to a section, consider where that section should END based on the NEXT SECTION information
- Do NOT assign content to a section if it clearly belongs to the next section's topic/scope
- Stop assigning content to a section when you encounter content that better fits the next section
- If content spans multiple sections, assign it to the most appropriate primary section

OUTLINE SECTIONS (with next section boundaries):
{chr(10).join(outline_sections_with_boundaries)}

DOCUMENT CHUNK ({chunk_idx + 1} of {len(ground_truth_chunks)}):
{chunk_content}
{truncation_note}

MAPPING INSTRUCTIONS:
1. Look for discrete sections in the chunk (titles, headers, paragraphs, topic changes)
2. For each section you identify, extract the ACTUAL TEXT CONTENT from the document
3. Determine which outline section (1-{len(section_descriptions)}) it best matches:
   - For [TOPIC/CONCEPT] sections: Match content that discusses the same topic/concept, even if wording differs
   - For [LITERAL TEXT] sections: Match content that contains the same or very similar text/wording
4. CRITICAL: When assigning content to a section, check if any content better belongs to the NEXT SECTION
5. STOP assigning content to a section when you encounter content that transitions to the next section's topic
6. Respond with valid JSON only

RESPONSE FORMAT (JSON only):
{{"mappings": [{{"section_content": "the actual text extracted from the document section", "outline_section": 1, "boundary_reasoning": "why this content belongs to this section and not the next"}}, {{"section_content": "the actual text from another document section", "outline_section": 2, "boundary_reasoning": "reasoning for section assignment and boundary decision"}}]}}

IMPORTANT: 
- "section_content" must contain the actual text from the document, not a description or summary
- "boundary_reasoning" should explain why this content belongs to the assigned section and not the next section
- Be conservative: if content could belong to either section, assign it to the earlier section but mention the ambiguity
"""

                # Log the exact prompt being sent to the LLM
                prompt_size = len(mapping_prompt)
                estimated_tokens = prompt_size // 4  # Rough estimate: 4 chars per token
                print("=" * 80)
                print(
                    f"SENDING MAPPING PROMPT TO LLM (Chunk {chunk_idx + 1}/{len(ground_truth_chunks)}):"
                )
                print(f"Prompt size: {prompt_size} chars (~{estimated_tokens} tokens)")
                if estimated_tokens > 30000:  # Warning threshold
                    print(f"WARNING: Prompt may exceed token limits!")
                print("=" * 80)
                print(mapping_prompt)
                print("=" * 80)

                # Check for potential LLM issues before calling
                try:
                    # Call LLM through rate limiter for proper token management
                    from app.services.universal_llm_wrapper import execute_llm_request_safely_sync
                    mapping_response = execute_llm_request_safely_sync(
                        llm, 
                        mapping_prompt, 
                        model_name=getattr(llm, 'model_name', 'gpt-4o')
                    ).content
                except Exception as llm_error:
                    print(f"LLM ERROR: {str(llm_error)}")
                    print(f"Prompt was {prompt_size} characters")
                    if "token" in str(llm_error).lower():
                        print("ERROR: Likely token limit exceeded!")
                    raise HTTPException(
                        status_code=413,
                        detail=f"Content too large for processing: {str(llm_error)}",
                    )

                # Log the raw response from the LLM
                print("=" * 80)
                print(
                    f"RAW LLM MAPPING RESPONSE (Chunk {chunk_idx + 1}/{len(ground_truth_chunks)}):"
                )
                print("=" * 80)
                print(mapping_response)
                print("=" * 80)

                # Save the mapping response to a .txt file for inspection
                try:
                    import os
                    from datetime import datetime

                    # Create a timestamp for unique filenames
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

                    # Create filename with chunk info
                    response_filename = f"llm_mapping_response_chunk_{chunk_idx + 1}_of_{len(ground_truth_chunks)}_{timestamp}.txt"

                    # Save to the backend directory (where this file is located)
                    backend_dir = os.path.dirname(os.path.abspath(__file__))
                    response_filepath = os.path.join(backend_dir, response_filename)

                    # Create the full content to save
                    file_content = f"""LLM MAPPING RESPONSE
Chunk: {chunk_idx + 1} of {len(ground_truth_chunks)}
Timestamp: {datetime.now().isoformat()}
Ground-truth document: {file.filename}
Chunk size: {len(chunk)} characters
Prompt size: {prompt_size} characters (~{estimated_tokens} tokens)

=== MAPPING PROMPT SENT TO LLM ===
{mapping_prompt}

=== RAW LLM RESPONSE ===
{mapping_response}

=== END OF RESPONSE ===
"""

                    with open(response_filepath, "w", encoding="utf-8") as f:
                        f.write(file_content)

                    print(f"✓ Saved mapping response to: {response_filename}")

                except Exception as save_error:
                    print(
                        f"Warning: Could not save mapping response to file: {save_error}"
                    )

                # Parse JSON response from LLM
                assigned_sections = []
                document_sections_identified = []

                try:
                    # Try to parse as JSON - handle markdown code blocks
                    response_text = mapping_response.strip()

                    # Remove markdown code blocks if present
                    if response_text.startswith("```json"):
                        response_text = response_text[7:]  # Remove ```json
                    elif response_text.startswith("```"):
                        response_text = response_text[3:]  # Remove ```

                    if response_text.endswith("```"):
                        response_text = response_text[:-3]  # Remove closing ```

                    response_text = response_text.strip()

                    # NEW: Additional sanitization of the response text
                    response_text = sanitize_text_for_json(response_text)

                    json_response = json.loads(response_text)

                    if "mappings" in json_response:
                        for mapping in json_response["mappings"]:
                            section_content = mapping.get("section_content", "").strip()
                            outline_section_num = mapping.get("outline_section", 0)
                            boundary_reasoning = mapping.get(
                                "boundary_reasoning", "No reasoning provided"
                            )

                            # Record the identified section content with boundary reasoning
                            if section_content:
                                content_preview = (
                                    section_content[:100] + "..."
                                    if len(section_content) > 100
                                    else section_content
                                )
                                reasoning_preview = (
                                    boundary_reasoning[:50] + "..."
                                    if len(boundary_reasoning) > 50
                                    else boundary_reasoning
                                )
                                document_sections_identified.append(
                                    f"{content_preview} [Boundary: {reasoning_preview}]"
                                )

                            # Map to actual section description and collect content
                            if 1 <= outline_section_num <= len(section_descriptions):
                                section_desc = section_descriptions[
                                    outline_section_num - 1
                                ]

                                # Debug: Log the mapping with original section index
                                original_section_index = section_index_mapping.get(
                                    section_desc, "unknown"
                                )
                                # print(
                                #    f"✓ LLM mapped section {outline_section_num} to description: '{section_desc[:50]}...'"
                                # )
                                # print(
                                #    f"✓ Original section index: {original_section_index}"
                                # )
                                # print(f"✓ Content length: {len(section_content)} chars")
                                # print(
                                #    f"✓ Section consultDocuments: {section_consult_flags[outline_section_num - 1]}"
                                # )

                                if section_desc not in assigned_sections:
                                    assigned_sections.append(section_desc)

                                # Store the actual section content (clean, without boundary reasoning)
                                if section_content:
                                    # Store only the clean content for optimization analysis
                                    section_to_content[section_desc].append(
                                        section_content
                                    )
                                    # print(
                                    #    f"✓ Stored content for section: '{section_desc[:30]}...' ({len(section_content)} chars)"
                                    # )

                    # print(
                    #    f"✓ Enhanced boundary-aware JSON parsing successful. Found {len(document_sections_identified)} document sections mapping to {len(assigned_sections)} outline sections"
                    # )

                    # Log extracted content summary with boundary awareness
                    total_extracted_chars = sum(
                        len(content)
                        for section_content_list in section_to_content.values()
                        for content in section_content_list
                    )
                    # print(
                    #    f"✓ Extracted {total_extracted_chars} characters of section content with boundary reasoning"
                    # )

                except json.JSONDecodeError as e:
                    print(f"✗ JSON parsing failed: {e}")
                    print(f"Raw response was: {mapping_response[:500]}...")
                    print(f"Cleaned response was: {response_text[:500]}...")

                    # NEW: Try to identify the problematic character
                    try:
                        # Find the character at the error position if available
                        if hasattr(e, "pos") and e.pos < len(response_text):
                            problem_char = response_text[e.pos]
                            print(
                                f"Problematic character at position {e.pos}: '{problem_char}' (ord: {ord(problem_char)})"
                            )
                    except:
                        pass

                    # Keep empty lists for fallback
                except Exception as e:
                    print(f"✗ Error processing JSON response: {e}")
                    if "response_text" in locals():
                        print(f"Cleaned response was: {response_text[:500]}...")
                    # Keep empty lists for fallback
                # Enhanced fallback logic if JSON parsing failed
                if not assigned_sections:
                    print(
                        f"No sections assigned by LLM for chunk {chunk_idx + 1}, using enhanced fallback..."
                    )

                    # Calculate chunk position and use content analysis for better fallback
                    chunk_position = chunk_idx / len(ground_truth_chunks)

                    # Analyze chunk content for keywords to help with mapping
                    chunk_lower = chunk[
                        :2000
                    ].lower()  # Analyze first 2000 chars for keywords

                    # Try to match chunk content to section keywords with consult documents consideration
                    section_scores = []
                    for i, (section_desc, consult_docs) in enumerate(
                        zip(section_descriptions, section_consult_flags)
                    ):
                        if consult_docs:
                            # For topic/concept sections, do semantic word overlap
                            section_words = set(section_desc.lower().split())
                            chunk_words = set(chunk_lower.split())
                            overlap = len(section_words.intersection(chunk_words))
                        else:
                            # For literal text sections, look for exact/near-exact substring matches
                            if section_desc.lower() in chunk_lower:
                                overlap = 100  # High score for exact match
                            elif any(
                                word in chunk_lower
                                for word in section_desc.lower().split()
                                if len(word) > 3
                            ):
                                overlap = 50  # Medium score for partial match
                            else:
                                overlap = 0  # No match

                        section_scores.append((i, overlap))

                    # Sort by overlap score
                    section_scores.sort(key=lambda x: x[1], reverse=True)
                    best_content_match = section_scores[0] if section_scores else (0, 0)

                    # Combine position and content analysis for better fallback
                    if best_content_match[1] > 0:  # If there's content overlap
                        assigned_sections = [
                            section_descriptions[best_content_match[0]]
                        ]
                        match_type = (
                            "exact text"
                            if section_consult_flags[best_content_match[0]] == False
                            else "topic"
                        )
                        reasoning = f"Fallback: Content analysis found {match_type} match (score: {best_content_match[1]}) with section {best_content_match[0] + 1}"
                    else:
                        # Use positional fallback with improved logic
                        if chunk_idx == 0:
                            # First chunk: likely maps to first section(s)
                            assigned_sections = [section_descriptions[0]]
                            reasoning = "Fallback: First chunk mapped to first section"
                        elif chunk_idx == len(ground_truth_chunks) - 1:
                            # Last chunk: likely maps to last section(s)
                            assigned_sections = [section_descriptions[-1]]
                            reasoning = "Fallback: Last chunk mapped to last section"
                        else:
                            # Middle chunks: proportional mapping
                            primary_section_index = min(
                                int(chunk_position * len(section_descriptions)),
                                len(section_descriptions) - 1,
                            )
                            assigned_sections = [
                                section_descriptions[primary_section_index]
                            ]
                            reasoning = f"Fallback: Positional mapping for chunk at {chunk_position:.1%} position"

                    document_sections_identified = [
                        "Fallback: No sections identified by LLM"
                    ]
                    print(
                        f"Fallback mapping: Assigned chunk {chunk_idx + 1} to {[section_descriptions.index(s) + 1 for s in assigned_sections]} - {reasoning}"
                    )

                    # IMPORTANT: Also populate section_to_content for fallback cases
                    # This ensures later sections get ground-truth content even if JSON parsing failed
                    for section_desc in assigned_sections:
                        section_to_content[section_desc].append(chunk)
                else:
                    reasoning = f"JSON parsing successful: {len(document_sections_identified)} sections identified"
                    print(
                        f"LLM mapping: Chunk {chunk_idx + 1} → sections {[section_descriptions.index(s) + 1 for s in assigned_sections]}"
                    )
                    print(
                        f"Document sections identified: {', '.join(document_sections_identified[:3])}{'...' if len(document_sections_identified) > 3 else ''}"
                    )

                # Update mapping statistics
                chunk_mapping_stats["total"] += 1
                # Note: No longer tracking confidence levels in simplified approach

                # Record this mapping decision for future context
                mapping_decision = {
                    "chunk_idx": chunk_idx + 1,
                    "assigned_sections": [
                        section_descriptions.index(s) + 1 for s in assigned_sections
                    ],
                    "reasoning": (
                        reasoning[:150] + "..." if len(reasoning) > 150 else reasoning
                    ),
                    "chunk_size": len(chunk),
                    "document_position": chunk_idx / len(ground_truth_chunks),
                    "fallback_used": "Fallback" in reasoning,
                    "document_sections_identified": (
                        document_sections_identified[:3]
                        if document_sections_identified
                        else []
                    ),
                }
                previous_mappings.append(mapping_decision)

                # Add chunk to assigned sections with comprehensive metadata
                for section in assigned_sections:
                    chunk_metadata = {
                        "content": chunk,
                        "chunk_index": chunk_idx,
                        "chunk_size": len(chunk),
                        "reasoning": reasoning,
                        "document_sections_identified": document_sections_identified,
                        "mapping_context": len(previous_mappings)
                        > 1,  # Whether context was available
                        "document_position": chunk_idx / len(ground_truth_chunks),
                        "fallback_used": "Fallback" in reasoning,
                        "section_assignment_count": len(
                            assigned_sections
                        ),  # How many sections this chunk maps to
                        "content_preview": (
                            chunk[:200] + "..." if len(chunk) > 200 else chunk
                        ),  # For debugging
                    }
                    section_to_chunks[section].append(chunk_metadata)

            print(
                "Ground-truth mapping complete. Generating optimization suggestions..."
            )

            # Debug: Print content extraction summary with section details
            total_content_extracted = sum(
                len(content_list) for content_list in section_to_content.values()
            )
            sections_with_content = sum(
                1 for content_list in section_to_content.values() if content_list
            )
            print(
                f"Content extraction summary: {total_content_extracted} total sections extracted, {sections_with_content}/{len(section_descriptions)} outline sections have mapped content"
            )

            # Debug: Show which sections have content mapped to them
            print("MAPPING RESULTS BY SECTION:")
            for i, section_desc in enumerate(section_descriptions):
                consult_flag = section_consult_flags[i]
                content_pieces = len(section_to_content.get(section_desc, []))
                chunk_pieces = len(section_to_chunks.get(section_desc, []))
                total_content_size = sum(
                    len(content) for content in section_to_content.get(section_desc, [])
                )

                section_type = "CONSULT" if consult_flag else "NO-CONSULT"
                print(
                    f"  Section {i+1} [{section_type}]: '{section_desc[:40]}...' - {content_pieces} content pieces ({total_content_size} chars), {chunk_pieces} chunks"
                )

                # Show a preview of the first content piece if available
                if section_to_content.get(section_desc):
                    first_content = section_to_content[section_desc][0]
                    preview = (
                        first_content[:100] + "..."
                        if len(first_content) > 100
                        else first_content
                    )
                    print(f"    Preview: {preview}")
                else:
                    print(f"    No content mapped to this section")
            print("END MAPPING RESULTS")
            
            # Complete matching stage and start comparing stage
            progress_tracker.complete_stage(task_id, "matching", "Document matching complete")
            progress_tracker.update_stage_progress(
                task_id, "comparing", 0, len(consulting_section_descriptions), "Starting section comparison..."
            )

            # 7. Compare each section's generated content to its mapped ground-truth chunks
            # Only process sections that consult documents to avoid numbering mismatches
            suggestions = []
            optimization_count = 0

            # Process only consulting sections to ensure consistent numbering
            for opt_idx, section_description in enumerate(consulting_section_descriptions):
                # Update comparing progress
                progress_tracker.update_stage_progress(
                    task_id, "comparing", opt_idx, len(consulting_section_descriptions),
                    f"Comparing section {opt_idx + 1}/{len(consulting_section_descriptions)}..."
                )
                
                if cancellation_requested:
                    print("Operation cancelled by client disconnect")
                    raise HTTPException(
                        status_code=408, detail="Operation cancelled by user"
                    )

                # Get the generated content for this section
                generated_content = generated_sections.get(section_description, "")

                # This section definitely consults documents (since it's in consulting_section_descriptions)
                print(f"Analyzing consulting section: {section_description[:50]}...")

                # Get the mapped ground-truth content for this section
                # Use the SAME section_description that was used as key during mapping
                mapped_chunks = section_to_chunks.get(section_description, [])
                mapped_content = section_to_content.get(section_description, [])

                # Debug: Log what we found
                print(
                    f"✓ Looking up content for section: '{section_description[:30]}...'"
                )
                print(
                    f"✓ Found {len(mapped_content)} content pieces, {len(mapped_chunks)} chunks"
                )

                # Use actual extracted content if available, otherwise fall back to chunks
                if mapped_content:
                    ground_truth_context = "\n\n".join(mapped_content)
                    print(
                        f"✓ Using mapped content: {len(ground_truth_context)} characters"
                    )
                else:
                    ground_truth_context = "\n\n".join(
                        [
                            chunk["content"] if isinstance(chunk, dict) else chunk
                            for chunk in mapped_chunks
                        ]
                    )
                    print(
                        f"✓ Using chunk content: {len(ground_truth_context)} characters"
                    )

                # Create an enhanced summary of the mapping for debugging and analysis
                mapping_summary = ""
                if mapped_chunks or mapped_content:
                    chunk_count = len(mapped_chunks)
                    content_count = len(mapped_content)
                    fallback_count = sum(
                        1
                        for chunk in mapped_chunks
                        if isinstance(chunk, dict) and chunk.get("fallback_used", False)
                    )

                    total_content_size = sum(
                        (
                            chunk.get("chunk_size", 0)
                            if isinstance(chunk, dict)
                            else len(chunk)
                        )
                        for chunk in mapped_chunks
                    )

                    extracted_content_size = sum(
                        len(content) for content in mapped_content
                    )

                    mapping_summary = f"Mapped {chunk_count} chunks → {content_count} extracted sections (total: {total_content_size} chars → {extracted_content_size} chars)"
                else:
                    mapping_summary = "No content mapped to this section"

                # If no chunks were mapped to this section, use a fallback
                if not ground_truth_context.strip():
                    ground_truth_context = "No specific content was mapped to this section from the ground-truth document."
                    print(
                        f"Warning: No ground-truth content mapped to section: {section_description[:30]}..."
                    )
                else:
                    print(
                        f"Mapped content to section: {section_description[:30]}... {mapping_summary}"
                    )

                # Generate optimization suggestion
                template_vars = {
                    "original_section": section_description,
                    "generated_content": generated_content[
                        :2000
                    ],  # Limit to avoid token limits
                    "ground_truth_content": ground_truth_context[
                        :2000
                    ],  # Limit to avoid token limits
                }

                # Add custom instructions if provided
                if custom_instructions:
                    template_vars["custom_instructions"] = (
                        f"\nADDITIONAL CUSTOM INSTRUCTIONS FOR OPTIMIZATION:\n{custom_instructions}\n"
                    )
                    print(
                        f"✓ Applying custom instructions to optimization analysis for section: {section_description[:30]}..."
                    )
                else:
                    template_vars["custom_instructions"] = ""

                suggestion_response = invoke_llm(
                    llm,
                    settings.REPORTGENIE_OPTIMIZE_OUTLINE_PROMPT_TEMPLATE,
                    template_vars,
                )

                # Parse the response with improved logic
                lines = suggestion_response.split("\n")

                # Extract needs revision (check for both "Yes" and case variations)
                needs_revision = False
                for line in lines:
                    if line.startswith("NEEDS_REVISION:"):
                        revision_value = (
                            line.replace("NEEDS_REVISION:", "").strip().lower()
                        )
                        needs_revision = revision_value in ["yes", "true", "y"]
                        break

                # Extract suggested section
                suggested_section = section_description  # Default to original
                for line in lines:
                    if line.startswith("SUGGESTED_SECTION:"):
                        suggested_section = line.replace(
                            "SUGGESTED_SECTION:", ""
                        ).strip()
                        break

                # Extract reason
                reason = "No specific reason provided"
                for line in lines:
                    if line.startswith("REASON:"):
                        reason = line.replace("REASON:", "").strip()
                        break

                # Extract quality gap severity (new field)
                quality_gap_severity = "unknown"
                for line in lines:
                    if line.startswith("QUALITY_GAP_SEVERITY:"):
                        quality_gap_severity = (
                            line.replace("QUALITY_GAP_SEVERITY:", "").strip().lower()
                        )
                        break

                # Log the analysis results for debugging
                print(
                    f"Optimization analysis for section '{section_description[:30]}...': "
                    f"needs_revision={needs_revision}, quality_gap={quality_gap_severity}"
                )

                if needs_revision:
                    optimization_count += 1

                suggestions.append(
                    OutlineSuggestion(
                        original_section=section_description,
                        suggested_section=suggested_section,
                        reason=reason,
                        current_output=generated_content[
                            :1000
                        ],  # Show first 1000 chars
                        ground_truth_content=ground_truth_context[
                            :1000
                        ],  # Show first 1000 chars
                        needs_revision=needs_revision,
                    )
                )

                print(
                    f"Section analysis complete: {'NEEDS OPTIMIZATION' if needs_revision else 'OK'}"
                )

            # 7. Compile results - Use the pre-defined consulting sections list
            # This ensures consistent section handling between mapping and optimization phases
            optimized_sections = [s.suggested_section for s in suggestions]

            # Calculate enhanced mapping statistics
            total_chunks = len(ground_truth_chunks)
            mapped_chunk_instances = sum(
                len(chunks) for chunks in section_to_chunks.values()
            )
            unique_chunks_mapped = len(
                set(
                    chunk.get("chunk_index", -1) if isinstance(chunk, dict) else -1
                    for chunks in section_to_chunks.values()
                    for chunk in chunks
                )
            )

            # Calculate coverage metrics
            coverage_percentage = (
                (unique_chunks_mapped / total_chunks * 100) if total_chunks > 0 else 0
            )

            # Calculate section type statistics using the original sections
            total_sections_in_outline = len(current_sections)
            sections_that_consult_docs = len(consulting_section_descriptions)
            sections_that_dont_consult_docs = (
                total_sections_in_outline - sections_that_consult_docs
            )
            sections_actually_optimized = len(
                [s for s in suggestions if s.needs_revision]
            )

            analysis_summary = f"""
Enhanced Content Extraction Analysis:
- Total outline sections: {total_sections_in_outline}
- Sections that consult documents (included in results): {sections_that_consult_docs}
- Sections that don't consult documents (excluded from results): {sections_that_dont_consult_docs}
- Sections needing optimization: {sections_actually_optimized}
- Sections working well: {sections_that_consult_docs - sections_actually_optimized}

Ground-truth Processing:
- Document: {file.filename}
- Total chunks processed: {total_chunks}
- Unique chunks mapped: {unique_chunks_mapped}
- Chunk instances mapped: {mapped_chunk_instances} (chunks can map to multiple sections)
- Coverage: {coverage_percentage:.1f}% of document mapped

Content Extraction:
- LLM mapped content to ALL {total_sections_in_outline} sections (including non-consulting ones)
- Only {sections_that_consult_docs} consulting sections analyzed for optimization
- Method: Complete section mapping with selective optimization analysis
- Note: Sections with 'Consult Documents' set to false are mapped but excluded from optimization results.
            """.strip()

            print(
                f"Optimization complete: {sections_actually_optimized}/{sections_that_consult_docs} document-consulting sections optimized ({sections_that_dont_consult_docs} non-consulting sections excluded from results)"
            )
            
            # Complete comparing stage and finalizing
            progress_tracker.complete_stage(task_id, "comparing", "Section comparison complete")
            progress_tracker.update_stage_progress(
                task_id, "finalizing", 0, 1, "Finalizing optimization results..."
            )
            
            progress_tracker.complete_stage(task_id, "finalizing", "Optimization complete!")

            return OptimizedOutlineResponse(
                original_sections=consulting_section_descriptions,  # Only return consulting sections
                suggestions=suggestions,  # Already filtered to only consulting sections
                optimized_sections=optimized_sections,
                analysis_summary=analysis_summary,
                task_id=task_id,
            )

    except Exception as e:
        print("Error in outline optimization:")
        print(str(e))
        traceback.print_exc()
        
        # Mark progress as failed if task_id exists
        if 'task_id' in locals() and task_id:
            progress_tracker.fail_task(task_id, f"Optimization failed: {str(e)}")
        
        raise HTTPException(
            status_code=500, detail=f"Error optimizing outline: {str(e)}"
        )
    finally:
        if disconnect_monitor:
            disconnect_monitor.cancel()


@router.post("/optimize-outline/csv", response_class=StreamingResponse)
async def generate_outline_optimization_csv(
    session: SessionDep, current_user: CurrentUser, request: DocxRequest
):
    """
    Generate a CSV file from outline optimization results with columns for:
    Section Number, Original Section, Suggested Section, Needs Revision, Reason, Current Output, Ground Truth Content.
    """
    print("Now generating CSV of outline optimization results...")
    try:
        # Get the content from the request - this should be the optimization results JSON
        if not request.content:
            raise HTTPException(
                status_code=400, detail="Optimization content is required"
            )

        # Create CSV content
        csv_buffer = StringIO()
        csv_writer = csv.writer(csv_buffer)

        # Write header with optimization-specific columns
        csv_writer.writerow(
            [
                "Section Number",
                "Original Section",
                "Suggested Section",
                "Needs Revision",
                "Reason",
                "Current Output",
                "Ground Truth Content",
                "Analysis Summary",
            ]
        )

        try:
            # Parse the content as JSON (optimization results data)
            data = json.loads(request.content)

            # Extract suggestions and analysis summary
            suggestions = data.get("suggestions", [])
            analysis_summary = data.get(
                "analysis_summary", "No analysis summary provided"
            )

            # Process each suggestion
            for index, suggestion in enumerate(suggestions, 1):
                original_section = suggestion.get("original_section", "")
                suggested_section = suggestion.get("suggested_section", "")
                needs_revision = suggestion.get("needs_revision", False)
                reason = suggestion.get("reason", "")
                current_output = suggestion.get("current_output", "")
                ground_truth_content = suggestion.get("ground_truth_content", "")

                # Clean up text fields (remove newlines and carriage returns for CSV)
                original_section_clean = original_section.replace("\n", " ").replace(
                    "\r", " "
                )
                suggested_section_clean = suggested_section.replace("\n", " ").replace(
                    "\r", " "
                )
                reason_clean = reason.replace("\n", " ").replace("\r", " ")
                current_output_clean = current_output.replace("\n", " ").replace(
                    "\r", " "
                )
                ground_truth_content_clean = ground_truth_content.replace(
                    "\n", " "
                ).replace("\r", " ")
                analysis_summary_clean = analysis_summary.replace("\n", " ").replace(
                    "\r", " "
                )

                # Write row
                csv_writer.writerow(
                    [
                        index,
                        original_section_clean,
                        suggested_section_clean,
                        "Yes" if needs_revision else "No",
                        reason_clean,
                        current_output_clean,
                        ground_truth_content_clean,
                        analysis_summary_clean,
                    ]
                )

        except json.JSONDecodeError:
            raise HTTPException(
                status_code=400,
                detail="Invalid content format. Expected JSON with outline optimization data.",
            )

        # Get CSV content
        csv_content = csv_buffer.getvalue()
        csv_buffer.close()

        # Create BytesIO object for the response
        csv_bytes = BytesIO(csv_content.encode("utf-8"))
        csv_bytes.seek(0)

        print("Outline optimization CSV file generated successfully.")

        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"outline_optimization_{timestamp}.csv"

        # Return the CSV as a downloadable file
        return StreamingResponse(
            csv_bytes,
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error generating optimization CSV: {str(e)}"
        )


@router.post("/generate/docx", response_class=StreamingResponse)
async def generate_docx(
    session: SessionDep, current_user: CurrentUser, request: DocxRequest
):
    """
    Generate a DOCX file from the report content.
    """
    print("Now generating DOCX of report...")
    try:
        # Get the markdown content from the request
        if not request.content:
            raise HTTPException(status_code=400, detail="Report content is required")

        # Convert markdown to HTML for parsing
        html_content = markdown.markdown(request.content, extensions=["tables"])
        soup = BeautifulSoup(html_content, "html.parser")

        print("Markdown content converted to HTML successfully.")
        # Create a new Document
        doc = Document()

        print("Adding title and date to the document...")
        # Add a title - hard-coding it for ReportGenie because it's using the service for Compare functionality with 'Document Comparison' as title
        title_text = (
            # request.title
            # if hasattr(request, "title") and request.title
            # else "Generated Document"
            "Generated Document"
        )
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_paragraph.add_run(
            f"Generated on: {datetime.now().strftime('%B %d, %Y')}"
        )
        date_run.italic = True

        # Add a separator
        doc.add_paragraph("─" * 50)

        print("Adding headers, paragraphs, lists, and tables...")
        # Process each element in the soup
        for element in soup.find_all():
            if element.name == "h1":
                doc.add_heading(element.get_text().strip(), level=1)
            elif element.name == "h2":
                doc.add_heading(element.get_text().strip(), level=2)
            elif element.name == "h3":
                doc.add_heading(element.get_text().strip(), level=3)
            elif element.name == "h4":
                doc.add_heading(element.get_text().strip(), level=4)
            elif element.name == "p":
                text = element.get_text().strip()
                if text:  # Only add non-empty paragraphs
                    paragraph = doc.add_paragraph(text)

                    # Handle formatting within paragraphs
                    for strong in element.find_all("strong"):
                        # Bold text formatting would need more complex handling
                        pass
                    for em in element.find_all("em"):
                        # Italic text formatting would need more complex handling
                        pass

            elif element.name == "table":
                # Handle tables
                rows = element.find_all("tr")
                if rows:
                    print(f"Adding table with {len(rows)} rows...")
                    table = doc.add_table(
                        rows=len(rows), cols=len(rows[0].find_all(["th", "td"]))
                    )
                    table.style = "Table Grid"

                    for i, row in enumerate(rows):
                        cells = row.find_all(["th", "td"])
                        for j, cell in enumerate(cells):
                            if j < len(table.rows[i].cells):
                                table.rows[i].cells[j].text = cell.get_text().strip()
                                # Make header row bold
                                if i == 0:
                                    for paragraph in table.rows[i].cells[j].paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True

            elif element.name == "ul":
                # Handle unordered lists
                for li in element.find_all("li", recursive=False):
                    text = li.get_text().strip()
                    if text:
                        paragraph = doc.add_paragraph(text, style="List Bullet")

            elif element.name == "ol":
                # Handle ordered lists
                for li in element.find_all("li", recursive=False):
                    text = li.get_text().strip()
                    if text:
                        paragraph = doc.add_paragraph(text, style="List Number")

        print("Saving the document to a BytesIO object...")
        # Save the document to a BytesIO object
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Create filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"report_{timestamp}.docx"

        print(f"DOCX file size: {len(doc_io.getvalue())} bytes")

        # Verify the document can be opened (basic integrity check)
        doc_io.seek(0)
        try:
            test_doc = Document(doc_io)
            print("DOCX file passed integrity check (can be opened by python-docx).")
        except Exception as e:
            print(f"DOCX integrity check failed: {e}")
            raise HTTPException(
                status_code=500, detail=f"Generated DOCX file is corrupted: {str(e)}"
            )

        doc_io.seek(0)
        print(
            "Document saved successfully. Preparing to return as a downloadable file."
        )

        # Return the document as a downloadable file
        return StreamingResponse(
            doc_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")


@router.post("/generate/csv", response_class=StreamingResponse)
async def generate_csv(
    session: SessionDep, current_user: CurrentUser, request: DocxRequest
):
    """
    Generate a CSV file from report content with columns for:
    Section Title, Content, Citations
    """
    print("Now generating CSV of report content...")
    try:
        # Get the content from the request
        if not request.content:
            raise HTTPException(status_code=400, detail="Report content is required")

        # Create CSV content
        csv_buffer = StringIO()
        csv_writer = csv.writer(csv_buffer)

        # Write header for report CSV
        csv_writer.writerow(["Section Title", "Content", "Citations"])

        try:
            # Parse the content as JSON (report sections data)
            data = json.loads(request.content)
            sections = data.get("sections", [])

            print(f"Processing {len(sections)} sections for CSV export...")

            # Process each section
            for section in sections:
                title = section.get("title", "")
                content = section.get("content", "")
                citations = section.get("source_citations", [])

                # Format citations as "filename: content | filename: content"
                citation_texts = []
                for citation in citations:
                    if isinstance(citation, dict):
                        source = citation.get("source", "unknown")
                        citation_content = citation.get("content", "")
                        citation_texts.append(f"{source}: {citation_content}")
                    elif isinstance(citation, str):
                        citation_texts.append(citation)

                citations_formatted = " | ".join(citation_texts)

                # Clean up text fields (remove newlines and carriage returns for CSV)
                title_clean = (
                    title.replace("\n", " ").replace("\r", " ").replace('"', '""')
                )
                content_clean = (
                    content.replace("\n", " ").replace("\r", " ").replace('"', '""')
                )
                citations_clean = (
                    citations_formatted.replace("\n", " ")
                    .replace("\r", " ")
                    .replace('"', '""')
                )

                # Write row
                csv_writer.writerow([title_clean, content_clean, citations_clean])

        except json.JSONDecodeError:
            # If not JSON, treat as plain markdown content
            print("Content is not JSON, treating as plain markdown content...")

            # Split content into sections based on headers
            lines = request.content.split("\n")
            current_section = ""
            current_content = []

            for line in lines:
                line = line.strip()
                if line.startswith("#"):
                    # Save previous section
                    if current_section:
                        content_text = " ".join(current_content).replace('"', '""')
                        csv_writer.writerow(
                            [
                                current_section.replace('"', '""'),
                                content_text,
                                "",  # No citations for markdown content
                            ]
                        )

                    # Start new section
                    current_section = line.lstrip("#").strip()
                    current_content = []
                elif line:
                    current_content.append(line)

            # Save final section
            if current_section:
                content_text = " ".join(current_content).replace('"', '""')
                csv_writer.writerow(
                    [current_section.replace('"', '""'), content_text, ""]
                )

        # Prepare CSV for download
        csv_content = csv_buffer.getvalue()
        csv_buffer.close()

        if not csv_content.strip() or csv_content.count("\n") <= 1:
            raise HTTPException(
                status_code=400, detail="No valid report data found to export"
            )

        print(f"CSV generated successfully with {csv_content.count(chr(10))} rows")

        # Create response
        csv_bytes = BytesIO(csv_content.encode("utf-8"))

        # Create filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"report_{timestamp}.csv"

        return StreamingResponse(
            csv_bytes,
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except HTTPException:
        raise
    except Exception as e:
        import traceback

        traceback.print_exc()
        print(f"Error generating CSV: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating CSV: {str(e)}")
