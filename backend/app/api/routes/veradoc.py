import uuid
import json
import traceback
import re  # Add missing import for regex operations
import tempfile
import zipfile
import os
from pathlib import Path
from app.models import (
    VeraDocResponse,
    VeraDocChecklist,
    RagChecklistRequest,
    EmbeddingModel,
    Source,
    KnowledgeBase,
    LlmInteraction,
    DocxRequest,
    VeraDocDetailResponse,
    Message,
    User,
    OptimizeChecklistRequest,
    ChecklistSuggestion,
    OptimizedChecklistResponse,
    GenerateQuestionsRequest,
    GenerateQuestionsResponse,
)

from app.api.deps import CurrentUser, SessionDep
from app.core.config import settings
from app.services.knowledgebases import get_embedding_model
from app.services.embeddings import load_embeddings_model
from app.services.llms import (
    get_default_llm,
    invoke_llm,
    invoke_llm_async,
    record_llm_interaction,
)
from app.services.translation import (
    translate_text_if_needed,
    translate_progress_message,
    translate,
)
from app.services.docx_translations import translate_docx_header
from app.services.retrievers import (
    create_ensemble_retriever,
)  # Import the ensemble retriever
from app.services.enhanced_retrieval import SmartRetrieverFactory
from app.services.progress_tracker import progress_tracker

# Cache for frontend translations
_frontend_translations_cache = None


def _load_frontend_translations():
    """
    Load frontend translations from common.json files.
    """
    global _frontend_translations_cache
    if _frontend_translations_cache is not None:
        return _frontend_translations_cache

    translations = {}
    # In Docker, frontend locales are copied to /app/frontend/src/locales
    locales_dir = Path("/app/frontend/src/locales")

    # Fallback for development (not in Docker)
    if not locales_dir.exists():
        # Get to project root: backend/app/api/routes -> backend/app/api -> backend/app -> backend -> project_root
        locales_dir = (
            Path(__file__).parent.parent.parent.parent.parent
            / "frontend"
            / "src"
            / "locales"
        )

    if locales_dir.exists():
        for lang_dir in locales_dir.iterdir():
            if lang_dir.is_dir() and (lang_dir / "common.json").exists():
                lang_code = lang_dir.name
                with open(lang_dir / "common.json", "r", encoding="utf-8") as f:
                    translations[lang_code] = json.load(f)
    else:
        print(f"WARNING: Frontend locales directory not found at {locales_dir}")

    _frontend_translations_cache = translations
    return translations


def translate_frontend(key, language="en", **kwargs):
    """
    Translate a key using frontend translations with parameter substitution.
    """
    translations = _load_frontend_translations()

    print(f"[translate_frontend] Translating key: {key}, language: {language}")
    print(f"[translate_frontend] Available languages: {list(translations.keys())}")

    # Get translations for the requested language, fallback to English
    lang_translations = translations.get(language, translations.get("en", {}))

    if not lang_translations:
        print(
            f"[translate_frontend] ERROR: No translations found for language: {language}"
        )
        return key

    # Navigate to the nested key
    keys = key.split(".")
    value = lang_translations
    for k in keys:
        if isinstance(value, dict) and k in value:
            value = value[k]
            print(f"[translate_frontend] Found key '{k}': {type(value)}")
        else:
            print(f"[translate_frontend] Key '{k}' not found in current level")
            # Fallback to English if key not found
            en_translations = translations.get("en", {})
            value = en_translations
            for k in keys:
                if isinstance(value, dict) and k in value:
                    value = value[k]
                else:
                    print(
                        f"[translate_frontend] Key '{k}' not found even in English, returning: {key}"
                    )
                    return key  # Return key if not found

    if not isinstance(value, str):
        print(f"[translate_frontend] ERROR: Final value is not a string: {type(value)}")
        return key

    print(f"[translate_frontend] Found translation: {value}")

    # Substitute parameters
    try:
        result = value.format(**kwargs)
        print(f"[translate_frontend] After substitution: {result}")
        return result
    except (KeyError, ValueError) as e:
        print(f"[translate_frontend] Error during substitution: {e}")
        return value


from sqlmodel import select
from fastapi import (
    APIRouter,
    UploadFile,
    File,
    Form,
    HTTPException,
    Depends,
    Request as FastAPIRequest,
    Query,
)
from fastapi.responses import StreamingResponse
from typing import List, Dict, Any, Optional
import asyncio
from concurrent.futures import ThreadPoolExecutor
from dotenv import load_dotenv
import json
import os
import re
from pathlib import Path
import csv
import zipfile
import traceback
from io import BytesIO, StringIO

from datetime import datetime
from starlette.requests import Request
import tempfile
import markdown
from langchain_community.vectorstores import Chroma

# from langchain_community.document_loaders import PyPDFLoader, TextLoader  # Removed - using pypdf instead
from langchain_community.document_loaders import TextLoader
from app.services.pdf_utils import load_pdf_with_pypdf
from langchain_core.documents import Document as LangchainDocument
from docx import Document  # For .docx file handling
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup

# Load environment variables from .env file
load_dotenv(dotenv_path="c:/miniconda/aibeniq-react/.env", override=False)

# Retrieve the OpenAI API key from the environment
openai_api_key = os.getenv("OPENAI_API_KEY")

# Initialize a flag to track API key status
is_openai_configured = False

if openai_api_key:
    # Set up OpenAI API key if available
    os.environ["OPENAI_API_KEY"] = openai_api_key
    is_openai_configured = True
    print("OpenAI API key configured successfully")
else:
    print(
        "WARNING: OPENAI_API_KEY is not set in environment variables. Some FormConnect features will be limited."
    )


async def prefetch_knowledge_base_context(
    retriever,
    question_list: List[Dict],
    llm,
    context_prompt_template: str,
    session,
    current_user,
    request: FastAPIRequest = None,
    task_id: str = None,
    user_language: str = "en",
) -> Dict[str, Dict[str, Any]]:
    """
    Pre-fetch knowledge base context for all questions to avoid redundant retrieval.
    Returns a dictionary mapping question text to its context and source citations.
    """
    print(f"Pre-fetching knowledge base context for {len(question_list)} questions...")

    question_contexts = {}

    for i, question_item in enumerate(question_list):
        # Yield to event loop at start of each question iteration
        await asyncio.sleep(0.05)  # Prevent connection timeouts during long operations

        # Update progress for context fetching
        if task_id:
            progress_tracker.update_stage_progress(
                task_id,
                "fetching_context",
                i,
                len(question_list),
                translate_progress_message(
                    "retrieving_policy_context",
                    user_language,
                    question_num=i + 1,
                    total_questions=len(question_list),
                ),
            )

        # Add delay between questions to prevent rate limit exhaustion
        if i > 0 and settings.VERADOC_ENABLE_PROCESSING_DELAYS:
            await asyncio.sleep(settings.PROCESSING_DELAY_BETWEEN_QUESTIONS)

        # Check for cancellation during context pre-fetching
        try:
            if request and await request.is_disconnected():
                print(
                    f"❌ CLIENT DISCONNECTED - Stopping context prefetch at question {i + 1}"
                )
                raise HTTPException(
                    status_code=408,
                    detail="Request cancelled - client disconnected during context prefetch",
                )
        except Exception as e:
            print(f"Warning: Could not check disconnect status during prefetch: {e}")

        question_text = question_item.get("text", "").strip()
        consult_documents = question_item.get("consultDocuments", True)

        if not question_text:
            continue

        print(
            f"Pre-fetching context for question {i+1}/{len(question_list)}: {question_text[:50]}..."
        )

        if consult_documents:
            try:
                # Step 1: Retrieve relevant context from the knowledge base
                docs = retriever.get_relevant_documents(question_text)

                if not docs:
                    print(
                        f"No documents retrieved for question: {question_text[:50]}..."
                    )
                    context = "No relevant documents found in the knowledge base for this question."
                    source_citations = []
                else:
                    print(
                        f"Retrieved {len(docs)} documents for question: {question_text[:50]}..."
                    )

                    # LLM-based relevance filtering for VeraDoc (similar to ReportGenie)
                    # This prevents irrelevant chunks from being included as citations
                    if docs and settings.RAG_ENABLE_LLM_RELEVANCE_FILTER:
                        print(
                            f"🔍 Filtering {len(docs)} retrieved chunks for relevance to question: {question_text[:50]}..."
                        )

                        # Check if this is a Full Document Scan (retriever returns ALL docs)
                        # We detect this by checking if we got more than RAG_NUM_CHUNKS documents
                        is_full_scan = len(docs) > settings.RAG_NUM_CHUNKS
                        if is_full_scan:
                            print(
                                "Full document scan detected - using batch processing for performance"
                            )

                        # Batch / concurrency settings with sensible defaults
                        BATCH_SIZE = getattr(
                            settings, "VERADOC_FULL_SCAN_FILTER_BATCH_SIZE", 10
                        )
                        REQUEST_DELAY = getattr(
                            settings, "PROCESSING_DELAY_BETWEEN_REQUESTS", 0.02
                        )

                        filtered_docs = []
                        print(
                            f"Starting batch processing with batch size {BATCH_SIZE} for {len(docs)} documents"
                        )

                        # Process docs in batches to reduce total runtime while still being rate-limit friendly
                        for start in range(0, len(docs), BATCH_SIZE):
                            batch = docs[start : start + BATCH_SIZE]
                            tasks = []

                            # Create async tasks that run the blocking invoke_llm_async in executor
                            for j, doc in enumerate(batch):
                                doc_idx = start + j

                                async def _check(doc=doc, doc_idx=doc_idx):
                                    try:
                                        # Yield to event loop more frequently for long-running operations
                                        # This prevents connection timeouts and allows progress polling to work
                                        await asyncio.sleep(0.05)

                                        # Update progress during relevance filtering (crucial for long operations)
                                        if task_id:
                                            # Show detailed progress: question X, analyzing chunk Y/Z
                                            progress_tracker.update_stage_progress(
                                                task_id,
                                                "fetching_context",
                                                i,
                                                len(question_list),
                                                translate_progress_message(
                                                    "question_analyzing_chunk",
                                                    user_language,
                                                    question_num=i + 1,
                                                    total_questions=len(question_list),
                                                    chunk_num=doc_idx + 1,
                                                    total_chunks=len(docs),
                                                ),
                                            )

                                        # Check for cancellation during filtering
                                        if request and await request.is_disconnected():
                                            print(
                                                f"❌ CLIENT DISCONNECTED - Stopping relevance filtering at chunk {doc_idx + 1}"
                                            )
                                            raise HTTPException(
                                                status_code=408,
                                                detail="Request cancelled during relevance filtering",
                                            )

                                        print(
                                            f"Analyzing chunk {doc_idx + 1}/{len(docs)} for relevance..."
                                        )

                                        # Use LLM to determine if this chunk is relevant
                                        relevance_check = await invoke_llm_async(
                                            llm,
                                            settings.VERADOC_RELEVANCE_FILTER_PROMPT_TEMPLATE,
                                            {
                                                "chunk": doc.page_content or "",
                                                "question": question_text,
                                            },
                                        )

                                        # Check for cancellation after LLM call
                                        if request and await request.is_disconnected():
                                            print(
                                                f"❌ CLIENT DISCONNECTED - Stopping after relevance check at chunk {doc_idx + 1}"
                                            )
                                            raise HTTPException(
                                                status_code=408,
                                                detail="Request cancelled during relevance filtering",
                                            )

                                        # Yield again after LLM call to prevent connection timeout
                                        await asyncio.sleep(0.02)

                                        return doc_idx, doc, relevance_check

                                    except Exception as filter_error:
                                        print(
                                            f"Error filtering chunk {doc_idx + 1}: {filter_error}"
                                        )
                                        # On error, include the chunk to be safe
                                        return doc_idx, doc, None

                                tasks.append(asyncio.create_task(_check()))

                            # Await this batch and handle results
                            print(
                                f"Awaiting batch {start // BATCH_SIZE + 1} with {len(tasks)} tasks"
                            )
                            results = await asyncio.gather(
                                *tasks, return_exceptions=True
                            )
                            print(f"Batch completed, got {len(results)} results")

                            for res in results:
                                if isinstance(res, Exception):
                                    # On error, include the chunk to be safe (preserve previous behavior)
                                    print(
                                        f"Warning: error during batch relevance check: {res}"
                                    )
                                    continue

                                doc_idx, doc, relevance_check = res
                                # Filter based on LLM response (same logic as before)
                                if (
                                    relevance_check
                                    and "No relevant information found"
                                    not in relevance_check
                                ):
                                    print(f"✅ Chunk {doc_idx + 1} is relevant")
                                    filtered_docs.append(doc)
                                else:
                                    print(
                                        f"❌ Chunk {doc_idx + 1} is not relevant - excluding from citations"
                                    )

                            # Small sleep between batches to help with rate limits and to yield loop
                            if REQUEST_DELAY:
                                await asyncio.sleep(REQUEST_DELAY)

                        print(
                            f"📊 Relevance filtering: {len(filtered_docs)}/{len(docs)} chunks are relevant"
                        )
                        docs = filtered_docs

                    # Build context from filtered documents
                    context = "\n\n".join(
                        [doc.page_content for doc in docs if doc.page_content]
                    )
                    print(
                        f"Final context length: {len(context)} characters from {len(docs)} documents"
                    )

                    # Store source documents for citation (now filtered if Full Document Scan)
                    source_citations = []
                    for doc in docs:
                        try:
                            # Process metadata and source citations
                            metadata = (
                                doc.metadata.copy()
                                if hasattr(doc, "metadata") and doc.metadata
                                else {}
                            )

                            # Source lookup logic (same as original)
                            if "source" in metadata and isinstance(
                                metadata["source"], str
                            ):
                                source_path = metadata["source"]
                                raw_filename = Path(source_path).name

                                match = re.search(r"^[^_]*_(.+)$", raw_filename)
                                if match:
                                    filename = match.group(1)
                                else:
                                    filename = raw_filename

                                try:
                                    source_entry = session.exec(
                                        select(Source).where(Source.name == filename)
                                    ).first()

                                    if not source_entry:
                                        source_entry = session.exec(
                                            select(Source).where(
                                                Source.name == raw_filename
                                            )
                                        ).first()

                                    if source_entry:
                                        metadata["source_data_id"] = str(
                                            source_entry.source_data_id
                                        )
                                except Exception as source_lookup_error:
                                    print(
                                        f"Error looking up source: {source_lookup_error}"
                                    )

                            source = {
                                "content": doc.page_content or "",
                                "metadata": metadata,
                            }
                            source_citations.append(source)
                        except Exception as citation_error:
                            print(f"Error processing citation: {citation_error}")
                            continue

            except Exception as retrieval_error:
                print(
                    f"Error retrieving documents for question '{question_text[:50]}...': {retrieval_error}"
                )
                context = "Error occurred while retrieving relevant documents from the knowledge base."
                source_citations = []

            try:
                # Step 2: Get the relevant policy context for this question with chunking for large contexts
                print("Generating context for question...")

                # Check if context is too large and needs chunking
                from app.services.text_processing import estimate_tokens, chunk_text

                context_tokens = estimate_tokens(context)

                if context_tokens > settings.VERADOC_KB_CHUNK_SIZE_LIMIT:
                    print(
                        f"⚠️ Large context detected ({context_tokens} tokens). Chunking for processing..."
                    )

                    # Chunk the context to prevent rate limit issues
                    context_chunks = chunk_text(
                        context, max_tokens=settings.VERADOC_KB_CHUNK_SIZE_LIMIT
                    )

                    # Process chunks and synthesize context
                    chunk_contexts = []
                    for chunk_idx, chunk in enumerate(context_chunks):
                        try:
                            # Update progress during context chunk processing
                            if task_id:
                                progress_tracker.update_stage_progress(
                                    task_id,
                                    "fetching_context",
                                    i,
                                    len(question_list),
                                    translate_progress_message(
                                        "question_processing_context_chunk",
                                        user_language,
                                        question_num=i + 1,
                                        total_questions=len(question_list),
                                        chunk_num=chunk_idx + 1,
                                        total_chunks=len(context_chunks),
                                    ),
                                )
                                await asyncio.sleep(
                                    0.01
                                )  # Allow progress API to respond

                            # Add delay between chunks
                            if (
                                chunk_idx > 0
                                and settings.VERADOC_ENABLE_PROCESSING_DELAYS
                            ):
                                await asyncio.sleep(
                                    settings.PROCESSING_DELAY_BETWEEN_CHUNKS
                                )

                            print(
                                f"Processing context chunk {chunk_idx+1}/{len(context_chunks)}"
                            )
                            chunk_context = await invoke_llm_async(
                                llm,
                                context_prompt_template,
                                {"context": chunk, "question": question_text},
                            )

                            # Check for cancellation after LLM call
                            if request and await request.is_disconnected():
                                print(
                                    f"❌ CLIENT DISCONNECTED - Stopping after context chunk processing at chunk {chunk_idx + 1}"
                                )
                                raise HTTPException(
                                    status_code=408,
                                    detail="Request cancelled during context chunk processing",
                                )

                            chunk_contexts.append(chunk_context)

                        except Exception as chunk_error:
                            print(
                                f"Error processing context chunk {chunk_idx+1}: {chunk_error}"
                            )
                            chunk_contexts.append(
                                f"Error processing part of context: {str(chunk_error)}"
                            )

                    # Combine chunk contexts
                    question_context = "\n\n".join(
                        [ctx for ctx in chunk_contexts if ctx]
                    )

                else:
                    # Context is small enough, process normally
                    question_context = await invoke_llm_async(
                        llm,
                        context_prompt_template,
                        {"context": context, "question": question_text},
                    )

                    # Check for cancellation after LLM call
                    if request and await request.is_disconnected():
                        print(
                            f"❌ CLIENT DISCONNECTED - Stopping after context generation"
                        )
                        raise HTTPException(
                            status_code=408,
                            detail="Request cancelled during context generation",
                        )

                print(f"Got context: {question_context[:100]}...")

                # Translate the question context if needed
                # question_context = await translate_text_if_needed(
                #     question_context, session, current_user, llm
                # )

            except Exception as context_error:
                print(f"Error generating context for question: {context_error}")

                # Implement fallback strategy for rate limit errors
                if (
                    "rate limiter" in str(context_error).lower()
                    or "rate limit" in str(context_error).lower()
                ):
                    print("🚨 Rate limit detected. Using fallback context strategy...")

                    # Fallback 1: Use a simplified context from document titles/metadata
                    if source_citations:
                        fallback_context = "Based on available documents: " + "; ".join(
                            [
                                f"Document: {citation.get('metadata', {}).get('source', 'Unknown')}"
                                for citation in source_citations[
                                    :5
                                ]  # Limit to first 5 documents
                            ]
                        )
                        question_context = f"Limited context due to processing constraints: {fallback_context}"
                    else:
                        # Fallback 2: Generic context about knowledge base availability
                        question_context = "Knowledge base documents are available but detailed context generation failed due to processing constraints. Please answer based on the question content."

                    print(f"Using fallback context: {question_context[:100]}...")
                else:
                    # For other errors, use error message
                    question_context = f"Error generating context: {str(context_error)}"

                # Translate the fallback context if needed
                # question_context = await translate_text_if_needed(
                #     question_context, session, current_user, llm
                # )
        else:
            # Skip knowledge base consultation
            question_context = (
                "No policy context consultation requested for this question."
            )
            # question_context = await translate_text_if_needed(
            #     question_context, session, current_user, llm
            # )
            source_citations = []
            print(
                f"Skipping document consultation for question: {question_text[:50]}..."
            )

        # Store the pre-fetched context and citations
        question_contexts[question_text] = {
            "context": question_context,
            "source_citations": source_citations,
            "consult_documents": consult_documents,
        }

    print(f"✅ Pre-fetched context for {len(question_contexts)} questions")
    return question_contexts


def extract_text_from_file(
    file_content: bytes, filename: str, current_user=None
) -> str:
    """Extract text from various file formats using unified document processing."""
    from app.services.document_utils import extract_text_from_file_unified

    if current_user:
        print(
            f"[VERADOC] extract_text_from_file called for {filename} with user {current_user.id}"
        )
    else:
        print(f"[VERADOC] extract_text_from_file called for {filename} with NO user")

    # Uses settings.PDF_PARSING_MODE by default, or user preference if current_user provided
    return extract_text_from_file_unified(
        file_content, filename, current_user=current_user
    )


async def extract_text_from_file_async(
    file_content: bytes, filename: str, current_user=None
) -> str:
    """
    Async wrapper for text extraction to prevent blocking on large files.
    Uses ThreadPoolExecutor for CPU-intensive text extraction operations.
    """
    # Always use thread pool for DOCX files since they can be slow to process
    # regardless of size
    is_docx = filename.lower().endswith((".docx", ".doc"))

    # Define size threshold for async processing
    if is_docx:
        SIZE_THRESHOLD = 10 * 1024  # Very low threshold for DOCX files (10KB)
    else:
        SIZE_THRESHOLD = 50 * 1024  # 50KB for other files

    if len(file_content) > SIZE_THRESHOLD or is_docx:
        print(
            f"File requires thread pool processing ({len(file_content)} bytes, is_docx: {is_docx})"
        )

        # Process in thread pool to avoid blocking the event loop
        loop = asyncio.get_event_loop()
        with ThreadPoolExecutor(max_workers=1) as executor:
            document_text = await loop.run_in_executor(
                executor, extract_text_from_file, file_content, filename, current_user
            )
        return document_text
    else:
        # For small non-DOCX files, process synchronously
        return extract_text_from_file(file_content, filename, current_user)


router = APIRouter(prefix="/veradoc", tags=["veradoc"])


@router.get("/progress/{task_id}")
async def get_veradoc_progress(
    task_id: str,
    current_user: CurrentUser,
    request: FastAPIRequest = None,
) -> Any:
    """
    Get progress information for a VeraDoc task (review).
    """
    # Log incoming request headers to help diagnose cross-host/CORS routing issues
    try:
        client_addr = (
            request.client.host
            if request and getattr(request, "client", None)
            else "unknown"
        )
        host_hdr = request.headers.get("host") if request else None
        origin_hdr = request.headers.get("origin") if request else None
        print(
            f"🔗 VERADOC PROGRESS REQUEST: task_id={task_id}, client={client_addr}, host={host_hdr}, origin={origin_hdr}"
        )
    except Exception as _:
        print("🔗 VERADOC PROGRESS: could not read request headers")

    # Make this async to prevent blocking during intensive operations
    progress_data = progress_tracker.get_progress(task_id)
    if not progress_data:
        # Don't immediately return 404 here. In high-latency flows (LLM calls, heavy processing)
        # the frontend may start polling before the backend has persisted progress or during
        # transient Redis/connectivity issues. Returning 404 causes the frontend to stop
        # polling and surface a "Task not found or expired" error to users.
        #
        # Instead, return a lightweight placeholder indicating the task is initializing.
        # The frontend will continue polling and pick up the real progress once available.
        print(
            f"⚠️ VERADOC PROGRESS: No progress found for task {task_id} - returning initializing placeholder"
        )
        placeholder = {
            "task_id": task_id,
            "operation": "Reviewing documents",
            "stages": {},
            "current_stage": "",
            "percentage": 0.0,
            "status": "pending",
            "message": "Initializing task - progress not yet available. Polling will continue.",
        }
        # Yield control back to the event loop before returning
        await asyncio.sleep(0)
        return placeholder

    # Debug logging to see what's actually being returned
    # print(f"🔍 VERADOC API RETURNING PROGRESS: task_id={task_id}")
    # print(
    #    f"🔍 PROGRESS DATA: status={progress_data.get('status')}, percentage={progress_data.get('percentage')}, current_stage={progress_data.get('current_stage')}"
    # )
    # print(f"🔍 PROGRESS MESSAGE: {progress_data.get('message')}")
    # print(f"🔍 PROGRESS STAGES: {list(progress_data.get('stages', {}).keys())}")

    # Check each stage completion status
    stages = progress_data.get("stages", {})
    for stage_name, stage_data in stages.items():
        completed = (
            stage_data.get("completed", False)
            if isinstance(stage_data, dict)
            else False
        )
        print(f"🔍 STAGE {stage_name}: completed={completed}")

    # Yield control to allow other async operations (like this API call) to run
    await asyncio.sleep(0)

    return progress_data


@router.get("/results/{task_id}")
async def get_veradoc_results(
    task_id: str,
    current_user: CurrentUser,
    request: FastAPIRequest = None,
) -> Any:
    """
    Get the results for a completed VeraDoc review task.
    This endpoint should be called after progress shows status='completed'.
    """
    # Log incoming request headers to help diagnose cross-host/CORS routing issues
    try:
        client_addr = (
            request.client.host
            if request and getattr(request, "client", None)
            else "unknown"
        )
        host_hdr = request.headers.get("host") if request else None
        origin_hdr = request.headers.get("origin") if request else None
        print(
            f"🔗 VERADOC RESULTS REQUEST: task_id={task_id}, client={client_addr}, host={host_hdr}, origin={origin_hdr}"
        )
    except Exception:
        print("🔗 VERADOC RESULTS: could not read request headers")

    # Retrieve results from task metadata
    results = progress_tracker.get_task_metadata(task_id)
    if not results:
        # Don't return a hard 404 here; callers (browsers) may be hitting a different host or race with metadata
        print(
            f"⚠️ VERADOC RESULTS: No metadata found for task {task_id} - returning placeholder to allow frontend to retry"
        )
        placeholder = {
            "task_id": task_id,
            "status": "pending",
            "message": "Results not yet available. Polling will continue.",
        }
        await asyncio.sleep(0)
        return placeholder

    return results


@router.post("/review/task")
async def create_review_task():
    """
    Create a progress tracking task for document review and return task_id immediately.
    This allows frontend to start progress polling before form submission.
    """
    task_id = progress_tracker.create_task(
        "Reviewing documents",
        {
            "setup": 0.05,
            "fetching_context": 0.60,
            "reviewing": 0.30,
            "finalizing": 0.05,
        },
    )
    progress_tracker.update_stage_progress(
        task_id, "setup", 0, 1, "Waiting to start document review..."
    )
    return {"task_id": task_id}


# Initialize the LLM
# llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.0)


def generate_template(questions: List[str]) -> Dict[str, str]:
    """
    Generate a JSON template from a list of questions.
    Each field will have a blank value.
    """
    return {field: "" for field in questions}


def parse_optimization_response(
    llm_response: str, original_qa: dict
) -> ChecklistSuggestion:
    """Parse the LLM optimization response into a structured suggestion."""
    try:
        lines = llm_response.strip().split("\n")
        revised_question = original_qa["question"]  # Default to original
        reason = "No changes suggested"
        needs_revision = False

        for line in lines:
            line = line.strip()
            if line.startswith("REVISED_QUESTION:"):
                revised_question = line.replace("REVISED_QUESTION:", "").strip()
            elif line.startswith("REASON:"):
                reason = line.replace("REASON:", "").strip()
            elif line.startswith("NEEDS_REVISION:"):
                needs_revision_str = line.replace("NEEDS_REVISION:", "").strip().lower()
                needs_revision = needs_revision_str in ["yes", "true", "1"]

        return ChecklistSuggestion(
            original_question=original_qa["question"],
            suggested_question=revised_question,
            reason=reason,
            current_answer=original_qa["answer"],
            needs_revision=needs_revision,
            policy_context=original_qa.get("context", ""),
        )
    except Exception as e:
        print(f"Error parsing optimization response: {e}")
        return ChecklistSuggestion(
            original_question=original_qa["question"],
            suggested_question=original_qa["question"],
            reason=f"Error parsing suggestion: {str(e)}",
            current_answer=original_qa["answer"],
            needs_revision=False,
            policy_context=original_qa.get("context", ""),
        )


def needs_optimization(answer: str) -> bool:
    """Determine if a checklist question needs optimization based on the answer."""
    answer_lower = answer.lower()
    negative_indicators = [
        "no",
        "not",
        "insufficient",
        "missing",
        "absent",
        "lacks",
        "does not",
        "doesn't",
        "cannot",
        "can't",
        "unable",
        "fails",
        "inadequate",
        "incomplete",
        "unclear",
        "vague",
        "poorly",
    ]

    # Check if any negative indicators are present
    for indicator in negative_indicators:
        if indicator in answer_lower:
            return True

    return False


# Add the new endpoint
@router.post("/process-rag", response_model=VeraDocResponse)
async def process_rag_checklist(
    session: SessionDep,
    current_user: CurrentUser,
    questions: Optional[str] = Form(None),
    knowledge_base_id: str = Form(...),
    files: List[UploadFile] = File(...),
    custom_instructions: Optional[str] = Form(None),
    search_mode: str = Form("vector"),
    task_id: Optional[str] = Form(None),
    request: FastAPIRequest = None,
):
    """
    Process the uploaded files using RAG with a knowledge base.
    Includes real-time progress tracking similar to ReportGenie.
    """
    print("process_rag_checklist function invoked!")
    print(f"Received search_mode: {search_mode}")
    print(
        f"Request data: knowledge_base_id={knowledge_base_id}, questions length={len(questions) if questions else 0}"
    )

    # Input validation
    if not knowledge_base_id:
        raise HTTPException(status_code=400, detail="Knowledge base ID is required")

    if not questions or not questions.strip():
        raise HTTPException(status_code=400, detail="Questions are required")

    # Check for at least one file
    total_files = len(files) if files else 0
    if total_files == 0:
        raise HTTPException(status_code=400, detail="At least one file is required")

    # Validate search mode
    if search_mode not in ["vector", "full_scan"]:
        print(f"Warning: Invalid search mode '{search_mode}', defaulting to 'vector'")
        search_mode = "vector"

    # Create or initialize progress tracking task
    if not task_id:
        # No task_id provided, create a new one
        task_id = progress_tracker.create_task(
            "Reviewing documents",
            {
                "setup": 0.05,
                "fetching_context": 0.60,
                "reviewing": 0.30,
                "finalizing": 0.05,
            },
        )
        print(f"Created new task_id: {task_id}")
    else:
        # Task_id provided by frontend, initialize it in the progress tracker
        print(f"Using provided task_id: {task_id}")
        progress_tracker.create_task_with_id(
            task_id,
            "Reviewing documents",
            {
                "setup": 0.05,
                "fetching_context": 0.60,
                "reviewing": 0.30,
                "finalizing": 0.05,
            },
        )

    # Create request_data object for backward compatibility with rest of the code
    class RequestData:
        pass

    request_data = RequestData()
    request_data.questions = questions
    request_data.knowledge_base_id = knowledge_base_id
    request_data.custom_instructions = custom_instructions
    request_data.search_mode = search_mode
    request_data.task_id = task_id

    # Get user's preferred language for progress message translation
    user_language = getattr(current_user, "preferred_language", "en") or "en"

    progress_tracker.update_stage_progress(
        task_id,
        "setup",
        0,
        1,
        translate_progress_message("initializing_document_review", user_language),
    )

    try:
        print("Processing RAG checklist...")

        # 1. Retrieve knowledge base from database
        kb = session.get(KnowledgeBase, request_data.knowledge_base_id)
        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        progress_tracker.complete_stage(
            task_id,
            "setup",
            translate_progress_message("setup_complete", user_language),
        )

        # Start fetching context stage
        progress_tracker.update_stage_progress(
            task_id,
            "fetching_context",
            0,
            1,
            translate_progress_message(
                "preparing_retrieve_policy_context", user_language
            ),
        )

        await asyncio.sleep(0.01)  # Allow progress API to respond

        # 2. Create a temporary directory for ChromaDB
        with tempfile.TemporaryDirectory() as temp_dir:
            # Extract the zipped ChromaDB into the temp directory
            if kb.storage_type == "file" and kb.file_path:
                # File-based storage: extract from file path
                if os.path.exists(kb.file_path):
                    with zipfile.ZipFile(kb.file_path, "r") as zip_ref:
                        zip_ref.extractall(temp_dir)
                else:
                    raise HTTPException(
                        status_code=400, detail="Knowledge base file not found on disk"
                    )
            elif kb.data:
                # Database storage: extract from data field
                with zipfile.ZipFile(BytesIO(kb.data), "r") as zip_ref:
                    zip_ref.extractall(temp_dir)
            else:
                raise HTTPException(
                    status_code=400, detail="Knowledge base has no vector database data"
                )

            # 3. Load the vector database with the SAME model used to create the knowledge base
            # Use the knowledge base's specific embedding model if available
            if kb.embedding_model_id:
                embedding_model = session.get(EmbeddingModel, kb.embedding_model_id)
                if embedding_model:
                    # Use the KB's original model
                    model_id = embedding_model.model_id
                    provider = embedding_model.provider
                    print(
                        f"Using knowledge base's original embedding model: {model_id}"
                    )
                else:
                    # Fallback if the model was deleted from the database
                    embedding_info = get_embedding_model(session, current_user)
                    model_id = embedding_info["model_id"]
                    provider = embedding_info["provider"]
                    print(
                        f"Original embedding model not found, using current default: {model_id}"
                    )
            else:
                # For knowledge bases created before tracking embedding models
                embedding_info = get_embedding_model(session, current_user)
                model_id = embedding_info["model_id"]
                provider = embedding_info["provider"]
                print(
                    f"Knowledge base has no embedding model record, using current default: {embedding_info}"
                )

            print(f"Initializing embedding model: {model_id} ({provider})")
            embeddings = load_embeddings_model(provider=provider, model_id=model_id)
            chroma_db = Chroma(
                persist_directory=temp_dir, embedding_function=embeddings
            )

            # Print all metadata in the vectorstore
            # print("======= CHROMA VECTORDB METADATA CONTENTS =======")
            # Get all documents with their metadata
            all_docs = chroma_db.get()
            if all_docs and "metadatas" in all_docs and False:
                for i, metadata in enumerate(all_docs["metadatas"]):
                    print(f"Document {i+1} Metadata: {metadata}")
                    # If you want to see document content as well
                    if "documents" in all_docs and i < len(all_docs["documents"]):
                        doc_preview = (
                            all_docs["documents"][i][:200] + "..."
                            if len(all_docs["documents"][i]) > 100
                            else all_docs["documents"][i]
                        )
                        print(f"Content preview: {doc_preview}")
                    print("-" * 50)
            else:
                print("No documents or metadata found in the vectorstore")
            print("================================================")

            # Create retriever based on search mode
            try:
                if request_data.search_mode == "full_scan":
                    # Full document scan: retrieve all documents from the knowledge base
                    print("Using full document scan mode")

                    # Test ChromaDB collection access first
                    try:
                        # Try multiple ways to get collection count
                        if hasattr(chroma_db, "_collection") and hasattr(
                            chroma_db._collection, "count"
                        ):
                            collection_count = chroma_db._collection.count()
                        else:
                            # Fallback: try to get a small sample to test access
                            test_data = chroma_db.get(limit=1)
                            collection_count = (
                                len(test_data.get("documents", [])) if test_data else 0
                            )

                        print(
                            f"Knowledge base collection has {collection_count} documents"
                        )

                        if collection_count == 0:
                            print(
                                f"Warning: Knowledge base '{request_data.knowledge_base_id}' appears to be empty"
                            )
                            # Don't raise error, let it proceed with empty results

                    except Exception as count_error:
                        print(f"Warning: Could not get collection count: {count_error}")
                        # Continue anyway - the get_relevant_documents method will handle empty collections

                    # Create a simple retriever that returns all documents
                    class FullScanRetriever:
                        def __init__(self, chroma_db):
                            self.chroma_db = chroma_db

                        def get_relevant_documents(self, query):
                            try:
                                print(
                                    f"FullScanRetriever: Processing query '{query[:50]}...'"
                                )

                                # Get all documents for true Full Document Scan
                                try:
                                    # Get all documents - no sampling to ensure nothing is missed
                                    all_data = self.chroma_db.get()

                                    if all_data and "documents" in all_data:
                                        print(
                                            f"📚 Full Document Scan: Processing ALL {len(all_data['documents'])} documents"
                                        )

                                except Exception as get_error:
                                    print(
                                        f"Error getting documents from ChromaDB: {get_error}"
                                    )
                                    return []

                                # Convert to Document objects similar to vector search
                                documents = []

                                if (
                                    all_data
                                    and "documents" in all_data
                                    and all_data["documents"]
                                ):
                                    print(
                                        f"FullScanRetriever: Found {len(all_data['documents'])} total documents"
                                    )

                                    for i, doc_content in enumerate(
                                        all_data["documents"]
                                    ):
                                        try:
                                            # Safely get metadata
                                            metadata = {}
                                            if (
                                                "metadatas" in all_data
                                                and isinstance(
                                                    all_data["metadatas"], list
                                                )
                                                and i < len(all_data["metadatas"])
                                                and all_data["metadatas"][i] is not None
                                            ):

                                                raw_metadata = all_data["metadatas"][i]
                                                if isinstance(raw_metadata, dict):
                                                    metadata = raw_metadata
                                                else:
                                                    print(
                                                        f"Warning: Metadata at index {i} is not a dict: {type(raw_metadata)}"
                                                    )
                                                    metadata = {}

                                            # Ensure doc_content is a string
                                            content = (
                                                str(doc_content)
                                                if doc_content is not None
                                                else ""
                                            )

                                            if (
                                                content
                                            ):  # Only add documents with content
                                                documents.append(
                                                    LangchainDocument(
                                                        page_content=content,
                                                        metadata=metadata,
                                                    )
                                                )
                                        except Exception as doc_error:
                                            print(
                                                f"Error processing document {i}: {doc_error}"
                                            )
                                            # Continue processing other documents instead of failing
                                            continue
                                else:
                                    print(
                                        "FullScanRetriever: No documents found in knowledge base"
                                    )

                                print(
                                    f"FullScanRetriever: Returning {len(documents)} documents"
                                )
                                return documents

                            except Exception as e:
                                print(
                                    f"Error in FullScanRetriever.get_relevant_documents: {e}"
                                )
                                import traceback

                                traceback.print_exc()
                                # Return empty list instead of raising exception
                                return []

                    retriever = FullScanRetriever(chroma_db)
                else:
                    # Vector search mode (default) with enhanced content filtering
                    print("Using enhanced vector search mode with content filtering")
                    try:
                        retriever = (
                            SmartRetrieverFactory.create_academic_paper_retriever(
                                chroma_db=chroma_db,
                                search_kwargs={"k": settings.RAG_NUM_CHUNKS},
                            )
                        )
                        print(
                            "Enhanced academic retriever created successfully - will filter bibliography content"
                        )
                    except Exception as retriever_error:
                        print(f"Error creating enhanced retriever: {retriever_error}")
                        # Fallback to general document retriever with smart filtering
                        retriever = (
                            SmartRetrieverFactory.create_general_document_retriever(
                                chroma_db=chroma_db,
                                search_kwargs={"k": settings.RAG_NUM_CHUNKS},
                            )
                        )
                        print("Using fallback smart retriever with content filtering")

            except Exception as retriever_setup_error:
                print(f"Error setting up retriever: {retriever_setup_error}")
                traceback.print_exc()
                raise HTTPException(
                    status_code=500,
                    detail=f"Failed to set up document retriever: {str(retriever_setup_error)}",
                )

            # 4. Initialize the LLM
            # llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.0)
            print("Now loading default LLM for session with following info:")
            print("Session:", session)
            llm = get_default_llm(session, current_user)
            print("LLM successfully loaded.")

            # Check if LLM supports vision
            from app.services.vision_service import VisionService

            vision_enabled = VisionService.is_vision_enabled(llm, current_user)

            # 5. Define the prompts for the different stages
            context_prompt_template = settings.VERADOC_CONTEXT_PROMPT_TEMPLATE
            qa_prompt_template = settings.VERADOC_QA_PROMPT_TEMPLATE
            final_prompt_template = settings.VERADOC_FINAL_PROMPT_TEMPLATE

            # 6. Parse questions first - support both legacy string format and new structured format
            try:
                # Try to parse as structured JSON format
                questions_data = json.loads(request_data.questions)
                if isinstance(questions_data, list) and all(
                    isinstance(item, dict)
                    and "text" in item
                    and "consultDocuments" in item
                    for item in questions_data
                ):
                    # New structured format
                    question_list = questions_data
                else:
                    raise ValueError("Not structured format")
            except (json.JSONDecodeError, ValueError):
                # Fallback to legacy string format
                question_texts = request_data.questions.strip().split("\n")
                question_list = [
                    {"text": q.strip(), "consultDocuments": True}
                    for q in question_texts
                    if q.strip()
                ]

            # 7. PRE-FETCH KNOWLEDGE BASE CONTEXT (OPTIMIZATION)
            # This step is the same regardless of which document is being reviewed
            print(
                "🚀 OPTIMIZATION: Pre-fetching knowledge base context for all questions..."
            )
            try:
                question_contexts = await prefetch_knowledge_base_context(
                    retriever=retriever,
                    question_list=question_list,
                    llm=llm,
                    context_prompt_template=context_prompt_template,
                    session=session,
                    current_user=current_user,
                    request=request,
                    task_id=task_id,
                    user_language=user_language,
                )
                print(f"✅ Pre-fetched context for {len(question_contexts)} questions")
            except HTTPException:
                # Re-raise cancellation errors
                raise
            except Exception as prefetch_error:
                print(f"Error during context pre-fetch: {prefetch_error}")
                # Fall back to processing without pre-fetch
                question_contexts = {}

            # Complete context fetching stage
            progress_tracker.complete_stage(
                task_id,
                "fetching_context",
                translate_progress_message("policy_context_retrieved", user_language),
            )

            # Start reviewing documents stage
            progress_tracker.update_stage_progress(
                task_id,
                "reviewing",
                0,
                total_files,
                translate_progress_message("beginning_document_review", user_language),
            )
            await asyncio.sleep(0.01)  # Allow progress API to respond

            # 8. Process each uploaded file using the pre-fetched context
            qa_pairs = []

            # 8. Process each uploaded file using the pre-fetched context
            all_files_results = []

            # Support multiple files - process each one
            for file_index, file in enumerate(files):
                # Update progress for each file
                file_preview = (
                    file.filename[:30] + "..."
                    if file.filename and len(file.filename) > 30
                    else file.filename
                )
                progress_tracker.update_stage_progress(
                    task_id,
                    "reviewing",
                    file_index,
                    len(files),
                    f"Reviewing file {file_index + 1}/{len(files)}: {file_preview}",
                )

                await asyncio.sleep(0.01)  # Allow progress API to respond

                if not file.filename:
                    print(f"Skipping file {file_index + 1}: No filename")
                    continue

                print(f"Processing file {file_index + 1}/{len(files)}: {file.filename}")

                # Check for cancellation before processing each file
                try:
                    if request and await request.is_disconnected():
                        print(
                            f"❌ CLIENT DISCONNECTED - Stopping at file {file_index + 1}"
                        )
                        return VeraDocResponse(
                            results={
                                "status": "cancelled",
                                "message": "Request cancelled - client disconnected",
                            }
                        )
                except Exception as e:
                    print(f"Warning: Could not check disconnect status: {e}")

                qa_pairs = []

                try:
                    content = await file.read()
                    if not content:
                        print(f"File {file.filename} appears to be empty, skipping")
                        continue

                    # Check if this is a potentially slow file to process
                    is_large_file = len(content) > 50000
                    is_docx_file = file.filename.lower().endswith((".docx", ".doc"))
                    needs_special_handling = is_large_file or is_docx_file

                    # Extract text using unified processing
                    print(
                        f"Processing file with unified text extraction: {file.filename}"
                    )
                    document_text = await extract_text_from_file_async(
                        content, file.filename, current_user
                    )
                    # Clean surrogates from document_text to prevent encoding issues
                    document_text = re.sub(r"[\ud800-\udfff]", "", document_text)

                    # Extract images if vision is enabled
                    document_images = []
                    if vision_enabled:
                        try:
                            from app.services.document_utils import (
                                extract_documents_and_images_from_file_unified,
                            )

                            _, document_images = (
                                extract_documents_and_images_from_file_unified(
                                    content, file.filename
                                )
                            )
                            print(
                                f"Extracted {len(document_images)} images from {file.filename}"
                            )
                        except Exception as img_error:
                            print(
                                f"Warning: Could not extract images from {file.filename}: {img_error}"
                            )

                    # Handle case where no text was extracted
                    if not document_text or document_text.strip() == "":
                        if vision_enabled and document_images:
                            print(
                                f"No text extracted from {file.filename}, but {len(document_images)} images found. Using vision analysis as fallback."
                            )
                            document_text = f"This document ({file.filename}) contains images but no extractable text. Vision analysis will be used to answer questions about the visual content."
                        else:
                            print(
                                f"Could not extract text from {file.filename}, skipping"
                            )
                            continue

                    print(
                        f"Extracted {len(document_text)} characters from {file.filename}"
                    )
                    await file.seek(0)

                except Exception as file_error:
                    print(f"Error processing file {file.filename}: {file_error}")
                    # Add error result for this file and continue with next file
                    all_files_results.append(
                        {
                            "filename": file.filename,
                            "final_evaluation": f"Error processing file {file.filename}: {str(file_error)}",
                            "qa_pairs": [],
                            "interaction_id": None,
                        }
                    )
                    continue

                # 9. Process each question using the PRE-FETCHED context
                for i, question_item in enumerate(question_list):
                    # Yield to event loop at start of each question
                    await asyncio.sleep(0.05)  # Prevent connection timeouts

                    # Update progress for answering questions
                    # Calculate overall progress: (files completed * questions per file + current question) / total work
                    total_questions = len(files) * len(question_list)
                    questions_completed = file_index * len(question_list) + i

                    question_preview = (
                        question_item.get("text", "")[:50] + "..."
                        if len(question_item.get("text", "")) > 50
                        else question_item.get("text", "")
                    )
                    progress_tracker.update_stage_progress(
                        task_id,
                        "reviewing",
                        questions_completed,
                        total_questions,
                        translate_progress_message(
                            "answering_question",
                            user_language,
                            question_num=i + 1,
                            total_questions=len(question_list),
                            file_num=file_index + 1,
                            total_files=len(files),
                            question_preview=question_preview,
                        ),
                    )

                    # Add delay between question processing to prevent rate limit exhaustion
                    if i > 0 and settings.VERADOC_ENABLE_PROCESSING_DELAYS:
                        await asyncio.sleep(settings.PROCESSING_DELAY_BETWEEN_REQUESTS)

                    try:
                        # Check if client has disconnected before processing each question
                        try:
                            if request and await request.is_disconnected():
                                print(
                                    f"❌ CLIENT DISCONNECTED - Stopping at question {i + 1} for file {file.filename}"
                                )
                                return VeraDocResponse(
                                    results={
                                        "status": "cancelled",
                                        "message": "Request cancelled - client disconnected",
                                    }
                                )
                        except Exception as e:
                            print(f"Warning: Could not check disconnect status: {e}")

                        question_text = re.sub(
                            r"[\ud800-\udfff]", "", question_item.get("text", "")
                        ).strip()

                        if not question_text:
                            print(f"Skipping empty question at index {i}")
                            continue

                        print(
                            f"Processing question {i+1}/{len(question_list)} for {file.filename}: {question_text[:50]}..."
                        )

                        # 🚀 OPTIMIZATION: Use pre-fetched context instead of retrieving again
                        # Debug: Show what keys are available
                        print(
                            f"🔍 DEBUG: Looking for question key: '{question_text[:50]}'"
                        )
                        print(
                            f"🔍 DEBUG: Available keys in question_contexts: {[key[:50] for key in question_contexts.keys()]}"
                        )

                        if question_text in question_contexts:
                            # Use pre-fetched context and citations
                            cached_context = question_contexts[question_text]
                            question_context = cached_context["context"]
                            # Clean surrogates from question_context
                            question_context = re.sub(
                                r"[\ud800-\udfff]", "", question_context
                            )
                            source_citations = cached_context["source_citations"]
                            consult_documents = cached_context["consult_documents"]
                            print(
                                f"✅ Using pre-fetched context for question: {question_text[:30]}..."
                            )
                        else:
                            # Fallback to original logic if pre-fetch failed for this question
                            print(
                                f"⚠️ No pre-fetched context found for question: {question_text[:30]}..., using fallback"
                            )
                            consult_documents = question_item.get(
                                "consultDocuments", True
                            )

                            if consult_documents:
                                try:
                                    docs = retriever.get_relevant_documents(
                                        question_text
                                    )
                                    if not docs:
                                        context = "No relevant documents found in the knowledge base for this question."
                                        source_citations = []
                                    else:
                                        context = "\n\n".join(
                                            [
                                                doc.page_content
                                                for doc in docs
                                                if doc.page_content
                                            ]
                                        )
                                        source_citations = []
                                        for doc in docs:
                                            metadata = (
                                                doc.metadata.copy()
                                                if doc.metadata
                                                else {}
                                            )

                                            # Source lookup logic (same as prefetch)
                                            if "source" in metadata and isinstance(
                                                metadata["source"], str
                                            ):
                                                source_path = metadata["source"]
                                                raw_filename = Path(source_path).name

                                                match = re.search(
                                                    r"^[^_]*_(.+)$", raw_filename
                                                )
                                                if match:
                                                    filename = match.group(1)
                                                else:
                                                    filename = raw_filename

                                                try:
                                                    source_entry = session.exec(
                                                        select(Source).where(
                                                            Source.name == filename
                                                        )
                                                    ).first()

                                                    if not source_entry:
                                                        source_entry = session.exec(
                                                            select(Source).where(
                                                                Source.name
                                                                == raw_filename
                                                            )
                                                        ).first()

                                                    if source_entry:
                                                        metadata["source_data_id"] = (
                                                            str(
                                                                source_entry.source_data_id
                                                            )
                                                        )
                                                except Exception as source_lookup_error:
                                                    print(
                                                        f"Error looking up source: {source_lookup_error}"
                                                    )

                                            source_citations.append(
                                                {
                                                    "content": doc.page_content or "",
                                                    "metadata": metadata,
                                                }
                                            )

                                    question_context = await invoke_llm_async(
                                        llm,
                                        context_prompt_template,
                                        {"context": context, "question": question_text},
                                    )

                                    # Check for cancellation after LLM call
                                    if request and await request.is_disconnected():
                                        print(
                                            f"❌ CLIENT DISCONNECTED - Stopping after fallback context generation"
                                        )
                                        raise HTTPException(
                                            status_code=408,
                                            detail="Request cancelled during fallback context generation",
                                        )

                                    # question_context = await translate_text_if_needed(
                                    #     question_context, session, current_user, llm
                                    # )
                                    # Clean surrogates from question_context
                                    question_context = re.sub(
                                        r"[\ud800-\udfff]", "", question_context
                                    )
                                except Exception as fallback_error:
                                    print(
                                        f"Error in fallback context generation: {fallback_error}"
                                    )
                                    question_context = f"Error generating context: {str(fallback_error)}"
                                    source_citations = []
                            else:
                                question_context = "No policy context consultation requested for this question."
                                # question_context = await translate_text_if_needed(
                                #     question_context, session, current_user, llm
                                # )
                                # Clean surrogates from question_context
                                question_context = re.sub(
                                    r"[\ud800-\udfff]", "", question_context
                                )
                                source_citations = []

                        print("Generating answer based on document and context...")

                        # Prepare custom instructions section
                        custom_instructions_section = ""
                        if (
                            hasattr(request_data, "custom_instructions")
                            and request_data.custom_instructions
                        ):
                            custom_instructions_section = f"\nADDITIONAL INSTRUCTIONS:\n{request_data.custom_instructions.strip()}\n"

                        # Prepare language instruction - ALWAYS include for consistency
                        language_instruction = ""
                        if user_language:
                            language_name = settings.SUPPORTED_LANGUAGES.get(
                                user_language, user_language
                            )
                            language_instruction = (
                                f"Respond in this language: {language_name}."
                            )

                        # DEBUG: Print the full prompt sent to the LLM
                        try:
                            rendered_prompt = qa_prompt_template.format(
                                document_text=document_text,
                                question=question_text,
                                question_context=question_context,
                                custom_instructions_section=custom_instructions_section,
                                language_instruction=language_instruction,
                            )
                        except Exception as e:
                            rendered_prompt = f"[ERROR rendering prompt: {e}]"
                        # Clean surrogates from rendered_prompt before printing to avoid UnicodeEncodeError
                        clean_prompt = re.sub(r"[\ud800-\udfff]", "", rendered_prompt)
                        # print(
                        #    "\n===== VERADOC_QA_PROMPT_TEMPLATE PROMPT SENT TO LLM =====\n"
                        # )
                        # print(clean_prompt)
                        # print(
                        #    "\n========================================================\n"
                        # )

                        try:
                            # Generate text-based answer
                            print(
                                f"DEBUG: language_instruction = '{language_instruction}'"
                            )
                            answer = await invoke_llm_async(
                                llm,
                                qa_prompt_template,
                                {
                                    "document_text": document_text,
                                    "question": question_text,
                                    "question_context": question_context,
                                    "custom_instructions_section": custom_instructions_section,
                                    "language_instruction": language_instruction,
                                },
                            )

                            # Check for cancellation after LLM call
                            if request and await request.is_disconnected():
                                print(
                                    f"❌ CLIENT DISCONNECTED - Stopping after question answering"
                                )
                                raise HTTPException(
                                    status_code=408,
                                    detail="Request cancelled during question answering",
                                )

                            # Yield after LLM call to prevent connection timeout
                            await asyncio.sleep(0.05)

                            print(f"Got text answer: {answer[:100]}...")

                            # Add vision analysis if images exist and LLM supports it
                            if vision_enabled and document_images:
                                print(
                                    f"Adding vision analysis for question: {question_text[:50]}..."
                                )

                                # Prepare images for processing
                                image_data_list = []
                                for idx, img_b64 in enumerate(document_images):
                                    image_data_list.append(
                                        {
                                            "image_data": img_b64,
                                            "source_file": file.filename,
                                            "image_index": idx,
                                            "metadata": {
                                                "extracted_from": file.filename
                                            },
                                        }
                                    )

                                try:
                                    vision_variables = {
                                        "question": question_text,
                                        "filename": file.filename,
                                        "custom_instructions": (
                                            request_data.custom_instructions
                                            if hasattr(
                                                request_data, "custom_instructions"
                                            )
                                            else ""
                                        ),
                                        "language_instruction": language_instruction,
                                    }

                                    print(
                                        f"DEBUG: vision language_instruction = '{vision_variables.get('language_instruction', '')}'"
                                    )
                                    vision_analysis = await VisionService.process_images_with_prompt(
                                        llm=llm,
                                        images=image_data_list,
                                        prompt_template=settings.VERADOC_VISION_PROMPT_TEMPLATE,
                                        variables=vision_variables,
                                    )

                                    # Combine text and vision analysis seamlessly (photogenic integration)
                                    if (
                                        "contains images but no extractable text"
                                        in document_text
                                        and len(document_text) < 200
                                    ):
                                        # For image-only documents, use vision-primary combination
                                        combined_answer = f"Based on visual analysis of the document: {vision_analysis}. This assessment relies on image content, as the document contains minimal extractable text."
                                    else:
                                        # Normal text + vision combination: Integrate narratively without markers
                                        # Clean the vision_analysis to remove any residual markers (if present)
                                        vision_analysis_clean = re.sub(
                                            r"## .*? ##|---.*?---", "", vision_analysis
                                        ).strip()
                                        # Combine into a flowing response
                                        combined_answer = f"Text Analysis: {answer} Visual Analysis: {vision_analysis_clean}"

                                    answer = combined_answer
                                    print(
                                        f"Combined answer with vision analysis: {answer[:100]}..."
                                    )

                                except Exception as vision_error:
                                    print(
                                        f"Vision analysis error for question '{question_text[:50]}...': {vision_error}"
                                    )
                                    # Continue with text-only answer

                            # Clean surrogates from answer
                            answer = re.sub(r"[\ud800-\udfff]", "", answer)

                        except Exception as answer_error:
                            print(
                                f"Error generating answer for question: {answer_error}"
                            )
                            answer = f"Error generating answer: {str(answer_error)}"

                        print(
                            "Source citations for question:", question_text
                        )  # Store the question-answer pair with context
                        qa_pairs.append(
                            {
                                "question": question_text,
                                "answer": answer,
                                "context": question_context,
                                "source_citations": source_citations,
                            }
                        )

                        # Update progress AFTER question is completed
                        total_questions = len(files) * len(question_list)
                        questions_completed = file_index * len(question_list) + (
                            i + 1
                        )  # +1 because we just completed this question
                        print(
                            f"📊 PROGRESS UPDATE: Completed {questions_completed}/{total_questions} questions"
                        )
                        progress_tracker.update_stage_progress(
                            task_id,
                            "reviewing",
                            questions_completed,
                            total_questions,
                            translate_progress_message(
                                "completed_question",
                                user_language,
                                question_num=i + 1,
                                total_questions=len(question_list),
                                file_num=file_index + 1,
                                total_files=len(files),
                            ),
                        )
                        # Give MORE time for progress polling to see the update, especially for Full Document Scan
                        await asyncio.sleep(
                            0.1
                        )  # Increased from 0.01 to ensure progress is visible

                    except Exception as question_processing_error:
                        print(
                            f"Error processing question '{question_text[:50]}...': {question_processing_error}"
                        )
                        import traceback

                        traceback.print_exc()
                        # Add error result instead of failing completely
                        qa_pairs.append(
                            {
                                "question": question_text,
                                "answer": f"Error processing this question: {str(question_processing_error)}",
                                "context": "Error occurred during processing",
                                "source_citations": [],
                            }
                        )
                        continue  # Store file-specific QA pairs and final evaluation
                qa_pairs_text = ""
                for i, qa in enumerate(qa_pairs):
                    qa_pairs_text += (
                        f"Question {i+1}: {qa['question']}\nAnswer: {qa['answer']}\n\n"
                    )

                # Generate final evaluation for this file
                try:
                    # Update progress for final evaluation
                    progress_tracker.update_stage_progress(
                        task_id,
                        "reviewing",
                        file_index * len(question_list) + len(question_list),
                        len(files) * len(question_list),
                        translate_progress_message(
                            "generating_final_evaluation",
                            user_language,
                            file_num=file_index + 1,
                            total_files=len(files),
                            file_preview=file_preview,
                        ),
                    )
                    await asyncio.sleep(0.01)  # Allow progress API to respond

                    print(f"Generating final evaluation for {file.filename}...")
                    print(
                        f"DEBUG: final evaluation language_instruction = '{language_instruction}'"
                    )
                    final_evaluation = invoke_llm(
                        llm,
                        final_prompt_template,
                        {
                            "qa_pairs": qa_pairs_text,
                            "language_instruction": language_instruction,
                        },
                    )
                    print(f"Got final evaluation: {final_evaluation[:100]}...")
                    # Clean surrogates from final_evaluation
                    final_evaluation = re.sub(r"[\ud800-\udfff]", "", final_evaluation)
                except Exception as final_eval_error:
                    print(
                        f"Error generating final evaluation for {file.filename}: {final_eval_error}"
                    )
                    final_evaluation = (
                        f"Error generating final evaluation: {str(final_eval_error)}"
                    )

                # Record interaction for this file
                try:
                    interaction_id = record_llm_interaction(
                        session=session,
                        user_id=current_user.id,
                        functionality="veradoc",
                        input_data={
                            "questions": request_data.questions,
                            "document_name": file.filename,
                            "kb_id": request_data.knowledge_base_id,
                            "search_mode": request_data.search_mode,
                            "multi_file_batch": len(files) > 1,
                            "file_index": file_index + 1,
                            "total_files": len(files),
                            "optimization_applied": True,
                        },
                        output_data={
                            "final_evaluation": final_evaluation,
                            "qa_count": len(qa_pairs),
                        },
                        metadata={"qa_pairs": qa_pairs},
                    )
                except Exception as interaction_error:
                    print(
                        f"Error recording interaction for {file.filename}: {interaction_error}"
                    )
                    interaction_id = None

                # Store results for this file
                file_result = {
                    "filename": file.filename,
                    "final_evaluation": final_evaluation,
                    "qa_pairs": qa_pairs,
                    "interaction_id": str(interaction_id) if interaction_id else None,
                }
                all_files_results.append(file_result)

            # 10. Return results based on number of files processed
            if len(all_files_results) == 0:
                raise HTTPException(
                    status_code=400, detail="No files were successfully processed"
                )

            # Complete reviewing stage and start finalizing
            print(f"📊 COMPLETING REVIEWING STAGE for task {task_id}")
            progress_tracker.complete_stage(task_id, "reviewing", "Review complete")
            await asyncio.sleep(
                0.1
            )  # Allow progress polling to see reviewing completion

            print(f"📊 STARTING FINALIZING STAGE for task {task_id}")
            progress_tracker.update_stage_progress(
                task_id, "finalizing", 0, 1, "Finalizing results..."
            )

            await asyncio.sleep(0.1)  # Allow progress API to respond

            # Return optimized multi-file response
            if len(all_files_results) > 1:
                # Multiple files: return optimized format with all results
                result = VeraDocResponse(
                    results={
                        "task_id": task_id,  # Include task_id for progress tracking
                        "multi_file_results": all_files_results,
                        "total_files_processed": len(all_files_results),
                        "optimization_applied": True,
                        "context_prefetch_count": len(question_contexts),
                        "search_mode": request_data.search_mode,
                        # For backward compatibility, include first file's data at root level
                        "filename": all_files_results[0]["filename"],
                        "final_evaluation": all_files_results[0]["final_evaluation"],
                        "qa_pairs": all_files_results[0]["qa_pairs"],
                        "interaction_id": all_files_results[0]["interaction_id"],
                    }
                )
            else:
                # Single file: return traditional format for compatibility
                result = VeraDocResponse(
                    results={
                        "task_id": task_id,  # Include task_id for progress tracking
                        **all_files_results[0],
                        "optimization_applied": True,
                        "context_prefetch_count": len(question_contexts),
                        "search_mode": request_data.search_mode,
                    }
                )

            # Complete finalizing stage
            print(f"📊 COMPLETING FINALIZING STAGE for task {task_id}")
            progress_tracker.complete_stage(
                task_id,
                "finalizing",
                translate_progress_message(
                    "review_completed_successfully", user_language
                ),
            )
            await asyncio.sleep(
                0.1
            )  # Allow progress polling to see finalizing completion

            print(f"📊 COMPLETING TASK for task {task_id}")
            progress_tracker.complete_task(
                task_id,
                translate_progress_message(
                    "review_completed_successfully", user_language
                ),
            )
            await asyncio.sleep(0.1)  # Allow progress polling to see task completion

            # Store results in task metadata for later retrieval
            print(f"📊 STORING RESULTS METADATA for task {task_id}")
            progress_tracker.update_task_metadata(task_id, result.results)

            return result

    except Exception as e:
        import traceback

        print("Error processing RAG checklist:")
        print(str(e))

        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error processing RAG checklist: {str(e)}"
        )


# Functions related to Checklists
@router.post("/checklists", response_model=VeraDocChecklist)
def create_checklist(
    checklist: VeraDocChecklist,
    session: SessionDep,
    current_user: CurrentUser,
):
    """
    Save a new checklist to the database.
    """
    existing_checklist = session.exec(
        select(VeraDocChecklist).where(VeraDocChecklist.name == checklist.name)
    ).first()
    if existing_checklist:
        raise HTTPException(
            status_code=400, detail="A checklist with this name already exists."
        )

    checklist.owner_id = current_user.id
    # Temporarily truncate description to fit current database constraint
    # TODO: Remove this after running migration to increase description length
    if checklist.description and len(checklist.description) > 255:
        checklist.description = checklist.description[:252] + "..."
    session.add(checklist)
    session.commit()
    session.refresh(checklist)
    return checklist


@router.get("/checklists", response_model=List[VeraDocChecklist])
def get_checklists(session: SessionDep, current_user: CurrentUser):
    """
    Retrieve all checklists from the database for this user.
    """
    return session.exec(
        select(VeraDocChecklist).where(VeraDocChecklist.owner_id == current_user.id)
    ).all()


@router.get("/checklists/{checklist_id}", response_model=VeraDocChecklist)
def get_checklist(checklist_id: uuid.UUID, session: SessionDep):
    """
    Retrieve a specific checklist by ID.
    """
    checklist = session.get(VeraDocChecklist, checklist_id)
    if not checklist:
        raise HTTPException(status_code=404, detail="Checklist not found.")
    return checklist


@router.put("/checklists/{checklist_id}", response_model=VeraDocChecklist)
def update_checklist(
    checklist_id: uuid.UUID,
    updated_checklist: VeraDocChecklist,
    session: SessionDep,
    current_user: CurrentUser,
):
    """
    Update an existing checklist.
    """
    checklist = session.get(VeraDocChecklist, checklist_id)
    if not checklist:
        raise HTTPException(status_code=404, detail="Checklist not found.")

    # Ensure the current user is the owner of the checklist
    if checklist.owner_id != current_user.id:
        raise HTTPException(
            status_code=403, detail="Not authorized to update this checklist."
        )

    checklist.name = updated_checklist.name
    # Temporarily truncate description to fit current database constraint
    # TODO: Remove this after running migration to increase description length
    description = updated_checklist.description
    if description and len(description) > 255:
        description = description[:252] + "..."
    checklist.description = description

    # Debug logging to see what questions data is being received
    print(
        f"Updating checklist {checklist_id} with questions: {updated_checklist.questions}"
    )

    checklist.questions = updated_checklist.questions
    checklist.date_modified = datetime.utcnow()

    session.add(checklist)
    session.commit()
    session.refresh(checklist)
    return checklist


@router.delete("/checklists/{checklist_id}", response_model=Message)
def delete_checklist(
    checklist_id: uuid.UUID, session: SessionDep, current_user: CurrentUser
):
    """
    Delete a checklist by ID.
    """
    checklist = session.get(VeraDocChecklist, checklist_id)
    if not checklist:
        raise HTTPException(status_code=404, detail="Checklist not found.")

    # Ensure the current user is the owner of the checklist
    if checklist.owner_id != current_user.id:
        raise HTTPException(
            status_code=403, detail="Not authorized to delete this checklist."
        )

    session.delete(checklist)
    session.commit()
    return Message(message="Checklist deleted successfully.")


@router.delete("/evaluations/{evaluation_id}", response_model=Message)
def delete_evaluation(
    evaluation_id: uuid.UUID, session: SessionDep, current_user: CurrentUser
):
    """
    Delete an evaluation/report by ID.
    """
    evaluation = session.get(LlmInteraction, evaluation_id)
    if not evaluation:
        raise HTTPException(status_code=404, detail="Evaluation not found.")

    # Ensure the current user is the owner of the evaluation
    if evaluation.user_id != current_user.id:
        raise HTTPException(
            status_code=403, detail="Not authorized to delete this evaluation."
        )

    # Only allow deletion of veradoc evaluations
    if evaluation.functionality != "veradoc":
        raise HTTPException(status_code=400, detail="Invalid evaluation type.")

    session.delete(evaluation)
    session.commit()
    return Message(message="Evaluation deleted successfully.")


@router.get("/history", response_model=List[Dict[str, Any]])
async def get_veradoc_history(
    session: SessionDep,
    current_user: CurrentUser,
    skip: int = Query(0, ge=0, le=10000),
    limit: int = Query(20, ge=1, le=100),
    show_all: bool = Query(False),
):
    """Retrieve past VeraDoc evaluation history for the current user or all users."""
    print("Retrieving VeraDoc history. Show all:", show_all)

    try:
        # Start with base query
        query = select(LlmInteraction).where(LlmInteraction.functionality == "veradoc")

        # Only filter by user if not showing all users
        if not show_all:
            query = query.where(LlmInteraction.user_id == current_user.id)

        # Add ordering and pagination
        reports = session.exec(
            query.order_by(LlmInteraction.date_created.desc()).offset(skip).limit(limit)
        ).all()

        print(f"Found {len(reports)} VeraDoc evaluations for user {current_user.id}")

        result = []
        for report in reports:
            # Initialize variables outside the try block
            input_data = {}
            output_data = {}
            extra_data = {}
            kb_name = "Unknown Knowledge Base"

            try:
                # Parse the input_data and output_data from string to dict
                input_data = json.loads(report.input_data) if report.input_data else {}
                output_data = (
                    json.loads(report.output_data) if report.output_data else {}
                )
                extra_data = report.extra_data or {}

                # Get KB name from input_data
                if input_data.get("kb_id"):
                    kb = session.get(KnowledgeBase, input_data.get("kb_id"))
                    kb_name = kb.title if kb else "Unknown Knowledge Base"

                # Create a user-friendly title
                document_name = input_data.get("document_name", "Unnamed Document")
                title = f"Evaluation of {document_name}"

                # Create result item
                result_item = {
                    "id": str(report.id),
                    "date_created": report.date_created,
                    "title": title,
                    "document_name": document_name,
                    "kb_name": kb_name,
                    "kb_id": input_data.get("kb_id", ""),
                    "questions": input_data.get("questions", ""),
                    "qa_count": output_data.get("qa_count", 0),
                    "final_evaluation": output_data.get("final_evaluation", ""),
                    "has_feedback": report.feedback is not None,
                }

                # Add feedback information if exists
                if report.feedback:
                    result_item["feedback"] = {
                        "feedback": report.feedback,
                        "feedbackText": report.feedback_text,
                    }

                # Add user info for all-users view
                if show_all:
                    user = session.get(User, report.user_id)
                    user_name = (
                        f"{user.full_name or 'User'} ({user.email})"
                        if user
                        else "Unknown User"
                    )
                    result_item["user_name"] = user_name

                result.append(result_item)
            except Exception as e:
                # If parsing fails, add a minimal entry
                print(f"Error processing report {report.id}: {e}")
                result.append(
                    {
                        "id": str(report.id),
                        "date_created": report.date_created,
                        "title": f"Evaluation from {report.date_created.strftime('%Y-%m-%d')}",
                    }
                )

        return result
    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error retrieving VeraDoc history: {str(e)}"
        )


@router.get("/history/{report_id}", response_model=Dict[str, Any])
async def get_veradoc_detail(
    report_id: uuid.UUID,
    session: SessionDep,
    current_user: CurrentUser,
    include_qa_pairs: bool = Query(
        default=True,
        description="If False, excludes the heavy qa_pairs data to improve performance",
    ),
):
    """
    Retrieve a specific VeraDoc evaluation by ID.

    Args:
        report_id: The UUID of the report to retrieve
        include_qa_pairs: If False, excludes the heavy qa_pairs data (default: True)

    Returns summary (without qa_pairs) when include_qa_pairs=False,
    or full detail (with qa_pairs) when include_qa_pairs=True.
    """
    try:
        report = session.get(LlmInteraction, report_id)
        if not report:
            raise HTTPException(status_code=404, detail="Report not found")

        # No longer need to check this as we may now view others' outputs
        # if report.user_id != current_user.id:
        #    raise HTTPException(
        #        status_code=403, detail="You don't have access to this report"
        #    )

        if report.functionality != "veradoc":
            raise HTTPException(
                status_code=400, detail="This is not a VeraDoc evaluation"
            )

        try:
            input_data = json.loads(report.input_data) if report.input_data else {}
            output_data = json.loads(report.output_data) if report.output_data else {}
            extra_data = report.extra_data or {}

            # For backward compatibility, try to reconstruct full results
            document_name = input_data.get("document_name", "Unknown Document")
            kb_name = "Unknown Knowledge Base"

            # Try to get KB name
            kb_id = input_data.get("kb_id")
            if kb_id:
                kb = session.get(KnowledgeBase, kb_id)
                kb_name = kb.title if kb else "Unknown Knowledge Base"

            qa_pairs = extra_data.get("qa_pairs", [])

            # Create a response that matches the structure expected by the frontend
            result = {
                "id": str(report.id),
                "date_created": report.date_created,
                "document_name": document_name,
                "kb_name": kb_name,
                "kb_id": kb_id,
                "questions": input_data.get("questions", ""),
                "results": {
                    "final_evaluation": output_data.get("final_evaluation", ""),
                    "interaction_id": str(report.id),
                },
                # Add feedback information
                "feedback": {
                    "feedback": report.feedback,
                    "feedbackText": report.feedback_text,
                    "feedbackDate": (
                        report.feedback_date.isoformat()
                        if report.feedback_date
                        else None
                    ),
                },
            }

            # Conditionally include qa_pairs based on parameter
            if include_qa_pairs:
                result["results"]["qa_pairs"] = qa_pairs
            else:
                # For summary view, include question headers (without answers/context/citations)
                result["results"]["qa_pairs_summary"] = [
                    {
                        "index": i,
                        "question": qa.get("question", ""),
                    }
                    for i, qa in enumerate(qa_pairs)
                ]
                result["results"]["qa_pairs_count"] = len(qa_pairs)

            return result

        except Exception as e:
            # Fallback if parsing fails
            fallback = {
                "id": str(report.id),
                "date_created": report.date_created,
                "results": {
                    "final_evaluation": f"Unable to reconstruct evaluation from {report.date_created}.\n\n"
                    f"This might be due to an older format or incomplete data.",
                    "interaction_id": str(report.id),
                },
            }

            if include_qa_pairs:
                fallback["results"]["qa_pairs"] = []
            else:
                fallback["results"]["qa_pairs_count"] = 0

            return fallback

    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error retrieving evaluation details: {str(e)}"
        )


@router.get("/history/{report_id}/qa-pair/{qa_index}", response_model=Dict[str, Any])
async def get_veradoc_qa_pair(
    report_id: uuid.UUID,
    qa_index: int,
    session: SessionDep,
    current_user: CurrentUser,
):
    """
    Retrieve a specific QA pair from a VeraDoc evaluation by index.

    This enables lazy loading of individual QA pairs for better performance.

    Args:
        report_id: The UUID of the report
        qa_index: The zero-based index of the QA pair to retrieve

    Returns the QA pair with question, answer, context, and source_citations.
    """
    try:
        report = session.get(LlmInteraction, report_id)
        if not report:
            raise HTTPException(status_code=404, detail="Report not found")

        if report.functionality != "veradoc":
            raise HTTPException(
                status_code=400, detail="This is not a VeraDoc evaluation"
            )

        try:
            extra_data = report.extra_data or {}
            qa_pairs = extra_data.get("qa_pairs", [])

            if qa_index < 0 or qa_index >= len(qa_pairs):
                raise HTTPException(
                    status_code=404,
                    detail=f"QA pair index {qa_index} not found. Valid range: 0-{len(qa_pairs)-1}",
                )

            qa_pair = qa_pairs[qa_index]

            return {
                "index": qa_index,
                "question": qa_pair.get("question", ""),
                "answer": qa_pair.get("answer", ""),
                "context": qa_pair.get("context", ""),
                "source_citations": qa_pair.get("source_citations", []),
            }

        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error retrieving QA pair: {str(e)}"
            )

    except HTTPException:
        raise
    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error retrieving QA pair: {str(e)}"
        )


@router.post("/optimize-checklist", response_model=OptimizedChecklistResponse)
async def optimize_checklist(
    session: SessionDep,
    current_user: CurrentUser,
    knowledge_base_id: str = Form(...),
    questions: str = Form(...),
    target_answers: str = Form("yes"),
    custom_instructions: Optional[str] = Form(None),
    search_mode: str = Form("vector"),
    files: List[UploadFile] = File(...),
    request: FastAPIRequest = None,
):
    """
    Optimize checklist questions by testing them against a document that should meet all requirements.
    Suggests revisions for questions that resulted in negative answers.
    """
    print("optimize_checklist function invoked!")
    # Disconnect monitoring disabled due to false positives

    # Create request_data object for backward compatibility with rest of the code
    class RequestData:
        pass

    request_data = RequestData()
    request_data.questions = questions
    request_data.knowledge_base_id = knowledge_base_id
    request_data.custom_instructions = custom_instructions
    request_data.search_mode = search_mode
    request_data.target_answers = target_answers

    try:
        print("Starting checklist optimization...")

        # Get user's preferred language
        user_language = getattr(current_user, "preferred_language", "en") or "en"

        # 1. Retrieve knowledge base
        kb = session.get(KnowledgeBase, request_data.knowledge_base_id)
        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # 2. Set up the same infrastructure as process_rag_checklist
        with tempfile.TemporaryDirectory() as temp_dir:
            # Extract ChromaDB
            if kb.storage_type == "file" and kb.file_path:
                # File-based storage: extract from file path
                if os.path.exists(kb.file_path):
                    with zipfile.ZipFile(kb.file_path, "r") as zip_ref:
                        zip_ref.extractall(temp_dir)
                else:
                    raise HTTPException(
                        status_code=400, detail="Knowledge base file not found on disk"
                    )
            elif kb.data:
                # Database storage: extract from data field
                with zipfile.ZipFile(BytesIO(kb.data), "r") as zip_ref:
                    zip_ref.extractall(temp_dir)
            else:
                raise HTTPException(
                    status_code=400, detail="Knowledge base has no vector database data"
                )

            # Load embeddings and vector database
            if kb.embedding_model_id:
                embedding_model = session.get(EmbeddingModel, kb.embedding_model_id)
                if embedding_model:
                    model_id = embedding_model.model_id
                    provider = embedding_model.provider
                else:
                    embedding_info = get_embedding_model(session, current_user)
                    model_id = embedding_info["model_id"]
                    provider = embedding_info["provider"]
            else:
                embedding_info = get_embedding_model(session, current_user)
                model_id = embedding_info["model_id"]
                provider = embedding_info["provider"]

            embeddings = load_embeddings_model(provider=provider, model_id=model_id)
            chroma_db = Chroma(
                persist_directory=temp_dir, embedding_function=embeddings
            )

            # Create retriever based on search mode
            if request_data.search_mode == "full_scan":
                print("Using full document scan mode for optimization")

                # Create a simple retriever that returns all documents
                class FullScanRetriever:
                    def __init__(self, chroma_db):
                        self.chroma_db = chroma_db

                    def get_relevant_documents(self, query):
                        try:
                            # Get all documents for full optimization scan
                            all_data = self.chroma_db.get()

                            if all_data and "documents" in all_data:
                                print(
                                    f"📚 Full Optimization Scan: Processing ALL {len(all_data['documents'])} documents"
                                )

                            documents = []

                            if (
                                all_data
                                and "documents" in all_data
                                and all_data["documents"]
                            ):
                                for i, doc_content in enumerate(all_data["documents"]):
                                    metadata = (
                                        all_data["metadatas"][i]
                                        if "metadatas" in all_data
                                        and i < len(all_data["metadatas"])
                                        else {}
                                    )
                                    documents.append(
                                        LangchainDocument(
                                            page_content=doc_content, metadata=metadata
                                        )
                                    )

                            return documents
                        except Exception as e:
                            print(f"Error in FullScanRetriever: {e}")
                            return []

                retriever = FullScanRetriever(chroma_db)
            else:
                print("Using enhanced vector search mode for optimization")
                retriever = SmartRetrieverFactory.create_academic_paper_retriever(
                    chroma_db=chroma_db,
                    search_kwargs={"k": settings.RAG_NUM_CHUNKS},
                )

            # Initialize LLM
            llm = get_default_llm(session, current_user)

            # 3. Process the test document
            file = files[0]
            content = await file.read()
            document_text = await extract_text_from_file_async(
                content, file.filename, current_user
            )
            print(
                f"Processing test document: {file.filename} ({len(document_text)} characters)"
            )

            # For large files or DOCX files, cancel disconnect monitoring to prevent false positives
            is_docx = file.filename.lower().endswith((".docx", ".doc"))
            large_document_threshold = (
                100000 if is_docx else 200000
            )  # Lower threshold for DOCX

            if len(document_text) > large_document_threshold or (
                is_docx and len(content) > 50000
            ):
                print(
                    f"Large document detected ({file.filename}), disabling disconnect monitoring to prevent false positives"
                )
                # Disconnect monitoring disabled due to false positives

            # 4. Run the review process with current questions
            question_list = request_data.questions.strip().split("\n")
            qa_results = []

            context_prompt_template = settings.VERADOC_CONTEXT_PROMPT_TEMPLATE
            qa_prompt_template = settings.VERADOC_QA_PROMPT_TEMPLATE
            optimize_prompt_template = settings.VERADOC_OPTIMIZE_PROMPT_TEMPLATE

            print(f"Evaluating {len(question_list)} questions...")

            for question in question_list:

                question = question.strip()
                if not question:
                    continue

                # Get relevant context
                docs = retriever.get_relevant_documents(question)
                context = "\n\n".join([doc.page_content for doc in docs])

                # Generate policy context
                question_context = invoke_llm(
                    llm,
                    context_prompt_template,
                    {"context": context, "question": question},
                )

                # Translate the question context if needed
                # question_context = await translate_text_if_needed(
                #     question_context, session, current_user, llm
                # )

                # Prepare custom instructions section if provided
                custom_instructions_section = ""
                if (
                    request_data.custom_instructions
                    and request_data.custom_instructions.strip()
                ):
                    custom_instructions_section = f"\nADDITIONAL INSTRUCTIONS:\n{request_data.custom_instructions.strip()}\n"

                # Prepare language instruction
                language_instruction = ""
                if user_language:
                    language_name = settings.SUPPORTED_LANGUAGES.get(
                        user_language, user_language
                    )
                    language_instruction = f"Respond in this language: {language_name}."

                # Generate answer
                print(f"DEBUG: language_instruction = '{language_instruction}'")
                answer = invoke_llm(
                    llm,
                    qa_prompt_template,
                    {
                        "document_text": document_text[:10000],
                        "question": question,
                        "question_context": question_context,
                        "custom_instructions_section": custom_instructions_section,
                        "language_instruction": language_instruction,
                    },
                )

                qa_results.append(
                    {
                        "question": question,
                        "answer": answer,
                        "context": question_context,
                    }
                )

                print(
                    f"Question: {question[:50]}... -> {'NEEDS OPTIMIZATION' if needs_optimization(answer) else 'OK'}"
                )

            # 5. Generate optimization suggestions
            suggestions = []
            optimization_count = 0

            for qa in qa_results:

                if needs_optimization(qa["answer"]):
                    optimization_count += 1
                    print(
                        f"Generating suggestion for question: {qa['question'][:50]}..."
                    )

                    # Generate optimization suggestion
                    suggestion_response = invoke_llm(
                        llm,
                        optimize_prompt_template,
                        {
                            "original_question": qa["question"],
                            "generated_answer": qa["answer"],
                            "document_context": qa["context"],
                        },
                    )

                    suggestion = parse_optimization_response(suggestion_response, qa)
                    # Add policy context to the suggestion
                    suggestion.policy_context = qa["context"]

                    # Translate current answer and analysis to user's language
                    # suggestion.current_answer = await translate_text_if_needed(
                    #     suggestion.current_answer, session, current_user, llm
                    # )
                    # suggestion.reason = await translate_text_if_needed(
                    #     suggestion.reason, session, current_user, llm
                    # )

                    suggestions.append(suggestion)
                else:
                    # Question is already working well
                    reason_text = "Question already generates positive responses"
                    # translated_reason = await translate_text_if_needed(
                    #     reason_text, session, current_user, llm
                    # )
                    translated_reason = reason_text
                    suggestion = ChecklistSuggestion(
                        original_question=qa["question"],
                        suggested_question=qa["question"],
                        reason=translated_reason,
                        current_answer=qa["answer"],
                        needs_revision=False,
                        policy_context=qa["context"],
                    )

                    # Translate current answer and analysis to user's language
                    # suggestion.current_answer = await translate_text_if_needed(
                    #     suggestion.current_answer, session, current_user, llm
                    # )
                    # suggestion.reason = await translate_text_if_needed(
                    #     suggestion.reason, session, current_user, llm
                    # )

                    suggestions.append(suggestion)

            # 6. Compile results
            original_questions = [qa["question"] for qa in qa_results]
            optimized_questions = [s.suggested_question for s in suggestions]

            analysis_summary = f"""
Checklist Optimization Analysis:
- Total questions evaluated: {len(original_questions)}
- Questions needing optimization: {optimization_count}
- Questions working well: {len(original_questions) - optimization_count}
- Test document: {file.filename}

The optimization process identified questions that resulted in negative responses when evaluating a document that should meet all requirements. Suggested revisions aim to make requirements more achievable while maintaining their intent.
            """.strip()

            print(
                f"Optimization complete: {optimization_count}/{len(original_questions)} questions optimized"
            )

            return OptimizedChecklistResponse(
                original_questions=original_questions,
                suggestions=suggestions,
                optimized_questions=optimized_questions,
                analysis_summary=analysis_summary,
            )

    except Exception as e:
        print("Error in checklist optimization:")
        print(str(e))
        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error optimizing checklist: {str(e)}"
        )


@router.post("/generate/docx", response_class=StreamingResponse)
async def generate_docx(
    session: SessionDep, current_user: CurrentUser, request: DocxRequest
):
    """
    Generate a DOCX file from the evaluation content.
    """
    print("Now generating DOCX of evaluation...")
    try:
        # Get the markdown content from the request
        if not request.content:
            raise HTTPException(status_code=400, detail="Report content is required")

        # Convert markdown to HTML for parsing
        html_content = markdown.markdown(request.content, extensions=["tables"])
        soup = BeautifulSoup(html_content, "html.parser")

        print("Markdown content converted to HTML successfully.")
        # Create a new Document
        doc = Document()

        print("Adding title and date to the document...")
        # Determine language - prioritize user's preferred language over request
        # because frontend i18n might not be properly initialized
        user_preferred = getattr(current_user, "preferred_language", None)
        language = user_preferred or request.language or "en"

        print(f"[DOCX Generation] Language determined: {language}")
        print(f"[DOCX Generation] request.language: {request.language}")
        print(f"[DOCX Generation] current_user.preferred_language: {user_preferred}")

        # Add a title
        title_text = (
            request.title
            if request.title
            else translate_docx_header("documentEvaluation", language)
        )
        print(f"[DOCX Generation] Title text: {title_text}")
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add translated subtitle with metadata
        date_str = datetime.now().strftime("%B %d, %Y at %H:%M")
        subtitle = translate_docx_header(
            "generatedOn",
            language,
            date=date_str,
            name=current_user.full_name or current_user.email,
            email=current_user.email,
        )
        subtitle_paragraph = doc.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.italic = True

        # Add a separator
        doc.add_paragraph("─" * 50)

        print("Adding headers, paragraphs, lists, and tables...")
        # Process all headers and paragraphs in the HTML
        for element in soup.find_all(
            ["h1", "h2", "h3", "h4", "p", "ul", "ol", "li", "table"]
        ):
            if element.name == "h1":
                doc.add_heading(element.text, level=1)
            elif element.name == "h2":
                doc.add_heading(element.text, level=2)
            elif element.name == "h3":
                doc.add_heading(element.text, level=3)
            elif element.name == "p":
                doc.add_paragraph(element.text)
            elif element.name == "ul":
                for li in element.find_all("li"):
                    paragraph = doc.add_paragraph(li.text)
                    paragraph.style = "List Bullet"
            elif element.name == "ol":
                for li in element.find_all("li"):
                    paragraph = doc.add_paragraph(li.text)
                    paragraph.style = "List Number"
            elif element.name == "table":
                table_rows = element.find_all("tr")
                if table_rows:
                    # Count the number of columns in the first row
                    first_row = table_rows[0]
                    columns = len(first_row.find_all(["th", "td"]))

                    # Create the table
                    table = doc.add_table(rows=0, cols=columns)
                    table.style = "Table Grid"

                    # Process header row
                    header_cells = first_row.find_all(["th", "td"])
                    if header_cells:
                        header_row = table.add_row().cells
                        for i, cell in enumerate(header_cells):
                            if i < len(header_row):
                                header_row[i].text = cell.text
                                run = header_row[i].paragraphs[0].runs[0]
                                run.bold = True

                    # Process data rows
                    for row in table_rows[1:]:
                        cells = row.find_all("td")
                        if cells:
                            row_cells = table.add_row().cells
                            for i, cell in enumerate(cells):
                                if i < len(row_cells):
                                    row_cells[i].text = cell.text

        # Save the document to a BytesIO object
        print("Saving the document to a BytesIO object...")
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        # --- CORRUPTION CHECKS ---
        # 1. Check file size
        size = docx_bytes.getbuffer().nbytes
        print(f"DOCX file size: {size} bytes")
        if size < 1000:
            print("Warning: DOCX file is very small and may be empty or corrupted.")

        # 2. Try to reload the DOCX to ensure it's readable
        try:
            docx_bytes.seek(0)
            _ = Document(docx_bytes)
            print("DOCX file passed integrity check (can be opened by python-docx).")
        except Exception as e:
            print(
                f"Integrity check failed: generated DOCX cannot be opened. Error: {e}"
            )
            raise HTTPException(
                status_code=500, detail="Generated DOCX file is corrupted."
            )

        docx_bytes.seek(0)

        print(
            "Document saved successfully. Preparing to return as a downloadable file."
        )

        # Create a filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"evaluation_{timestamp}.docx"

        # Return the document as a downloadable file
        return StreamingResponse(
            docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")


@router.post("/generate/csv", response_class=StreamingResponse)
async def generate_csv(
    session: SessionDep, current_user: CurrentUser, request: DocxRequest
):
    """
    Generate a CSV file from VeraDoc review results with columns for:
    Checklist Question, Policy Context, Citations, Answer, and Final Evaluation.
    """
    print("Now generating CSV of VeraDoc review...")
    try:
        # Get the content from the request - this should be the full review JSON
        if not request.content:
            raise HTTPException(status_code=400, detail="Review content is required")

        # Create CSV content
        csv_buffer = StringIO()
        csv_writer = csv.writer(csv_buffer)

        # Write header with VeraDoc-specific columns
        csv_writer.writerow(
            [
                "Checklist Question",
                "Policy Context",
                "Citations",
                "Answer",
                "Final Evaluation",
            ]
        )

        try:
            # Parse the content as JSON (review data)
            data = json.loads(request.content)

            # Extract QA pairs and final evaluation
            qa_pairs = data.get("qa_pairs", [])
            final_evaluation = data.get(
                "final_evaluation", "No final evaluation provided"
            )

            # Process each QA pair
            for qa in qa_pairs:
                question = qa.get("question", "")
                answer = qa.get("answer", "")
                context = qa.get("context", "")

                # Extract citations
                citations = []
                if "source_citations" in qa:
                    for citation in qa["source_citations"]:
                        source_name = "Unknown"
                        if citation.get("metadata", {}).get("source"):
                            source_path = citation["metadata"]["source"]
                            # Extract filename from path
                            source_name = source_path.split("/")[-1].split("\\")[-1]
                            # Remove prefixes using more robust logic
                            if "_" in source_name:
                                parts = source_name.split("_")
                                if len(parts) > 1:
                                    first_part = parts[0]
                                    # If first part is short and alphanumeric (likely a prefix), remove it
                                    if len(first_part) <= 10 and first_part.isalnum():
                                        source_name = "_".join(parts[1:])
                                    # Otherwise keep the original filename
                                    # This prevents removing legitimate parts of the filename

                        citation_text = (
                            re.sub(r"[\ud800-\udfff]", "", citation.get("content", ""))
                            .replace("\n", " ")
                            .replace("\r", " ")
                        )
                        citations.append(f"{source_name}: {citation_text}")

                citations_text = " | ".join(citations) if citations else "No citations"

                # Clean up text fields
                question_clean = (
                    re.sub(r"[\ud800-\udfff]", "", question)
                    .replace("\n", " ")
                    .replace("\r", " ")
                )
                answer_clean = (
                    re.sub(r"[\ud800-\udfff]", "", answer)
                    .replace("\n", " ")
                    .replace("\r", " ")
                )
                context_clean = (
                    re.sub(r"[\ud800-\udfff]", "", context)
                    .replace("\n", " ")
                    .replace("\r", " ")
                )

                # For the final evaluation, we'll include it for each row
                # (since it's a summary of all QA pairs)
                final_eval_clean = (
                    re.sub(r"[\ud800-\udfff]", "", final_evaluation)
                    .replace("\n", " ")
                    .replace("\r", " ")
                )

                # Write row
                csv_writer.writerow(
                    [
                        question_clean,
                        context_clean,
                        citations_text,
                        answer_clean,
                        final_eval_clean,
                    ]
                )

        except json.JSONDecodeError:
            raise HTTPException(
                status_code=400,
                detail="Invalid content format. Expected JSON with VeraDoc review data.",
            )

        # Get CSV content
        csv_content = csv_buffer.getvalue()
        csv_buffer.close()

        # Create BytesIO object for the response
        csv_bytes = BytesIO(csv_content.encode("utf-8"))
        csv_bytes.seek(0)

        print("VeraDoc CSV file generated successfully.")

        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"veradoc_review_{timestamp}.csv"

        # Return the CSV as a downloadable file
        return StreamingResponse(
            csv_bytes,
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating CSV: {str(e)}")


@router.post("/optimization/csv", response_class=StreamingResponse)
async def generate_optimization_csv(
    session: SessionDep, current_user: CurrentUser, request: DocxRequest
):
    """
    Generate a CSV file from checklist optimization results with columns for:
    Question Number, Original Question, Suggested Question, Needs Revision, Reason, Current Answer.
    """
    print("Now generating CSV of checklist optimization results...")
    try:
        # Get the content from the request - this should be the optimization results JSON
        if not request.content:
            raise HTTPException(
                status_code=400, detail="Optimization content is required"
            )

        # Create CSV content
        csv_buffer = StringIO()
        csv_writer = csv.writer(csv_buffer)

        # Write header with optimization-specific columns
        csv_writer.writerow(
            [
                "Question Number",
                "Original Question",
                "Suggested Question",
                "Needs Revision",
                "Policy Context",
                "Current Answer",
                "Reason",
                "Analysis Summary",
            ]
        )

        try:
            # Parse the content as JSON (optimization results data)
            data = json.loads(request.content)

            # Extract suggestions and analysis summary
            suggestions = data.get("suggestions", [])
            analysis_summary = data.get(
                "analysis_summary", "No analysis summary provided"
            )

            # Process each suggestion
            for index, suggestion in enumerate(suggestions, 1):
                original_question = suggestion.get("original_question", "")
                suggested_question = suggestion.get("suggested_question", "")
                needs_revision = suggestion.get("needs_revision", False)
                reason = suggestion.get("reason", "")
                current_answer = suggestion.get("current_answer", "")
                policy_context = suggestion.get("policy_context", "")

                # Clean up text fields (remove newlines and carriage returns for CSV)
                original_question_clean = original_question.replace("\n", " ").replace(
                    "\r", " "
                )
                suggested_question_clean = suggested_question.replace(
                    "\n", " "
                ).replace("\r", " ")
                reason_clean = reason.replace("\n", " ").replace("\r", " ")
                current_answer_clean = current_answer.replace("\n", " ").replace(
                    "\r", " "
                )
                policy_context_clean = policy_context.replace("\n", " ").replace(
                    "\r", " "
                )
                analysis_summary_clean = analysis_summary.replace("\n", " ").replace(
                    "\r", " "
                )

                # Write row
                csv_writer.writerow(
                    [
                        index,
                        original_question_clean,
                        suggested_question_clean,
                        "Yes" if needs_revision else "No",
                        policy_context_clean,
                        current_answer_clean,
                        reason_clean,
                        analysis_summary_clean,
                    ]
                )

        except json.JSONDecodeError:
            raise HTTPException(
                status_code=400,
                detail="Invalid content format. Expected JSON with checklist optimization data.",
            )

        # Get CSV content
        csv_content = csv_buffer.getvalue()
        csv_buffer.close()

        # Create BytesIO object for the response
        csv_bytes = BytesIO(csv_content.encode("utf-8"))
        csv_bytes.seek(0)

        print("Checklist optimization CSV file generated successfully.")

        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"checklist_optimization_{timestamp}.csv"

        # Return the CSV as a downloadable file
        return StreamingResponse(
            csv_bytes,
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error generating optimization CSV: {str(e)}"
        )


@router.post("/generate-questions-with-files", response_model=GenerateQuestionsResponse)
async def generate_questions_with_files(
    session: SessionDep,
    current_user: CurrentUser,
    description: str = Form(...),
    checklist_type: str = Form(default="general"),
    num_questions: Optional[int] = Form(default=None),
    files: List[UploadFile] = File(default=[]),
):
    """
    Generate checklist questions based on a description using LLM, with optional reference documents.
    """
    from app.services.text_processing import chunk_text
    from app.core.config import settings

    try:
        # Get the default LLM
        llm = get_default_llm(session, current_user)

        # Process uploaded files to extract reference document content
        reference_document_content = ""
        if files:
            print(f"Processing {len(files)} uploaded files for question generation")

            for file in files:
                if file.size > 0:
                    try:
                        # Read and process the file content with visual enhancement
                        file_content = await file.read()

                        # Use enhanced processing with vision capabilities
                        from app.services.document_utils import (
                            extract_text_with_vision_enhancement,
                        )

                        file_text = await extract_text_with_vision_enhancement(
                            file_content,
                            file.filename or "unknown",
                            llm,
                            purpose="checklist question generation",
                            current_user=current_user,
                        )

                        if file_text.strip():
                            reference_document_content += f"\n\n--- Content from {file.filename} ---\n{file_text.strip()}"
                            print(
                                f"Extracted {len(file_text)} characters from {file.filename} (with vision enhancement)"
                            )

                    except Exception as e:
                        print(f"Error processing file {file.filename}: {str(e)}")
                        # Add the error info to the document content so user knows what happened
                        reference_document_content += f"\n\n--- Error processing {file.filename} ---\nError: {str(e)}\n"
                        continue

        # Check if content exceeds token limits and chunk if necessary
        if reference_document_content:
            print(
                f"Total document content: {len(reference_document_content)} characters"
            )

            # Using conservative chunking similar to TWINCHECK settings
            max_chunk_size = 80000  # Conservative chunk size for 128K context limit

            if len(reference_document_content) > max_chunk_size:
                print(
                    f"Document too large ({len(reference_document_content)} chars), chunking for processing"
                )

                # Chunk the document content
                chunks = chunk_text(
                    reference_document_content, max_tokens=max_chunk_size
                )

                # Process each chunk to generate questions
                all_chunk_questions = []

                for i, chunk in enumerate(chunks):
                    print(f"Processing chunk {i+1}/{len(chunks)}")

                    # Generate questions for this chunk
                    chunk_prompt_variables = {
                        "description": description,
                        "checklist_type": checklist_type,
                        "reference_documents_instruction": "You can find additional requirements in the reference documents provided below.",
                        "reference_documents_content": chunk,
                        "additional_instructions": "\n11. Use the reference documents provided below to identify additional requirements that should be included in the checklist questions",
                    }

                    try:
                        chunk_response = invoke_llm(
                            llm,
                            settings.VERADOC_GENERATE_QUESTIONS_PROMPT_TEMPLATE,
                            chunk_prompt_variables,
                        )

                        # Parse questions from chunk response
                        chunk_questions = []
                        lines = chunk_response.strip().split("\n")
                        in_questions_section = False

                        for line in lines:
                            line = line.strip()
                            if line.startswith("QUESTIONS:"):
                                in_questions_section = True
                                continue
                            elif line.startswith("ANALYSIS:"):
                                in_questions_section = False
                                continue

                            if in_questions_section:
                                if re.match(r"^\d+\.\s+", line):
                                    question = re.sub(r"^\d+\.\s+", "", line)
                                    if question.strip():
                                        chunk_questions.append(question.strip())

                        # If parsing failed, try simpler approach
                        if not chunk_questions:
                            for line in lines:
                                line = line.strip()
                                if re.match(r"^\d+\.\s+", line):
                                    question = re.sub(r"^\d+\.\s+", "", line)
                                    if question.strip():
                                        chunk_questions.append(question.strip())

                        all_chunk_questions.extend(chunk_questions)

                    except Exception as e:
                        print(f"Error processing chunk {i+1}: {e}")
                        continue

                # Deduplicate and refine questions across all chunks
                if all_chunk_questions:
                    # Remove duplicates while preserving order
                    seen = set()
                    unique_questions = []
                    for q in all_chunk_questions:
                        if q.lower() not in seen:
                            seen.add(q.lower())
                            unique_questions.append(q)

                    # If we have too many questions, synthesize and prioritize
                    if len(unique_questions) > (num_questions or 50):
                        synthesis_prompt_variables = {
                            "description": description,
                            "checklist_type": checklist_type,
                            "questions_list": "\n".join(
                                [f"{i+1}. {q}" for i, q in enumerate(unique_questions)]
                            ),
                            "num_questions": num_questions or 20,
                        }

                        synthesis_prompt = f"""From the following list of checklist questions, select and refine the {num_questions or 20} most important and relevant questions for {checklist_type} verification based on: {description}

Questions to review:
{chr(10).join([f"{i+1}. {q}" for i, q in enumerate(unique_questions)])}

Requirements:
1. Select the most critical and actionable questions
2. Ensure no redundancy
3. Maintain clarity and specificity
4. Focus on questions most relevant to the description

Return only the final selected questions, one per line, numbered."""

                        try:
                            refined_response = invoke_llm(llm, synthesis_prompt, {})
                            questions = []
                            for line in refined_response.strip().split("\n"):
                                line = line.strip()
                                if line and (
                                    line[0].isdigit()
                                    or line.startswith("-")
                                    or line.startswith("*")
                                ):
                                    question = re.sub(r"^\d+\.\s+", "", line)
                                    question = re.sub(r"^[-*]\s+", "", question)
                                    if question.strip():
                                        questions.append(question.strip())
                        except Exception as e:
                            print(f"Error in question synthesis: {e}")
                            questions = unique_questions[: num_questions or 20]
                    else:
                        questions = unique_questions[: num_questions or 20]

                    # For analysis, show chunked processing was used
                    analysis = f"Generated {len(questions)} questions from chunked document analysis ({len(chunks)} chunks processed) based on the provided description to ensure comprehensive evaluation coverage."

                    # Record the interaction
                    record_llm_interaction(
                        session=session,
                        user_id=current_user.id,
                        functionality="generate_checklist_questions",
                        input_data={
                            "description": description,
                            "requested_questions": num_questions,
                            "checklist_type": checklist_type,
                            "chunked_processing": True,
                            "chunk_count": len(chunks),
                        },
                        output_data={
                            "questions_count": len(questions),
                            "analysis": analysis,
                        },
                        metadata={},
                    )

                    return GenerateQuestionsResponse(
                        questions=questions, description_analysis=analysis
                    )
                else:
                    # Fallback to description-only generation if chunk processing failed
                    reference_document_content = ""

        # Continue with existing logic for small documents or when no files provided
        if reference_document_content:
            reference_documents_instruction = "You can find additional requirements in the reference documents provided below."
            additional_instructions = "\n11. Use the reference documents provided below to identify additional requirements that should be included in the checklist questions"
        else:
            reference_documents_instruction = ""
            reference_document_content = ""
            additional_instructions = ""

        prompt_variables = {
            "description": description,
            "checklist_type": checklist_type,
            "reference_documents_instruction": reference_documents_instruction,
            "reference_documents_content": reference_document_content,
            "additional_instructions": additional_instructions,
            "language_instruction": f"Respond in this language: {settings.SUPPORTED_LANGUAGES.get(getattr(current_user, 'preferred_language', 'en'), 'English')}.",
        }

        # Generate questions using the LLM
        questions_response = invoke_llm(
            llm,
            settings.VERADOC_GENERATE_QUESTIONS_PROMPT_TEMPLATE,
            prompt_variables,
        )

        # Parse the response to extract questions and analysis
        questions = []
        analysis = ""

        lines = questions_response.strip().split("\n")
        in_questions_section = False
        in_analysis_section = False

        for line in lines:
            line = line.strip()
            if line.startswith("QUESTIONS:"):
                in_questions_section = True
                in_analysis_section = False
                continue
            elif line.startswith("ANALYSIS:"):
                in_questions_section = False
                in_analysis_section = True
                continue

            if in_questions_section:
                # Extract questions (numbered list)
                if re.match(r"^\d+\.\s+", line):
                    question = re.sub(r"^\d+\.\s+", "", line)
                    if question.strip():
                        questions.append(question.strip())
            elif in_analysis_section:
                if line:
                    if analysis:
                        analysis += " " + line
                    else:
                        analysis = line

        # If parsing failed, try simpler approach
        if not questions:
            # Split by lines and look for numbered items
            for line in lines:
                line = line.strip()
                if re.match(r"^\d+\.\s+", line):
                    question = re.sub(r"^\d+\.\s+", "", line)
                    if question.strip():
                        questions.append(question.strip())

        # Ensure we have some questions
        if not questions:
            raise HTTPException(
                status_code=500,
                detail="Failed to generate questions from the description. Please try with a more detailed description.",
            )

        # Apply user-specified limit if provided, otherwise use all generated questions
        if request.num_questions:
            questions = questions[: request.num_questions]

        if not analysis:
            search_method = (
                "vector search"
                if request.search_mode == "vector"
                else "full document scan"
            )
            analysis = f"Generated {len(questions)} checklist questions based on the provided description using {search_method}"
            if request.knowledge_base_id:
                analysis += " with knowledge base reference."

        # Record the interaction
        record_llm_interaction(
            session=session,
            user_id=current_user.id,
            functionality="generate_checklist_questions",
            input_data={
                "description": request.description,
                "checklist_type": request.checklist_type,
                "requested_questions": request.num_questions,
                "knowledge_base_id": request.knowledge_base_id,
                "search_mode": request.search_mode,
            },
            output_data={
                "questions_count": len(questions),
                "analysis": analysis,
            },
            metadata={},
        )

        return GenerateQuestionsResponse(
            questions=questions, description_analysis=analysis
        )

    except Exception as e:
        print(f"Error generating questions: {e}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error generating questions: {str(e)}"
        )
