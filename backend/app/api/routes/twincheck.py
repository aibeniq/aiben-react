import uuid
import difflib
import re
import csv
from typing import List, Dict, Any, Optional
from datetime import datetime
import json
import traceback
import tempfile
import os
import docx
import io
from io import BytesIO, StringIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup
import tiktoken

from fastapi import APIRouter, Depends, File, UploadFile, HTTPException, Form, Request as FastAPIRequest
from fastapi.responses import StreamingResponse
from sqlmodel import select
import asyncio
import threading

from app.api.deps import CurrentUser, SessionDep
from app.models import (
    TwinCheckRequest,
    TwinCheckResponse,
    TwinCheckTopicList,
    TwinCheckDetailResponse,
    LlmInteraction,
    DocxRequest,
    Message,
    GenerateTopicsRequest,
    GenerateTopicsResponse,
)
from app.core.config import settings
from app.services.llms import get_default_llm, invoke_llm, record_llm_interaction
from app.services.translation import translate_text_if_needed
from app.services.pdf_utils import load_pdf_with_pypdf

# Async wrapper for invoke_llm that respects cancellation
async def invoke_llm_async(llm, prompt, variables=None):
    """Async wrapper for invoke_llm that properly handles cancellation"""
    # Simple approach: just call invoke_llm directly since we're checking
    # disconnection at the caller level anyway
    try:
        # Run in a thread to avoid blocking the event loop
        loop = asyncio.get_running_loop()
        result = await loop.run_in_executor(None, invoke_llm, llm, prompt, variables)
        return result
    except Exception as e:
        print(f"Error in invoke_llm_async: {e}")
        raise

# No longer using global cancellation tracking - relying on asyncio cancellation

# from langchain_community.document_loaders import PyPDFLoader  # Removed - using pypdf instead
import mimetypes
import logging

# Configure logging for TwinCheck
logger = logging.getLogger(__name__)

router = APIRouter(prefix="/twincheck", tags=["twincheck"])


def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text content from uploaded files using unified document processing.
    """
    from app.services.document_utils import extract_text_from_file_unified

    print(f"Processing file: {file.filename}")

    # Read the file content
    file_content = file.file.read()
    if not file_content:
        raise HTTPException(
            status_code=400,
            detail=f"Uploaded file {file.filename} appears to be empty",
        )

    try:
        return extract_text_from_file_unified(file_content, file.filename or "unknown")
    except Exception as e:
        print(f"Error processing file {file.filename}: {str(e)}")
        raise HTTPException(
            status_code=400, detail=f"Error processing file {file.filename}: {str(e)}"
        )


# Estimate the number of tokens in a text string.
# Uses tiktoken for accurate token counting.
def estimate_tokens(text: str, model: str = "gpt-4") -> int:
    try:
        # Try to get the encoding for the specific model
        if "gpt-4" in model.lower():
            encoding = tiktoken.encoding_for_model("gpt-4")
        elif "gpt-3.5" in model.lower():
            encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
        else:
            # Default to cl100k_base encoding used by most modern models
            encoding = tiktoken.get_encoding("cl100k_base")

        return len(encoding.encode(text))
    except Exception as e:
        # Fallback: rough estimation (1 token ≈ 4 characters)
        print(f"Token estimation error: {e}, using fallback method")
        return len(text) // 4


def chunk_diff_text(diff_text: str, max_tokens: int = None) -> List[str]:
    """
    Split diff text into chunks that don't exceed the token limit.
    Tries to preserve diff context by splitting at natural boundaries.
    """
    if max_tokens is None:
        max_tokens = settings.TWINCHECK_MAX_TOKENS_PER_CHUNK

    if estimate_tokens(diff_text) <= max_tokens:
        return [diff_text]

    chunks = []
    lines = diff_text.split("\n")
    current_chunk = []
    current_tokens = 0

    # Reserve tokens for prompt template and overhead
    chunk_token_limit = max_tokens - settings.TWINCHECK_PROMPT_RESERVE_TOKENS

    for line in lines:
        line_tokens = estimate_tokens(line + "\n")

        # If adding this line would exceed the limit, start a new chunk
        if current_tokens + line_tokens > chunk_token_limit and current_chunk:
            chunks.append("\n".join(current_chunk))
            current_chunk = []
            current_tokens = 0

        current_chunk.append(line)
        current_tokens += line_tokens

    # Add the last chunk if it has content
    if current_chunk:
        chunks.append("\n".join(current_chunk))

    return chunks


def create_synthesis_prompt(
    chunk_results: List[Dict[str, Any]], doc1_name: str, doc2_name: str, topics: str
) -> str:
    """
    Create a prompt for synthesizing multiple chunk analysis results.
    """
    return f"""
    You are synthesizing analysis results from multiple chunks of a document comparison.

    Documents compared:
    - Document 1: {doc1_name}
    - Document 2: {doc2_name}

    Topics of interest: {topics}

    Below are the analysis results from each chunk:

    {"=" * 50}

    {chr(10).join(
        [f"CHUNK {i+1} ANALYSIS:{chr(10)}{result['analysis']}{chr(10)}{chr(10)}"
         for i, result in enumerate(chunk_results)]
    )}

    {"=" * 50}

    Please provide a comprehensive synthesis that:
    1. Combines all the chunk analyses into a coherent overview
    2. Identifies patterns and themes across all chunks
    3. Highlights the most significant differences between the documents
    4. Removes any redundancy or overlap between chunk analyses
    5. Provides clear, actionable insights about the document differences

    Your synthesis should be well-structured and comprehensive while avoiding repetition.
    """


# Process documents for comparison
@router.post("/compare", response_model=TwinCheckResponse)
async def compare_documents(
    session: SessionDep,
    current_user: CurrentUser,
    request_data: TwinCheckRequest = Depends(),
    document1: UploadFile = File(...),
    document2: UploadFile = File(...),
    request: FastAPIRequest = None,
):
    """
    Compare two documents based on the provided comparison topics.
    Supports PDF, DOCX, and plain text files.
    """
    # Generate unique operation ID for logging
    operation_id = str(uuid.uuid4())
    print(f"Starting comparison operation {operation_id}")
    
    try:
        # Reset file pointers (in case they were read elsewhere)
        document1.file.seek(0)
        document2.file.seek(0)

        # Load the LLM model first to check vision capabilities
        llm = get_default_llm(session, current_user)

        # Check if LLM supports vision
        from app.services.vision_service import VisionService

        vision_enabled = VisionService.is_vision_enabled(llm)

        # Extract text from both documents
        doc1_text = extract_text_from_file(document1)

        # Reset file pointer for document2
        document2.file.seek(0)
        doc2_text = extract_text_from_file(document2)

        # Extract images if vision is enabled
        doc1_images = []
        doc2_images = []
        if vision_enabled:
            # Reset file pointers for image extraction
            document1.file.seek(0)
            document2.file.seek(0)

            # Extract images from both documents
            from app.services.document_utils import (
                extract_documents_and_images_from_file_unified,
            )

            doc1_content = document1.file.read()
            _, doc1_images = extract_documents_and_images_from_file_unified(
                doc1_content, document1.filename
            )

            document2.file.seek(0)
            doc2_content = document2.file.read()
            _, doc2_images = extract_documents_and_images_from_file_unified(
                doc2_content, document2.filename
            )

            print(f"Extracted {len(doc1_images)} images from {document1.filename}")
            print(f"Extracted {len(doc2_images)} images from {document2.filename}")

        # Split files into lines for diffing
        doc1_lines = doc1_text.splitlines()
        doc2_lines = doc2_text.splitlines()

        # Generate diff using difflib
        differ = difflib.Differ()
        diff_result = list(differ.compare(doc1_lines, doc2_lines))
        diff_text = "\n".join(diff_result)

        print(f"Generated diff text with {estimate_tokens(diff_text)} estimated tokens")

        # Parse comparison topics
        topic_list = request_data.comparison_topics.strip().split("\n")
        topic_analysis = []

        # Check if we need to chunk the diff text
        diff_chunks = chunk_diff_text(diff_text)
        is_chunked = len(diff_chunks) > 1

        print(f"Split diff into {len(diff_chunks)} chunks")

        # Process each topic with the LLM
        for topic_idx, topic in enumerate(topic_list):
            # CRITICAL: Check if client has disconnected before processing each topic
            try:
                if request and await request.is_disconnected():
                    print(f"❌ CLIENT DISCONNECTED - Stopping at topic {topic_idx + 1}")
                    return TwinCheckResponse(
                        results={
                            "status": "cancelled",
                            "message": "Request cancelled - client disconnected"
                        }
                    )
            except Exception as e:
                print(f"Warning: Could not check disconnect status: {e}")
            
            if not topic.strip():
                continue

            print(f"Processing topic {topic_idx + 1}/{len(topic_list)}: {topic}")

            # Simplified topic processing without Knowledge Base
            source_citations = []  # Keep this empty array for consistency

            if is_chunked:
                # Process each chunk for this topic
                chunk_results = []

                for i, chunk in enumerate(diff_chunks):
                    # CRITICAL: Check if client has disconnected before each chunk
                    try:
                        if request and await request.is_disconnected():
                            print(f"❌ CLIENT DISCONNECTED - Stopping at chunk {i + 1}")
                            return TwinCheckResponse(
                                results={
                                    "status": "cancelled",
                                    "message": "Request cancelled - client disconnected during chunk processing"
                                }
                            )
                    except Exception as e:
                        print(f"Warning: Could not check disconnect status: {e}")
                    
                    print(
                        f"  Processing chunk {i+1}/{len(diff_chunks)} for topic: {topic}"
                    )

                    try:
                        # Simplified prompt without knowledge base context
                        prompt_variables = {
                            "diff_text": chunk,
                            "topic": topic,
                            "doc1_name": document1.filename,
                            "doc2_name": document2.filename,
                            "knowledge_base_context": "",  # Always empty
                        }

                        chunk_result = invoke_llm(
                            llm,
                            settings.TWINCHECK_ANALYSIS_PROMPT_TEMPLATE,
                            prompt_variables,
                        )
                        
                        # CRITICAL: Check if client disconnected after LLM call
                        try:
                            if request and await request.is_disconnected():
                                print(f"❌ CLIENT DISCONNECTED - After LLM call for chunk {i + 1}")
                                return TwinCheckResponse(
                                    results={
                                        "status": "cancelled",
                                        "message": "Request cancelled - client disconnected after LLM call"
                                    }
                                )
                        except Exception as e:
                            print(f"Warning: Could not check disconnect status: {e}")

                        chunk_results.append(
                            {"chunk_index": i + 1, "analysis": chunk_result}
                        )

                    except Exception as e:
                        chunk_results.append(
                            {
                                "chunk_index": i + 1,
                                "analysis": f"Error analyzing chunk {i+1}: {str(e)}",
                            }
                        )

                # Synthesize the chunk results for this topic
                try:
                    synthesis_prompt = create_synthesis_prompt(
                        chunk_results, document1.filename, document2.filename, topic
                    )

                    print(
                        f"  Synthesizing {len(chunk_results)} chunk results for topic: {topic}"
                    )

                    synthesized_result = invoke_llm(llm, synthesis_prompt, {})

                    # Translate the synthesized result if needed
                    synthesized_result = await translate_text_if_needed(
                        synthesized_result, session, current_user, llm
                    )

                    # Add vision analysis if images exist and LLM supports it
                    if vision_enabled and (doc1_images or doc2_images):
                        print(f"  Adding vision analysis for chunked topic: {topic}")

                        # Prepare images for comparison
                        combined_images = VisionService.prepare_images_for_comparison(
                            doc1_images,
                            doc2_images,
                            document1.filename,
                            document2.filename,
                        )

                        vision_variables = {
                            "topic": topic,
                            "doc1_image_count": len(doc1_images),
                            "doc2_image_count": len(doc2_images),
                        }

                        try:
                            vision_analysis = await VisionService.process_images_with_prompt(
                                llm=llm,
                                images=combined_images,
                                prompt_template=settings.TWINCHECK_VISION_COMPARISON_PROMPT_TEMPLATE,
                                variables=vision_variables,
                            )

                            # Translate vision analysis if needed
                            vision_analysis = await translate_text_if_needed(
                                vision_analysis, session, current_user, llm
                            )

                            # Combine text and vision analysis
                            final_analysis = (
                                VisionService.combine_text_and_vision_analysis(
                                    synthesized_result, vision_analysis
                                )
                            )

                            topic_analysis.append(
                                {
                                    "topic": topic,
                                    "analysis": final_analysis,
                                    "chunk_count": len(diff_chunks),
                                    "has_vision_analysis": True,
                                    "image_count": {
                                        "doc1": len(doc1_images),
                                        "doc2": len(doc2_images),
                                    },
                                    "source_citations": source_citations,
                                }
                            )

                        except Exception as vision_error:
                            print(
                                f"Vision analysis error for chunked topic {topic}: {vision_error}"
                            )
                            # Fall back to text-only analysis
                            topic_analysis.append(
                                {
                                    "topic": topic,
                                    "analysis": synthesized_result,
                                    "chunk_count": len(diff_chunks),
                                    "vision_error": str(vision_error),
                                    "source_citations": source_citations,
                                }
                            )
                    else:
                        # Text-only analysis for chunked processing
                        topic_analysis.append(
                            {
                                "topic": topic,
                                "analysis": synthesized_result,
                                "chunk_count": len(diff_chunks),
                                "source_citations": source_citations,
                            }
                        )

                except Exception as e:
                    # Fallback: combine chunk results manually
                    combined_analysis = (
                        f"Analysis from {len(chunk_results)} chunks:\n\n"
                    )
                    for result in chunk_results:
                        combined_analysis += (
                            f"Chunk {result['chunk_index']}:\n{result['analysis']}\n\n"
                        )

                    topic_analysis.append(
                        {
                            "topic": topic,
                            "analysis": combined_analysis,
                            "chunk_count": len(diff_chunks),
                            "synthesis_error": str(e),
                            "source_citations": source_citations,
                        }
                    )
            else:
                # Single chunk processing (original behavior)
                try:
                    # Simplified prompt without knowledge base context
                    prompt_variables = {
                        "diff_text": diff_text,
                        "topic": topic,
                        "doc1_name": document1.filename,
                        "doc2_name": document2.filename,
                        "knowledge_base_context": "",  # Always empty
                    }

                    topic_result = invoke_llm(
                        llm,
                        settings.TWINCHECK_ANALYSIS_PROMPT_TEMPLATE,
                        prompt_variables,
                    )

                    # Add vision analysis if images exist and LLM supports it
                    if vision_enabled and (doc1_images or doc2_images):
                        # Check for cancellation before vision processing
                        await asyncio.sleep(0.01)
                        print(f"  Adding vision analysis for topic: {topic}")

                        # Prepare images for comparison
                        combined_images = VisionService.prepare_images_for_comparison(
                            doc1_images,
                            doc2_images,
                            document1.filename,
                            document2.filename,
                        )

                        vision_variables = {
                            "topic": topic,
                            "doc1_image_count": len(doc1_images),
                            "doc2_image_count": len(doc2_images),
                        }

                        try:
                            vision_analysis = await VisionService.process_images_with_prompt(
                                llm=llm,
                                images=combined_images,
                                prompt_template=settings.TWINCHECK_VISION_COMPARISON_PROMPT_TEMPLATE,
                                variables=vision_variables,
                            )

                            # Translate vision analysis if needed
                            vision_analysis = await translate_text_if_needed(
                                vision_analysis, session, current_user, llm
                            )

                            # Combine text and vision analysis
                            combined_analysis = (
                                VisionService.combine_text_and_vision_analysis(
                                    topic_result, vision_analysis
                                )
                            )

                            topic_analysis.append(
                                {
                                    "topic": topic,
                                    "analysis": combined_analysis,
                                    "has_vision_analysis": True,
                                    "image_count": {
                                        "doc1": len(doc1_images),
                                        "doc2": len(doc2_images),
                                    },
                                    "source_citations": source_citations,
                                }
                            )

                        except Exception as vision_error:
                            print(
                                f"Vision analysis error for topic {topic}: {vision_error}"
                            )
                            # Fall back to text-only analysis
                            topic_analysis.append(
                                {
                                    "topic": topic,
                                    "analysis": topic_result,
                                    "vision_error": str(vision_error),
                                    "source_citations": source_citations,
                                }
                            )
                    else:
                        # Text-only analysis
                        # Translate the topic result if needed
                        await asyncio.sleep(0.01)  # Allow cancellation during translation
                        topic_result = await translate_text_if_needed(
                            topic_result, session, current_user, llm
                        )

                        topic_analysis.append(
                            {
                                "topic": topic,
                                "analysis": topic_result,
                                "source_citations": source_citations,
                            }
                        )

                except Exception as e:
                    topic_analysis.append(
                        {
                            "topic": topic,
                            "analysis": f"Error analyzing this topic: {str(e)}",
                            "source_citations": source_citations,
                        }
                    )

        # Create a comprehensive summary
        if is_chunked:
            # For chunked processing, create summary from topic analyses
            print("Creating summary from topic analyses (chunked mode)")

            topic_summaries = "\n\n".join(
                [f"Topic: {ta['topic']}\n{ta['analysis']}" for ta in topic_analysis]
            )

            summary_prompt = f"""
            You are creating a comprehensive summary of a document comparison that was processed in chunks due to size.
            
            Documents compared:
            - Document 1: {document1.filename}
            - Document 2: {document2.filename}
            
            The comparison was processed in {len(diff_chunks)} chunks and analyzed across the following topics:
            {request.comparison_topics}
            
            Below are the detailed topic analyses:
            
            {topic_summaries}
            
            Please provide a comprehensive executive summary that:
            1. Highlights the most significant overall differences between the documents
            2. Synthesizes patterns across all topic analyses
            3. Provides clear, actionable insights about the document comparison
            4. Is well-structured and avoids repetition
            
            Focus on the big picture and most important differences.
            """

            try:
                summary = await invoke_llm_async(llm, summary_prompt, {})

                # Translate the summary if needed
                summary = await translate_text_if_needed(
                    summary, session, current_user, llm
                )

            except Exception as e:
                summary = f"Summary generation error: {str(e)}\n\nPlease refer to the individual topic analyses below for detailed insights."
        else:
            # Single chunk processing (original behavior)
            print("Creating summary from diff text (single chunk mode)")
            summary = invoke_llm(
                llm,
                settings.TWINCHECK_SUMMARY_PROMPT_TEMPLATE,
                {
                    "diff_text": diff_text,
                    "doc1_name": document1.filename,
                    "doc2_name": document2.filename,
                    "topics": request_data.comparison_topics,
                },
            )

            # Translate the summary if needed
            summary = await translate_text_if_needed(
                summary, session, current_user, llm
            )

        # Record this interaction for history
        interaction_id = record_llm_interaction(
            session=session,
            user_id=current_user.id,
            functionality="twincheck",
            input_data={
                "comparison_topics": request_data.comparison_topics,
                "document1_name": document1.filename,
                "document2_name": document2.filename,
            },
            output_data={"summary": summary, "topic_count": len(topic_analysis)},
            metadata={
                "topic_analysis": topic_analysis,  # Store detailed analysis for retrieval
                "diff_stats": {
                    "total_tokens": estimate_tokens(diff_text),
                    "chunk_count": len(diff_chunks),
                    "was_chunked": is_chunked,
                    "additions": diff_text.count("\n+ "),
                    "deletions": diff_text.count("\n- "),
                    "changes": diff_text.count("\n? "),
                },
            },
        )

        # Return the results
        result = {
            "summary": summary,
            "topic_analysis": topic_analysis,
            "interaction_id": str(interaction_id) if interaction_id else None,
            "processing_info": {
                "was_chunked": is_chunked,
                "chunk_count": len(diff_chunks),
                "estimated_tokens": estimate_tokens(diff_text),
            },
        }

        return TwinCheckResponse(results=result)

    except asyncio.CancelledError:
        print(f"Comparison operation {operation_id} was cancelled")
        raise  # Re-raise to properly handle the cancellation
    except Exception as e:
        print(f"Error in comparison operation {operation_id}: {str(e)}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error comparing documents: {str(e)}"
        )



# Get history of comparison operations
@router.get("/history", response_model=List[Dict[str, Any]])
async def get_comparison_history(
    session: SessionDep,
    current_user: CurrentUser,
    skip: int = 0,
    limit: int = 20,
    show_all: bool = False,
):
    """Retrieve past document comparison history for the current user or all users."""
    print("Retrieving TwinCheck history. Show all:", show_all)

    try:
        # Start with base query
        query = select(LlmInteraction).where(
            LlmInteraction.functionality == "twincheck"
        )

        # Only filter by user if not showing all users
        if not show_all:
            query = query.where(LlmInteraction.user_id == current_user.id)

        # Add ordering and pagination
        comparisons = session.exec(
            query.order_by(LlmInteraction.date_created.desc()).offset(skip).limit(limit)
        ).all()

        result = []
        for comparison in comparisons:
            try:
                input_data = (
                    json.loads(comparison.input_data) if comparison.input_data else {}
                )
                output_data = (
                    json.loads(comparison.output_data) if comparison.output_data else {}
                )

                # Create result item
                result_item = {
                    "id": str(comparison.id),
                    "date_created": comparison.date_created,
                    "document1_name": input_data.get(
                        "document1_name", "Unknown Document 1"
                    ),
                    "document2_name": input_data.get(
                        "document2_name", "Unknown Document 2"
                    ),
                    "comparison_topics": input_data.get("comparison_topics", ""),
                    "topic_count": output_data.get("topic_count", 0),
                    "has_feedback": comparison.feedback is not None,
                }

                # Add feedback information if exists
                if comparison.feedback:
                    result_item["feedback"] = {
                        "feedback": comparison.feedback,
                        "feedbackText": comparison.feedback_text,
                    }

                # Add user info for all-users view
                if show_all:
                    from app.models import User  # Import here to avoid circular imports

                    user = session.get(User, comparison.user_id)
                    user_name = (
                        f"{user.full_name or 'User'} ({user.email})"
                        if user
                        else "Unknown User"
                    )
                    result_item["user_name"] = user_name

                result.append(result_item)
            except json.JSONDecodeError:
                # If JSON parsing fails, use minimal information
                # Create result item with minimal info
                result_item = {
                    "id": str(comparison.id),
                    "date_created": comparison.date_created,
                    "document1_name": "Unknown Document 1",
                    "document2_name": "Unknown Document 2",
                    "topic_count": 0,
                    "has_feedback": comparison.feedback is not None,
                }

                # Add feedback information if exists
                if comparison.feedback:
                    result_item["feedback"] = {
                        "feedback": comparison.feedback,
                        "feedbackText": comparison.feedback_text,
                    }

                # Add user info for all-users view
                if show_all:
                    from app.models import User  # Import here to avoid circular imports

                    user = session.get(User, comparison.user_id)
                    user_name = (
                        f"{user.full_name or 'User'} ({user.email})"
                        if user
                        else "Unknown User"
                    )
                    result_item["user_name"] = user_name

                result.append(result_item)

        return result
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error retrieving comparison history: {str(e)}"
        )


# Get details of a specific comparison
@router.get("/history/{comparison_id}", response_model=TwinCheckDetailResponse)
async def get_comparison_detail(
    comparison_id: uuid.UUID,
    session: SessionDep,
    current_user: CurrentUser,
):
    """Retrieve a specific comparison's full content by ID."""
    try:
        comparison = session.get(LlmInteraction, comparison_id)
        if not comparison:
            raise HTTPException(status_code=404, detail="Comparison not found")

        # No longer need to check this since we now allow viewing others' outputs
        # if comparison.user_id != current_user.id:
        #    raise HTTPException(
        #        status_code=403, detail="You don't have access to this comparison"
        #    )

        if comparison.functionality != "twincheck":
            raise HTTPException(
                status_code=400, detail="This is not a TwinCheck comparison"
            )

        try:
            input_data = (
                json.loads(comparison.input_data) if comparison.input_data else {}
            )
            output_data = (
                json.loads(comparison.output_data) if comparison.output_data else {}
            )
            extra_data = comparison.extra_data or {}

            # Create a response that matches the structure expected by the frontend
            result = {
                "id": str(comparison.id),
                "date_created": comparison.date_created,
                "document1_name": input_data.get(
                    "document1_name", "Unknown Document 1"
                ),
                "document2_name": input_data.get(
                    "document2_name", "Unknown Document 2"
                ),
                "comparison_topics": input_data.get("comparison_topics", ""),
                "results": {
                    "summary": output_data.get("summary", ""),
                    "topic_analysis": extra_data.get("topic_analysis", []),
                    "interaction_id": str(comparison.id),
                },
                # Add feedback information
                "feedback": {
                    "feedback": comparison.feedback,
                    "feedbackText": comparison.feedback_text,
                    "feedbackDate": (
                        comparison.feedback_date.isoformat()
                        if comparison.feedback_date
                        else None
                    ),
                },
            }

            return result

        except json.JSONDecodeError:
            # Fallback if JSON parsing fails
            return {
                "id": str(comparison.id),
                "date_created": comparison.date_created,
                "document1_name": "Unknown Document 1",
                "document2_name": "Unknown Document 2",
                "results": {
                    "summary": "Unable to reconstruct comparison from this record. This might be due to an older format or incomplete data.",
                    "topic_analysis": [],
                },
                # Add empty feedback object for consistency
                "feedback": {
                    "feedback": None,
                    "feedbackText": None,
                    "feedbackDate": None,
                },
            }

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error retrieving comparison details: {str(e)}"
        )


# Functions related to Comparisons (saved comparison topics)
@router.post("/comparisons", response_model=TwinCheckTopicList)
def create_comparison(
    comparison: TwinCheckTopicList, session: SessionDep, current_user: CurrentUser
):
    """
    Save a new comparison topic set to the database.
    """
    print(f"🐛 DEBUG: Received comparison data: {comparison}")
    print(
        f"🐛 DEBUG: Name: '{comparison.name}', Topics: '{comparison.topics}', Description: '{comparison.description}'"
    )

    # Validate required fields
    if not comparison.name or not comparison.name.strip():
        print(f"🐛 DEBUG: Name validation failed - name is empty or None")
        raise HTTPException(
            status_code=400, detail="Comparison name is required and cannot be empty"
        )

    if not comparison.topics or not comparison.topics.strip():
        print(f"🐛 DEBUG: Topics validation failed - topics is empty or None")
        raise HTTPException(
            status_code=400, detail="Comparison topics are required and cannot be empty"
        )

    existing_comparison = session.exec(
        select(TwinCheckTopicList).where(TwinCheckTopicList.name == comparison.name)
    ).first()

    if existing_comparison:
        print(
            f"🐛 DEBUG: Duplicate name found - existing comparison: {existing_comparison.name}"
        )
        raise HTTPException(
            status_code=400, detail="A comparison with this name already exists."
        )

    comparison.owner_id = current_user.id
    session.add(comparison)
    session.commit()
    session.refresh(comparison)
    print(f"🐛 DEBUG: Successfully created comparison with ID: {comparison.id}")
    return comparison


@router.get("/comparisons", response_model=List[TwinCheckTopicList])
def get_comparisons(session: SessionDep, current_user: CurrentUser):
    """
    Retrieve all saved comparison topic sets from the database for this user.
    """
    return session.exec(
        select(TwinCheckTopicList).where(TwinCheckTopicList.owner_id == current_user.id)
    ).all()


@router.get("/comparisons/{comparison_id}", response_model=TwinCheckTopicList)
def get_comparison(comparison_id: uuid.UUID, session: SessionDep):
    """
    Retrieve a specific comparison topic set by ID.
    """
    comparison = session.get(TwinCheckTopicList, comparison_id)
    if not comparison:
        raise HTTPException(status_code=404, detail="Comparison not found.")
    return comparison


@router.put("/comparisons/{comparison_id}", response_model=TwinCheckTopicList)
def update_comparison(
    comparison_id: uuid.UUID,
    updated_comparison: TwinCheckTopicList,
    session: SessionDep,
    current_user: CurrentUser,
):
    """
    Update an existing comparison topic set.
    """
    comparison = session.get(TwinCheckTopicList, comparison_id)
    if not comparison:
        raise HTTPException(status_code=404, detail="Comparison not found.")

    # Ensure the current user is the owner of the comparison
    if comparison.owner_id != current_user.id:
        raise HTTPException(
            status_code=403, detail="Not authorized to update this comparison."
        )

    comparison.name = updated_comparison.name
    comparison.description = updated_comparison.description
    comparison.topics = updated_comparison.topics
    comparison.date_modified = datetime.utcnow()

    session.add(comparison)
    session.commit()
    session.refresh(comparison)
    return comparison


@router.delete("/comparisons/{comparison_id}", response_model=Message)
def delete_comparison(
    comparison_id: uuid.UUID, session: SessionDep, current_user: CurrentUser
):
    """
    Delete a comparison topic set by ID.
    """
    comparison = session.get(TwinCheckTopicList, comparison_id)
    if not comparison:
        raise HTTPException(status_code=404, detail="Comparison not found.")

    # Ensure the current user is the owner of the comparison
    if comparison.owner_id != current_user.id:
        raise HTTPException(
            status_code=403, detail="Not authorized to delete this comparison."
        )

    session.delete(comparison)
    session.commit()
    return Message(message="Comparison deleted successfully.")


@router.post("/generate/docx", response_class=StreamingResponse)
async def generate_docx(
    session: SessionDep, current_user: CurrentUser, request: DocxRequest
):
    """
    Generate a DOCX file from the comparison content.
    """
    print("Now generating DOCX of comparison...")
    try:
        # Get the markdown content from the request
        if not request.content:
            raise HTTPException(status_code=400, detail="Report content is required")

        # Convert markdown to HTML for parsing
        html_content = markdown.markdown(request.content, extensions=["tables"])
        soup = BeautifulSoup(html_content, "html.parser")

        print("Markdown content converted to HTML successfully.")
        # Create a new Document
        doc = Document()

        print("Adding title and date to the document...")
        # Add a title
        title_text = (
            request.title
            if hasattr(request, "title") and request.title
            else "Document Comparison"
        )
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_paragraph.add_run(
            f"Generated on: {datetime.now().strftime('%B %d, %Y')}"
        )
        date_run.italic = True

        # Add a separator
        doc.add_paragraph("─" * 50)

        print("Adding headers, paragraphs, lists, and tables...")
        # Process all headers and paragraphs in the HTML
        for element in soup.find_all(
            ["h1", "h2", "h3", "h4", "p", "ul", "ol", "li", "table"]
        ):
            if element.name == "h1":
                doc.add_heading(element.text, level=1)
            elif element.name == "h2":
                doc.add_heading(element.text, level=2)
            elif element.name == "h3":
                doc.add_heading(element.text, level=3)
            elif element.name == "p":
                doc.add_paragraph(element.text)
            elif element.name == "ul":
                for li in element.find_all("li"):
                    paragraph = doc.add_paragraph(li.text)
                    paragraph.style = "List Bullet"
            elif element.name == "ol":
                for li in element.find_all("li"):
                    paragraph = doc.add_paragraph(li.text)
                    paragraph.style = "List Number"
            elif element.name == "table":
                table_rows = element.find_all("tr")
                if table_rows:
                    # Count the number of columns in the first row
                    first_row = table_rows[0]
                    columns = len(first_row.find_all(["th", "td"]))

                    # Create the table
                    table = doc.add_table(rows=0, cols=columns)
                    table.style = "Table Grid"

                    # Process header row
                    header_cells = first_row.find_all(["th", "td"])
                    if header_cells:
                        header_row = table.add_row().cells
                        for i, cell in enumerate(header_cells):
                            if i < len(header_row):
                                header_row[i].text = cell.text
                                run = header_row[i].paragraphs[0].runs[0]
                                run.bold = True

                    # Process data rows
                    for row in table_rows[1:]:
                        cells = row.find_all("td")
                        if cells:
                            row_cells = table.add_row().cells
                            for i, cell in enumerate(cells):
                                if i < len(row_cells):
                                    row_cells[i].text = cell.text

        # Save the document to a BytesIO object
        print("Saving the document to a BytesIO object...")
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        # --- CORRUPTION CHECKS ---
        # 1. Check file size
        size = docx_bytes.getbuffer().nbytes
        print(f"DOCX file size: {size} bytes")
        if size < 1000:
            print("Warning: DOCX file is very small and may be empty or corrupted.")

        # 2. Try to reload the DOCX to ensure it's readable
        try:
            docx_bytes.seek(0)
            _ = Document(docx_bytes)
            print("DOCX file passed integrity check (can be opened by python-docx).")
        except Exception as e:
            print(
                f"Integrity check failed: generated DOCX cannot be opened. Error: {e}"
            )
            raise HTTPException(
                status_code=500, detail="Generated DOCX file is corrupted."
            )

        docx_bytes.seek(0)

        print(
            "Document saved successfully. Preparing to return as a downloadable file."
        )

        # Create a filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"comparison_{timestamp}.docx"

        # Return the document as a downloadable file
        return StreamingResponse(
            docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")


@router.post("/generate/csv", response_class=StreamingResponse)
async def generate_csv(
    session: SessionDep, current_user: CurrentUser, request: DocxRequest
):
    """
    Generate a CSV file from the comparison content.
    """
    print("Now generating CSV of comparison...")
    try:
        # Get the content from the request
        if not request.content:
            raise HTTPException(status_code=400, detail="Report content is required")

        # Try to parse the content as JSON first (for structured data)
        try:
            content_data = json.loads(request.content)
            summary = content_data.get("summary", "")
            topic_results = content_data.get("topic_results", [])
            doc1_name = content_data.get("doc1_name", "Document 1")
            doc2_name = content_data.get("doc2_name", "Document 2")
        except json.JSONDecodeError:
            # If it's not JSON, treat it as markdown content
            content_lines = request.content.split("\n")
            summary = ""
            topic_results = []
            doc1_name = "Document 1"
            doc2_name = "Document 2"

            # Extract summary and topics from markdown
            current_section = ""
            current_content = []

            for line in content_lines:
                if line.startswith("# Summary"):
                    current_section = "summary"
                    current_content = []
                elif line.startswith("# Topic Analysis") or line.startswith(
                    "## Topic:"
                ):
                    if current_section == "summary":
                        summary = "\n".join(current_content).strip()
                    current_section = "topic"
                    if line.startswith("## Topic:"):
                        topic_name = line.replace("## Topic:", "").strip()
                        current_content = [topic_name]
                elif line.strip() and current_section:
                    current_content.append(line.strip())

            # Handle last section
            if current_section == "summary":
                summary = "\n".join(current_content).strip()

        # Create CSV content
        output = StringIO()
        writer = csv.writer(output)

        # Write headers
        headers = ["Comparison Topic", "Analysis", "Document 1", "Document 2"]
        writer.writerow(headers)

        # Write summary row
        if summary:
            writer.writerow(
                [
                    "Overall Summary",
                    summary.replace("\n", " ").replace("\r", "").replace('"', '""'),
                    doc1_name,
                    doc2_name,
                ]
            )

        # Write topic analysis rows
        for topic_result in topic_results:
            if isinstance(topic_result, dict):
                topic = topic_result.get("topic", "Unknown Topic")
                analysis = topic_result.get("analysis", "No analysis available")
            else:
                topic = str(topic_result)
                analysis = "No analysis available"

            # Clean analysis text for CSV
            cleaned_analysis = (
                analysis.replace("\n", " ").replace("\r", "").replace('"', '""')
            )

            writer.writerow([topic, cleaned_analysis, doc1_name, doc2_name])

        # If no structured data was found, create a simple row with the content
        if not summary and not topic_results:
            writer.writerow(
                [
                    "Comparison Analysis",
                    request.content.replace("\n", " ")
                    .replace("\r", "")
                    .replace('"', '""'),
                    doc1_name,
                    doc2_name,
                ]
            )

        # Prepare the CSV for download
        csv_content = output.getvalue()
        output.close()

        # Convert to bytes
        csv_bytes = csv_content.encode("utf-8")

        # Generate timestamp for filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"twincheck_comparison_{timestamp}.csv"

        print(
            "CSV file generated successfully. Preparing to return as a downloadable file."
        )

        return StreamingResponse(
            BytesIO(csv_bytes),
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating CSV: {str(e)}")


@router.post("/generate-topics", response_model=GenerateTopicsResponse)
async def generate_topics(
    session: SessionDep,
    current_user: CurrentUser,
    description: str = Form(...),
    comparison_type: str = Form("general"),
    num_topics: Optional[int] = Form(None),
    search_mode: str = "vector",
    knowledge_base_id: Optional[str] = Form(None),
    files: List[UploadFile] = File(default=[]),
):
    """
    Generate comparison topics based on a description using LLM, with optional example document.
    """
    print("generate_topics function invoked!")
    print(f"Received search_mode: {search_mode}")
    print(f"Request data: description={description[:50]}...")

    # Validate search mode
    if search_mode not in ["vector", "full_scan"]:
        print(f"Warning: Invalid search mode '{search_mode}', defaulting to 'vector'")
        search_mode = "vector"

    try:
        # Get the default LLM
        llm = get_default_llm(session, current_user)

        # Extract text from example document if provided
        example_document = ""
        example_instruction = ""
        example_analysis_instruction = ""

        if files and len(files) > 0:
            file = files[0]
            if file.size > 0:
                # Reset file pointer to beginning before processing
                file.file.seek(0)
                file_content = file.file.read()

                # Use enhanced processing with vision capabilities
                from app.services.document_utils import (
                    extract_text_with_vision_enhancement,
                )

                example_document = await extract_text_with_vision_enhancement(
                    file_content,
                    file.filename or "unknown",
                    llm,
                    purpose="comparison topic generation",
                )
                example_instruction = f" and use the uploaded example document ({file.filename}) as a reference for the appropriate scope and depth of comparison topics"
                example_analysis_instruction = f" and explain how they align with the scope shown in the example document ({file.filename})"

        # Check if example document exceeds token limits and chunk if necessary
        if example_document:
            print(f"Total example document content: {len(example_document)} characters")

            # Using conservative chunking similar to TWINCHECK settings
            max_chunk_size = 80000  # Conservative chunk size for 128K context limit

            if len(example_document) > max_chunk_size:
                print(
                    f"Example document too large ({len(example_document)} chars), chunking for processing"
                )

                from app.services.text_processing import chunk_text

                # Chunk the document content
                chunks = chunk_text(example_document, max_tokens=max_chunk_size)

                # Process each chunk to generate topics
                all_chunk_topics = []

                for i, chunk in enumerate(chunks):
                    print(f"Processing chunk {i+1}/{len(chunks)}")

                    # Generate topics for this chunk
                    chunk_prompt_variables = {
                        "description": description,
                        "comparison_type": comparison_type,
                        "example_document": f"EXAMPLE DOCUMENT: {file.filename}\n{chunk}\n",
                        "example_instruction": example_instruction,
                        "example_analysis_instruction": example_analysis_instruction,
                        "knowledge_base_content": "",
                        "knowledge_base_instruction": "",
                    }

                    # If knowledge base is specified, retrieve content using selected search mode
                    if knowledge_base_id:
                        print(f"\n=== KB INTEGRATION FOR CHUNK {i+1} ====")
                        try:
                            from app.services.content_retrieval import (
                                retrieve_knowledge_base_content,
                            )

                            print(
                                f"Retrieving knowledge base content for chunk {i+1}, KB ID: {knowledge_base_id}, search mode: {search_mode}"
                            )
                            content, instruction = (
                                await retrieve_knowledge_base_content(
                                    session=session,
                                    current_user=current_user,
                                    knowledge_base_id=knowledge_base_id,
                                    search_mode=search_mode,
                                    query=description,
                                )
                            )

                            if content:
                                print(
                                    f"✅ Successfully retrieved KB content for chunk {i+1}: {len(content)} characters"
                                )
                                chunk_prompt_variables["knowledge_base_content"] = (
                                    f"REFERENCE DOCUMENTS FROM KNOWLEDGE BASE:\n{content}"
                                )
                                chunk_prompt_variables["knowledge_base_instruction"] = (
                                    f"\n12. {instruction} Use them as inspiration for the type of comparison topics and scope, "
                                    f"adapting the topics to match the specific requirements in the description. "
                                    f"Search mode used: {search_mode}"
                                )
                                chunk_prompt_variables[
                                    "example_analysis_instruction"
                                ] += f". Briefly mention how the knowledge base content (using {search_mode}) influenced the topic selection"
                            else:
                                print(
                                    f"No content retrieved from knowledge base for chunk {i+1}"
                                )
                        except Exception as e:
                            print(
                                f"Error retrieving knowledge base content for chunk {i+1}: {e}"
                            )
                            import traceback

                            traceback.print_exc()

                    try:
                        chunk_response = invoke_llm(
                            llm,
                            settings.TWINCHECK_GENERATE_TOPICS_PROMPT_TEMPLATE,
                            chunk_prompt_variables,
                        )

                        # Parse topics from chunk response
                        chunk_topics = []
                        lines = chunk_response.strip().split("\n")
                        in_topics_section = False

                        for line in lines:
                            line = line.strip()
                            if line.startswith("TOPICS:"):
                                in_topics_section = True
                                continue
                            elif line.startswith("ANALYSIS:"):
                                in_topics_section = False
                                continue

                            if in_topics_section:
                                if re.match(r"^\d+\.\s+", line):
                                    topic = re.sub(r"^\d+\.\s+", "", line)
                                    if topic.strip():
                                        chunk_topics.append(topic.strip())

                        # If parsing failed, try simpler approach
                        if not chunk_topics:
                            for line in lines:
                                line = line.strip()
                                if re.match(r"^\d+\.\s+", line):
                                    topic = re.sub(r"^\d+\.\s+", "", line)
                                    if topic.strip():
                                        chunk_topics.append(topic.strip())

                        all_chunk_topics.extend(chunk_topics)

                    except Exception as e:
                        print(f"Error processing chunk {i+1}: {e}")
                        continue

                # Deduplicate and refine topics across all chunks
                if all_chunk_topics:
                    # Remove duplicates while preserving order
                    seen = set()
                    unique_topics = []
                    for t in all_chunk_topics:
                        if t.lower() not in seen:
                            seen.add(t.lower())
                            unique_topics.append(t)

                    # If we have too many topics, synthesize and prioritize
                    if len(unique_topics) > (num_topics or 20):
                        synthesis_prompt = f"""From the following list of comparison topics, select and refine the {num_topics or 10} most important and relevant topics for {comparison_type} comparison based on: {description}

Topics to review:
{chr(10).join([f"{i+1}. {t}" for i, t in enumerate(unique_topics)])}

Requirements:
1. Select the most critical and comprehensive topics
2. Ensure good coverage of comparison aspects
3. Maintain clarity and specificity
4. Focus on topics most relevant to the description

Return only the final selected topics, one per line, numbered."""

                        try:
                            refined_response = invoke_llm(llm, synthesis_prompt, {})
                            topics = []
                            for line in refined_response.strip().split("\n"):
                                line = line.strip()
                                if line and (
                                    line[0].isdigit()
                                    or line.startswith("-")
                                    or line.startswith("*")
                                ):
                                    topic = re.sub(r"^\d+\.\s+", "", line)
                                    topic = re.sub(r"^[-*]\s+", "", topic)
                                    if topic.strip():
                                        topics.append(topic.strip())
                        except Exception as e:
                            print(f"Error in topic synthesis: {e}")
                            topics = unique_topics[: num_topics or 10]
                    else:
                        topics = unique_topics[: num_topics or 20]

                    # For analysis, show chunked processing was used
                    analysis = f"Generated {len(topics)} topics from chunked example document analysis ({len(chunks)} chunks processed) based on the provided description to ensure comprehensive document comparison coverage."

                    # Record the interaction
                    record_llm_interaction(
                        session=session,
                        user_id=current_user.id,
                        functionality="generate_comparison_topics",
                        input_data={
                            "description": description,
                            "requested_topics": num_topics,
                            "comparison_type": comparison_type,
                            "has_example_document": True,
                            "chunked_processing": True,
                            "chunk_count": len(chunks),
                            "search_mode": search_mode,
                        },
                        output_data={
                            "topics_count": len(topics),
                            "analysis": analysis,
                        },
                        metadata={},
                    )

                    return GenerateTopicsResponse(
                        topics=topics, description_analysis=analysis
                    )
                else:
                    # Fallback to description-only generation if chunk processing failed
                    example_document = ""
                    example_instruction = ""
                    example_analysis_instruction = ""

        # Continue with existing logic for small documents or when no files provided
        prompt_variables = {
            "description": description,
            "comparison_type": comparison_type,
            "example_document": (
                f"EXAMPLE DOCUMENT: {file.filename}\n{example_document}\n"
                if example_document
                else ""
            ),
            "example_instruction": example_instruction,
            "example_analysis_instruction": example_analysis_instruction,
            "knowledge_base_content": "",
            "knowledge_base_instruction": "",
        }

        # If knowledge base is specified, retrieve content using selected search mode
        if knowledge_base_id:
            print(f"\n=== KNOWLEDGE BASE INTEGRATION START ===")
            print(f"Knowledge Base ID: {knowledge_base_id}")
            print(f"Search Mode: {search_mode}")
            print(f"Query (description): {description[:100]}...")
            try:
                from app.services.content_retrieval import (
                    retrieve_knowledge_base_content,
                )

                print(
                    f"Retrieving knowledge base content for KB ID: {knowledge_base_id}, search mode: {search_mode}"
                )
                content, instruction = await retrieve_knowledge_base_content(
                    session=session,
                    current_user=current_user,
                    knowledge_base_id=knowledge_base_id,
                    search_mode=search_mode,
                    query=description,
                )

                if content:
                    print(
                        f"✅ Successfully retrieved KB content: {len(content)} characters"
                    )
                    print(f"Instruction: {instruction}")
                    print(f"Content preview: {content[:200]}...")
                    prompt_variables["knowledge_base_content"] = (
                        f"REFERENCE DOCUMENTS FROM KNOWLEDGE BASE:\n{content}"
                    )
                    prompt_variables["knowledge_base_instruction"] = (
                        f"\n12. {instruction} Use them as inspiration for the type of comparison topics and scope, "
                        f"adapting the topics to match the specific requirements in the description. "
                        f"Search mode used: {search_mode}"
                    )
                    prompt_variables[
                        "example_analysis_instruction"
                    ] += f". Briefly mention how the knowledge base content (using {search_mode}) influenced the topic selection"
                    print("✅ Knowledge base content added to prompt variables")
                else:
                    print("❌ No content retrieved from knowledge base")
                    print("This means either:")
                    print("  - No sources found in the knowledge base")
                    print("  - No extractable content in the sources")
                    print("  - Knowledge base doesn't exist or access denied")
                print(f"=== KNOWLEDGE BASE INTEGRATION END ===\n")
            except Exception as e:
                print(f"❌ Error retrieving knowledge base content: {e}")
                import traceback

                traceback.print_exc()
                print(f"=== KNOWLEDGE BASE INTEGRATION FAILED ===\n")

        # Generate topics using the LLM
        print(f"\n=== PROMPT VARIABLES SUMMARY ===")
        print(f"Description: {prompt_variables['description'][:100]}...")
        print(f"KB Content Length: {len(prompt_variables['knowledge_base_content'])}")
        print(
            f"KB Instruction: {prompt_variables['knowledge_base_instruction'][:100]}..."
        )
        print(f"Example Document Length: {len(prompt_variables['example_document'])}")
        print(f"=== CALLING LLM ===\n")

        topics_response = await invoke_llm_async(
            llm,
            settings.TWINCHECK_GENERATE_TOPICS_PROMPT_TEMPLATE,
            prompt_variables,
        )

        # Parse the response to extract topics and analysis
        topics = []
        analysis = ""

        lines = topics_response.strip().split("\n")
        in_topics_section = False
        in_analysis_section = False

        for line in lines:
            line = line.strip()
            if line.startswith("TOPICS:"):
                in_topics_section = True
                in_analysis_section = False
                continue
            elif line.startswith("ANALYSIS:"):
                in_topics_section = False
                in_analysis_section = True
                continue

            if in_topics_section:
                # Extract topics (numbered list)
                if re.match(r"^\d+\.\s+", line):
                    topic = re.sub(r"^\d+\.\s+", "", line)
                    if topic.strip():
                        topics.append(topic.strip())
            elif in_analysis_section:
                if line:
                    if analysis:
                        analysis += " " + line
                    else:
                        analysis = line

        # If parsing failed, try simpler approach
        if not topics:
            # Split by lines and look for numbered items
            for line in lines:
                line = line.strip()
                if re.match(r"^\d+\.\s+", line):
                    topic = re.sub(r"^\d+\.\s+", "", line)
                    if topic.strip():
                        topics.append(topic.strip())

        # Ensure we have some topics
        if not topics:
            raise HTTPException(
                status_code=500,
                detail="Failed to generate topics from the description. Please try with a more detailed description.",
            )

        # Apply user-specified limit if provided, otherwise use all generated topics
        if num_topics:
            topics = topics[:num_topics]

        if not analysis:
            analysis = f"Generated {len(topics)} comparison topics based on the provided description to ensure comprehensive document comparison coverage."

        # Record the interaction
        record_llm_interaction(
            session=session,
            user_id=current_user.id,
            functionality="generate_comparison_topics",
            input_data={
                "description": description,
                "requested_topics": num_topics,
                "comparison_type": comparison_type,
                "has_example_document": len(files) > 0 and files[0].size > 0,
                "search_mode": search_mode,
            },
            output_data={
                "topics_count": len(topics),
                "analysis": analysis,
            },
            metadata={},
        )

        return GenerateTopicsResponse(topics=topics, description_analysis=analysis)

    except Exception as e:
        print(f"Error generating topics: {e}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error generating topics: {str(e)}"
        )


@router.post("/generate-topics-json", response_model=GenerateTopicsResponse)
async def generate_topics_json(
    session: SessionDep, current_user: CurrentUser, request: GenerateTopicsRequest
):
    """
    Generate comparison topics based on a description using LLM, with optional knowledge base reference (JSON version).
    """
    try:
        # Get the default LLM
        llm = get_default_llm(session, current_user)

        # Handle optional description
        description = request.description or ""

        # Prepare variables for the prompt
        prompt_variables = {
            "description": description,
            "comparison_type": request.comparison_type or "general",
            "example_document": "",
            "example_instruction": "",
            "example_analysis_instruction": "",
            "knowledge_base_instruction": "",
            "knowledge_base_content": "",
        }

        # If knowledge base is specified, retrieve content using selected search mode
        if request.knowledge_base_id:
            try:
                from app.services.content_retrieval import (
                    retrieve_knowledge_base_content,
                )

                content, instruction = await retrieve_knowledge_base_content(
                    session=session,
                    current_user=current_user,
                    knowledge_base_id=request.knowledge_base_id,
                    search_mode=request.search_mode,
                    query=description,
                )

                if content:
                    prompt_variables["knowledge_base_content"] = (
                        f"REFERENCE DOCUMENTS FROM KNOWLEDGE BASE:\n{content}"
                    )
                    prompt_variables["knowledge_base_instruction"] = (
                        f"\n12. {instruction} Use them as inspiration for the type of comparison topics and scope, "
                        f"adapting the topics to match the specific requirements in the description. "
                        f"Search mode used: {request.search_mode}"
                    )
                    prompt_variables["example_analysis_instruction"] = (
                        f". Briefly mention how the knowledge base content (using {request.search_mode}) influenced the topic selection"
                    )

            except Exception as e:
                logger.warning(f"Error retrieving from knowledge base: {str(e)}")
                # Continue without knowledge base content rather than failing
                pass

        # Generate topics using the LLM
        topics_response = await invoke_llm_async(
            llm,
            settings.TWINCHECK_GENERATE_TOPICS_PROMPT_TEMPLATE,
            prompt_variables,
        )

        # Parse the response to extract topics and analysis
        topics = []
        analysis = ""

        lines = topics_response.strip().split("\n")
        in_topics_section = False
        in_analysis_section = False

        for line in lines:
            line = line.strip()
            if line.startswith("TOPICS:"):
                in_topics_section = True
                in_analysis_section = False
                continue
            elif line.startswith("ANALYSIS:"):
                in_topics_section = False
                in_analysis_section = True
                continue

            if in_topics_section:
                # Extract topics (numbered list)
                if re.match(r"^\d+\.\s+", line):
                    topic = re.sub(r"^\d+\.\s+", "", line)
                    if topic.strip():
                        topics.append(topic.strip())
            elif in_analysis_section:
                if line:
                    if analysis:
                        analysis += " " + line
                    else:
                        analysis = line

        # If parsing failed, try simpler approach
        if not topics:
            # Split by lines and look for numbered items
            for line in lines:
                line = line.strip()
                if re.match(r"^\d+\.\s+", line):
                    topic = re.sub(r"^\d+\.\s+", "", line)
                    if topic.strip():
                        topics.append(topic.strip())

        # Ensure we have some topics
        if not topics:
            raise HTTPException(
                status_code=500,
                detail="Failed to generate topics from the description. Please try with a more detailed description.",
            )

        # Apply user-specified limit if provided, otherwise use all generated topics
        if request.num_topics:
            topics = topics[: request.num_topics]

        if not analysis:
            search_method = (
                "vector search"
                if request.search_mode == "vector"
                else "full document scan"
            )
            analysis = f"Generated {len(topics)} comparison topics based on the provided description using {search_method}"
            if request.knowledge_base_id:
                analysis += " with knowledge base reference."

        # Record the interaction
        record_llm_interaction(
            session=session,
            user_id=current_user.id,
            functionality="generate_topics",
            input_data={
                "description": request.description,
                "comparison_type": request.comparison_type,
                "requested_topics": request.num_topics,
                "knowledge_base_id": request.knowledge_base_id,
                "search_mode": request.search_mode,
            },
            output_data={
                "topics_count": len(topics),
                "analysis": analysis,
            },
            metadata={},
        )

        return GenerateTopicsResponse(topics=topics, description_analysis=analysis)

    except Exception as e:
        logger.error(f"Error generating topics: {e}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error generating topics: {str(e)}"
        )
