import uuid
import json
import csv
import tempfile
import os
import asyncio
import re
import traceback
import base64
import markdown
from pathlib import Path
from io import BytesIO, StringIO
from datetime import datetime
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from app.models import (
    FormConnectRequest,
    FormConnectResponse,
    FormConnectForm,
    FormConnectDetailResponse,
    GenerateFormFieldsRequest,
    GenerateFormFieldsResponse,
    DocxRequest,
    LlmInteraction,
    Message,
    KnowledgeBase,
    EmbeddingModel,
)
from app.services.llms import (
    get_default_llm,
    invoke_llm,
    invoke_llm_with_image,
    record_llm_interaction,
)
from app.services.translation import translate_text_if_needed, translate
from app.api.deps import CurrentUser, SessionDep
from app.core.config import settings
from app.services.progress_tracker import progress_tracker

from sqlmodel import Session, select
from fastapi import (
    APIRouter,
    UploadFile,
    File,
    Form,
    HTTPException,
    Depends,
    Request as FastAPIRequest,
    Query,
)
from typing import List, Dict, Any, Literal, Optional

from langchain.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI
from langchain.schema import AIMessage
from dotenv import load_dotenv

from app.services.embeddings import load_embeddings_model
from app.services.knowledgebases import get_embedding_model
from app.services.retrievers import create_ensemble_retriever
from app.services.enhanced_retrieval import SmartRetrieverFactory
from tempfile import NamedTemporaryFile

# import fitz  # PyMuPDF - Removed for commercial licensing
from app.services.pdf_utils import load_pdf_with_pypdf

from datetime import datetime

# Load environment variables from .env file
load_dotenv(dotenv_path="c:/miniconda/aibeniq-react/.env", override=False)

# Retrieve the OpenAI API key from the environment
openai_api_key = os.getenv("OPENAI_API_KEY")
# Initialize a flag to track API key status
is_openai_configured = False

if openai_api_key:
    # Set up OpenAI API key if available
    os.environ["OPENAI_API_KEY"] = openai_api_key
    is_openai_configured = True
    print("OpenAI API key configured successfully")
else:
    print(
        "WARNING: OPENAI_API_KEY is not set in environment variables. Some FormConnect features will be limited."
    )

router = APIRouter(prefix="/formconnect", tags=["formconnect"])


@router.post("/process/task")
async def create_process_task():
    """
    Create a new FormConnect processing task for progress tracking.
    Returns a task_id that can be used to track progress.
    """
    # Define stages for FormConnect processing
    # Format: Dict[stage_name, weight] where weights are relative
    stages = {
        "setup": 10,
        "loading": 15,
        "extracting": 60,
        "comparing": 10,
        "finalizing": 5,
    }

    # Initialize progress tracker with stages for FormConnect processing
    task_id = progress_tracker.create_task(
        operation="FormConnect Processing", stages=stages
    )

    print(f"📋 Created FormConnect task with ID: {task_id}")
    return {"task_id": task_id}


@router.get("/progress/{task_id}")
async def get_formconnect_progress(
    task_id: str,
    current_user: CurrentUser,
    request: FastAPIRequest = None,
) -> Any:
    """
    Get the progress of a FormConnect processing task.
    """
    progress = progress_tracker.get_progress(task_id)

    if not progress:
        raise HTTPException(status_code=404, detail="Task not found")

    print(
        f"🔍 PROGRESS DATA: status={progress.get('status')}, percentage={progress.get('percentage')}, current_stage={progress.get('current_stage')}"
    )
    print(f"🔍 PROGRESS MESSAGE: {progress.get('message')}")

    return progress


def generate_template(fields: List[str]) -> Dict[str, str]:
    """
    Generate a JSON template from a list of fields.
    Each field will have a blank value.
    """
    return {field: "" for field in fields}


async def extract_fields_from_digitized_document(
    file: UploadFile,
    template: Dict[str, str],
    llm=None,
    search_mode: str = "full_scan",
    current_user=None,
) -> Dict[str, str]:
    """
    Extract fields from a document using the LLM.
    Supports both full text processing and vector search modes.
    """
    # Read the file content
    content = await file.read()

    if search_mode == "vector":
        # TRUE VECTOR SEARCH IMPLEMENTATION
        return await extract_fields_using_vector_search(file, content, template, llm)
    else:
        # FULL TEXT MODE (existing implementation)
        return await extract_fields_using_full_text(
            content, file.filename, template, llm, current_user
        )


async def extract_fields_using_vector_search(
    file: UploadFile,
    content: bytes,
    template: Dict[str, str],
    llm=None,
    current_user=None,
) -> Dict[str, str]:
    """
    Extract fields using vector search to find relevant document sections.
    Uses a temporary ChromaDB instance to perform semantic search.
    """
    try:
        import tempfile
        import shutil
        from langchain_community.vectorstores import Chroma
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        from app.services.embeddings import load_embeddings_model
        from app.services.knowledgebases import get_embedding_model
        from app.services.retrievers import create_ensemble_retriever
        from app.services.enhanced_retrieval import SmartRetrieverFactory

        print(
            f"🔍 Using enhanced vector search mode for field extraction from {file.filename}"
        )

        # Get embedding model
        from app.api.deps import get_db

        session = next(get_db())
        embedding_info = get_embedding_model(session, current_user)

        if not embedding_info:
            print("❌ No embedding model available, falling back to full text mode")
            return await extract_fields_using_full_text(
                content, file.filename, template, llm, current_user
            )

        print(
            f"Using embedding model: {embedding_info['model_id']} ({embedding_info['provider']})"
        )

        # Extract text and images from the document using enhanced unified processing
        from app.services.document_utils import (
            extract_text_from_file_unified,
            extract_documents_and_images_from_file_unified,
            ensure_documents_for_vector_search,
        )
        from app.services.vision_service import VisionService

        # First try regular text extraction
        text = extract_text_from_file_unified(
            content, file.filename or "unknown", current_user=current_user
        )

        # Also extract images for vision-enhanced processing (for PDFs)
        document_images = []
        vision_enabled = VisionService.is_vision_enabled(llm, current_user)
        file_ext = Path(file.filename or "").suffix.lower()

        if vision_enabled and file_ext == ".pdf":
            try:
                print(
                    f"🖼️ Extracting images from PDF for enhanced vector search: {file.filename}"
                )
                _, document_images = extract_documents_and_images_from_file_unified(
                    content, file.filename or "unknown"
                )
                if document_images:
                    print(
                        f"✅ Found {len(document_images)} images to enhance vector search"
                    )
                else:
                    print(f"ℹ️ No images found in {file.filename}")
            except Exception as e:
                print(f"⚠️ Failed to extract images from {file.filename}: {str(e)}")
                document_images = []

        # If no text, try vision extraction for image-only documents
        if not text.strip():
            print("No text found, checking for images...")
            documents, images = extract_documents_and_images_from_file_unified(
                content, file.filename or "unknown"
            )

            # Check if we have vision capabilities for image-only processing
            if images and VisionService.is_vision_enabled(llm, current_user):
                print(f"Found {len(images)} images, using vision extraction")
                # Use VisionService for image-only extraction
                try:
                    # Convert base64 images to the format expected by VisionService
                    vision_images = []
                    for img_b64 in images:
                        vision_images.append(
                            {
                                "image_data": img_b64,
                                "metadata": {"source": file.filename or "unknown"},
                            }
                        )

                    vision_result = VisionService.safe_vision_analysis(
                        llm=llm,
                        prompt_template=settings.FORMCONNECT_VISION_PROMPT_TEMPLATE,
                        variables={
                            "template_fields": list(template.keys()),
                            "image_count": len(images),
                        },
                        images=vision_images,
                    )

                    # Parse the vision result as JSON if possible
                    import json
                    import re

                    try:
                        # Try to extract JSON from the response
                        json_match = re.search(
                            r"```(?:json)?\s*\n?({.*?})\s*\n?```",
                            vision_result,
                            re.DOTALL | re.IGNORECASE,
                        )
                        if json_match:
                            vision_result = json_match.group(1)

                        extracted_data = json.loads(vision_result)
                        print(
                            f"Successfully extracted data using vision: {extracted_data}"
                        )
                        return extracted_data

                    except (json.JSONDecodeError, AttributeError) as e:
                        print(f"Could not parse vision result as JSON: {e}")
                        # Fall back to text processing with the vision result
                        text = vision_result

                except Exception as e:
                    print(f"Vision extraction failed: {e}")
                    return {
                        k: "Could not extract: Document contains only images and vision extraction failed"
                        for k in template.keys()
                    }
            else:
                return {
                    k: "Could not extract: Empty document or no vision support available"
                    for k in template.keys()
                }

        # Create temporary directory for ChromaDB
        temp_dir = tempfile.mkdtemp()

        try:
            # Load embedding model
            embeddings = load_embeddings_model(
                provider=embedding_info["provider"], model_id=embedding_info["model_id"]
            )

            # Split text into chunks for vector storage
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=settings.RAG_DOCUMENT_CHUNK_SIZE,
                chunk_overlap=settings.RAG_DOCUMENT_CHUNK_OVERLAP,
                length_function=len,
            )
            chunks = text_splitter.split_text(text)

            # Ensure we have at least one chunk for vector search
            if not chunks:
                chunks = [text or f"Fallback content for {file.filename or 'document'}"]

            print(f"📄 Split document into {len(chunks)} chunks for vector search")

            # Create ChromaDB vector store with error handling
            try:
                chroma_db = Chroma.from_texts(
                    texts=chunks,
                    embedding=embeddings,
                    persist_directory=temp_dir,
                    metadatas=[
                        {"source": file.filename, "chunk_id": i}
                        for i in range(len(chunks))
                    ],
                )
            except Exception as e:
                print(f"Error creating vector store: {e}")
                # Fall back to full text extraction if vector store fails
                return await extract_fields_using_full_text(
                    content, file.filename, template, llm, current_user
                )

            # Create retriever for semantic search
            retriever = create_ensemble_retriever(
                chroma_db=chroma_db,
                vector_weight=0.8,  # Higher weight for vector search in FormConnect
                keyword_weight=0.2,
                search_kwargs={"k": settings.FORMCONNECT_VECTOR_SEARCH_CHUNKS},
            )

            # Extract fields using vector search
            extracted_data = {}
            field_count = len(template)

            print(f"🔎 Extracting {field_count} fields using vector search...")

            for i, (field_name, field_description) in enumerate(template.items(), 1):
                print(f"[{i}/{field_count}] Searching for: {field_name}")

                # Create search query combining field name and description
                search_query = f"{field_name} {field_description}".strip()

                try:
                    # Retrieve relevant chunks
                    relevant_docs = retriever.invoke(search_query)

                    if relevant_docs:
                        # Combine relevant text from retrieved chunks
                        relevant_chunks = []
                        total_tokens = 0

                        for doc in relevant_docs:
                            chunk_text = doc.page_content
                            chunk_tokens = count_tokens(chunk_text)

                            # Limit total context to avoid token limits
                            if (
                                total_tokens + chunk_tokens > 10000
                            ):  # Conservative limit
                                break

                            relevant_chunks.append(chunk_text)
                            total_tokens += chunk_tokens

                        relevant_text = "\n\n".join(relevant_chunks)

                        if relevant_text.strip():
                            # Create field-specific extraction prompt
                            field_prompt = f"""Extract the value for "{field_name}" from the following text.

Field description: {field_description}

Relevant text:
{relevant_text}

Instructions:
1. Look for the specific information related to "{field_name}"
2. If found, return only the extracted value
3. If not found, return "Not found"
4. Be precise and concise

Extracted value:"""

                            # Extract field value using LLM
                            field_response = invoke_llm(llm, field_prompt, {})
                            extracted_value = (
                                field_response.content
                                if hasattr(field_response, "content")
                                else str(field_response)
                            )

                            # Check if we got a good result or should try vision
                            if (
                                extracted_value.strip().lower()
                                in ["not found", "n/a", "none", "null", ""]
                                and vision_enabled
                                and document_images
                            ):
                                print(
                                    f"   🔍 Text search failed for {field_name}, trying vision analysis..."
                                )
                                vision_value = await extract_field_from_images(
                                    field_name,
                                    field_description,
                                    document_images,
                                    llm,
                                    file.filename,
                                )
                                if (
                                    vision_value
                                    and vision_value.strip().lower()
                                    not in ["not found", "n/a", "none", "null"]
                                ):
                                    extracted_data[field_name] = vision_value.strip()
                                    print(
                                        f"   ✅ Vision found: {vision_value.strip()[:50]}..."
                                    )
                                else:
                                    extracted_data[field_name] = extracted_value.strip()
                                    print(
                                        f"   ⚠️ Both text and vision failed for {field_name}"
                                    )
                            else:
                                extracted_data[field_name] = extracted_value.strip()
                                print(f"   ✅ Found: {extracted_value.strip()[:50]}...")
                        else:
                            # No relevant text found, try vision if available
                            if vision_enabled and document_images:
                                print(
                                    f"   🔍 No relevant text for {field_name}, trying vision analysis..."
                                )
                                vision_value = await extract_field_from_images(
                                    field_name,
                                    field_description,
                                    document_images,
                                    llm,
                                    file.filename,
                                )
                                if (
                                    vision_value
                                    and vision_value.strip().lower()
                                    not in ["not found", "n/a", "none", "null"]
                                ):
                                    extracted_data[field_name] = vision_value.strip()
                                    print(
                                        f"   ✅ Vision found: {vision_value.strip()[:50]}..."
                                    )
                                else:
                                    extracted_data[field_name] = "Not found"
                                    print(
                                        f"   ❌ No relevant content found (text or vision)"
                                    )
                            else:
                                extracted_data[field_name] = "Not found"
                                print(f"   ❌ No relevant content found")
                    else:
                        # No search results, try vision if available
                        if vision_enabled and document_images:
                            print(
                                f"   🔍 No search results for {field_name}, trying vision analysis..."
                            )
                            vision_value = await extract_field_from_images(
                                field_name,
                                field_description,
                                document_images,
                                llm,
                                file.filename,
                            )
                            if vision_value and vision_value.strip().lower() not in [
                                "not found",
                                "n/a",
                                "none",
                                "null",
                            ]:
                                extracted_data[field_name] = vision_value.strip()
                                print(
                                    f"   ✅ Vision found: {vision_value.strip()[:50]}..."
                                )
                            else:
                                extracted_data[field_name] = "Not found"
                                print(f"   ❌ No search results (text or vision)")
                        else:
                            extracted_data[field_name] = "Not found"
                            print(f"   ❌ No search results")

                except Exception as e:
                    print(f"   ❌ Error searching for {field_name}: {str(e)}")
                    extracted_data[field_name] = f"Error: {str(e)}"

            print("✅ Vector search extraction completed")
            return extracted_data

        finally:
            # Cleanup temporary directory
            try:
                shutil.rmtree(temp_dir)
                print(f"🧹 Cleaned up temporary directory: {temp_dir}")
            except Exception as e:
                print(f"Warning: Could not cleanup temporary directory: {str(e)}")

    except Exception as e:
        print(f"❌ Vector search failed: {str(e)}. Falling back to full text mode.")
        return await extract_fields_using_full_text(
            content, file.filename, template, llm, current_user
        )


async def extract_fields_using_full_text(
    content: bytes, filename: str, template: Dict[str, str], llm=None, current_user=None
) -> Dict[str, str]:
    """
    Extract fields using full text processing with enhanced visual processing.
    Now includes image extraction and vision analysis for PDFs with embedded images.
    """
    print(f"📄 Using full text mode for field extraction from {filename}")

    # Check file extension to determine processing method
    file_ext = Path(filename).suffix.lower() if filename else ""

    # Import vision service for image processing
    from app.services.vision_service import VisionService

    # Check if vision is enabled for this LLM
    vision_enabled = VisionService.is_vision_enabled(llm, current_user)
    print(f"🔍 Vision processing enabled: {vision_enabled}")

    try:
        # Use unified document processing for all file types
        from app.services.document_utils import (
            extract_text_from_file_unified,
            extract_documents_and_images_from_file_unified,
        )

        text = extract_text_from_file_unified(
            content, filename, current_user=current_user
        )
        document_images = []

        # Extract images if vision is enabled and we're processing a PDF
        if vision_enabled and file_ext == ".pdf":
            try:
                print(f"🖼️ Extracting images from PDF: {filename}")
                _, document_images = extract_documents_and_images_from_file_unified(
                    content, filename
                )
                if document_images:
                    print(f"✅ Extracted {len(document_images)} images from {filename}")
                else:
                    print(f"ℹ️ No images found in {filename}")
            except Exception as e:
                print(f"⚠️ Failed to extract images from {filename}: {str(e)}")
                document_images = []

        # Handle case where document has no text but has images
        if not text.strip() and document_images and vision_enabled:
            print(f"📷 Processing image-only document: {filename}")
            return await process_images_only(document_images, template, llm, filename)

        # If no text and no images, return error
        elif not text.strip():
            return {
                k: f"Could not extract: Empty document {filename}"
                for k in template.keys()
            }

        # Check token count and implement chunking if needed
        token_count = count_tokens(text + json.dumps(template))

        if token_count > 150000:  # Token limit safety
            print(
                f"⚠️ Large document detected ({token_count} tokens), implementing chunking..."
            )
            return await extract_fields_with_chunking(text, template, llm, current_user)

    except Exception as e:
        print(f"Error processing file {filename}: {str(e)}")
        return {
            k: f"Could not extract: Error processing {filename} - {str(e)}"
            for k in template.keys()
        }

    # Process with or without vision depending on available images
    if vision_enabled and document_images:
        print(
            f"🔍 Processing {filename} with both text and {len(document_images)} images"
        )
        # Use vision-enhanced processing
        response = await process_with_text_and_images(
            text, document_images, template, llm, filename, current_user
        )
    else:
        print(f"📄 Processing {filename} with text only (no images or vision disabled)")
        # Use standard text-only processing
        # Get user language and create language instruction
        user_language = (
            getattr(current_user, "preferred_language", "en") or "en"
            if current_user
            else "en"
        )
        language_name = settings.SUPPORTED_LANGUAGES.get(user_language, "English")
        language_instruction = f"Respond in this language: {language_name}."

        prompt_template = settings.FORMCONNECT_DIGITIZED_PROMPT_TEMPLATE
        variables = {
            "template": json.dumps(template),
            "document_text": text,
            "language_instruction": language_instruction,
        }
        response = invoke_llm(llm, prompt_template, variables)

    # Try to parse JSON from the response
    try:
        # The output might already be a dictionary
        if isinstance(response, dict):
            return response
        content_dict = json.loads(response)
        return content_dict
    except Exception:
        return {"raw_content": str(response)}


def count_tokens(text: str, model: str = "gpt-4o-mini") -> int:
    """Count tokens in text for the given model."""
    try:
        import tiktoken

        encoding = tiktoken.encoding_for_model(model)
        return len(encoding.encode(text))
    except KeyError:
        # Fallback to cl100k_base for unknown models
        import tiktoken

        encoding = tiktoken.get_encoding("cl100k_base")
        return len(encoding.encode(text))


async def process_images_only(
    document_images: List[str], template: Dict[str, str], llm, filename: str
) -> Dict[str, str]:
    """
    Process a document that contains only images (no extractable text).
    Uses vision analysis to extract form fields from the images.
    """
    try:
        from app.services.vision_service import VisionService

        # Convert base64 images to the format expected by VisionService
        vision_images = []
        for i, img_b64 in enumerate(document_images):
            vision_images.append(
                {
                    "image_data": img_b64,
                    "metadata": {"source": filename, "page": i + 1},
                }
            )

        vision_result = VisionService.safe_vision_analysis(
            llm=llm,
            prompt_template=settings.FORMCONNECT_VISION_PROMPT_TEMPLATE,
            variables={
                "template_fields": list(template.keys()),
                "image_count": len(document_images),
                "filename": filename,
            },
            images=vision_images,
        )

        # Parse the vision result as JSON if possible
        import json
        import re

        try:
            # Try to extract JSON from the response
            json_match = re.search(
                r"```(?:json)?\s*\n?({.*?})\s*\n?```",
                vision_result,
                re.DOTALL | re.IGNORECASE,
            )
            if json_match:
                vision_result = json_match.group(1)

            extracted_data = json.loads(vision_result)
            print(
                f"✅ Successfully extracted data using vision from {filename}: {extracted_data}"
            )
            return extracted_data

        except (json.JSONDecodeError, AttributeError) as e:
            print(f"⚠️ Could not parse vision result as JSON: {e}")
            # Return structured response indicating vision processing was attempted
            return {
                k: f"Vision analysis completed but data format unclear: {vision_result[:100]}..."
                for k in template.keys()
            }

    except Exception as e:
        print(f"❌ Vision processing failed for {filename}: {e}")
        return {
            k: f"Could not extract: Vision processing failed - {str(e)}"
            for k in template.keys()
        }


async def process_with_text_and_images(
    text: str,
    document_images: List[str],
    template: Dict[str, str],
    llm,
    filename: str,
    current_user=None,
) -> Dict[str, str]:
    """
    Process a document that has both text and images using enhanced vision analysis.
    Combines text extraction with vision analysis for comprehensive field extraction.
    """
    try:
        from app.services.vision_service import VisionService

        print(f"🔄 Processing {filename} with combined text and vision analysis")

        # Convert base64 images to the format expected by VisionService
        vision_images = []
        for i, img_b64 in enumerate(document_images):
            vision_images.append(
                {
                    "image_data": img_b64,
                    "metadata": {"source": filename, "page": i + 1},
                }
            )

        # Use a specialized prompt that combines text and vision analysis
        prompt_template = settings.FORMCONNECT_VISION_PROMPT_TEMPLATE
        vision_variables = {
            "template_fields": list(template.keys()),
            "document_text": text[:10000],  # Limit text to avoid token overflow
            "image_count": len(document_images),
            "filename": filename,
            "has_text": True,
        }

        vision_result = VisionService.safe_vision_analysis(
            llm=llm,
            prompt_template=prompt_template,
            variables=vision_variables,
            images=vision_images,
        )

        # Parse the combined result
        import json
        import re

        try:
            # Try to extract JSON from the response
            json_match = re.search(
                r"```(?:json)?\s*\n?({.*?})\s*\n?```",
                vision_result,
                re.DOTALL | re.IGNORECASE,
            )
            if json_match:
                vision_result = json_match.group(1)

            extracted_data = json.loads(vision_result)
            print(
                f"✅ Successfully extracted data using combined text+vision from {filename}"
            )
            return extracted_data

        except (json.JSONDecodeError, AttributeError) as e:
            print(
                f"⚠️ JSON parsing failed for combined analysis, falling back to text-only"
            )
            # Fallback to text-only processing
            # Get user language and create language instruction
            user_language = (
                getattr(current_user, "preferred_language", "en") or "en"
                if current_user
                else "en"
            )
            language_name = settings.SUPPORTED_LANGUAGES.get(user_language, "English")
            language_instruction = f"Respond in this language: {language_name}."

            prompt_template = settings.FORMCONNECT_DIGITIZED_PROMPT_TEMPLATE
            variables = {
                "template": json.dumps(template),
                "document_text": text,
                "language_instruction": language_instruction,
            }
            response = invoke_llm(llm, prompt_template, variables)

            if isinstance(response, dict):
                return response
            try:
                return json.loads(response)
            except:
                return {"raw_content": str(response)}

    except Exception as e:
        print(
            f"⚠️ Combined processing failed for {filename}, falling back to text-only: {e}"
        )
        # Fallback to text-only processing
        prompt_template = settings.FORMCONNECT_DIGITIZED_PROMPT_TEMPLATE
        variables = {"template": json.dumps(template), "document_text": text}
        response = invoke_llm(llm, prompt_template, variables)

        if isinstance(response, dict):
            return response
        try:
            return json.loads(response)
        except:
            return {"raw_content": str(response)}


async def extract_field_from_images(
    field_name: str,
    field_description: str,
    document_images: List[str],
    llm,
    filename: str,
) -> str:
    """
    Extract a specific field value from document images using vision analysis.
    Used as a fallback when vector/text search doesn't find the field.
    """
    try:
        from app.services.vision_service import VisionService

        # Convert base64 images to the format expected by VisionService
        vision_images = []
        for i, img_b64 in enumerate(document_images):
            vision_images.append(
                {
                    "image_data": img_b64,
                    "metadata": {"source": filename, "page": i + 1},
                }
            )

        # Create a focused prompt for this specific field
        field_specific_prompt = f"""You are analyzing document images to extract a specific piece of information.

FIELD TO FIND: "{field_name}"
FIELD DESCRIPTION: {field_description}

Please look through all the provided images and find the value for "{field_name}".

Instructions:
1. Carefully examine each image for information related to "{field_name}"
2. If you find the information, return ONLY the value (no explanations)
3. If you cannot find the information, return "Not found"
4. Be precise and extract only the specific value requested

The images are from: {filename}

Value for "{field_name}":"""

        vision_result = VisionService.safe_vision_analysis(
            llm=llm,
            prompt_template=field_specific_prompt,
            variables={},
            images=vision_images,
        )

        # Clean up the response
        if isinstance(vision_result, str):
            result = vision_result.strip()
            # Remove common wrapper text that might appear
            if result.lower().startswith("the value"):
                # Try to extract just the value part
                import re

                value_match = re.search(r"value.*?is:?\s*(.+?)(?:\.|$)", result.lower())
                if value_match:
                    result = value_match.group(1).strip()
            return result
        else:
            return str(vision_result).strip()

    except Exception as e:
        print(f"❌ Vision field extraction failed for {field_name}: {e}")
        return "Not found"


async def extract_fields_with_chunking(
    document_text: str, template: Dict[str, str], llm=None, current_user=None
) -> Dict[str, str]:
    """
    Extract fields from large documents using chunking.
    """
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    print("🔄 Processing large document with chunking...")

    # Create chunks
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=50000,  # Conservative chunk size
        chunk_overlap=200,
        length_function=len,
    )

    chunks = splitter.split_text(document_text)
    print(f"📄 Split document into {len(chunks)} chunks")

    all_extractions = []

    # Process each chunk
    for i, chunk in enumerate(chunks):
        print(f"[{i+1}/{len(chunks)}] Processing chunk ({count_tokens(chunk)} tokens)")

        try:
            # Get user language and create language instruction
            user_language = (
                getattr(current_user, "preferred_language", "en") or "en"
                if current_user
                else "en"
            )
            language_name = settings.SUPPORTED_LANGUAGES.get(user_language, "English")
            language_instruction = f"Respond in this language: {language_name}."

            prompt_template = settings.FORMCONNECT_DIGITIZED_PROMPT_TEMPLATE
            variables = {
                "template": json.dumps(template),
                "document_text": chunk,
                "language_instruction": language_instruction,
            }
            response = invoke_llm(llm, prompt_template, variables)

            try:
                if isinstance(response, dict):
                    chunk_extraction = response
                else:
                    chunk_extraction = json.loads(response)
                all_extractions.append(chunk_extraction)
            except json.JSONDecodeError as e:
                print(f"Error parsing JSON from chunk {i+1}: {e}")
                continue

        except Exception as e:
            print(f"Error processing chunk {i+1}: {str(e)}")
            continue

    # Merge extractions from all chunks
    return merge_field_extractions(all_extractions, template)


def merge_field_extractions(
    extractions: list, template: Dict[str, str]
) -> Dict[str, str]:
    """
    Merge field extractions from multiple document chunks.
    """
    if not extractions:
        return {k: "Not found" for k in template.keys()}

    if len(extractions) == 1:
        return extractions[0]

    # Start with the template structure
    merged = {k: "Not found" for k in template.keys()}

    # For each field, take the first non-empty value found
    for extraction in extractions:
        for field, value in extraction.items():
            if field in merged and value and str(value).strip():
                # Only update if current value is empty/placeholder
                current_value = str(merged[field]).strip()
                if current_value.lower() in ["not found", "", "n/a", "null", "none"]:
                    merged[field] = value

    return merged


async def convert_pdf_to_images(content: bytes) -> List[str]:
    """
    Convert PDF pages to base64-encoded images.
    """
    try:
        import fitz  # PyMuPDF - optional dependency

        # Open PDF from bytes
        pdf_doc = fitz.open(stream=content, filetype="pdf")
        page_images = []

        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            # Render page as image (higher DPI for better OCR)
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))  # 2x scaling
            img_data = pix.tobytes("png")
            img_base64 = base64.b64encode(img_data).decode("utf-8")
            page_images.append(img_base64)

        pdf_doc.close()
        print(f"✅ Successfully converted PDF to {len(page_images)} images")
        return page_images

    except ImportError:
        print("❌ PyMuPDF (fitz) not available - cannot convert PDF to images")
        print(
            "💡 To fix this: Install PyMuPDF with 'pip install PyMuPDF' in the container"
        )
        return []
    except Exception as e:
        print(f"❌ Error converting PDF to images: {str(e)}")
        # Try to detect specific error types for better user guidance
        error_str = str(e).lower()
        if "password" in error_str or "encrypted" in error_str:
            print("💡 PDF appears to be password protected or encrypted")
        elif "corrupt" in error_str or "invalid" in error_str:
            print("💡 PDF file may be corrupted or invalid")
        else:
            print("💡 PDF conversion failed - try uploading as separate image files")
        return []


async def convert_docx_to_images(content: bytes, file_ext: str) -> List[str]:
    """
    Convert DOCX/DOC pages to base64-encoded images.
    Note: This is a placeholder implementation. In practice, you might need
    additional libraries like python-docx2pdf + pdf2image or similar.
    """
    try:
        # For now, fall back to text extraction and create a simple image
        # In a full implementation, you'd use libraries to convert DOCX to images
        print(f"⚠️ DOCX/DOC to image conversion not fully implemented for {file_ext}")

        return []  # Return empty to trigger fallback

    except Exception as e:
        print(f"❌ Error converting DOCX to images: {str(e)}")
        return []


def merge_page_extractions(
    extractions: List[Dict[str, str]], template: Dict[str, str]
) -> Dict[str, str]:
    """
    Merge field extractions from multiple pages, prioritizing non-empty values.
    """
    print(f"🔀 Merging extractions from {len(extractions)} page(s)")
    for i, extraction in enumerate(extractions):
        print(f"📄 Page {i+1} extraction: {extraction}")

    merged = {k: "" for k in template.keys()}

    for extraction in extractions:
        # Handle raw_content responses (unparsed JSON)
        if "raw_content" in extraction and len(extraction) == 1:
            print(
                f"⚠️ Found raw_content, attempting to parse: {extraction['raw_content'][:100]}..."
            )
            try:
                import json
                import re

                raw_content = extraction["raw_content"]

                # Remove markdown code block wrappers if present
                if "```json" in raw_content or "```" in raw_content:
                    json_match = re.search(
                        r"```(?:json)?\s*\n?({.*?})\s*\n?```", raw_content, re.DOTALL
                    )
                    if json_match:
                        raw_content = json_match.group(1)
                        print(
                            f"🎯 Extracted JSON from raw_content: {raw_content[:100]}..."
                        )

                parsed_content = json.loads(raw_content.strip())
                print(f"✅ Successfully parsed raw_content: {parsed_content}")
                extraction = parsed_content

            except Exception as e:
                print(f"❌ Failed to parse raw_content: {str(e)}")
                continue

        # Merge field values
        for field, value in extraction.items():
            if field in merged and value and str(value).strip():
                value_str = str(value).strip()

                # Skip obviously bad values
                if value_str.lower() in [
                    "not found",
                    "n/a",
                    "null",
                    "none",
                    "",
                ] or value_str.startswith("Error"):
                    continue

                # If we don't have a value yet, or current value is better
                current_value = str(merged[field]).strip()
                if (
                    not current_value
                    or current_value.lower() in ["not found", "n/a", "null", "none"]
                    or current_value.startswith("Error")
                    or len(value_str) > len(current_value)
                ):
                    merged[field] = value_str
                    print(f"✅ Updated field '{field}': '{value_str}'")

    # Mark fields not found across all pages
    for field in merged:
        if not merged[field] or not str(merged[field]).strip():
            merged[field] = "Not found in document"
            print(f"❌ Field '{field}' not found in any page")

    print(f"🎯 Final merged result: {merged}")
    return merged


async def format_single_document_result(
    extracted_data: Dict[str, str], file_name: str, llm, current_user=None
) -> str:
    """
    Format the results from a single document into a clear presentation.
    """
    # Use original filename
    clean_filename = file_name

    # Convert dict to string, escaping any curly braces for the formatter
    data_str = str(extracted_data).replace("{", "{{").replace("}", "}}")

    # Use the single document template from config
    # Get user language and create language instruction
    user_language = (
        getattr(current_user, "preferred_language", "en") or "en"
        if current_user
        else "en"
    )
    language_name = settings.SUPPORTED_LANGUAGES.get(user_language, "English")
    language_instruction = f"Respond in this language: {language_name}."

    prompt_template = settings.FORMCONNECT_SINGLE_DOCUMENT_PROMPT_TEMPLATE
    variables = {
        "document_name": clean_filename,
        "extracted_data": data_str,
        "language_instruction": language_instruction,
    }

    print(f"Formatting single document result for: {clean_filename}")

    # Invoke the LLM to format the results
    response = invoke_llm(llm, prompt_template, variables)

    return response


async def compare_multiple_documents(
    documents: List[Dict[str, str]], file_names: List[str], llm, current_user, session
) -> str:
    """
    Compare fields across multiple documents using the LLM.
    """
    # Create a combined representation of all documents WITH ACTUAL FILENAMES
    documents_str = ""
    clean_filenames = []

    for i, (doc, name) in enumerate(zip(documents, file_names)):
        # Use original filename
        clean_filename = name
        clean_filenames.append(clean_filename)

        # Convert dict to string, escaping any curly braces for the formatter
        doc_str = str(doc).replace("{", "{{").replace("}", "}}")
        documents_str += f"Document: {clean_filename}\nExtracted Data: {doc_str}\n\n"

    print(
        "documents_str for comparison:", documents_str[:500]
    )  # Print first 500 chars for debugging

    # Create an enhanced prompt that explicitly instructs the LLM to use actual filenames
    enhanced_prompt_template = """Compare the extracted fields across the following documents and provide a detailed analysis.

IMPORTANT: When referring to documents in your analysis and tables, use the actual document filenames provided below, NOT generic labels like "Document 1", "Document 2", etc.

Document Filenames:
{filename_list}

Documents to compare:
{documents_str}

Instructions:
1. Create a comparison table showing field values across all documents
2. Use the actual document filenames as column headers in any tables
3. Identify discrepancies and highlight the most likely correct values
4. Provide a summary of findings
5. If creating markdown tables, use the document filenames as column headers

Format your response in markdown with clear tables and analysis."""

    variables = {
        "documents_str": documents_str,
        "filename_list": "\n".join([f"- {name}" for name in clean_filenames]),
    }
    response = invoke_llm(llm, enhanced_prompt_template, variables)
    # Translate the response if needed
    # translated_response = await translate_text_if_needed(
    #     response, session, current_user, llm
    # )
    return response


@router.post("/process", response_model=FormConnectResponse)
async def process_form(
    session: SessionDep,
    current_user: CurrentUser,
    fields: str = Form(...),
    search_mode: Literal["vector", "full_scan"] = Form("vector"),
    form_name: Optional[str] = Form(None),
    digitized_files: List[UploadFile] = File(None),
    handwritten_files: List[UploadFile] = File(None),
    request: FastAPIRequest = None,
    task_id: Optional[str] = Form(None),
):
    """
    Process the uploaded files and fields with unified visual processing.
    All files now benefit from automatic visual enhancement for embedded images.

    Args:
        task_id: Optional task ID for progress tracking
    """
    print("process_form function invoked!")
    print(f"Received search_mode: {search_mode}")
    print(f"Received task_id: {task_id}")
    print(f"Request data: fields={fields[:50]}...")

    # Get the default LLM
    llm = get_default_llm(session, current_user)

    # Log the model type being used
    if hasattr(llm, "__class__") and "ReplicateWrapper" in llm.__class__.__name__:
        print(
            f"Using Replicate model for FormConnect: {getattr(llm, 'model_id', 'unknown')}"
        )
    else:
        print(f"Using LangChain model for FormConnect: {type(llm).__name__}")

    # Combine digitized and handwritten files
    files = []
    if digitized_files:
        files.extend(digitized_files)
    if handwritten_files:
        files.extend(handwritten_files)

    total_files = len(files)
    print(f"Now processing {total_files} files with unified visual processing...")
    print(f"  - Digitized files: {len(digitized_files) if digitized_files else 0}")
    print(
        f"  - Handwritten files: {len(handwritten_files) if handwritten_files else 0}"
    )

    # Check if we have at least one file
    if total_files < 1:
        raise HTTPException(
            status_code=400, detail="At least one file must be uploaded."
        )
    # Parse the fields into a list
    field_list = fields.splitlines()

    if not field_list:
        raise HTTPException(status_code=400, detail="No fields provided.")

    # Update progress: Setup complete
    if task_id:
        progress_tracker.complete_stage(
            task_id, "setup", message_key="common.progress.initializing"
        )
        progress_tracker.update_stage_progress(
            task_id, "loading", 0, 1, message_key="common.progress.processing"
        )
        await asyncio.sleep(0.01)  # Yield to event loop

    # Generate the JSON template
    template = generate_template(field_list)

    # Update progress: Loading complete
    if task_id:
        progress_tracker.complete_stage(
            task_id, "loading", message_key="common.progress.processing"
        )
        progress_tracker.update_stage_progress(
            task_id,
            "extracting",
            0,
            total_files,
            message_key="common.progress.extracting",
        )
        await asyncio.sleep(0.01)  # Yield to event loop

    # Extract fields from all documents using unified visual processing
    extracted_results = []
    file_names = []

    # Process all files with unified visual enhancement
    if files:
        for i, file in enumerate(files):
            # Update progress for this file
            if task_id:
                progress_tracker.update_stage_progress(
                    task_id,
                    "extracting",
                    i,
                    total_files,
                    message_key="common.progress.extracting",
                )
                await asyncio.sleep(0.01)  # Yield to event loop

            # Add delay between file processing to prevent rate limit exhaustion
            if i > 0:
                await asyncio.sleep(settings.PROCESSING_DELAY_BETWEEN_DOCUMENTS)

            # CRITICAL: Check if client has disconnected before processing each file
            try:
                if request and await request.is_disconnected():
                    print(f"❌ CLIENT DISCONNECTED - Stopping at file {i + 1}")
                    return FormConnectResponse(
                        results={
                            "status": "cancelled",
                            "message": "Request cancelled - client disconnected",
                        }
                    )
            except Exception as e:
                print(f"Warning: Could not check disconnect status: {e}")

            # Read file content
            file_content = await file.read()
            filename = file.filename

            if search_mode == "vector":
                # Use vector search with visual enhancement
                extracted = await extract_fields_using_vector_search(
                    file, file_content, template, llm, current_user
                )
            else:
                # Use full text processing with visual enhancement
                extracted = await extract_fields_using_full_text(
                    file_content, filename, template, llm, current_user
                )

            # CRITICAL: Check if client disconnected after field extraction
            try:
                if request and await request.is_disconnected():
                    print(f"❌ CLIENT DISCONNECTED - After processing file {i + 1}")
                    return FormConnectResponse(
                        results={
                            "status": "cancelled",
                            "message": "Request cancelled - client disconnected after file processing",
                        }
                    )
            except Exception as e:
                print(f"Warning: Could not check disconnect status: {e}")

            print("Results for file name:", filename)
            print("Extracted fields:", extracted)
            extracted_results.append(extracted)
            file_names.append(filename)

            # Reset file position for potential future reads
            await file.seek(0)

    # Update progress: Extraction complete, starting comparison
    if task_id:
        progress_tracker.complete_stage(
            task_id, "extracting", message_key="common.progress.extracting"
        )
        progress_tracker.update_stage_progress(
            task_id, "comparing", 0, 1, message_key="match.progress.formatting"
        )
        await asyncio.sleep(0.01)  # Yield to event loop

    # If there's only one file, format the results nicely instead of just showing raw data
    if total_files == 1:
        # Format the single document result for better presentation
        formatted_result = await format_single_document_result(
            extracted_results[0], file_names[0], llm, current_user
        )
        result = {
            "message_key": "match.singleDocumentSuccess",
            "comparison": formatted_result,
            "extracted_data": extracted_results[0],
        }
    else:
        # Compare the extracted fields
        comparison_result = await compare_multiple_documents(
            extracted_results, file_names, llm, current_user, session
        )
        result = {
            "message_key": "match.matchSuccess",
            "comparison": comparison_result,
            "extracted_data": extracted_results,
        }

    # Update progress: Comparison complete, finalizing
    if task_id:
        progress_tracker.complete_stage(
            task_id, "comparing", message_key="match.progress.formatting"
        )
        progress_tracker.update_stage_progress(
            task_id, "finalizing", 0, 1, message_key="common.progress.processing"
        )
        await asyncio.sleep(0.01)  # Yield to event loop

    interaction_id = record_llm_interaction(
        session=session,
        user_id=current_user.id,
        functionality="formconnect",
        input_data={
            "fields": fields,
            "files": file_names,
            "search_mode": search_mode,
            "form_name": form_name,
        },
        output_data=result,
        metadata={
            "file_count": total_files,
            "field_count": len(field_list),
            "document_count": total_files,
            "digitized_files": (
                [f.filename for f in digitized_files] if digitized_files else []
            ),
            "handwritten_files": (
                [f.filename for f in handwritten_files] if handwritten_files else []
            ),
            "fields": field_list,
            "search_mode": search_mode,
        },
    )

    print(f"[DEBUG] FormConnect interaction_id returned: {interaction_id}")
    # Add interaction_id to the result
    result["interaction_id"] = str(interaction_id) if interaction_id else None
    print(
        f"[DEBUG] FormConnect result with interaction_id: {result.get('interaction_id')}"
    )

    # Update progress: Complete
    if task_id:
        progress_tracker.complete_stage(task_id, "finalizing", "Processing complete")
        # Mark the entire task as complete using the ProgressTracker API
        progress_tracker.complete_task(
            task_id, "Form processing completed successfully"
        )
        await asyncio.sleep(0.01)  # Yield to event loop

    # Return the comparison results as a dictionary
    return FormConnectResponse(results=result)


# Functions related to Forms
@router.post("/forms", response_model=FormConnectForm)
def create_form(
    form: FormConnectForm,
    session: SessionDep,
    current_user: CurrentUser,
):
    """
    Save a new form to the database.
    """
    existing_form = session.exec(
        select(FormConnectForm).where(FormConnectForm.name == form.name)
    ).first()
    if existing_form:
        raise HTTPException(
            status_code=400, detail="A form with this name already exists."
        )

    form.owner_id = current_user.id
    session.add(form)
    session.commit()
    session.refresh(form)
    return form


@router.get("/forms", response_model=List[FormConnectForm])
def get_forms(session: SessionDep, current_user: CurrentUser):
    """
    Retrieve all forms from the database for this user.
    """
    return session.exec(
        select(FormConnectForm).where(FormConnectForm.owner_id == current_user.id)
    ).all()


@router.get("/forms/{form_id}", response_model=FormConnectForm)
def get_form(form_id: uuid.UUID, session: SessionDep):
    """
    Retrieve a specific form by ID.
    """
    form = session.get(FormConnectForm, form_id)
    if not form:
        raise HTTPException(status_code=404, detail="Form not found.")
    return form


@router.put("/forms/{form_id}", response_model=FormConnectForm)
def update_form(
    form_id: uuid.UUID,
    updated_form: FormConnectForm,
    session: SessionDep,
    current_user: CurrentUser,
):
    """
    Update an existing form.
    """
    form = session.get(FormConnectForm, form_id)
    if not form:
        raise HTTPException(status_code=404, detail="Form not found.")

    # Ensure the current user is the owner of the form
    if form.owner_id != current_user.id:
        raise HTTPException(
            status_code=403, detail="Not authorized to update this form."
        )

    form.name = updated_form.name
    form.description = updated_form.description
    form.fields = updated_form.fields
    form.date_modified = datetime.utcnow()

    session.add(form)
    session.commit()
    session.refresh(form)
    return form


@router.delete("/forms/{form_id}", response_model=Message)
def delete_form(form_id: uuid.UUID, session: SessionDep, current_user: CurrentUser):
    """
    Delete a form by ID.
    """
    form = session.get(FormConnectForm, form_id)
    if not form:
        raise HTTPException(status_code=404, detail="Form not found.")

    # Ensure the current user is the owner of the form
    if form.owner_id != current_user.id:
        raise HTTPException(
            status_code=403, detail="Not authorized to delete this form."
        )

    session.delete(form)
    session.commit()
    return Message(message="Form deleted successfully.")


# Add this new endpoint to get history details for a specific form processing
@router.get("/history/{interaction_id}", response_model=FormConnectDetailResponse)
async def get_form_detail(
    interaction_id: uuid.UUID,
    session: SessionDep,
    current_user: CurrentUser,
):
    """Retrieve a specific form processing's full content by ID."""
    print("Received interaction ID:", interaction_id)
    try:
        report = session.get(LlmInteraction, interaction_id)
        if not report:
            raise HTTPException(
                status_code=404, detail="Form processing result not found"
            )

        # No longer need to check this as we now allow viewing other users' outputs
        # if report.user_id != current_user.id:
        #    raise HTTPException(
        #        status_code=403, detail="You don't have access to this form processing"
        #    )

        if report.functionality != "formconnect":
            raise HTTPException(
                status_code=400, detail="This is not a FormConnect processing"
            )

        # Try to reconstruct the original form processing structure
        try:
            input_data = json.loads(report.input_data) if report.input_data else {}
            output_data = json.loads(report.output_data) if report.output_data else {}

            # Create a response that matches the structure expected by the frontend
            result = {
                "id": str(report.id),
                "date_created": report.date_created,
                "fields": input_data.get("fields", ""),
                "file_names": input_data.get("files", []),
                "results": output_data,
                # Add feedback information
                "feedback": {
                    "feedback": report.feedback,
                    "feedbackText": report.feedback_text,
                    "feedbackDate": (
                        report.feedback_date.isoformat()
                        if report.feedback_date
                        else None
                    ),
                },
            }

            return result

        except json.JSONDecodeError:
            # Fallback if JSON parsing fails
            return {
                "id": str(report.id),
                "date_created": report.date_created,
                "results": {
                    "message": f"Unable to reconstruct form processing from {report.date_created}.\n\n"
                    f"This might be due to an older format or incomplete data."
                },
                # Add empty feedback object for consistency
                "feedback": {
                    "feedback": None,
                    "feedbackText": None,
                    "feedbackDate": None,
                },
            }

    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Error retrieving form processing details: {str(e)}",
        )


# Also add a history endpoint to get a list of past form processing operations
@router.get("/history", response_model=List[Dict[str, Any]])
async def get_form_history(
    session: SessionDep,
    current_user: CurrentUser,
    skip: int = Query(0, ge=0, le=10000),
    limit: int = Query(20, ge=1, le=100),
    show_all: bool = Query(False),
):
    """Retrieve past form processing history for the current user or all users."""
    print("Retrieving FormConnect history. Show all:", show_all)

    try:
        # Start with base query
        query = select(LlmInteraction).where(
            LlmInteraction.functionality == "formconnect"
        )

        # Only filter by user if not showing all users
        if not show_all:
            query = query.where(LlmInteraction.user_id == current_user.id)

        # Add ordering and pagination
        interactions = session.exec(
            query.order_by(LlmInteraction.date_created.desc()).offset(skip).limit(limit)
        ).all()

        result = []
        for interaction in interactions:
            # Parse the input_data and output_data
            try:
                input_data = (
                    json.loads(interaction.input_data) if interaction.input_data else {}
                )
                output_data = (
                    json.loads(interaction.output_data)
                    if interaction.output_data
                    else {}
                )
                # Fix: Use extra_data instead of metadata
                metadata = interaction.extra_data if interaction.extra_data else {}

                file_count = len(input_data.get("files", []))
                fields = input_data.get("fields", "").split("\n")
                field_count = len([f for f in fields if f.strip()])

                # Use stored form_name if available, otherwise create from first field
                stored_form_name = input_data.get("form_name")
                if stored_form_name:
                    form_name = stored_form_name
                else:
                    # Fallback: Create a display name from the first field
                    form_name = "Unnamed Form"
                    if fields and len(fields) > 0:
                        first_field = fields[0].strip()
                        if first_field:
                            # Truncate if too long
                            form_name = first_field[:50] + (
                                "..." if len(first_field) > 50 else ""
                            )
                        else:
                            form_name = "Custom Form"

                # Create result item
                result_item = {
                    "id": str(interaction.id),
                    "date_created": interaction.date_created,
                    "file_names": input_data.get("files", []),
                    "file_count": file_count,
                    "field_count": field_count,
                    "fields": fields,
                    "form_name": form_name,
                    "has_feedback": interaction.feedback is not None,
                    # Add metadata information for enhanced display
                    "metadata": metadata,
                    "files": metadata.get("files", metadata.get("digitized_files", [])),
                    "document_count": metadata.get("document_count", file_count),
                    "search_mode": metadata.get("search_mode", "unknown"),
                }

                # Add feedback information if exists
                if interaction.feedback:
                    result_item["feedback"] = {
                        "feedback": interaction.feedback,
                        "feedbackText": interaction.feedback_text,
                    }

                # Add user info for all-users view
                if show_all:
                    from app.models import User  # Import here to avoid circular imports

                    user = session.get(User, interaction.user_id)
                    user_name = (
                        f"{user.full_name or 'User'} ({user.email})"
                        if user
                        else "Unknown User"
                    )
                    result_item["user_name"] = user_name

                result.append(result_item)
            except json.JSONDecodeError:
                # If JSON parsing fails, use minimal information
                # Create result item with minimal info
                result_item = {
                    "id": str(interaction.id),
                    "date_created": interaction.date_created,
                    "file_names": [],
                    "file_count": 0,
                    "field_count": 0,
                    "fields": [],
                    "has_feedback": interaction.feedback is not None,
                    # Add empty metadata for consistency
                    "metadata": {},
                    "files": [],
                    "document_count": 0,
                    "search_mode": "unknown",
                }

                # Add user info for all-users view
                if show_all:
                    from app.models import User  # Import here to avoid circular imports

                    user = session.get(User, interaction.user_id)
                    user_name = (
                        f"{user.full_name or 'User'} ({user.email})"
                        if user
                        else "Unknown User"
                    )
                    result_item["user_name"] = user_name

                result.append(result_item)

        return result
    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Error retrieving form processing history: {str(e)}",
        )


@router.post("/generate-fields", response_model=GenerateFormFieldsResponse)
async def generate_form_fields(
    session: SessionDep, current_user: CurrentUser, request: GenerateFormFieldsRequest
):
    """
    Generate form fields based on a description with optional knowledge base reference.
    """
    try:
        # Get the default LLM
        llm = get_default_llm(session, current_user)

        # Handle optional description
        description = request.description or ""

        # Prepare variables for the prompt
        # Get user language and create language instruction
        user_language = getattr(current_user, "preferred_language", "en") or "en"
        language_name = settings.SUPPORTED_LANGUAGES.get(user_language, "English")
        language_instruction = f"Respond in this language: {language_name}."

        prompt_variables = {
            "description": description,
            "example_instruction": "",
            "analysis_instruction": "",
            "analysis_note": "",
            "knowledge_base_instruction": "",
            "knowledge_base_content": "",
            "language_instruction": language_instruction,
        }

        # If knowledge base is specified, retrieve content using selected search mode
        if request.knowledge_base_id:
            try:
                from app.services.content_retrieval import (
                    retrieve_knowledge_base_content,
                )

                content, instruction = await retrieve_knowledge_base_content(
                    session=session,
                    current_user=current_user,
                    knowledge_base_id=str(request.knowledge_base_id),
                    search_mode=request.search_mode,
                    query=description,
                )

                if content:
                    prompt_variables["knowledge_base_content"] = (
                        f"REFERENCE DOCUMENTS FROM KNOWLEDGE BASE:\n{content}"
                    )
                    prompt_variables["knowledge_base_instruction"] = (
                        f"\n11. {instruction} Use them as examples to understand the types of fields "
                        f"that are typically found in similar documents. Search mode used: {request.search_mode}"
                    )
                    prompt_variables["analysis_instruction"] = (
                        f". Briefly mention how the knowledge base content (using {request.search_mode}) influenced the field selection"
                    )

            except Exception as e:
                print(f"Error retrieving from knowledge base: {str(e)}")
                # Continue without knowledge base content rather than failing
                pass

        # Generate fields using the LLM
        fields_response = invoke_llm(
            llm,
            settings.FORMCONNECT_GENERATE_FIELDS_PROMPT_TEMPLATE,
            prompt_variables,
        )

        # Parse the response to extract fields and analysis
        fields = []
        analysis = ""

        lines = fields_response.strip().split("\n")
        in_fields_section = False
        in_analysis_section = False

        for line in lines:
            line = line.strip()
            if line.startswith("FIELDS:"):
                in_fields_section = True
                in_analysis_section = False
                continue
            elif line.startswith("ANALYSIS:"):
                in_fields_section = False
                in_analysis_section = True
                continue

            if in_fields_section:
                # Extract fields (numbered list)
                if re.match(r"^\d+\.\s+", line):
                    field = re.sub(r"^\d+\.\s+", "", line)
                    if field.strip():
                        fields.append(field.strip())
            elif in_analysis_section:
                if line:
                    if analysis:
                        analysis += " " + line
                    else:
                        analysis = line

        # If parsing failed, try simpler approach
        if not fields:
            # Split by lines and look for numbered items
            for line in lines:
                line = line.strip()
                if re.match(r"^\d+\.\s+", line):
                    field = re.sub(r"^\d+\.\s+", "", line)
                    if field.strip():
                        fields.append(field.strip())

        # Ensure we have some fields
        if not fields:
            raise HTTPException(
                status_code=500,
                detail="Failed to generate fields from the description. Please try with a more detailed description.",
            )

        # Apply user-specified limit if provided, otherwise use all generated fields
        if request.num_fields:
            fields = fields[: request.num_fields]

        if not analysis:
            search_method = (
                "vector search"
                if request.search_mode == "vector"
                else "full document scan"
            )
            analysis = f"Generated {len(fields)} form fields based on the provided description using {search_method}"
            if request.knowledge_base_id:
                analysis += " with knowledge base reference."

        # Translate the analysis if needed
        # translated_analysis = await translate_text_if_needed(
        #     analysis, session, current_user, llm
        # )
        translated_analysis = analysis

        # Record the interaction
        record_llm_interaction(
            session=session,
            user_id=current_user.id,
            functionality="generate_form_fields",
            input_data={
                "description": request.description,
                "requested_fields": request.num_fields,
                "knowledge_base_id": (
                    str(request.knowledge_base_id)
                    if request.knowledge_base_id
                    else None
                ),
                "search_mode": request.search_mode,
            },
            output_data={
                "fields_count": len(fields),
                "analysis": translated_analysis,
            },
            metadata={},
        )

        return GenerateFormFieldsResponse(
            fields=fields, description_analysis=translated_analysis
        )

    except Exception as e:
        print(f"Error generating form fields: {e}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error generating form fields: {str(e)}"
        )


@router.post("/generate-fields-json", response_model=GenerateFormFieldsResponse)
async def generate_form_fields_json(
    session: SessionDep, current_user: CurrentUser, request: GenerateFormFieldsRequest
):
    """
    Generate form fields based on a description with optional knowledge base reference (JSON version).
    """
    # This is the same as generate_form_fields but ensures JSON request/response
    return await generate_form_fields(session, current_user, request)


@router.post("/generate-fields-with-files", response_model=GenerateFormFieldsResponse)
async def generate_form_fields_with_files(
    session: SessionDep,
    current_user: CurrentUser,
    description: str = Form(...),
    num_fields: int = Form(15),
    search_mode: Literal["vector", "full_scan"] = Form("vector"),
    files: List[UploadFile] = File(...),
):
    """
    Generate form fields based on a description and uploaded reference documents.
    """
    try:
        if not files or len(files) == 0:
            raise HTTPException(
                status_code=400, detail="At least one reference file must be uploaded."
            )

        # Get the default LLM
        llm = get_default_llm(session, current_user)

        # Extract text from all uploaded files
        extracted_documents = []
        file_names = []

        print(f"🔍 Processing {len(files)} reference files for field suggestion...")

        for i, file in enumerate(files):
            # Add delay between reference file processing to prevent rate limit exhaustion
            if i > 0:
                await asyncio.sleep(settings.PROCESSING_DELAY_BETWEEN_DOCUMENTS)

            try:
                # Read file content
                content = await file.read()

                # Extract text using enhanced processing with vision capabilities
                from app.services.document_utils import (
                    extract_text_with_vision_enhancement,
                )

                text = await extract_text_with_vision_enhancement(
                    content,
                    file.filename or "unknown",
                    llm,
                    purpose="form field generation",
                    current_user=current_user,
                )

                if text.strip():
                    extracted_documents.append(
                        {"filename": file.filename, "content": text}
                    )
                    file_names.append(file.filename)
                    print(
                        f"✅ Successfully processed {file.filename} (with vision enhancement)"
                    )
                else:
                    print(f"⚠️ No content extracted from {file.filename}")

            except Exception as e:
                print(f"❌ Error processing {file.filename}: {str(e)}")
                # Continue with other files instead of failing completely
                continue

        if not extracted_documents:
            raise HTTPException(
                status_code=400,
                detail="No text could be extracted from any of the uploaded files. Please check your files and try again.",
            )

        # Combine all document content for analysis
        combined_content = ""
        for doc in extracted_documents:
            combined_content += f"\n\n--- DOCUMENT: {doc['filename']} ---\n"
            combined_content += doc["content"]

        # Prepare variables for the prompt
        # Get user language and create language instruction
        user_language = getattr(current_user, "preferred_language", "en") or "en"
        language_name = settings.SUPPORTED_LANGUAGES.get(user_language, "English")
        language_instruction = f"Respond in this language: {language_name}."

        prompt_variables = {
            "description": (
                description.strip()
                if description
                else "Form template based on reference documents"
            ),
            "example_instruction": f"\n12. Use the following {len(extracted_documents)} reference document(s) as examples to understand the types of fields that are typically found in similar documents:",
            "analysis_instruction": f". Briefly mention how the {len(extracted_documents)} reference document(s) influenced the field selection using {search_mode} analysis",
            "analysis_note": f" (analyzed {len(extracted_documents)} reference documents)",
            "knowledge_base_instruction": "",
            "knowledge_base_content": f"REFERENCE DOCUMENTS:\n{combined_content}",
            "language_instruction": language_instruction,
        }

        # Generate fields using the LLM
        fields_response = invoke_llm(
            llm,
            settings.FORMCONNECT_GENERATE_FIELDS_PROMPT_TEMPLATE,
            prompt_variables,
        )

        # Parse the response to extract fields and analysis
        fields = []
        analysis = ""

        lines = fields_response.strip().split("\n")
        in_fields_section = False
        in_analysis_section = False

        for line in lines:
            line = line.strip()
            if line.startswith("FIELDS:"):
                in_fields_section = True
                in_analysis_section = False
                continue
            elif line.startswith("ANALYSIS:"):
                in_fields_section = False
                in_analysis_section = True
                continue

            if in_fields_section:
                # Extract fields (numbered list)
                if re.match(r"^\d+\.\s+", line):
                    field = re.sub(r"^\d+\.\s+", "", line)
                    if field.strip():
                        fields.append(field.strip())
            elif in_analysis_section:
                if line:
                    if analysis:
                        analysis += " " + line
                    else:
                        analysis = line

        # If parsing failed, try simpler approach
        if not fields:
            # Split by lines and look for numbered items
            for line in lines:
                line = line.strip()
                if re.match(r"^\d+\.\s+", line):
                    field = re.sub(r"^\d+\.\s+", "", line)
                    if field.strip():
                        fields.append(field.strip())

        # Ensure we have some fields
        if not fields:
            raise HTTPException(
                status_code=500,
                detail="Failed to generate fields from the reference documents. Please try with different documents or add a more detailed description.",
            )

        # Apply user-specified limit if provided
        if num_fields and num_fields > 0:
            fields = fields[:num_fields]

        if not analysis:
            analysis = f"Generated {len(fields)} form fields based on analysis of {len(extracted_documents)} reference document(s) using {search_mode} method"

        # Translate the analysis if needed
        # translated_analysis = await translate_text_if_needed(
        #     analysis, session, current_user, llm
        # )
        translated_analysis = analysis

        # Record the interaction
        record_llm_interaction(
            session=session,
            user_id=current_user.id,
            functionality="generate_form_fields_with_files",
            input_data={
                "description": description,
                "requested_fields": num_fields,
                "file_count": len(extracted_documents),
                "file_names": file_names,
                "search_mode": search_mode,
            },
            output_data={
                "fields_count": len(fields),
                "analysis": translated_analysis,
            },
            metadata={
                "processed_files": len(extracted_documents),
                "total_files": len(files),
            },
        )

        return GenerateFormFieldsResponse(
            fields=fields, description_analysis=translated_analysis
        )

    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        print(f"Error generating form fields with files: {e}")
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Error generating form fields: {str(e)}"
        )


@router.post("/generate/docx", response_class=StreamingResponse)
async def generate_docx(
    session: SessionDep, current_user: CurrentUser, request: DocxRequest
):
    """
    Generate a DOCX file from the FormConnect results content.
    Handles markdown tables with extra care for proper rendering.
    """
    print("Now generating DOCX of FormConnect results...")
    try:
        # Get the markdown content from the request
        if not request.content:
            raise HTTPException(
                status_code=400, detail="FormConnect content is required"
            )

        # Convert markdown to HTML for parsing with tables extension
        html_content = markdown.markdown(
            request.content, extensions=["tables", "extra"]
        )
        soup = BeautifulSoup(html_content, "html.parser")

        print("Markdown content converted to HTML successfully.")
        # Create a new Document
        doc = Document()

        print("Adding title and date to the document...")
        # Determine language
        language = (
            request.language
            or getattr(current_user, "preferred_language", "en")
            or "en"
        )

        # Add a title
        title_text = (
            request.title if request.title else translate("matching_results", language)
        )
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add translated subtitle with metadata
        date_str = datetime.now().strftime("%B %d, %Y at %H:%M")
        subtitle_template = translate("generated_on", language)
        subtitle = subtitle_template.format(
            date=date_str,
            name=current_user.full_name or current_user.email,
            email=current_user.email,
        )
        subtitle_paragraph = doc.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.italic = True

        # Add a separator
        doc.add_paragraph("─" * 50)

        print("Processing content elements with special attention to tables...")
        # Process each element in the soup
        for element in soup.find_all():
            if element.name == "h1":
                doc.add_heading(element.get_text().strip(), level=1)
            elif element.name == "h2":
                doc.add_heading(element.get_text().strip(), level=2)
            elif element.name == "h3":
                doc.add_heading(element.get_text().strip(), level=3)
            elif element.name == "h4":
                doc.add_heading(element.get_text().strip(), level=4)
            elif element.name == "p":
                text = element.get_text().strip()
                if text:  # Only add non-empty paragraphs
                    paragraph = doc.add_paragraph(text)

            elif element.name == "table":
                # Handle tables with extra care for FormConnect markdown tables
                rows = element.find_all("tr")
                if rows:
                    print(f"Adding table with {len(rows)} rows...")

                    # Count maximum columns across all rows
                    max_cols = 0
                    for row in rows:
                        cells = row.find_all(["th", "td"])
                        max_cols = max(max_cols, len(cells))

                    if max_cols > 0:
                        table = doc.add_table(rows=len(rows), cols=max_cols)
                        table.style = "Table Grid"

                        # Set consistent column widths
                        for col in table.columns:
                            col.width = Inches(6.0 / max_cols)

                        for i, row in enumerate(rows):
                            cells = row.find_all(["th", "td"])
                            for j, cell in enumerate(cells):
                                if j < len(table.rows[i].cells):
                                    cell_text = cell.get_text().strip()
                                    table.rows[i].cells[j].text = cell_text

                                    # Make header row bold and centered
                                    if i == 0 or cell.name == "th":
                                        for paragraph in (
                                            table.rows[i].cells[j].paragraphs
                                        ):
                                            paragraph.alignment = (
                                                WD_ALIGN_PARAGRAPH.CENTER
                                            )
                                            for run in paragraph.runs:
                                                run.bold = True

                                    # Handle cell alignment for data rows
                                    elif (
                                        cell_text.isdigit()
                                        or cell_text.replace(".", "")
                                        .replace("-", "")
                                        .isdigit()
                                    ):
                                        # Right-align numeric content
                                        for paragraph in (
                                            table.rows[i].cells[j].paragraphs
                                        ):
                                            paragraph.alignment = (
                                                WD_ALIGN_PARAGRAPH.RIGHT
                                            )

            elif element.name == "ul":
                # Handle unordered lists
                for li in element.find_all("li", recursive=False):
                    text = li.get_text().strip()
                    if text:
                        paragraph = doc.add_paragraph(text, style="List Bullet")

            elif element.name == "ol":
                # Handle ordered lists
                for li in element.find_all("li", recursive=False):
                    text = li.get_text().strip()
                    if text:
                        paragraph = doc.add_paragraph(text, style="List Number")

        print("Saving the document to a BytesIO object...")
        # Save the document to a BytesIO object
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Create filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"formconnect_results_{timestamp}.docx"

        print(f"DOCX file size: {len(doc_io.getvalue())} bytes")

        # Verify the document can be opened (basic integrity check)
        doc_io.seek(0)
        try:
            test_doc = Document(doc_io)
            print("DOCX file passed integrity check (can be opened by python-docx).")
        except Exception as e:
            print(f"DOCX integrity check failed: {e}")
            raise HTTPException(
                status_code=500, detail=f"Generated DOCX file is corrupted: {str(e)}"
            )

        doc_io.seek(0)
        print(
            "Document saved successfully. Preparing to return as a downloadable file."
        )

        # Return the document as a downloadable file
        return StreamingResponse(
            doc_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")


@router.post("/generate/csv", response_class=StreamingResponse)
async def generate_csv(
    session: SessionDep, current_user: CurrentUser, request: DocxRequest
):
    """
    Generate a CSV file from the FormConnect results content using LLM formatting.
    """
    print("Now generating CSV of FormConnect results...")
    try:
        # Get the content from the request
        if not request.content:
            raise HTTPException(status_code=400, detail="Results content is required")

        # Get the default LLM for formatting
        llm = get_default_llm(session, current_user)

        # Try to parse the content as JSON first (for structured data)
        try:
            content_data = json.loads(request.content)
            extracted_data = content_data.get("extracted_data", [])
            comparison_data = content_data.get("comparison", "")
            message_data = content_data.get("message", "")
        except json.JSONDecodeError:
            # If it's not JSON, treat it as raw content and use LLM to format
            extracted_data = []
            comparison_data = request.content
            message_data = ""

        # Create LLM prompt to format the data as CSV
        csv_prompt_template = """You are tasked with converting FormConnect comparison results into a well-structured CSV format.

Input Data:
{input_data}

Instructions:
1. Analyze the extracted data and comparison results
2. Create a CSV table with the following structure:
   - First column: "Filename" (document names)
   - Subsequent columns: Field names from the form template
3. Each row should represent one document with its extracted field values
4. If a field value is missing for a document, leave the cell empty
5. Clean up any formatting issues and ensure values are CSV-safe
6. Return ONLY the CSV content, no additional text or formatting

Expected format:
Filename,Field1,Field2,Field3,...
Document1.pdf,Value1,Value2,Value3,...
Document2.pdf,Value1,Value2,Value3,...

Return the CSV content:"""

        # Prepare the input data for the LLM
        input_data = {
            "extracted_data": extracted_data,
            "comparison": comparison_data,
            "message": message_data,
        }

        input_data_str = json.dumps(input_data, indent=2)

        variables = {"input_data": input_data_str}

        # Use LLM to generate the CSV content
        print("Calling LLM to format FormConnect results as CSV...")
        csv_content = invoke_llm(llm, csv_prompt_template, variables)

        # Clean up the response - remove any markdown formatting or extra text
        csv_content = csv_content.strip()
        if csv_content.startswith("```csv"):
            csv_content = csv_content[6:]
        if csv_content.startswith("```"):
            csv_content = csv_content[3:]
        if csv_content.endswith("```"):
            csv_content = csv_content[:-3]
        csv_content = csv_content.strip()

        # Ensure we have valid CSV content
        if not csv_content or "Filename" not in csv_content:
            # Fallback to basic formatting if LLM didn't produce good results
            output = StringIO()
            writer = csv.writer(output)

            # Create basic headers
            headers = ["Filename", "Field", "Value"]
            writer.writerow(headers)

            # Add extracted data if available
            if extracted_data:
                for i, doc_data in enumerate(extracted_data):
                    filename = f"Document_{i+1}"
                    if isinstance(doc_data, dict):
                        for field_name, field_value in doc_data.items():
                            cleaned_value = (
                                str(field_value)
                                .replace("\n", " ")
                                .replace("\r", "")
                                .replace('"', '""')
                            )
                            writer.writerow([filename, field_name, cleaned_value])

            # Add comparison data if available
            if comparison_data:
                writer.writerow(
                    [
                        "Comparison Results",
                        "Analysis",
                        comparison_data.replace("\n", " ")
                        .replace("\r", "")
                        .replace('"', '""'),
                    ]
                )

            csv_content = output.getvalue()
            output.close()

        # Convert to bytes
        csv_bytes = csv_content.encode("utf-8")

        # Generate timestamp for filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"formconnect_results_{timestamp}.csv"

        print(
            "CSV file generated successfully using LLM formatting. Preparing to return as a downloadable file."
        )

        return StreamingResponse(
            BytesIO(csv_bytes),
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating CSV: {str(e)}")
