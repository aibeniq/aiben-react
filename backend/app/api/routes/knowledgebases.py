import uuid
from typing import Any, List, Optional

from fastapi import APIRouter, HTTPException, UploadFile, File, Depends
from sqlmodel import func, select, delete

import zipfile
import io
import os
import shutil
from io import BytesIO

from app.api.deps import CurrentUser, SessionDep
from app.models import (
    KnowledgeBase,
    KnowledgeBaseCreate,
    KnowledgeBasePublic,
    KnowledgeBasesPublic,
    KnowledgeBaseUpdate,
    Message,
    Source,
    SourceData,
    EmbeddingModel,
)
from app.services.embeddings import load_embeddings_model
from app.core.config import settings
import hashlib

from app.services.knowledgebases import KnowledgeBaseService

from sqlalchemy.sql import func

import tempfile

from datetime import datetime

from langchain_community.document_loaders import TextLoader, PyPDFLoader
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain.schema.document import Document
import mimetypes

router = APIRouter(prefix="/knowledge-bases", tags=["knowledge-bases"])


def load_correct_embeddings_model(
    session: SessionDep,
    current_user: Any,  # Pass the current user object here
    embedding_model_id: Optional[uuid.UUID] = None,
) -> Any:
    """
    Load the correct embeddings model based on the provided embedding_model_id or user's default model.

    Args:
        session (SessionDep): The database session.
        current_user: The current user object.
        embedding_model_id (Optional[uuid.UUID]): The ID of the embedding model to load.

    Returns:
        embeddings: The loaded embeddings model.
        model_id: The ID of the embedding model.
        provider: The provider of the embedding model.
    """
    if embedding_model_id:
        print("Using provided embedding model ID:", embedding_model_id)
        # Verify the embedding model exists
        model = session.get(EmbeddingModel, embedding_model_id)
        if not model:
            raise HTTPException(
                status_code=404,
                detail=f"Embedding model with ID {embedding_model_id} not found",
            )
        model_id = model.model_id
        provider = model.provider
    else:
        # Try to get the user's default embedding model
        if current_user and current_user.default_embedding_model:
            model = session.get(EmbeddingModel, current_user.default_embedding_model)
            if model:
                print("Using user's default embedding model:", model)
                model_id = model.model_id
                provider = model.provider
            else:
                model_id = None
                provider = None
        else:
            model_id = None
            provider = None

        # Fallback to system default if user has no default
        if not model_id or not provider:
            default_model = session.exec(
                select(EmbeddingModel).where(EmbeddingModel.owner_id.is_(None))
            ).first()
            print("System default embedding model:", default_model)
            if default_model:
                model_id = default_model.model_id
                provider = default_model.provider
            else:
                print("No default embedding model found, using hardcoded value.")
                model_id = "all-MiniLM-L6-v2"
                provider = "huggingface"

    # Initialize embeddings with the selected model
    try:
        embeddings = load_embeddings_model(provider=provider, model_id=model_id)
        print("Embeddings model loaded successfully.")
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Error initializing embedding model {model_id}: {str(e)}",
        )

    print(f"Using embedding model: {model_id}")
    return embeddings, model_id, provider


def extract_text_from_docx(file_path: str, filename: str) -> List[Any]:
    doc = docx.Document(file_path)

    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():  # Skip empty paragraphs
            full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                full_text.append(" | ".join(row_text))

    combined_text = "\n\n".join(full_text)

    metadata = {
        "source": filename,
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    # Try to get document properties
    try:
        core_properties = doc.core_properties
        if core_properties.title:
            metadata["title"] = core_properties.title
        if core_properties.author:
            metadata["author"] = core_properties.author
        if core_properties.created:
            metadata["created"] = str(core_properties.created)
        if core_properties.modified:
            metadata["modified"] = str(core_properties.modified)
    except Exception as e:
        print(f"Could not extract document properties: {str(e)}")

    # Create a Document object compatible with langchain
    return [Document(page_content=combined_text, metadata=metadata)]


def load_uploaded_file(file: UploadFile) -> List[Any]:
    """
    Load an uploaded file based on its type (e.g., PDF, text file).

    Args:
        file (UploadFile): The uploaded file to process.

    Returns:
        List[Any]: A list of loaded documents from the file.
    """
    print(f"Processing file: {file.filename}")
    content_type = file.content_type or mimetypes.guess_type(file.filename)[0]
    print(f"Detected content type: {content_type}")

    with tempfile.NamedTemporaryFile(
        delete=False, suffix=f"_{file.filename}"
    ) as temp_file:
        temp_file.write(
            file.file.read()
        )  # Write the file content to the temporary file
        temp_file_path = temp_file.name

    try:
        if content_type == "application/pdf" or file.filename.lower().endswith(".pdf"):
            print("Loading PDF with PyPDFLoader...")
            loader = PyPDFLoader(temp_file_path)
            loaded_documents = loader.load()
        elif (
            content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or file.filename.lower().endswith(".docx")
        ):
            print("Loading DOCX with python-docx library...")
            loaded_documents = extract_text_from_docx(temp_file_path, file.filename)
        else:
            print("Loading text with TextLoader...")
            # Try with different encodings if utf-8 fails
            try:
                loader = TextLoader(temp_file_path, encoding="utf-8")
                loaded_documents = loader.load()
            except UnicodeDecodeError:
                print("UTF-8 decoding failed. Retrying with Latin-1 encoding...")
                loader = TextLoader(temp_file_path, encoding="latin-1")
                loaded_documents = loader.load()

        return loaded_documents
    except Exception as e:
        print(f"Error processing file {file.filename}: {str(e)}")
        raise HTTPException(
            status_code=400, detail=f"Error processing file {file.filename}: {str(e)}"
        )
    finally:
        # Clean up the temporary file
        os.unlink(temp_file_path)


@router.get("/", response_model=KnowledgeBasesPublic)
def read_knowledge_bases(
    session: SessionDep, current_user: CurrentUser, skip: int = 0, limit: int = 100
) -> Any:
    """
    Retrieve knowledge bases with additional metadata: Number of Sources, Date Created, and Date Modified.
    """
    # Base query to count sources and retrieve metadata
    query = (
        session.query(
            KnowledgeBase,
            func.count(Source.id).label("number_of_sources"),
            KnowledgeBase.date_created,
            KnowledgeBase.date_modified,
        )
        .join(Source, Source.knowledge_base_id == KnowledgeBase.id, isouter=True)
        .group_by(KnowledgeBase.id)
    )

    # Apply filters based on user permissions
    if current_user.is_superuser:
        count_statement = select(func.count()).select_from(KnowledgeBase)
        count = session.exec(count_statement).one()
        query = query.offset(skip).limit(limit)
    else:
        count_statement = (
            select(func.count())
            .select_from(KnowledgeBase)
            .where(KnowledgeBase.owner_id == current_user.id)
        )
        count = session.exec(count_statement).one()
        query = (
            query.filter(KnowledgeBase.owner_id == current_user.id)
            .offset(skip)
            .limit(limit)
        )

    # Execute the query
    results = query.all()

    # Format the response
    knowledge_bases = []
    for kb in results:
        # Get embedding model name if it exists
        embedding_model_name = None
        if kb.KnowledgeBase.embedding_model_id:
            model = session.get(EmbeddingModel, kb.KnowledgeBase.embedding_model_id)
            if model:
                embedding_model_name = model.name

        kb_public = KnowledgeBasePublic(
            id=kb.KnowledgeBase.id,
            owner_id=kb.KnowledgeBase.owner_id,
            title=kb.KnowledgeBase.title,
            description=kb.KnowledgeBase.description,
            files=[],  # Files can be populated separately if needed
            number_of_sources=kb.number_of_sources,
            date_created=kb.date_created,
            date_modified=kb.date_modified,
            embedding_model_id=kb.KnowledgeBase.embedding_model_id,
            embedding_model_name=embedding_model_name,
        )
        knowledge_bases.append(kb_public)

    print("Knowledge Bases Response:", knowledge_bases)

    return KnowledgeBasesPublic(data=knowledge_bases, count=count)


@router.get("/{id}", response_model=KnowledgeBasePublic)
def read_knowledge_base(
    session: SessionDep, current_user: CurrentUser, id: uuid.UUID
) -> Any:
    """
    Get knowledge base by ID.
    """
    knowledge_base = session.get(KnowledgeBase, id)
    if not knowledge_base:
        raise HTTPException(status_code=404, detail="Knowledge base not found")

    # Get embedding model name if it exists
    embedding_model_name = None
    if knowledge_base.embedding_model_id:
        model = session.get(EmbeddingModel, knowledge_base.embedding_model_id)
        if model:
            embedding_model_name = model.name

    # Get all sources for this knowledge base
    sources = session.exec(select(Source).where(Source.knowledge_base_id == id)).all()

    files = []
    for source in sources:
        # Only include metadata, not the actual file content
        files.append(
            {
                "id": str(source.source_data_id),
                "name": source.name,
                "date_created": source.date_created,
                # Don't include data_base64 here
            }
        )

    # Construct the response model
    knowledge_base_public = KnowledgeBasePublic(
        **knowledge_base.model_dump(),
        files=files,
        embedding_model_name=embedding_model_name,
    )
    return knowledge_base_public


@router.post("/", response_model=KnowledgeBasePublic)
def create_knowledge_base(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    knowledge_base_in: KnowledgeBaseCreate = Depends(),
    files: List[UploadFile] = File(...),
) -> Any:
    """
    Create new knowledge base with a compressed folder with the Chroma VectorDB.
    """

    print("Received the following metadata for the knowledge base:")
    print(knowledge_base_in)

    # Check if a knowledge base with this title already exists for this user
    existing_kb = session.exec(
        select(KnowledgeBase).where(
            KnowledgeBase.title == knowledge_base_in.title,
            KnowledgeBase.owner_id == current_user.id,
        )
    ).first()

    print("Checking for existing knowledge base")

    if existing_kb:
        raise HTTPException(
            status_code=409,  # Using 409 Conflict for duplicate resource
            detail=f"A knowledge base with the title '{knowledge_base_in.title}' already exists",
        )

    # Initialize variables for Chroma
    documents = []

    # Process each uploaded file
    for file in files:
        loaded_documents = load_uploaded_file(file)
        documents.extend(loaded_documents)

        # Reset the file pointer before passing to create_source_entries
        file.file.seek(0)

    # Clean up temporary files
    for root, dirs, files_in_dir in os.walk(tempfile.gettempdir()):
        for filename in files_in_dir:
            if any(uploaded_file.filename in filename for uploaded_file in files):
                try:
                    os.unlink(os.path.join(root, filename))
                except:
                    pass

    print("Splitting documents...")

    # Split documents into chunks using RecursiveCharacterTextSplitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.DOCUMENT_CHUNK_SIZE,
        chunk_overlap=settings.DOCUMENT_CHUNK_OVERLAP,
    )
    splits = text_splitter.split_documents(documents)

    print("Initializing embeddings...")

    embeddings, model_id, provider = load_correct_embeddings_model(
        session=session,
        current_user=current_user,  # Pass the current user to load the correct model
        embedding_model_id=knowledge_base_in.embedding_model_id,
    )

    print(f"Using embedding model: {model_id}")

    # Clear out any existing chroma_db directory
    chroma_dir = tempfile.mkdtemp()
    try:
        # Create Chroma VectorDB from the document splits
        Chroma.from_documents(
            documents=splits, embedding=embeddings, persist_directory=chroma_dir
        )
    except Exception as e:
        print(f"Error creating Chroma VectorDB: {str(e)}")
        # Clean up the directory on error
        if os.path.exists(chroma_dir):
            shutil.rmtree(chroma_dir)
        raise HTTPException(
            status_code=500, detail=f"Error creating vector database: {str(e)}"
        )

    print("Zipping Chroma database...")

    # Compress the Chroma database directory into a zip file
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for root, _, filenames in os.walk(chroma_dir):
            for filename in filenames:
                file_path = os.path.join(root, filename)
                arcname = os.path.relpath(
                    file_path, chroma_dir
                )  # Preserve directory structure
                zip_file.write(file_path, arcname)
    zip_buffer.seek(0)

    print("Validating knowledge base...")

    # Use model_validate to create and validate the knowledge base
    knowledge_base = KnowledgeBase.model_validate(
        knowledge_base_in,
        update={
            "owner_id": current_user.id,
            "data": zip_buffer.read(),
            "embedding_model_id": knowledge_base_in.embedding_model_id,
            "date_created": datetime.utcnow(),
            "date_modified": datetime.utcnow(),
        },
    )

    print("Adding knowledge base to session...")

    session.add(knowledge_base)

    session.flush()  # This ensures the knowledge_base.id is available

    # Process each file
    for file in files:
        print(f"Received file: {file}")
        print(f"Type of file: {type(file)}")

        KnowledgeBaseService.create_source_entries(
            session=session,
            current_user=current_user,
            knowledge_base_id=knowledge_base.id,
            file=file,
        )

    session.commit()
    session.refresh(knowledge_base)
    return knowledge_base


@router.put("/{id}", response_model=KnowledgeBasePublic)
def update_knowledge_base(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    id: uuid.UUID,
    knowledge_base_in: KnowledgeBaseUpdate = Depends(),
    files: Optional[List[UploadFile]] = None,
) -> Any:
    """
    Update a knowledge base.
    """
    print("Now updating knowledge base...")
    if knowledge_base_in.removed_file_ids is None:
        knowledge_base_in.removed_file_ids = []
    print("Source IDs to remove:", knowledge_base_in.removed_file_ids)
    knowledge_base = session.get(KnowledgeBase, id)
    if not knowledge_base:
        raise HTTPException(status_code=404, detail="Knowledge base not found")
    if not current_user.is_superuser and (knowledge_base.owner_id != current_user.id):
        raise HTTPException(status_code=400, detail="Not enough permissions")

    update_dict = knowledge_base_in.model_dump(exclude_unset=True)
    knowledge_base.sqlmodel_update(update_dict)

    # Retrieve the compressed folder from the database
    if files or knowledge_base_in.removed_file_ids:
        print("Retrieving existing VectorDB...")

        # Clear out any existing chroma_db directory
        chroma_dir = tempfile.mkdtemp()
        try:
            # Extract the zipped ChromaDB into the temp directory
            if knowledge_base.data:
                with zipfile.ZipFile(io.BytesIO(knowledge_base.data), "r") as zip_ref:
                    zip_ref.extractall(chroma_dir)
            else:
                raise HTTPException(
                    status_code=400, detail="Knowledge base has no vector database data"
                )

            # Load the vector database with the SAME model used to create the knowledge base
            embeddings, model_id, provider = load_correct_embeddings_model(
                session=session,
                current_user=current_user,  # Pass the current
                embedding_model_id=knowledge_base_in.embedding_model_id,
            )

            # Initialize Chroma with the existing database
            chroma_vector_database = Chroma(
                persist_directory=chroma_dir, embedding_function=embeddings
            )

            # Process new files (if any)
            if files:
                print("Adding new files to VectorDB...")
                documents = []
                for file in files:
                    loaded_documents = load_uploaded_file(file)
                    documents.extend(loaded_documents)

                    # Reset the file pointer before passing to create_source_entries
                    file.file.seek(0)

                    # Add source entries for the new files
                    KnowledgeBaseService.create_source_entries(
                        session=session,
                        current_user=current_user,
                        knowledge_base_id=knowledge_base.id,
                        file=file,
                    )

                # Add the new documents to the VectorDB only if we have new files
                if documents:
                    chroma_vector_database.add_documents(documents)

            # Handle file deletions
            if knowledge_base_in.removed_file_ids:
                print("Removing files from VectorDB...")
                sources_to_remove = session.exec(
                    select(Source).where(
                        Source.source_data_id.in_(knowledge_base_in.removed_file_ids)
                    )
                ).all()

                print("Sources to remove:", sources_to_remove)

                for source in sources_to_remove:
                    chroma_vector_database.delete(ids=[str(source.source_data_id)])

                # 1. First, only delete Source entries for THIS specific knowledge base
                sources_to_delete = session.exec(
                    select(Source).where(
                        (Source.source_data_id.in_(knowledge_base_in.removed_file_ids))
                        & (Source.knowledge_base_id == id)
                    )
                ).all()

                # 2. Remember which source_data_ids we're removing
                source_data_ids_to_check = [
                    source.source_data_id for source in sources_to_delete
                ]

                # 3. Delete the specific Source entries (just the associations)
                session.exec(
                    delete(Source).where(
                        (Source.source_data_id.in_(knowledge_base_in.removed_file_ids))
                        & (Source.knowledge_base_id == id)
                    )
                )

                # 4. For each source_data_id, check if it's still referenced by any Source
                for source_data_id in source_data_ids_to_check:
                    # Count how many Sources still reference this source_data_id
                    remaining_references = session.exec(
                        select(func.count()).where(
                            Source.source_data_id == source_data_id
                        )
                    ).one()

                    # If no Sources reference this source_data_id anymore, delete the SourceData
                    if remaining_references == 0:
                        session.exec(
                            delete(SourceData).where(SourceData.id == source_data_id)
                        )

            # Zip the updated VectorDB
            print("Zipping updated VectorDB...")
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                for root, _, filenames in os.walk(chroma_dir):
                    for filename in filenames:
                        file_path = os.path.join(root, filename)
                        arcname = os.path.relpath(file_path, chroma_dir)
                        zip_file.write(file_path, arcname)
            zip_buffer.seek(0)

            # Update the knowledge base data in the database
            print("Updating knowledge base data in the database...")
            knowledge_base.data = zip_buffer.read()

        except Exception as e:
            # Clean up on error
            if os.path.exists(chroma_dir):
                shutil.rmtree(chroma_dir)
            raise HTTPException(
                status_code=500, detail=f"Error updating vector database: {str(e)}"
            )
        finally:
            # Clean up regardless of success or failure
            if os.path.exists(chroma_dir):
                print(f"Cleaning up {chroma_dir} directory...")
                shutil.rmtree(chroma_dir)

    # Update the date_modified field
    knowledge_base.date_modified = datetime.utcnow()

    session.add(knowledge_base)

    session.commit()
    session.refresh(knowledge_base)
    return knowledge_base


@router.delete("/{id}", response_model=Message)
def delete_knowledge_base(
    session: SessionDep, current_user: CurrentUser, id: uuid.UUID
) -> Message:
    """
    Delete a knowledge base.
    """
    knowledge_base = session.get(KnowledgeBase, id)
    if not knowledge_base:
        raise HTTPException(status_code=404, detail="Knowledge base not found")
    if not current_user.is_superuser and (knowledge_base.owner_id != current_user.id):
        raise HTTPException(status_code=400, detail="Not enough permissions")
    session.delete(knowledge_base)
    session.commit()
    return Message(message="Knowledge base deleted successfully")
