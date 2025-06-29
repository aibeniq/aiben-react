#!/usr/bin/env python3
"""
Test script to verify .docx text extraction is working correctly.
"""

import sys
import os
from pathlib import Path

# Add the backend app to the Python path
backend_path = Path(__file__).parent / "backend"
sys.path.insert(0, str(backend_path))

try:
    from app.api.routes.veradoc import extract_text_from_file
    from docx import Document
    import tempfile

    def test_docx_creation_and_extraction():
        """Create a test .docx file and extract text from it."""
        print("Testing .docx creation and text extraction...")

        # Create a test .docx file
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test paragraph in a Word document.")
        doc.add_paragraph("This is another paragraph with some text.")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1,1"
        table.cell(0, 1).text = "Cell 1,2"
        table.cell(1, 0).text = "Cell 2,1"
        table.cell(1, 1).text = "Cell 2,2"

        # Save to a temporary file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            doc.save(temp_file.name)
            temp_file_path = temp_file.name

        try:
            # Read the file content
            with open(temp_file_path, "rb") as f:
                file_content = f.read()

            # Extract text using our function
            extracted_text = extract_text_from_file(file_content, "test_document.docx")

            print("Extracted text:")
            print("-" * 50)
            print(extracted_text)
            print("-" * 50)

            # Check if the extraction was successful
            expected_texts = [
                "Test Document",
                "This is a test paragraph",
                "This is another paragraph",
                "Cell 1,1",
                "Cell 2,2",
            ]

            success = True
            for expected in expected_texts:
                if expected not in extracted_text:
                    print(
                        f"ERROR: Expected text '{expected}' not found in extracted text!"
                    )
                    success = False
                else:
                    print(f"✓ Found expected text: '{expected}'")

            if success:
                print("\n✅ .docx text extraction test PASSED!")
            else:
                print("\n❌ .docx text extraction test FAILED!")

            return success

        finally:
            # Clean up the temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)

    if __name__ == "__main__":
        test_docx_creation_and_extraction()

except ImportError as e:
    print(f"Import error: {e}")
    print(
        "Make sure you're running this from the correct directory and all dependencies are installed."
    )
except Exception as e:
    print(f"Error: {e}")
    import traceback

    traceback.print_exc()
